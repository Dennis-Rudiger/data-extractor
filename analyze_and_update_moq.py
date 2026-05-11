import json
import pandas as pd
from docx import Document

def format_detailed_json(df_all, branch_name):
    # Load supporting data for categories and prices if available
    try:
        with open('item_category_map.json', 'r') as f:
            cat_map = json.load(f)
    except:
        cat_map = {}
        
    try:
        with open('inventory_valuation.json', 'r') as f:
            val_data = json.load(f)
            price_map = {}
            for c in val_data:
                for p in c.get('products', []):
                    price_map[p.get('item_code', '')] = p.get('buying_price', p.get('unit_cost', 0))
    except:
        price_map = {}

    working_days = 96 # Roughly Jan-Apr
    df_all['category'] = df_all['code'].map(lambda x: cat_map.get(x, 'UNKNOWN'))
    df_all['buying_price'] = df_all['code'].map(lambda x: price_map.get(x, 0))
    df_all['weekly_value'] = df_all['new_moq'] * df_all['buying_price']
    
    # Calculate Summary
    total_items = len(df_all)
    fast_movers = len(df_all[df_all['qty_out_jan_apr'] > 0])
    slow_movers = total_items - fast_movers
    total_qty_out = df_all['qty_out_jan_apr'].sum()
    total_moq = df_all['new_moq'].sum()
    total_value = df_all['weekly_value'].sum()
    
    output = {
        "period": "January - April, 2026",
        "working_days": working_days,
        "moq_multiplier": 6,
        "formula": "MOQ = 6 x Daily Consumption (Based on Peak Q1 vs April trend)",
        "summary": {
            "total_items": int(total_items),
            "fast_movers": int(fast_movers),
            "slow_movers": int(slow_movers),
            "total_qty_out": round(float(total_qty_out), 2),
            "total_weekly_moq": round(float(total_moq), 2),
            "total_weekly_value": round(float(total_value), 2)
        },
        "categories": {}
    }
    
    grouped = df_all.groupby('category')
    for cat_name, group in grouped:
        if pd.isna(cat_name) or not cat_name:
            cat_name = 'UNKNOWN'
            
        cat_fast = len(group[group['qty_out_jan_apr'] > 0])
        
        products = []
        for _, row in group.iterrows():
            products.append({
                "item_code": str(row['code']),
                "item_description": str(row['description']),
                "qty_in": 0.0,
                "qty_out": round(float(row['qty_out_jan_apr']), 2),
                "closing_balance": 0.0,
                "daily_average": round(float(row[['avg_monthly_q1', 'april_qty']].max() / 24), 2),
                "weekly_moq": round(float(row['new_moq']), 2),
                "buying_price": round(float(row['buying_price']), 2),
                "weekly_value": round(float(row['weekly_value']), 2),
                "category": str(cat_name)
            })
            
        output["categories"][str(cat_name)] = {
            "total_items": len(group),
            "fast_movers": cat_fast,
            "slow_movers": len(group) - cat_fast,
            "total_qty_out": round(float(group['qty_out_jan_apr'].sum()), 2),
            "total_weekly_moq": round(float(group['new_moq'].sum()), 2),
            "total_weekly_value": round(float(group['weekly_value'].sum()), 2),
            "products": products
        }
        
    return output

def analyze_branch(branch_name, q1_file, apr_file):
    df_q1 = pd.read_csv(q1_file)
    df_apr = pd.read_csv(apr_file)
    
    # Calculate Q1 (Jan-Mar) by subtracting April from Jan-Apr (assuming the file Jan-Apr covers all 4 months, thus Q1 = (Jan_Apr - Apr))
    df_merged = pd.merge(df_q1, df_apr, on='code', how='outer', suffixes=('_jan_apr', '_apr'))
    df_merged.fillna(0, inplace=True)
    df_merged['description'] = df_merged['description_jan_apr'].combine_first(df_merged['description_apr'])
    
    # Q1 is Jan-Mar, so Total - April
    df_merged['qty_q1'] = df_merged['qty_out_jan_apr'] - df_merged['qty_out_apr']
    df_merged['qty_q1'] = df_merged['qty_q1'].clip(lower=0) # ensure no negative
    
    # Averages
    df_merged['avg_monthly_q1'] = df_merged['qty_q1'] / 3.0
    df_merged['april_qty'] = df_merged['qty_out_apr']
    
    # Find increased movement
    df_merged['increase'] = df_merged['april_qty'] - df_merged['avg_monthly_q1']
    # Add a small buffer to avoid division by zero
    df_merged['pct_increase'] = (df_merged['increase'] / (df_merged['avg_monthly_q1'] + 0.001)) * 100
    
    # Let's say MOQs are based on daily consumption. A month is roughly 24 working days. 
    # Weekly MOQ buffer = Daily * 6 (which is monthly / 4)
    # We set the new MOQ matching the max trend (either the Q1 average or April spike)
    df_merged['old_moq'] = (df_merged['avg_monthly_q1'] / 24) * 6
    df_merged['new_moq'] = (df_merged[['april_qty', 'avg_monthly_q1']].max(axis=1) / 24) * 6
    
    increased = df_merged[(df_merged['increase'] > 0) & (df_merged['april_qty'] > 5)].sort_values('increase', ascending=False)
    return df_merged, increased

def generate_report():
    doc = Document()
    doc.add_heading('MOQ Update & M-o-M Increase Report (April 2026)', 0)
    
    branches = [
        ('Bomas', 'bomas_q1_extracted.csv', 'bomas_apr_extracted.csv'), 
        ('Karen', 'karen_q1_extracted.csv', 'karen_apr_extracted.csv')
    ]
    
    for branch, f_q1, f_apr in branches:
        branch_doc = Document()
        branch_doc.add_heading(f'{branch} Branch - MOQ Update & April Trends Report', 0)

        df_all, df_inc = analyze_branch(branch, f_q1, f_apr)
        doc.add_heading(f'{branch} Branch - April Surges', 1)
        doc.add_paragraph(f"Analyzed {len(df_all)} total items. Found {len(df_inc)} items with significant movement increases in April compared to Q1 (Jan-Mar) averages.")
        
        branch_doc.add_heading('Summary', 1)
        branch_doc.add_paragraph(f"Analyzed {len(df_all)} total items across Jan-April. Found {len(df_inc)} items with significant demand increases in April compared to Q1 (Jan-Mar) averages.")

        # Create Top items table for main combined doc
        doc.add_heading('Top 20 Items Trending Upwards (Action: Increase MOQ)', 2)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Code'
        hdr[1].text = 'Item'
        hdr[2].text = 'Q1 Monthly Avg'
        hdr[3].text = 'April Qty'
        hdr[4].text = 'Increase (%)'
        hdr[5].text = 'New Wkly MOQ'

        # Create same table for branch-specific doc but with Top 50
        branch_doc.add_heading('Top Items Requiring MOQ Adjustments', 2)
        b_table = branch_doc.add_table(rows=1, cols=6)
        b_table.style = 'Table Grid'
        b_hdr = b_table.rows[0].cells
        b_hdr[0].text = 'Code'
        b_hdr[1].text = 'Item'
        b_hdr[2].text = 'Q1 Monthly Avg'
        b_hdr[3].text = 'April Qty'
        b_hdr[4].text = 'Increase (%)'
        b_hdr[5].text = 'New Wkly MOQ'

        # Fill tables
        for idx, row in df_inc.head(20).iterrows():
            r_cells = table.add_row().cells
            r_cells[0].text = str(row['code'])
            r_cells[1].text = str(row['description'])
            r_cells[2].text = f"{row['avg_monthly_q1']:.1f}"
            r_cells[3].text = f"{row['april_qty']:.1f}"
            pct_str = "NEW" if row['avg_monthly_q1'] <= 0 else f"{row['pct_increase']:.1f}%"
            r_cells[4].text = pct_str
            r_cells[5].text = f"{row['new_moq']:.1f}"

        for idx, row in df_inc.head(50).iterrows():
            r_cells = b_table.add_row().cells
            r_cells[0].text = str(row['code'])
            r_cells[1].text = str(row['description'])
            r_cells[2].text = f"{row['avg_monthly_q1']:.1f}"
            r_cells[3].text = f"{row['april_qty']:.1f}"
            pct_str = "NEW" if row['avg_monthly_q1'] <= 0 else f"{row['pct_increase']:.1f}%"
            r_cells[4].text = pct_str
            r_cells[5].text = f"{row['new_moq']:.1f}"

        # Export individual branch JSON format and CSV
        export_cols = ['code', 'description', 'avg_monthly_q1', 'april_qty', 'old_moq', 'new_moq']
        df_all[export_cols].to_csv(f'{branch.lower()}_updated_moqs.csv', index=False)
        
        # Build the specialized detailed JSON output
        detailed_json = format_detailed_json(df_all, branch)
        with open(f'{branch.lower()}_updated_moqs.json', 'w') as f:
            json.dump(detailed_json, f, indent=2)
        
        branch_doc.save(f'{branch}_MOQ_Report_Q1_vs_April.docx')
        
    doc.save('MOQ_Updates_and_April_Trends.docx')
    print('Analysis, JSON data, and Branch Reports Generated.')

if __name__ == '__main__':
    generate_report()