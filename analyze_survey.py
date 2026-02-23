"""
Solid Waste Management Survey Analysis
Generates Chapter 4 Findings Report with Charts and PDF/Word
Professional formatting with proper pagination and content flow
"""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from collections import Counter
import openpyxl
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle, Paragraph, 
                                 Spacer, PageBreak, Image, KeepTogether, 
                                 ListFlowable, ListItem, CondPageBreak)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY, TA_RIGHT
from reportlab.lib.units import inch, cm
from reportlab.platypus.flowables import HRFlowable
import os
from datetime import datetime

# Word document imports
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create charts directory
os.makedirs('survey_charts', exist_ok=True)

# Professional color palette
CHART_COLORS = ['#2E86AB', '#A23B72', '#F18F01', '#C73E1D', '#3B1F2B', 
                '#95C623', '#5C4D7D', '#E07A5F', '#81B29A', '#F2CC8F']

# PDF styling colors
PRIMARY_COLOR = colors.HexColor('#2E86AB')
SECONDARY_COLOR = colors.HexColor('#34495e')
ACCENT_COLOR = colors.HexColor('#3498db')
LIGHT_BG = colors.HexColor('#f8f9fa')
TABLE_HEADER_BG = colors.HexColor('#2E86AB')

PAGE_WIDTH, PAGE_HEIGHT = A4

def load_survey_data(filepath):
    """Load survey data from Excel file"""
    wb = openpyxl.load_workbook(filepath)
    ws = wb.active
    
    headers = [ws.cell(row=1, column=col).value for col in range(1, ws.max_column + 1)]
    
    data = []
    for row in range(2, ws.max_row + 1):
        row_data = {}
        for col, header in enumerate(headers, 1):
            row_data[header] = ws.cell(row=row, column=col).value
        data.append(row_data)
    
    return headers, data

def analyze_column(data, column_name):
    """Analyze a single column and return frequency counts"""
    values = [row.get(column_name) for row in data if row.get(column_name) is not None]
    values = [str(v).strip() for v in values if str(v).strip()]
    
    counter = Counter(values)
    total = len(values)
    
    results = []
    for value, count in counter.most_common():
        percentage = (count / total * 100) if total > 0 else 0
        results.append({
            'value': value,
            'count': count,
            'percentage': round(percentage, 1)
        })
    
    return results, total

def categorize_age(age_value):
    """Categorize age into groups"""
    try:
        age = int(age_value)
        if age < 18:
            return 'Below 18'
        elif 18 <= age <= 25:
            return '18-25'
        elif 26 <= age <= 35:
            return '26-35'
        elif 36 <= age <= 45:
            return '36-45'
        elif 46 <= age <= 55:
            return '46-55'
        else:
            return '56 and above'
    except:
        return None

def clean_label(text, max_len=30):
    """Clean and truncate label text, removing emojis"""
    import re
    # Remove emojis and special unicode characters
    text = re.sub(r'[^\x00-\x7F]+', '', str(text))
    text = text.strip()
    if len(text) > max_len:
        return text[:max_len-3] + '...'
    return text

def create_pie_chart(data, title, filename):
    """Create a professional pie chart"""
    # Clean labels
    labels = [clean_label(d['value'], 25) for d in data]
    sizes = [d['percentage'] for d in data]
    
    fig, ax = plt.subplots(figsize=(7, 3.5))
    
    colors_to_use = CHART_COLORS[:len(data)]
    
    # Create pie with better styling
    explode = [0.02] * len(data)  # Slight separation
    
    wedges, texts, autotexts = ax.pie(
        sizes, 
        labels=None, 
        autopct=lambda pct: f'{pct:.1f}%' if pct > 5 else '',
        colors=colors_to_use, 
        startangle=90,
        explode=explode,
        pctdistance=0.75,
        wedgeprops=dict(width=0.7, edgecolor='white', linewidth=2)
    )
    
    for autotext in autotexts:
        autotext.set_fontsize(8)
        autotext.set_fontweight('bold')
        autotext.set_color('white')
    
    # Professional legend
    legend_labels = [f"{labels[i]} ({sizes[i]:.1f}%)" for i in range(len(data))]
    ax.legend(wedges, legend_labels, 
              loc="center left", 
              bbox_to_anchor=(1.02, 0.5),
              fontsize=8,
              frameon=True,
              fancybox=True,
              shadow=True)
    
    ax.set_title(title, fontsize=10, fontweight='bold', pad=10, color='#2c3e50')
    
    plt.tight_layout()
    filepath = f'survey_charts/{filename}'
    plt.savefig(filepath, dpi=120, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close()
    return filepath

def create_bar_chart(data, title, filename):
    """Create a professional horizontal bar chart"""
    labels = [clean_label(d['value'], 35) for d in data[:8]]  # Top 8 max
    values = [d['percentage'] for d in data[:8]]
    
    fig, ax = plt.subplots(figsize=(8, max(2.5, len(labels) * 0.45)))
    
    y_pos = np.arange(len(labels))
    colors_to_use = CHART_COLORS[:len(labels)]
    
    # Create horizontal bars with gradient effect
    bars = ax.barh(y_pos, values, color=colors_to_use, height=0.6, edgecolor='white', linewidth=1)
    
    ax.set_yticks(y_pos)
    ax.set_yticklabels(labels, fontsize=9)
    ax.set_xlabel('Percentage (%)', fontsize=10, fontweight='bold')
    ax.set_title(title, fontsize=10, fontweight='bold', pad=10, color='#2c3e50')
    
    # Add value labels on bars
    for bar, val in zip(bars, values):
        width = bar.get_width()
        ax.text(width + 1, bar.get_y() + bar.get_height()/2,
                f'{val:.1f}%', va='center', fontsize=9, fontweight='bold', color='#2c3e50')
    
    # Style improvements
    ax.set_xlim(0, max(values) * 1.25)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_color('#bdc3c7')
    ax.spines['left'].set_color('#bdc3c7')
    ax.tick_params(colors='#7f8c8d')
    ax.set_facecolor('#fafafa')
    
    plt.tight_layout()
    filepath = f'survey_charts/{filename}'
    plt.savefig(filepath, dpi=120, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close()
    return filepath

def create_donut_chart(data, title, filename, center_text=""):
    """Create a donut chart for yes/no or binary questions"""
    labels = [clean_label(d['value'], 20) for d in data]
    sizes = [d['percentage'] for d in data]
    
    fig, ax = plt.subplots(figsize=(6, 3.5))
    
    colors_to_use = ['#2E86AB', '#E07A5F'] if len(data) == 2 else CHART_COLORS[:len(data)]
    
    wedges, texts, autotexts = ax.pie(
        sizes, 
        labels=None,
        autopct='%1.1f%%',
        colors=colors_to_use,
        startangle=90,
        pctdistance=0.75,
        wedgeprops=dict(width=0.5, edgecolor='white', linewidth=2)
    )
    
    for autotext in autotexts:
        autotext.set_fontsize(9)
        autotext.set_fontweight('bold')
        autotext.set_color('white')
    
    # Center text
    if center_text:
        ax.text(0, 0, center_text, ha='center', va='center', fontsize=11, fontweight='bold', color='#2c3e50')
    
    # Legend
    legend_labels = [f"{labels[i]}: {data[i]['count']} ({sizes[i]:.1f}%)" for i in range(len(data))]
    ax.legend(wedges, legend_labels, loc="center left", bbox_to_anchor=(1, 0.5), fontsize=9)
    
    ax.set_title(title, fontsize=10, fontweight='bold', pad=10, color='#2c3e50')
    
    plt.tight_layout()
    filepath = f'survey_charts/{filename}'
    plt.savefig(filepath, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return filepath

def generate_interpretation(question_type, data, total_responses):
    """Generate professional interpretation text"""
    if not data:
        return "No responses were recorded for this question."
    
    top = data[0]
    
    if len(data) == 2:
        majority = "majority" if top['percentage'] > 60 else "larger portion"
        return (f"The findings indicate that the {majority} of respondents ({top['percentage']}%, n={top['count']}) "
                f"responded '{top['value']}', while {data[1]['percentage']}% (n={data[1]['count']}) indicated "
                f"'{data[1]['value']}'. Out of {total_responses} total responses, this demonstrates a "
                f"{'clear consensus' if top['percentage'] > 75 else 'notable trend'} in the community.")
    
    elif len(data) <= 5:
        secondary = f", followed by '{data[1]['value']}' ({data[1]['percentage']}%)" if len(data) > 1 else ""
        tertiary = f" and '{data[2]['value']}' ({data[2]['percentage']}%)" if len(data) > 2 else ""
        return (f"Analysis of the {total_responses} responses reveals that '{top['value']}' was the most "
                f"prevalent response at {top['percentage']}% (n={top['count']}){secondary}{tertiary}. "
                f"These findings provide valuable insights into the patterns within the surveyed population.")
    
    else:
        return (f"The data analysis shows considerable diversity in responses. The leading response was "
                f"'{clean_label(top['value'], 40)}' representing {top['percentage']}% of respondents (n={top['count']}). "
                f"This was followed by '{clean_label(data[1]['value'], 40)}' at {data[1]['percentage']}%. "
                f"The variety of responses across {total_responses} participants indicates diverse perspectives "
                f"on this matter within the community.")

def get_styles():
    """Create professional paragraph styles"""
    styles = getSampleStyleSheet()
    
    custom_styles = {
        'ChapterTitle': ParagraphStyle(
            'ChapterTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=25,
            spaceBefore=0,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold',
            textColor=PRIMARY_COLOR,
            borderPadding=10,
        ),
        'SectionHeading': ParagraphStyle(
            'SectionHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=20,
            spaceAfter=12,
            fontName='Helvetica-Bold',
            textColor=SECONDARY_COLOR,
            borderColor=PRIMARY_COLOR,
            borderWidth=0,
            borderPadding=5,
            leftIndent=0,
        ),
        'SubsectionHeading': ParagraphStyle(
            'SubsectionHeading',
            parent=styles['Heading3'],
            fontSize=12,
            spaceBefore=15,
            spaceAfter=8,
            fontName='Helvetica-Bold',
            textColor=colors.HexColor('#34495e'),
            leftIndent=10,
        ),
        'BodyText': ParagraphStyle(
            'BodyText',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            spaceBefore=6,
            alignment=TA_JUSTIFY,
            leading=16,
            firstLineIndent=20,
            textColor=colors.HexColor('#2c3e50'),
        ),
        'Caption': ParagraphStyle(
            'Caption',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=18,
            spaceBefore=8,
            alignment=TA_CENTER,
            fontName='Helvetica-Oblique',
            textColor=colors.HexColor('#7f8c8d'),
        ),
        'IntroText': ParagraphStyle(
            'IntroText',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=15,
            alignment=TA_JUSTIFY,
            leading=16,
            textColor=colors.HexColor('#2c3e50'),
        ),
    }
    
    return custom_styles

def create_styled_table(data_rows, col_widths=None):
    """Create a professionally styled table"""
    if col_widths is None:
        col_widths = [3.2*inch, 1.2*inch, 1.3*inch]
    
    table = Table(data_rows, colWidths=col_widths)
    table.setStyle(TableStyle([
        # Header styling
        ('BACKGROUND', (0, 0), (-1, 0), TABLE_HEADER_BG),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TOPPADDING', (0, 0), (-1, 0), 12),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        
        # Body styling - alternating rows
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
        ('ALIGN', (1, 1), (-1, -1), 'CENTER'),
        ('ALIGN', (0, 1), (0, -1), 'LEFT'),
        ('LEFTPADDING', (0, 1), (0, -1), 10),
        
        # Grid
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#bdc3c7')),
        ('LINEBELOW', (0, 0), (-1, 0), 2, PRIMARY_COLOR),
        
        # Alternating row colors
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, LIGHT_BG]),
    ]))
    
    return table

def create_pdf_report(findings, output_file='survey_findings_report.pdf'):
    """Generate professional PDF report with proper pagination"""
    
    doc = SimpleDocTemplate(
        output_file, 
        pagesize=A4,
        leftMargin=0.8*inch, 
        rightMargin=0.8*inch,
        topMargin=0.75*inch, 
        bottomMargin=0.75*inch
    )
    
    styles = get_styles()
    story = []
    
    # ============ TITLE PAGE ============
    story.append(Spacer(1, 2*inch))
    story.append(Paragraph("CHAPTER 4", styles['ChapterTitle']))
    story.append(Spacer(1, 0.3*inch))
    
    # Decorative line
    story.append(HRFlowable(width="60%", thickness=3, color=PRIMARY_COLOR, spaceAfter=15, spaceBefore=15))
    
    story.append(Paragraph("FINDINGS OF THE STUDY", styles['ChapterTitle']))
    story.append(Spacer(1, 0.5*inch))
    
    # Subtitle
    subtitle_style = ParagraphStyle(
        'Subtitle',
        fontSize=12,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#7f8c8d'),
        spaceAfter=30,
    )
    story.append(Paragraph("Solid Waste Management Practices Among Households", subtitle_style))
    story.append(Paragraph("Dandora Ward", subtitle_style))
    
    story.append(Spacer(1, 1*inch))
    
    # Survey info box
    info_data = [
        ['Survey Information', ''],
        ['Total Respondents', '132'],
        ['Location', 'Dandora Ward'],
        ['Survey Type', 'Household Survey'],
        ['Analysis Date', datetime.now().strftime('%B %d, %Y')],
    ]
    info_table = Table(info_data, colWidths=[2.5*inch, 2.5*inch])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), PRIMARY_COLOR),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('SPAN', (0, 0), (-1, 0)),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('TOPPADDING', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
        ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#bdc3c7')),
        ('BACKGROUND', (0, 1), (-1, -1), LIGHT_BG),
        ('ALIGN', (1, 1), (1, -1), 'CENTER'),
    ]))
    story.append(info_table)
    
    story.append(PageBreak())
    
    # ============ 4.0 INTRODUCTION ============
    story.append(Paragraph("4.0 INTRODUCTION", styles['SectionHeading']))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#ecf0f1'), spaceAfter=12))
    
    intro_text = """This chapter presents the findings and analysis of the study on solid waste management 
    practices among households in Dandora Ward. The data was collected through a structured questionnaire 
    administered to 132 respondents from households within the study area. The findings are organized into 
    five main sections: Social Demographic Factors, Knowledge and Awareness, Waste Management Practices, 
    Challenges and Barriers, and Attitudes and Perceptions."""
    story.append(Paragraph(intro_text, styles['IntroText']))
    
    intro_text2 = """The collected data was analyzed using frequency distribution and percentages to establish 
    patterns and trends. Results are presented through frequency tables, pie charts, and bar graphs, each 
    accompanied by detailed interpretations to provide meaningful insights into the current state of solid 
    waste management in the community."""
    story.append(Paragraph(intro_text2, styles['IntroText']))
    
    story.append(Spacer(1, 20))
    
    # ============ PROCESS EACH SECTION ============
    section_num = 1
    table_num = 1
    figure_num = 1
    
    for section_name, section_findings in findings.items():
        subsection_num = 1
        
        for idx, finding in enumerate(section_findings):
            # Keep heading and content together using KeepTogether
            content_block = []
            
            # For the FIRST subsection, include the section heading in the same KeepTogether block
            if idx == 0:
                # Add section heading to content block so it stays with first subsection
                content_block.append(CondPageBreak(5*inch))  # Ensure enough space for section + first item
                content_block.append(Paragraph(f"4.{section_num} {section_name.upper()}", styles['SectionHeading']))
                content_block.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#ecf0f1'), spaceAfter=15))
            
            # Subsection heading
            content_block.append(Paragraph(
                f"4.{section_num}.{subsection_num} {finding['title']}", 
                styles['SubsectionHeading']
            ))
            
            # Add chart or table
            if finding.get('chart_path') and os.path.exists(finding['chart_path']):
                content_block.append(Spacer(1, 6))
                img = Image(finding['chart_path'], width=4.8*inch, height=2.6*inch)
                content_block.append(img)
                content_block.append(Paragraph(f"Figure {figure_num}: {finding['title']}", styles['Caption']))
                figure_num += 1
            
            elif finding.get('table_data'):
                content_block.append(Spacer(1, 6))
                table_data = [['Response', 'Frequency', 'Percentage']]
                for row in finding['table_data'][:10]:  # Limit rows
                    clean_value = clean_label(row['value'], 45)
                    table_data.append([clean_value, str(row['count']), f"{row['percentage']}%"])
                
                table = create_styled_table(table_data)
                content_block.append(table)
                content_block.append(Paragraph(f"Table {table_num}: {finding['title']}", styles['Caption']))
                table_num += 1
            
            # Interpretation
            content_block.append(Paragraph(finding['interpretation'], styles['BodyText']))
            content_block.append(Spacer(1, 12))
            
            # Use KeepTogether to prevent orphaned headings
            story.append(KeepTogether(content_block))
            
            subsection_num += 1
        
        section_num += 1
    
    # Build PDF
    doc.build(story)
    print(f"PDF Report generated: {output_file}")

def set_cell_shading(cell, color_hex):
    """Set background color for a Word table cell"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex.replace('#', ''))
    cell._tc.get_or_add_tcPr().append(shading)

def create_word_report(findings, output_file='survey_findings_report.docx'):
    """Generate professional Word document report"""
    
    doc = Document()
    
    # Set up document styles
    styles = doc.styles
    
    # Modify Normal style
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    
    # ============ TITLE PAGE ============
    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("CHAPTER 4")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(46, 134, 171)
    
    doc.add_paragraph()
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("FINDINGS OF THE STUDY")
    sub_run.bold = True
    sub_run.font.size = Pt(20)
    sub_run.font.color.rgb = RGBColor(46, 134, 171)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Topic
    topic = doc.add_paragraph()
    topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    topic_run = topic.add_run("Solid Waste Management Practices Among Households")
    topic_run.font.size = Pt(14)
    topic_run.font.color.rgb = RGBColor(127, 140, 141)
    
    location = doc.add_paragraph()
    location.alignment = WD_ALIGN_PARAGRAPH.CENTER
    loc_run = location.add_run("Dandora Ward")
    loc_run.font.size = Pt(14)
    loc_run.font.color.rgb = RGBColor(127, 140, 141)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Info table
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    info_data = [
        ('Survey Information', ''),
        ('Total Respondents', '132'),
        ('Location', 'Dandora Ward'),
        ('Survey Type', 'Household Survey'),
        ('Analysis Date', datetime.now().strftime('%B %d, %Y'))
    ]
    
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        if i == 0:
            # Merge first row for header
            row.cells[0].merge(row.cells[1])
            cell = row.cells[0]
            cell.text = label
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            set_cell_shading(cell, '2E86AB')
        else:
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = value
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(row.cells[0], 'f8f9fa')
            set_cell_shading(row.cells[1], 'f8f9fa')
    
    # Set table borders
    for row in info_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(6)
            cell.paragraphs[0].paragraph_format.space_after = Pt(6)
    
    doc.add_page_break()
    
    # ============ 4.0 INTRODUCTION ============
    intro_heading = doc.add_paragraph()
    intro_run = intro_heading.add_run("4.0 INTRODUCTION")
    intro_run.bold = True
    intro_run.font.size = Pt(14)
    intro_run.font.color.rgb = RGBColor(52, 73, 94)
    
    intro_para = doc.add_paragraph(
        "This chapter presents the findings and analysis of the study on solid waste management "
        "practices among households in Dandora Ward. The data was collected through a structured "
        "questionnaire administered to 132 respondents from households within the study area. The "
        "findings are organized into five main sections: Social Demographic Factors, Knowledge and "
        "Awareness, Waste Management Practices, Challenges and Barriers, and Attitudes and Perceptions."
    )
    intro_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    intro_para2 = doc.add_paragraph(
        "The collected data was analyzed using frequency distribution and percentages to establish "
        "patterns and trends. Results are presented through frequency tables, pie charts, and bar graphs, "
        "each accompanied by detailed interpretations to provide meaningful insights into the current "
        "state of solid waste management in the community."
    )
    intro_para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    # ============ PROCESS EACH SECTION ============
    section_num = 1
    figure_num = 1
    
    for section_name, section_findings in findings.items():
        # Section heading
        section_heading = doc.add_paragraph()
        section_run = section_heading.add_run(f"4.{section_num} {section_name.upper()}")
        section_run.bold = True
        section_run.font.size = Pt(14)
        section_run.font.color.rgb = RGBColor(52, 73, 94)
        
        subsection_num = 1
        
        for finding in section_findings:
            # Subsection heading
            subsection = doc.add_paragraph()
            subsection_run = subsection.add_run(f"4.{section_num}.{subsection_num} {finding['title']}")
            subsection_run.bold = True
            subsection_run.font.size = Pt(12)
            subsection_run.font.color.rgb = RGBColor(52, 73, 94)
            
            # Add chart if exists
            if finding.get('chart_path') and os.path.exists(finding['chart_path']):
                doc.add_paragraph()
                doc.add_picture(finding['chart_path'], width=Inches(5.5))
                
                # Caption
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption.add_run(f"Figure {figure_num}: {finding['title']}")
                caption_run.italic = True
                caption_run.font.size = Pt(10)
                caption_run.font.color.rgb = RGBColor(127, 140, 141)
                figure_num += 1
            
            # Interpretation
            doc.add_paragraph()
            interp = doc.add_paragraph(finding['interpretation'])
            interp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            interp.paragraph_format.first_line_indent = Inches(0.3)
            
            doc.add_paragraph()
            
            subsection_num += 1
        
        section_num += 1
    
    # Save document
    doc.save(output_file)
    print(f"Word Report generated: {output_file}")

def main():
    print("=" * 60)
    print("SOLID WASTE MANAGEMENT SURVEY ANALYSIS")
    print("Dandora Ward - Chapter 4 Findings Report")
    print("=" * 60)
    print()
    
    # Load data
    print("Loading survey data...")
    headers, data = load_survey_data('research proposal.xlsx')
    print(f"Loaded {len(data)} responses with {len(headers)} questions")
    print()
    
    findings = {}
    
    # ========== SECTION 1: SOCIAL DEMOGRAPHIC FACTORS ==========
    print("Analyzing Social Demographic Factors...")
    demographic_findings = []
    
    # 1.1 Gender
    gender_col = 'What is your gender'
    # Normalize gender values
    for d in data:
        val = d.get(gender_col)
        if val:
            val_lower = str(val).strip().lower()
            if 'female' in val_lower or val_lower == 'f':
                d[gender_col] = 'Female'
            elif 'male' in val_lower or val_lower == 'm':
                d[gender_col] = 'Male'
    gender_results, gender_total = analyze_column(data, gender_col)
    gender_chart = create_donut_chart(gender_results, "Gender Distribution of Respondents", "gender_donut.png", "Gender")
    demographic_findings.append({
        'title': 'Gender of the Respondents',
        'chart_path': gender_chart,
        'data': gender_results,
        'interpretation': generate_interpretation('gender', gender_results, gender_total)
    })
    print(f"  - Gender analysis complete: {gender_total} responses")
    
    # 1.2 Age
    age_col = 'what is your age'
    for d in data:
        age = d.get(age_col)
        d['age_group'] = categorize_age(age)
    
    age_results, age_total = analyze_column(data, 'age_group')
    # Sort age groups
    age_order = ['Below 18', '18-25', '26-35', '36-45', '46-55', '56 and above']
    age_results_sorted = []
    for age_group in age_order:
        for r in age_results:
            if r['value'] == age_group:
                age_results_sorted.append(r)
                break
    age_results_final = age_results_sorted if age_results_sorted else age_results
    
    age_chart = create_bar_chart(age_results_final, "Age Distribution of Respondents", "age_bar.png")
    demographic_findings.append({
        'title': 'Age of the Respondents',
        'chart_path': age_chart,
        'data': age_results_final,
        'interpretation': f"The age distribution analysis reveals that the majority of respondents ({age_results_final[0]['percentage']}%, n={age_results_final[0]['count']}) fall within the {age_results_final[0]['value']} age bracket. This is followed by the {age_results_final[1]['value']} group at {age_results_final[1]['percentage']}%. The demographic composition suggests a relatively {'young' if '18-25' in age_results_final[0]['value'] or '26-35' in age_results_final[0]['value'] else 'mature'} population participating in the survey, which may influence their waste management practices and awareness levels."
    })
    print(f"  - Age analysis complete: {age_total} responses")
    
    # 1.3 Education Level
    edu_col = 'what is your educational level'
    edu_results, edu_total = analyze_column(data, edu_col)
    edu_chart = create_pie_chart(edu_results, "Educational Level of Respondents", "education_pie.png")
    demographic_findings.append({
        'title': 'Educational Level of the Respondents',
        'chart_path': edu_chart,
        'data': edu_results,
        'interpretation': generate_interpretation('education', edu_results, edu_total)
    })
    print(f"  - Education analysis complete: {edu_total} responses")
    
    # 1.4 Occupation
    occ_col = 'what is your occupation'
    occ_results, occ_total = analyze_column(data, occ_col)
    occ_chart = create_pie_chart(occ_results, "Occupation of Respondents", "occupation_pie.png")
    demographic_findings.append({
        'title': 'Occupation of the Respondents',
        'chart_path': occ_chart,
        'data': occ_results,
        'interpretation': generate_interpretation('occupation', occ_results, occ_total)
    })
    print(f"  - Occupation analysis complete: {occ_total} responses")
    
    findings['Social Demographic Factors'] = demographic_findings
    
    # ========== SECTION 2: KNOWLEDGE AND AWARENESS ==========
    print("Analyzing Knowledge and Awareness...")
    knowledge_findings = []
    
    # 2.1 Understanding of solid waste
    understanding_col = 'what do you understand by the term "solid waste"'
    understanding_results, understanding_total = analyze_column(data, understanding_col)
    understanding_chart = create_bar_chart(understanding_results[:6], "Understanding of Solid Waste Term", "understanding_bar.png")
    knowledge_findings.append({
        'title': 'Understanding of the Term "Solid Waste"',
        'chart_path': understanding_chart,
        'data': understanding_results,
        'interpretation': generate_interpretation('understanding', understanding_results, understanding_total)
    })
    print(f"  - Understanding analysis complete: {understanding_total} responses")
    
    # 2.2 Importance of proper waste management
    importance_col = 'Do you know the importance of proper solid waste management'
    importance_results, importance_total = analyze_column(data, importance_col)
    importance_chart = create_donut_chart(importance_results, "Knowledge of Waste Management Importance", "importance_donut.png", "Awareness")
    knowledge_findings.append({
        'title': 'Knowledge of Importance of Proper Waste Management',
        'chart_path': importance_chart,
        'data': importance_results,
        'interpretation': f"The survey findings reveal that {importance_results[0]['percentage']}% of respondents (n={importance_results[0]['count']}) indicated '{importance_results[0]['value']}' regarding their knowledge of the importance of proper solid waste management. This demonstrates a {'high' if importance_results[0]['percentage'] > 70 else 'moderate'} level of awareness about waste management significance within the Dandora Ward community, suggesting that {'awareness campaigns have been effective' if importance_results[0]['percentage'] > 70 else 'there is room for improvement in awareness programs'}."
    })
    print(f"  - Importance awareness analysis complete: {importance_total} responses")
    
    # 2.3 Types of solid waste familiar with
    types_col = 'what are the main types of solid wastes you are familiar with'
    types_results, types_total = analyze_column(data, types_col)
    types_chart = create_pie_chart(types_results[:6], "Types of Solid Waste Familiar With", "waste_types_pie.png")
    knowledge_findings.append({
        'title': 'Types of Solid Waste Respondents are Familiar With',
        'chart_path': types_chart,
        'data': types_results,
        'interpretation': generate_interpretation('waste_types', types_results, types_total)
    })
    print(f"  - Waste types analysis complete: {types_total} responses")
    
    # 2.4 Sources of information
    sources_col = 'What sources of information do you rely on for solid waste management?'
    sources_results, sources_total = analyze_column(data, sources_col)
    sources_chart = create_bar_chart(sources_results[:6], "Sources of Information on Waste Management", "info_sources_bar.png")
    knowledge_findings.append({
        'title': 'Sources of Information on Waste Management',
        'chart_path': sources_chart,
        'data': sources_results,
        'interpretation': generate_interpretation('sources', sources_results, sources_total)
    })
    print(f"  - Information sources analysis complete: {sources_total} responses")
    
    findings['Knowledge and Awareness'] = knowledge_findings
    
    # ========== SECTION 3: WASTE MANAGEMENT PRACTICES ==========
    print("Analyzing Waste Management Practices...")
    practices_findings = []
    
    # 3.1 Frequency of waste generation
    freq_col = 'How often do you generate waste in your household'
    freq_results, freq_total = analyze_column(data, freq_col)
    freq_chart = create_pie_chart(freq_results, "Frequency of Household Waste Generation", "waste_frequency_pie.png")
    practices_findings.append({
        'title': 'Frequency of Waste Generation in Households',
        'chart_path': freq_chart,
        'data': freq_results,
        'interpretation': generate_interpretation('frequency', freq_results, freq_total)
    })
    print(f"  - Waste generation frequency analysis complete: {freq_total} responses")
    
    # 3.2 Waste separation
    sep_col = 'Do you separate waste into different categories (e.g., organics, recyclable, non recyclable)?'
    sep_results, sep_total = analyze_column(data, sep_col)
    sep_chart = create_donut_chart(sep_results, "Waste Separation Practices", "separation_donut.png", "Separation")
    practices_findings.append({
        'title': 'Waste Separation Practices',
        'chart_path': sep_chart,
        'data': sep_results,
        'interpretation': f"When examining waste separation practices, the data shows that {sep_results[0]['percentage']}% of respondents (n={sep_results[0]['count']}) indicated '{sep_results[0]['value']}'. This finding suggests that {'a significant portion of households actively practice' if 'yes' in sep_results[0]['value'].lower() and sep_results[0]['percentage'] > 50 else 'there is considerable room for improvement in'} waste separation at the household level in Dandora Ward, which has implications for recycling and waste processing efficiency."
    })
    print(f"  - Waste separation analysis complete: {sep_total} responses")
    
    # 3.3 Proper waste disposal
    disposal_col = 'do you dispose off waste in a proper way?'
    disposal_results, disposal_total = analyze_column(data, disposal_col)
    disposal_chart = create_donut_chart(disposal_results, "Proper Waste Disposal Practices", "disposal_donut.png", "Disposal")
    practices_findings.append({
        'title': 'Proper Waste Disposal Practices',
        'chart_path': disposal_chart,
        'data': disposal_results,
        'interpretation': generate_interpretation('disposal', disposal_results, disposal_total)
    })
    print(f"  - Waste disposal analysis complete: {disposal_total} responses")
    
    # 3.4 Recycling
    recycle_col = 'do you recycle any materials?'
    recycle_results, recycle_total = analyze_column(data, recycle_col)
    recycle_chart = create_donut_chart(recycle_results, "Recycling Practices Among Households", "recycling_donut.png", "Recycling")
    practices_findings.append({
        'title': 'Recycling Practices',
        'chart_path': recycle_chart,
        'data': recycle_results,
        'interpretation': f"Regarding recycling practices among households, {recycle_results[0]['percentage']}% of respondents (n={recycle_results[0]['count']}) answered '{recycle_results[0]['value']}'. {'The data indicates that a significant portion of the community actively engages in recycling activities, which contributes positively to waste reduction efforts' if 'yes' in recycle_results[0]['value'].lower() and recycle_results[0]['percentage'] > 50 else 'These results highlight the potential for increased recycling adoption through targeted community education and improved access to recycling facilities'}."
    })
    print(f"  - Recycling analysis complete: {recycle_total} responses")
    
    findings['Waste Management Practices'] = practices_findings
    
    # ========== SECTION 4: CHALLENGES AND BARRIERS ==========
    print("Analyzing Challenges and Barriers...")
    challenges_findings = []
    
    # 4.1 Barriers to waste management
    barriers_col = 'what barriers do you face in managing solid waste'
    barriers_results, barriers_total = analyze_column(data, barriers_col)
    barriers_chart = create_pie_chart(barriers_results[:6], "Barriers to Effective Waste Management", "barriers_pie.png")
    challenges_findings.append({
        'title': 'Barriers Faced in Managing Solid Waste',
        'chart_path': barriers_chart,
        'data': barriers_results,
        'interpretation': generate_interpretation('barriers', barriers_results, barriers_total)
    })
    print(f"  - Barriers analysis complete: {barriers_total} responses")
    
    # 4.2 Factors influencing practices
    factors_col = 'which factors influence your solid waste management practices?'
    factors_results, factors_total = analyze_column(data, factors_col)
    factors_chart = create_bar_chart(factors_results[:6], "Factors Influencing Waste Management", "factors_bar.png")
    challenges_findings.append({
        'title': 'Factors Influencing Waste Management Practices',
        'chart_path': factors_chart,
        'data': factors_results,
        'interpretation': generate_interpretation('factors', factors_results, factors_total)
    })
    print(f"  - Influencing factors analysis complete: {factors_total} responses")
    
    # 4.3 Barriers to awareness
    awareness_barriers_col = 'what barriers prevent you from being more aware of solid waste management?'
    awareness_barriers_results, awareness_barriers_total = analyze_column(data, awareness_barriers_col)
    awareness_barriers_chart = create_pie_chart(awareness_barriers_results[:6], "Barriers to Waste Management Awareness", "awareness_barriers_pie.png")
    challenges_findings.append({
        'title': 'Barriers Preventing Greater Awareness',
        'chart_path': awareness_barriers_chart,
        'data': awareness_barriers_results,
        'interpretation': generate_interpretation('awareness_barriers', awareness_barriers_results, awareness_barriers_total)
    })
    print(f"  - Awareness barriers analysis complete: {awareness_barriers_total} responses")
    
    findings['Challenges and Barriers'] = challenges_findings
    
    # ========== SECTION 5: ATTITUDES AND PERCEPTIONS ==========
    print("Analyzing Attitudes and Perceptions...")
    attitudes_findings = []
    
    # 5.1 General attitude
    attitude_col = 'how would describe your general attitude toward solid waste management?'
    attitude_results, attitude_total = analyze_column(data, attitude_col)
    attitude_chart = create_pie_chart(attitude_results, "General Attitude Toward Waste Management", "attitude_pie.png")
    attitudes_findings.append({
        'title': 'General Attitude Toward Solid Waste Management',
        'chart_path': attitude_chart,
        'data': attitude_results,
        'interpretation': generate_interpretation('attitude', attitude_results, attitude_total)
    })
    print(f"  - Attitude analysis complete: {attitude_total} responses")
    
    # 5.2 Individual responsibility
    responsibility_col = 'do you think that individuals have a responsibility to manage waste?'
    responsibility_results, responsibility_total = analyze_column(data, responsibility_col)
    responsibility_chart = create_donut_chart(responsibility_results, "Perception of Individual Responsibility", "responsibility_donut.png", "Responsibility")
    attitudes_findings.append({
        'title': 'Perception of Individual Responsibility',
        'chart_path': responsibility_chart,
        'data': responsibility_results,
        'interpretation': f"A significant {responsibility_results[0]['percentage']}% of respondents (n={responsibility_results[0]['count']}) believe that individuals {'have' if 'yes' in responsibility_results[0]['value'].lower() else 'do not have'} a responsibility to manage waste. This indicates a {'strong' if responsibility_results[0]['percentage'] > 80 else 'moderate'} sense of personal accountability toward waste management in the community, which is {'encouraging for community-based waste management initiatives' if responsibility_results[0]['percentage'] > 70 else 'an area that requires further community engagement'}."
    })
    print(f"  - Responsibility perception analysis complete: {responsibility_total} responses")
    
    # 5.3 Environmental impact
    impact_col = 'Do you think that improper waste disposal affects the environment and public health?'
    impact_results, impact_total = analyze_column(data, impact_col)
    impact_chart = create_donut_chart(impact_results, "Awareness of Environmental and Health Impact", "impact_donut.png", "Impact")
    attitudes_findings.append({
        'title': 'Awareness of Environmental and Health Impact',
        'chart_path': impact_chart,
        'data': impact_results,
        'interpretation': f"When asked about the impact of improper waste disposal on environment and public health, {impact_results[0]['percentage']}% of respondents (n={impact_results[0]['count']}) indicated '{impact_results[0]['value']}'. This demonstrates a {'high' if impact_results[0]['percentage'] > 80 else 'moderate'} level of awareness regarding the environmental and public health consequences of improper waste management, suggesting that the community {'understands the urgency of proper waste disposal' if impact_results[0]['percentage'] > 80 else 'would benefit from more education on environmental health'}."
    })
    print(f"  - Environmental impact awareness analysis complete: {impact_total} responses")
    
    # 5.4 Main problems
    problems_col = 'What do you think are the main problems with solid waste management in your community?'
    problems_results, problems_total = analyze_column(data, problems_col)
    problems_chart = create_bar_chart(problems_results[:6], "Main Problems Identified by Respondents", "problems_bar.png")
    attitudes_findings.append({
        'title': 'Perceived Main Problems with Solid Waste Management',
        'chart_path': problems_chart,
        'data': problems_results,
        'interpretation': generate_interpretation('problems', problems_results, problems_total)
    })
    print(f"  - Problems perception analysis complete: {problems_total} responses")
    
    findings['Attitudes and Perceptions'] = attitudes_findings
    
    # Generate PDF Report
    print()
    print("=" * 60)
    print("Generating Professional PDF Report...")
    create_pdf_report(findings, 'survey_findings_report.pdf')
    
    # Generate Word Report
    print("Generating Professional Word Report...")
    create_word_report(findings, 'survey_findings_report.docx')
    
    print()
    print("=" * 60)
    print("ANALYSIS COMPLETE!")
    print("=" * 60)
    print()
    print("Generated files:")
    print("  - survey_findings_report.pdf (PDF Chapter 4 Report)")
    print("  - survey_findings_report.docx (Word Chapter 4 Report)")
    print("  - survey_charts/ (All charts and figures)")
    print()

if __name__ == "__main__":
    main()
