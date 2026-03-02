import json
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import (BaseDocTemplate, PageTemplate, Frame, Table, TableStyle,
                                Paragraph, Spacer, PageBreak, KeepTogether)
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.units import inch, cm
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ===== COLOR PALETTE =====
NAVY = '#1a3a5c'
DARK_BLUE = '#2c3e50'
TEAL = '#16a085'
TEAL_LIGHT = '#d5f5e3'
ORANGE = '#e67e22'
RED = '#c0392b'
GREY_HEADER = '#5d6d7e'
GREY_TEXT = '#566573'
GREY_LIGHT = '#7f8c8d'
ROW_ALT = '#f2f6fa'
BORDER = '#d5d8dc'
WHITE = '#ffffff'
CAT_HEADER_BG = '#1a3a5c'
ACCENT_GREEN = '#27ae60'
ACCENT_BLUE = '#2980b9'
INFO_BG = '#eaf2f8'

def load_data():
    with open('moq_analysis.json', 'r') as f:
        return json.load(f)


def build_insights(data):
    """Generate data-driven narrative insights from the MOQ analysis."""
    s = data['summary']
    cats = sorted(data['categories'].items(), key=lambda x: x[1]['total_qty_out'], reverse=True)
    fast = data['fast_movers']
    slow = data['slow_movers']

    total_qty = s['total_qty_out']
    top5_qty = sum(f['qty_out'] for f in fast[:5])
    top30_qty = sum(f['qty_out'] for f in fast[:30])
    top5_pct = top5_qty / total_qty * 100 if total_qty else 0
    top30_pct = top30_qty / total_qty * 100 if total_qty else 0
    fast_pct = s['fast_movers'] / s['total_items'] * 100 if s['total_items'] else 0
    slow_pct = s['slow_movers'] / s['total_items'] * 100 if s['total_items'] else 0

    # Top 3 categories by value
    cats_by_value = sorted(cats, key=lambda x: x[1].get('total_weekly_value', 0), reverse=True)
    top3_val = [(n, c.get('total_weekly_value', 0)) for n, c in cats_by_value[:3]]
    top3_val_total = sum(v for _, v in top3_val)
    top3_val_pct = top3_val_total / s.get('total_weekly_value', 1) * 100

    # Categories with highest slow-mover ratio
    high_slow = [(n, c['slow_movers'], c['total_items'],
                  c['slow_movers'] / c['total_items'] * 100 if c['total_items'] else 0)
                 for n, c in cats if c['total_items'] >= 10]
    high_slow.sort(key=lambda x: x[3], reverse=True)

    insights = {
        'executive_summary': (
            f"In February 2026, BOMAS Hardware tracked {s['total_items']:,} SKUs across 23 categories "
            f"with total outward movement of {total_qty:,.0f} units over 24 working days. "
            f"Of these, {s['fast_movers']:,} items ({fast_pct:.0f}%) recorded active sales, while "
            f"{s['slow_movers']:,} items ({slow_pct:.0f}%) had zero outward movement — a significant "
            f"portion of inventory sitting idle on shelves. The calculated weekly MOQ stands at "
            f"{s['total_weekly_moq']:,} units, representing a weekly restocking investment of approximately "
            f"KES {s.get('total_weekly_value', 0):,.0f}."
        ),
        'demand_concentration': (
            f"Demand is heavily concentrated: the top 5 items alone account for {top5_pct:.0f}% of all "
            f"quantity moved, and the top 30 items drive {top30_pct:.0f}% of total movement. This Pareto "
            f"pattern is typical in hardware retail and means that a relatively small number of SKUs "
            f"generate the bulk of turnover. Ensuring these high-velocity items are always in stock "
            f"should be the top priority — a single day of stockout on a fast mover like Simba Cement "
            f"(~1,120 units/day) can result in substantial lost revenue."
        ),
        'category_value': (
            f"By weekly restocking value, the top three categories are {top3_val[0][0]} "
            f"(KES {top3_val[0][1]:,.0f}/wk), {top3_val[1][0]} (KES {top3_val[1][1]:,.0f}/wk), and "
            f"{top3_val[2][0]} (KES {top3_val[2][1]:,.0f}/wk) — together accounting for "
            f"{top3_val_pct:.0f}% of the total weekly MOQ value of KES {s.get('total_weekly_value', 0):,.0f}. "
            f"These categories require the largest capital allocation and should be reviewed "
            f"weekly to balance cash flow against stockout risk."
        ),
        'dead_stock': (
            f"Dead stock — inventory that has been received but shows zero customer demand — represents "
            f"{s['slow_movers']:,} items ({slow_pct:.0f}% of tracked SKUs). In inventory management, "
            f"dead stock ties up capital, occupies valuable shelf and warehouse space, increases carrying "
            f"costs (estimated at 20-30% of inventory value annually), and risks obsolescence. "
            f"The categories with the highest proportion of non-moving items are "
            f"{high_slow[0][0]} ({high_slow[0][1]}/{high_slow[0][2]} items, {high_slow[0][3]:.0f}% idle) "
            f"and {high_slow[1][0]} ({high_slow[1][1]}/{high_slow[1][2]} items, {high_slow[1][3]:.0f}% idle). "
            f"These items should be reviewed for clearance sales, bundling with fast movers, return to "
            f"supplier, or discontinuation to free up working capital."
        ),
        'moq_importance': (
            "Minimum Order Quantities (MOQs) serve as a critical planning tool that bridges the gap "
            "between customer demand and procurement decisions. By calculating MOQ as 6 times the daily "
            "average consumption (equivalent to one week's supply), BOMAS ensures: (1) sufficient stock "
            "to cover a full ordering cycle without risking stockouts, (2) a data-driven basis for "
            "purchase orders that eliminates guesswork, (3) better cash flow management by ordering only "
            "what is needed rather than over-purchasing, and (4) reduced carrying costs by avoiding "
            "excess inventory that depreciates or becomes obsolete. Regularly reviewing MOQs against "
            "actual sales trends ensures the business adapts to seasonal shifts and changing customer "
            "preferences — what moves fast in February may slow down in other months."
        ),
        'cash_flow': (
            f"With a weekly restocking value of KES {s.get('total_weekly_value', 0):,.0f}, prudent cash "
            f"allocation is essential. Prioritise spending on the top-performing categories and fast movers "
            f"that generate the highest turnover. For categories with high slow-mover ratios, reduce "
            f"reorder quantities or pause ordering altogether until existing stock is cleared. "
            f"Consider negotiating extended payment terms with suppliers of high-value categories "
            f"(Cement, Steel, Timber) to smooth out weekly cash outflows. The goal is to maintain "
            f"a lean inventory — enough stock to serve customers without tying up excessive capital "
            f"in slow-moving or dead items."
        ),
        'recommendations': [
            f"Prioritise restocking the top 30 fast movers weekly — they drive {top30_pct:.0f}% of all movement.",
            f"Conduct a dead stock clearance exercise on the {s['slow_movers']:,} non-moving items — "
            f"consider markdowns, bundle deals, or supplier returns.",
            f"Review the {high_slow[0][0]} category ({high_slow[0][3]:.0f}% idle) for possible SKU rationalisation.",
            "Set up weekly MOQ-based purchase orders to replace ad-hoc buying, ensuring consistent stock levels.",
            "Track MOQ fulfilment rates monthly to identify categories where actual orders deviate from calculated needs.",
            "Negotiate better terms with Cement and Steel suppliers given the high weekly spend (combined KES 11.4M/wk).",
        ],
    }
    return insights


def build_jan_feb_comparison():
    """Compare January and February stock movement to identify trends."""
    import os
    if not os.path.exists('stock_movement_jan.json') or not os.path.exists('stock_movement_feb.json'):
        return None

    with open('stock_movement_jan.json') as f:
        jan = json.load(f)
    with open('stock_movement_feb.json') as f:
        feb = json.load(f)

    jan_wd = jan.get('working_days', 24)
    feb_wd = feb.get('working_days', 24)
    jan_items = {i['code']: i for i in jan['items']}
    feb_items = {i['code']: i for i in feb['items']}
    common = set(jan_items) & set(feb_items)

    jan_total_out = sum(i['qty_out'] for i in jan['items'])
    feb_total_out = sum(i['qty_out'] for i in feb['items'])
    jan_fast = sum(1 for i in jan['items'] if i['qty_out'] > 0)
    feb_fast = sum(1 for i in feb['items'] if i['qty_out'] > 0)
    jan_slow = len(jan['items']) - jan_fast
    feb_slow = len(feb['items']) - feb_fast
    jan_moq = round(jan_total_out / jan_wd * 6)
    feb_moq = round(feb_total_out / feb_wd * 6)

    # Build per-item comparison for common items
    movers = []
    for code in common:
        j, f = jan_items[code], feb_items[code]
        j_out, f_out = j['qty_out'], f['qty_out']
        diff = f_out - j_out
        j_moq = round(j_out / jan_wd * 6)
        f_moq = round(f_out / feb_wd * 6)
        movers.append({
            'code': code, 'description': j['description'],
            'jan_qty': j_out, 'feb_qty': f_out, 'qty_change': diff,
            'jan_moq': j_moq, 'feb_moq': f_moq, 'moq_change': f_moq - j_moq,
        })

    gains = sorted(movers, key=lambda x: x['qty_change'], reverse=True)[:15]
    drops = sorted(movers, key=lambda x: x['qty_change'])[:15]

    # Status shifts
    jan_active_feb_dead = [c for c in common if jan_items[c]['qty_out'] > 0 and feb_items[c]['qty_out'] == 0]
    jan_dead_feb_active = [c for c in common if jan_items[c]['qty_out'] == 0 and feb_items[c]['qty_out'] > 0]
    feb_only = set(feb_items) - set(jan_items)

    qty_change = feb_total_out - jan_total_out
    qty_change_pct = qty_change / jan_total_out * 100 if jan_total_out else 0
    moq_change = feb_moq - jan_moq

    # Narrative
    direction = 'increased' if qty_change > 0 else 'decreased'
    narrative = (
        f"Total stock movement {direction} by {abs(qty_change):,.0f} units ({abs(qty_change_pct):.1f}%) "
        f"from January ({jan_total_out:,.0f} units) to February ({feb_total_out:,.0f} units), "
        f"while the item catalogue expanded from {len(jan['items']):,} to {len(feb['items']):,} tracked SKUs. "
        f"The aggregate weekly MOQ rose from {jan_moq:,} to {feb_moq:,} units (+{moq_change:,}). "
        f"However, {len(jan_active_feb_dead)} items that were active in January recorded zero sales in "
        f"February, suggesting seasonal demand shifts or possible stockout issues. Conversely, "
        f"{len(jan_dead_feb_active)} previously dormant items came to life in February, and "
        f"{len(feb_only)} entirely new items were introduced to the catalogue."
    )

    return {
        'jan_period': jan.get('period', 'January 2026'),
        'feb_period': feb.get('period', 'February 2026'),
        'jan_items': len(jan['items']), 'feb_items': len(feb['items']),
        'jan_fast': jan_fast, 'feb_fast': feb_fast,
        'jan_slow': jan_slow, 'feb_slow': feb_slow,
        'jan_total_out': jan_total_out, 'feb_total_out': feb_total_out,
        'jan_moq': jan_moq, 'feb_moq': feb_moq,
        'qty_change': qty_change, 'qty_change_pct': qty_change_pct,
        'moq_change': moq_change,
        'common_items': len(common),
        'jan_active_feb_dead': len(jan_active_feb_dead),
        'jan_dead_feb_active': len(jan_dead_feb_active),
        'new_items': len(feb_only),
        'top_gains': gains,
        'top_drops': drops,
        'narrative': narrative,
    }


class MOQReport:
    def __init__(self, filename='moq_report.pdf'):
        self.filename = filename
        self.page_width, self.page_height = A4
        self.margin = 0.55 * inch
        self.usable_width = self.page_width - 2 * self.margin
        self.setup_styles()

    def setup_styles(self):
        self.styles = {
            'title': ParagraphStyle('Title', fontSize=16, textColor=colors.HexColor(NAVY),
                                    alignment=TA_CENTER, fontName='Helvetica-Bold', spaceAfter=4),
            'subtitle': ParagraphStyle('Subtitle', fontSize=9, textColor=colors.HexColor(GREY_TEXT),
                                       alignment=TA_CENTER, spaceAfter=8),
            'heading': ParagraphStyle('Heading', fontSize=12, textColor=colors.HexColor(NAVY),
                                      fontName='Helvetica-Bold', spaceAfter=6, spaceBefore=10),
            'subheading': ParagraphStyle('SubHeading', fontSize=10, textColor=colors.HexColor(DARK_BLUE),
                                          fontName='Helvetica-Bold', spaceAfter=4, spaceBefore=6),
            'body': ParagraphStyle('Body', fontSize=8, textColor=colors.HexColor(GREY_TEXT),
                                    spaceAfter=6, leading=12),
            'small': ParagraphStyle('Small', fontSize=7, textColor=colors.HexColor(GREY_LIGHT),
                                    spaceAfter=2),
            'note': ParagraphStyle('Note', fontSize=7.5, textColor=colors.HexColor(GREY_TEXT),
                                    spaceAfter=4, fontName='Helvetica-Oblique'),
            'bullet': ParagraphStyle('Bullet', fontSize=8, textColor=colors.HexColor(GREY_TEXT),
                                      spaceAfter=3, leftIndent=18, bulletIndent=8, leading=11),
            'insight_heading': ParagraphStyle('InsightHead', fontSize=9.5,
                                              textColor=colors.HexColor(DARK_BLUE),
                                              fontName='Helvetica-Bold', spaceAfter=3, spaceBefore=8),
        }

    def header_footer(self, canvas, doc):
        canvas.saveState()
        canvas.setFillColor(colors.HexColor(NAVY))
        canvas.rect(0, self.page_height - 0.5 * inch, self.page_width, 0.5 * inch, fill=True, stroke=False)
        canvas.setFont('Helvetica-Bold', 10)
        canvas.setFillColor(colors.white)
        canvas.drawString(self.margin, self.page_height - 0.33 * inch,
                          'BOMAS HARDWARE — MOQ REPORT  |  FEBRUARY 2026  |  24 Working Days  |  Formula: 6 × Daily Avg')
        canvas.setFont('Helvetica', 8)
        canvas.drawRightString(self.page_width - self.margin, self.page_height - 0.33 * inch,
                               f'Page {doc.page}')
        canvas.setStrokeColor(colors.HexColor(BORDER))
        canvas.setLineWidth(0.3)
        canvas.line(self.margin, 0.35 * inch, self.page_width - self.margin, 0.35 * inch)
        canvas.setFont('Helvetica', 6.5)
        canvas.setFillColor(colors.HexColor(GREY_LIGHT))
        canvas.drawString(self.margin, 0.22 * inch, f'Generated: {datetime.now().strftime("%d %b %Y %H:%M")}')
        canvas.drawRightString(self.page_width - self.margin, 0.22 * inch, 'Confidential — BOMAS Hardware Store')
        canvas.restoreState()

    def make_table_style(self, header_color=DARK_BLUE, pad_top=4, pad_bot=4, font_header=8, font_body=7):
        return TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(header_color)),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), font_header),
            ('FONTSIZE', (0, 1), (-1, -1), font_body),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('ALIGN', (0, 0), (1, -1), 'LEFT'),
            ('ALIGN', (2, 0), (-1, -1), 'RIGHT'),
            ('TOPPADDING', (0, 0), (-1, 0), pad_top + 1),
            ('BOTTOMPADDING', (0, 0), (-1, 0), pad_bot + 1),
            ('TOPPADDING', (0, 1), (-1, -1), pad_top),
            ('BOTTOMPADDING', (0, 1), (-1, -1), pad_bot),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, 0), 0.5, colors.HexColor(header_color)),
            ('LINEBELOW', (0, 0), (-1, 0), 1, colors.HexColor(header_color)),
            ('LINEBELOW', (0, 1), (-1, -2), 0.3, colors.HexColor(BORDER)),
            ('LINEBELOW', (0, -1), (-1, -1), 0.5, colors.HexColor(BORDER)),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor(ROW_ALT)]),
        ])

    def make_insight_box(self, text):
        """Create a shaded insight paragraph block."""
        t = Table([[Paragraph(text, ParagraphStyle('ins', fontSize=8,
                    textColor=colors.HexColor(GREY_TEXT), leading=12))]],
                  colWidths=[self.usable_width - 12])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(INFO_BG)),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
            ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor(ACCENT_BLUE)),
        ]))
        return t

    def generate(self):
        data = load_data()
        insights = build_insights(data)

        doc = BaseDocTemplate(self.filename, pagesize=A4,
                              topMargin=0.6 * inch, bottomMargin=0.45 * inch,
                              leftMargin=self.margin, rightMargin=self.margin)

        frame = Frame(self.margin, 0.45 * inch, self.usable_width,
                      self.page_height - 1.05 * inch, id='main')
        template = PageTemplate(id='main', frames=frame, onPage=self.header_footer)
        doc.addPageTemplates([template])

        elements = []
        summary = data['summary']
        sorted_cats = sorted(data['categories'].items(),
                             key=lambda x: x[1]['total_qty_out'], reverse=True)

        # ========== PAGE 1: SUMMARY ==========
        elements.append(Spacer(1, 4))
        elements.append(Paragraph('INVENTORY MOVEMENT & MOQ ANALYSIS', self.styles['title']))
        elements.append(Paragraph(
            'Period: February 1-28, 2026  •  24 Working Days (4 weeks × 6 days)  •  '
            'Daily Avg = Qty Out ÷ 24  •  MOQ = Daily Avg × 6',
            self.styles['subtitle']))
        elements.append(Spacer(1, 6))

        # KPI Cards
        kpi_data = [
            ['Total Items', 'Fast Movers', 'Slow Movers', 'Total Qty Out', 'Weekly MOQ', 'Weekly Value'],
            [f"{summary['total_items']:,}", f"{summary['fast_movers']:,}", f"{summary['slow_movers']:,}",
             f"{summary['total_qty_out']:,.0f}", f"{summary['total_weekly_moq']:,}",
             f"KES {summary.get('total_weekly_value', 0):,.0f}"]
        ]
        kpi_widths = [self.usable_width / 6] * 6
        t = Table(kpi_data, colWidths=kpi_widths)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(TEAL)),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('BACKGROUND', (0, 1), (-1, 1), colors.HexColor(TEAL_LIGHT)),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTNAME', (0, 1), (-1, 1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 7.5),
            ('FONTSIZE', (0, 1), (-1, 1), 9),
            ('TEXTCOLOR', (0, 1), (-1, 1), colors.HexColor(DARK_BLUE)),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('TOPPADDING', (0, 0), (-1, -1), 7),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 7),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.white),
            ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor(TEAL)),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 10))

        # Executive Summary
        elements.append(Paragraph('EXECUTIVE SUMMARY', self.styles['heading']))
        elements.append(Paragraph(insights['executive_summary'], self.styles['body']))
        elements.append(Spacer(1, 6))

        # Category Summary Table
        elements.append(Paragraph('MOQ BY CATEGORY', self.styles['heading']))
        cat_header = ['Category', 'Items', 'Fast', 'Slow', 'Qty Out', 'Wkly MOQ', 'Wkly Value (KES)']
        cat_rows = [cat_header]
        for cat_name, cat_info in sorted_cats:
            if cat_info['total_qty_out'] > 0 or cat_info['total_items'] > 0:
                cat_rows.append([
                    cat_name[:25],
                    str(cat_info['total_items']),
                    str(cat_info['fast_movers']),
                    str(cat_info['slow_movers']),
                    f"{cat_info['total_qty_out']:,.0f}",
                    f"{cat_info['total_weekly_moq']:,}",
                    f"{cat_info.get('total_weekly_value', 0):,.0f}"
                ])

        cat_rows.append([
            'TOTAL', str(summary['total_items']), str(summary['fast_movers']),
            str(summary['slow_movers']), f"{summary['total_qty_out']:,.0f}",
            f"{summary['total_weekly_moq']:,}",
            f"{summary.get('total_weekly_value', 0):,.0f}"
        ])

        col_w = [2.1 * inch, 0.5 * inch, 0.45 * inch, 0.45 * inch, 0.85 * inch, 0.8 * inch, 1.2 * inch]
        t = Table(cat_rows, colWidths=col_w)
        style = self.make_table_style(DARK_BLUE)
        style.add('BACKGROUND', (0, -1), (-1, -1), colors.HexColor(ACCENT_GREEN))
        style.add('TEXTCOLOR', (0, -1), (-1, -1), colors.white)
        style.add('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold')
        t.setStyle(style)
        elements.append(t)
        elements.append(PageBreak())

        # ========== PAGE 2: DEMAND & CASH FLOW ANALYSIS ==========
        elements.append(Paragraph('DEMAND ANALYSIS & INSIGHTS', self.styles['title']))
        elements.append(Spacer(1, 8))

        elements.append(Paragraph('Demand Concentration', self.styles['insight_heading']))
        elements.append(self.make_insight_box(insights['demand_concentration']))
        elements.append(Spacer(1, 6))

        elements.append(Paragraph('Category Value Distribution', self.styles['insight_heading']))
        elements.append(self.make_insight_box(insights['category_value']))
        elements.append(Spacer(1, 6))

        elements.append(Paragraph('Dead Stock & Slow Movers', self.styles['insight_heading']))
        elements.append(self.make_insight_box(insights['dead_stock']))
        elements.append(Spacer(1, 6))

        elements.append(Paragraph('Why MOQs Matter', self.styles['insight_heading']))
        elements.append(self.make_insight_box(insights['moq_importance']))
        elements.append(Spacer(1, 6))

        elements.append(Paragraph('Cash Flow & Capital Allocation', self.styles['insight_heading']))
        elements.append(self.make_insight_box(insights['cash_flow']))
        elements.append(PageBreak())

        # ========== PAGE 3: RECOMMENDATIONS ==========
        elements.append(Paragraph('KEY RECOMMENDATIONS', self.styles['heading']))
        elements.append(Paragraph(
            'Based on the February 2026 movement data, the following actions are recommended:',
            self.styles['note']))
        elements.append(Spacer(1, 4))
        for i, rec in enumerate(insights['recommendations'], 1):
            elements.append(Paragraph(
                f'<b>{i}.</b>  {rec}', self.styles['bullet']))
        elements.append(Spacer(1, 10))

        # ========== JAN vs FEB COMPARISON ==========
        comparison = build_jan_feb_comparison()
        if comparison:
            elements.append(PageBreak())
            elements.append(Paragraph('JANUARY vs FEBRUARY — MOQ COMPARISON', self.styles['title']))
            elements.append(Paragraph(
                f'{comparison["jan_period"]}  →  {comparison["feb_period"]}  |  '
                f'Both months: 24 working days  |  {comparison["common_items"]:,} common items compared',
                self.styles['subtitle']))
            elements.append(Spacer(1, 6))

            # Overview narrative
            elements.append(Paragraph('Overview', self.styles['insight_heading']))
            elements.append(self.make_insight_box(comparison['narrative']))
            elements.append(Spacer(1, 8))

            # Summary comparison KPIs
            cmp_header = ['Metric', 'January', 'February', 'Change']
            c = comparison
            arrow_up = '▲'
            arrow_down = '▼'
            def fmt_chg(val, pct=None):
                sign = '+' if val >= 0 else ''
                pct_str = f' ({sign}{pct:.1f}%)' if pct is not None else ''
                return f'{sign}{val:,.0f}{pct_str}'

            cmp_rows = [cmp_header,
                ['Total Items Tracked', f"{c['jan_items']:,}", f"{c['feb_items']:,}",
                 fmt_chg(c['feb_items'] - c['jan_items'])],
                ['Active Items (Fast Movers)', f"{c['jan_fast']:,}", f"{c['feb_fast']:,}",
                 fmt_chg(c['feb_fast'] - c['jan_fast'])],
                ['Inactive Items (Slow/Dead)', f"{c['jan_slow']:,}", f"{c['feb_slow']:,}",
                 fmt_chg(c['feb_slow'] - c['jan_slow'])],
                ['Total Qty Out', f"{c['jan_total_out']:,.0f}", f"{c['feb_total_out']:,.0f}",
                 fmt_chg(c['qty_change'], c['qty_change_pct'])],
                ['Weekly MOQ (units)', f"{c['jan_moq']:,}", f"{c['feb_moq']:,}",
                 fmt_chg(c['moq_change'])],
            ]
            cw = [2.0 * inch, 1.2 * inch, 1.2 * inch, 1.6 * inch]
            t = Table(cmp_rows, colWidths=cw)
            style = self.make_table_style(ACCENT_BLUE, pad_top=5, pad_bot=5, font_body=8)
            t.setStyle(style)
            elements.append(t)
            elements.append(Spacer(1, 6))

            # Status shifts
            shifts_text = (
                f"Between the two months, {c['jan_active_feb_dead']} items that had active sales in January "
                f"went dormant in February (potential demand loss or stock issues), while {c['jan_dead_feb_active']} "
                f"previously inactive items started selling. Additionally, {c['new_items']} new SKUs were "
                f"added to the February catalogue that did not exist in January."
            )
            elements.append(Paragraph('Item Status Shifts', self.styles['insight_heading']))
            elements.append(Paragraph(shifts_text, self.styles['body']))
            elements.append(Spacer(1, 6))

            # Top 15 Gains table
            elements.append(Paragraph('TOP 15 ITEMS — BIGGEST DEMAND INCREASE', self.styles['subheading']))
            gain_header = ['#', 'Code', 'Description', 'Jan Qty', 'Feb Qty', 'Change', 'Jan MOQ', 'Feb MOQ']
            gain_rows = [gain_header]
            for idx, g in enumerate(c['top_gains'][:15], 1):
                gain_rows.append([
                    str(idx), g['code'],
                    Paragraph(g['description'][:32], ParagraphStyle('gt', fontSize=6.5)),
                    f"{g['jan_qty']:,.0f}", f"{g['feb_qty']:,.0f}",
                    f"{g['qty_change']:+,.0f}",
                    f"{g['jan_moq']:,}", f"{g['feb_moq']:,}",
                ])
            gw = [0.25*inch, 0.65*inch, 1.55*inch, 0.65*inch, 0.65*inch, 0.65*inch, 0.6*inch, 0.6*inch]
            t = Table(gain_rows, colWidths=gw, repeatRows=1)
            style = self.make_table_style(ACCENT_GREEN, pad_top=3, pad_bot=3, font_body=6.5)
            style.add('BACKGROUND', (0, 1), (-1, 3), colors.HexColor('#e8f8f5'))
            t.setStyle(style)
            elements.append(t)
            elements.append(Spacer(1, 8))

            # Top 15 Drops table
            elements.append(Paragraph('TOP 15 ITEMS — BIGGEST DEMAND DECLINE', self.styles['subheading']))
            drop_header = ['#', 'Code', 'Description', 'Jan Qty', 'Feb Qty', 'Change', 'Jan MOQ', 'Feb MOQ']
            drop_rows = [drop_header]
            for idx, d_item in enumerate(c['top_drops'][:15], 1):
                drop_rows.append([
                    str(idx), d_item['code'],
                    Paragraph(d_item['description'][:32], ParagraphStyle('dt', fontSize=6.5)),
                    f"{d_item['jan_qty']:,.0f}", f"{d_item['feb_qty']:,.0f}",
                    f"{d_item['qty_change']:+,.0f}",
                    f"{d_item['jan_moq']:,}", f"{d_item['feb_moq']:,}",
                ])
            t = Table(drop_rows, colWidths=gw, repeatRows=1)
            style = self.make_table_style(RED, pad_top=3, pad_bot=3, font_body=6.5)
            style.add('BACKGROUND', (0, 1), (-1, 3), colors.HexColor('#fdedec'))
            t.setStyle(style)
            elements.append(t)

        # ========== TOP 30 FAST MOVERS ==========
        elements.append(Paragraph('TOP 30 FAST MOVING ITEMS', self.styles['heading']))
        elements.append(Paragraph(
            'Items ranked by total quantity moved out in February 2026. '
            'These items require consistent restocking to meet demand.',
            self.styles['note']))

        fast_header = ['#', 'Code', 'Description', 'Category', 'Qty Out', 'Daily', 'Wkly MOQ']
        fast_rows = [fast_header]
        for i, item in enumerate(data['fast_movers'][:30], 1):
            fast_rows.append([
                str(i), item['item_code'],
                Paragraph(item['item_description'][:38], ParagraphStyle('td', fontSize=6.5)),
                item.get('category', '')[:16],
                f"{item['qty_out']:,.0f}",
                f"{item['daily_average']:.1f}",
                f"{item['weekly_moq']:,}"
            ])

        fw = [0.3 * inch, 0.7 * inch, 2.0 * inch, 1.1 * inch, 0.7 * inch, 0.55 * inch, 0.7 * inch]
        t = Table(fast_rows, colWidths=fw, repeatRows=1)
        style = self.make_table_style(ORANGE, pad_top=3, pad_bot=3, font_body=6.5)
        style.add('BACKGROUND', (0, 1), (-1, 5), colors.HexColor('#fef9e7'))
        t.setStyle(style)
        elements.append(t)
        elements.append(PageBreak())

        # ========== SLOW MOVERS ==========
        slow_count = summary['slow_movers']
        elements.append(Paragraph(f'SLOW MOVING ITEMS ({slow_count} items with ZERO sales)', self.styles['heading']))
        elements.append(Paragraph(
            'Items that received stock but recorded no outward movement during February. '
            'These items are potential dead stock and should be reviewed for clearance, '
            'bundling, or discontinuation to free up capital and shelf space.',
            self.styles['note']))

        slow_header = ['#', 'Code', 'Description', 'Category', 'Qty In', 'Closing']
        slow_rows = [slow_header]
        for i, item in enumerate(data['slow_movers'][:30], 1):
            cat_label = item.get('category', '')
            if not cat_label:
                for cat_name, cat_info in data['categories'].items():
                    if any(p['item_code'] == item['item_code'] for p in cat_info['products']):
                        cat_label = cat_name
                        break
            slow_rows.append([
                str(i), item['item_code'],
                Paragraph(item['item_description'][:38], ParagraphStyle('td', fontSize=6.5)),
                cat_label[:18],
                f"{item['qty_in']:,.0f}",
                f"{item.get('closing_balance', 0):,.0f}"
            ])

        sw = [0.3 * inch, 0.7 * inch, 2.0 * inch, 1.2 * inch, 0.7 * inch, 0.7 * inch]
        t = Table(slow_rows, colWidths=sw, repeatRows=1)
        t.setStyle(self.make_table_style(RED, pad_top=3, pad_bot=3, font_body=6.5))
        elements.append(t)
        if slow_count > 30:
            remaining = slow_count - 30
            elements.append(Paragraph(f'+ {remaining} more slow-moving items not shown', self.styles['small']))
        elements.append(PageBreak())

        # ========== DETAILED CATEGORY BREAKDOWN ==========
        elements.append(Paragraph('DETAILED MOQ BY CATEGORY', self.styles['title']))
        elements.append(Paragraph(
            'Each category shows its top 5 moving items with calculated MOQ values. '
            'Use this section for weekly ordering decisions.',
            self.styles['subtitle']))
        elements.append(Spacer(1, 8))

        for cat_name, cat_info in sorted_cats:
            if cat_info['total_qty_out'] == 0 and cat_info['fast_movers'] == 0:
                continue

            total_v = cat_info.get('total_weekly_value', 0)
            value_str = f'  |  Wkly Value: KES {total_v:,.0f}' if total_v > 0 else ''
            header_text = (
                f"{cat_name}  |  {cat_info['total_items']} Items  |  "
                f"Fast: {cat_info['fast_movers']}  |  Slow: {cat_info['slow_movers']}  |  "
                f"Wkly MOQ: {cat_info['total_weekly_moq']:,}{value_str}"
            )
            header_table = Table([[Paragraph(header_text,
                ParagraphStyle('ch', fontSize=8, textColor=colors.white,
                               fontName='Helvetica-Bold'))]], colWidths=[self.usable_width])
            header_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(CAT_HEADER_BG)),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ]))

            products = [p for p in cat_info['products'] if p['qty_out'] > 0][:5]
            if not products:
                elements.append(KeepTogether([header_table, Spacer(1, 8)]))
                continue

            prod_header = ['Code', 'Description', 'Qty In', 'Qty Out', 'Closing', 'Daily', 'Wkly MOQ']
            prod_rows = [prod_header]
            for p in products:
                prod_rows.append([
                    p['item_code'],
                    Paragraph(p['item_description'][:35],
                              ParagraphStyle('pd', fontSize=6.5)),
                    f"{p['qty_in']:,.0f}",
                    f"{p['qty_out']:,.0f}",
                    f"{p.get('closing_balance', 0):,.0f}",
                    f"{p['daily_average']:.1f}",
                    f"{p['weekly_moq']:,}"
                ])

            pw = [0.7 * inch, 2.1 * inch, 0.65 * inch, 0.65 * inch, 0.65 * inch, 0.5 * inch, 0.7 * inch]
            prod_table = Table(prod_rows, colWidths=pw, repeatRows=1)
            prod_table.setStyle(self.make_table_style(GREY_HEADER, pad_top=3, pad_bot=3, font_body=6.5))

            remaining = len([p for p in cat_info['products'] if p['qty_out'] > 0]) - len(products)
            group = [header_table, prod_table]
            if remaining > 0:
                group.append(Paragraph(f'+ {remaining} more items with movement', self.styles['small']))
            group.append(Spacer(1, 10))

            elements.append(KeepTogether(group[:3]))
            if len(group) > 3:
                for g in group[3:]:
                    elements.append(g)

        doc.build(elements)
        print(f'Report generated: {self.filename}')


class MOQReportWord:
    """Generates a Word (.docx) version of the MOQ report."""

    def __init__(self, filename='moq_report.docx'):
        self.filename = filename

    def hex_to_rgb(self, hex_color):
        h = hex_color.lstrip('#')
        return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

    def set_cell_shading(self, cell, hex_color):
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), hex_color.lstrip('#'))
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    def set_row_shading(self, row, hex_color):
        for cell in row.cells:
            self.set_cell_shading(cell, hex_color)

    def add_insight_block(self, doc, title, text, title_color=DARK_BLUE):
        """Add a titled insight paragraph with shaded background."""
        heading = doc.add_paragraph()
        run = heading.add_run(title)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*self.hex_to_rgb(title_color))

        # Use a single-cell table for the shaded box
        tbl = doc.add_table(rows=1, cols=1)
        cell = tbl.rows[0].cells[0]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(*self.hex_to_rgb(GREY_TEXT))
        self.set_cell_shading(cell, INFO_BG)

    def generate(self):
        data = load_data()
        insights = build_insights(data)
        summary = data['summary']
        sorted_cats = sorted(data['categories'].items(),
                             key=lambda x: x[1]['total_qty_out'], reverse=True)

        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.8)
            section.right_margin = Cm(1.8)

        # ========== TITLE ==========
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('BOMAS HARDWARE — MOQ REPORT')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*self.hex_to_rgb(NAVY))

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(
            'Period: February 1-28, 2026  •  24 Working Days (4 weeks × 6 days)  •  '
            'Daily Avg = Qty Out ÷ 24  •  MOQ = Daily Avg × 6'
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*self.hex_to_rgb(GREY_TEXT))

        # ========== KPI SUMMARY ==========
        heading = doc.add_paragraph()
        run = heading.add_run('SUMMARY')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*self.hex_to_rgb(NAVY))

        kpi_table = doc.add_table(rows=1, cols=6)
        kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        kpi_headers = ['Total Items', 'Fast Movers', 'Slow Movers', 'Total Qty Out', 'Weekly MOQ', 'Weekly Value']
        kpi_values = [
            f"{summary['total_items']:,}", f"{summary['fast_movers']:,}", f"{summary['slow_movers']:,}",
            f"{summary['total_qty_out']:,.0f}", f"{summary['total_weekly_moq']:,}",
            f"KES {summary.get('total_weekly_value', 0):,.0f}"
        ]

        for i, h in enumerate(kpi_headers):
            cell = kpi_table.rows[0].cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            self.set_cell_shading(cell, TEAL)

        val_row = kpi_table.add_row()
        for i, v in enumerate(kpi_values):
            cell = val_row.cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(v)
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(*self.hex_to_rgb(DARK_BLUE))
            self.set_cell_shading(cell, TEAL_LIGHT)

        doc.add_paragraph()

        # ========== EXECUTIVE SUMMARY ==========
        heading = doc.add_paragraph()
        run = heading.add_run('EXECUTIVE SUMMARY')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*self.hex_to_rgb(NAVY))

        p = doc.add_paragraph()
        run = p.add_run(insights['executive_summary'])
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*self.hex_to_rgb(GREY_TEXT))

        # ========== CATEGORY SUMMARY ==========
        heading = doc.add_paragraph()
        run = heading.add_run('MOQ BY CATEGORY')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*self.hex_to_rgb(NAVY))

        cat_table = doc.add_table(rows=1, cols=7)
        cat_headers = ['Category', 'Items', 'Fast', 'Slow', 'Qty Out', 'Wkly MOQ', 'Wkly Value (KES)']
        for i, h in enumerate(cat_headers):
            cell = cat_table.rows[0].cells[i]
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            if i >= 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self.set_cell_shading(cell, DARK_BLUE)

        for idx, (cat_name, cat_info) in enumerate(sorted_cats):
            if cat_info['total_qty_out'] > 0 or cat_info['total_items'] > 0:
                row = cat_table.add_row()
                vals = [
                    cat_name[:25],
                    str(cat_info['total_items']),
                    str(cat_info['fast_movers']),
                    str(cat_info['slow_movers']),
                    f"{cat_info['total_qty_out']:,.0f}",
                    f"{cat_info['total_weekly_moq']:,}",
                    f"{cat_info.get('total_weekly_value', 0):,.0f}"
                ]
                for i, v in enumerate(vals):
                    cell = row.cells[i]
                    p = cell.paragraphs[0]
                    run = p.add_run(v)
                    run.font.size = Pt(7.5)
                    if i >= 1:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if idx % 2 == 1:
                    self.set_row_shading(row, ROW_ALT)

        total_row = cat_table.add_row()
        total_vals = [
            'TOTAL', str(summary['total_items']), str(summary['fast_movers']),
            str(summary['slow_movers']), f"{summary['total_qty_out']:,.0f}",
            f"{summary['total_weekly_moq']:,}",
            f"{summary.get('total_weekly_value', 0):,.0f}"
        ]
        for i, v in enumerate(total_vals):
            cell = total_row.cells[i]
            p = cell.paragraphs[0]
            run = p.add_run(v)
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            if i >= 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self.set_cell_shading(cell, ACCENT_GREEN)

        doc.add_page_break()

        # ========== DEMAND ANALYSIS & INSIGHTS ==========
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run('DEMAND ANALYSIS & INSIGHTS')
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*self.hex_to_rgb(NAVY))

        self.add_insight_block(doc, 'Demand Concentration', insights['demand_concentration'])
        self.add_insight_block(doc, 'Category Value Distribution', insights['category_value'])
        self.add_insight_block(doc, 'Dead Stock & Slow Movers', insights['dead_stock'])
        self.add_insight_block(doc, 'Why MOQs Matter', insights['moq_importance'])
        self.add_insight_block(doc, 'Cash Flow & Capital Allocation', insights['cash_flow'])

        doc.add_page_break()

        # ========== RECOMMENDATIONS ==========
        heading = doc.add_paragraph()
        run = heading.add_run('KEY RECOMMENDATIONS')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*self.hex_to_rgb(NAVY))

        note = doc.add_paragraph()
        run = note.add_run('Based on the February 2026 movement data, the following actions are recommended:')
        run.font.size = Pt(8.5)
        run.font.italic = True
        run.font.color.rgb = RGBColor(*self.hex_to_rgb(GREY_TEXT))

        for i, rec in enumerate(insights['recommendations'], 1):
            p = doc.add_paragraph()
            run = p.add_run(f'{i}.  ')
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(*self.hex_to_rgb(DARK_BLUE))
            run = p.add_run(rec)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(*self.hex_to_rgb(GREY_TEXT))

        doc.add_paragraph()

        # ========== JAN vs FEB COMPARISON ==========
        comparison = build_jan_feb_comparison()
        if comparison:
            doc.add_page_break()

            heading = doc.add_paragraph()
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = heading.add_run('JANUARY vs FEBRUARY — MOQ COMPARISON')
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(*self.hex_to_rgb(NAVY))

            sub = doc.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = sub.add_run(
                f'{comparison["jan_period"]}  →  {comparison["feb_period"]}  |  '
                f'Both months: 24 working days  |  {comparison["common_items"]:,} common items compared'
            )
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(*self.hex_to_rgb(GREY_TEXT))

            # Overview narrative
            self.add_insight_block(doc, 'Overview', comparison['narrative'])

            # Summary comparison KPIs
            c = comparison
            def fmt_chg(val, pct=None):
                sign = '+' if val >= 0 else ''
                pct_str = f' ({sign}{pct:.1f}%)' if pct is not None else ''
                return f'{sign}{val:,.0f}{pct_str}'

            heading = doc.add_paragraph()
            run = heading.add_run('Key Metrics Comparison')
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(*self.hex_to_rgb(DARK_BLUE))

            cmp_table = doc.add_table(rows=1, cols=4)
            cmp_headers = ['Metric', 'January', 'February', 'Change']
            for i, h in enumerate(cmp_headers):
                cell = cmp_table.rows[0].cells[i]
                p = cell.paragraphs[0]
                run = p.add_run(h)
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                if i >= 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                self.set_cell_shading(cell, ACCENT_BLUE)

            cmp_data = [
                ['Total Items Tracked', f"{c['jan_items']:,}", f"{c['feb_items']:,}",
                 fmt_chg(c['feb_items'] - c['jan_items'])],
                ['Active Items (Fast)', f"{c['jan_fast']:,}", f"{c['feb_fast']:,}",
                 fmt_chg(c['feb_fast'] - c['jan_fast'])],
                ['Inactive Items (Slow)', f"{c['jan_slow']:,}", f"{c['feb_slow']:,}",
                 fmt_chg(c['feb_slow'] - c['jan_slow'])],
                ['Total Qty Out', f"{c['jan_total_out']:,.0f}", f"{c['feb_total_out']:,.0f}",
                 fmt_chg(c['qty_change'], c['qty_change_pct'])],
                ['Weekly MOQ (units)', f"{c['jan_moq']:,}", f"{c['feb_moq']:,}",
                 fmt_chg(c['moq_change'])],
            ]
            for idx, vals in enumerate(cmp_data):
                row = cmp_table.add_row()
                for i, v in enumerate(vals):
                    cell = row.cells[i]
                    p = cell.paragraphs[0]
                    run = p.add_run(v)
                    run.font.size = Pt(8)
                    if i == 0:
                        run.font.bold = True
                    if i >= 1:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if idx % 2 == 1:
                    self.set_row_shading(row, ROW_ALT)

            doc.add_paragraph()

            # Status shifts
            shifts_text = (
                f"Between the two months, {c['jan_active_feb_dead']} items that had active sales in January "
                f"went dormant in February (potential demand loss or stock issues), while {c['jan_dead_feb_active']} "
                f"previously inactive items started selling. Additionally, {c['new_items']} new SKUs were "
                f"added to the February catalogue that did not exist in January."
            )
            self.add_insight_block(doc, 'Item Status Shifts', shifts_text)

            # Top 15 Gains
            heading = doc.add_paragraph()
            run = heading.add_run('TOP 15 ITEMS — BIGGEST DEMAND INCREASE')
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(*self.hex_to_rgb(ACCENT_GREEN))

            gain_table = doc.add_table(rows=1, cols=8)
            gain_headers = ['#', 'Code', 'Description', 'Jan Qty', 'Feb Qty', 'Change', 'Jan MOQ', 'Feb MOQ']
            for i, h in enumerate(gain_headers):
                cell = gain_table.rows[0].cells[i]
                p = cell.paragraphs[0]
                run = p.add_run(h)
                run.font.size = Pt(7.5)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                if i >= 3:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                self.set_cell_shading(cell, ACCENT_GREEN)

            for idx, g in enumerate(c['top_gains'][:15], 1):
                row = gain_table.add_row()
                vals = [
                    str(idx), g['code'], g['description'][:32],
                    f"{g['jan_qty']:,.0f}", f"{g['feb_qty']:,.0f}",
                    f"{g['qty_change']:+,.0f}",
                    f"{g['jan_moq']:,}", f"{g['feb_moq']:,}",
                ]
                for i, v in enumerate(vals):
                    cell = row.cells[i]
                    p = cell.paragraphs[0]
                    run = p.add_run(v)
                    run.font.size = Pt(7)
                    if i >= 3:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if idx <= 3:
                    self.set_row_shading(row, '#e8f8f5')
                elif idx % 2 == 0:
                    self.set_row_shading(row, ROW_ALT)

            doc.add_paragraph()

            # Top 15 Drops
            heading = doc.add_paragraph()
            run = heading.add_run('TOP 15 ITEMS — BIGGEST DEMAND DECLINE')
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(*self.hex_to_rgb(RED))

            drop_table = doc.add_table(rows=1, cols=8)
            drop_headers = ['#', 'Code', 'Description', 'Jan Qty', 'Feb Qty', 'Change', 'Jan MOQ', 'Feb MOQ']
            for i, h in enumerate(drop_headers):
                cell = drop_table.rows[0].cells[i]
                p = cell.paragraphs[0]
                run = p.add_run(h)
                run.font.size = Pt(7.5)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                if i >= 3:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                self.set_cell_shading(cell, RED)

            for idx, d_item in enumerate(c['top_drops'][:15], 1):
                row = drop_table.add_row()
                vals = [
                    str(idx), d_item['code'], d_item['description'][:32],
                    f"{d_item['jan_qty']:,.0f}", f"{d_item['feb_qty']:,.0f}",
                    f"{d_item['qty_change']:+,.0f}",
                    f"{d_item['jan_moq']:,}", f"{d_item['feb_moq']:,}",
                ]
                for i, v in enumerate(vals):
                    cell = row.cells[i]
                    p = cell.paragraphs[0]
                    run = p.add_run(v)
                    run.font.size = Pt(7)
                    if i >= 3:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if idx <= 3:
                    self.set_row_shading(row, '#fdedec')
                elif idx % 2 == 0:
                    self.set_row_shading(row, ROW_ALT)

        # ========== TOP 30 FAST MOVERS ==========
        heading = doc.add_paragraph()
        run = heading.add_run('TOP 30 FAST MOVING ITEMS')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*self.hex_to_rgb(NAVY))

        note = doc.add_paragraph()
        run = note.add_run(
            'Items ranked by total quantity moved out in February 2026. '
            'These items require consistent restocking to meet demand.'
        )
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = RGBColor(*self.hex_to_rgb(GREY_TEXT))

        fast_table = doc.add_table(rows=1, cols=7)
        fast_headers = ['#', 'Code', 'Description', 'Category', 'Qty Out', 'Daily', 'Wkly MOQ']
        for i, h in enumerate(fast_headers):
            cell = fast_table.rows[0].cells[i]
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            if i >= 4:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self.set_cell_shading(cell, ORANGE)

        for idx, item in enumerate(data['fast_movers'][:30], 1):
            row = fast_table.add_row()
            vals = [
                str(idx), item['item_code'],
                item['item_description'][:38],
                item.get('category', '')[:16],
                f"{item['qty_out']:,.0f}",
                f"{item['daily_average']:.1f}",
                f"{item['weekly_moq']:,}"
            ]
            for i, v in enumerate(vals):
                cell = row.cells[i]
                p = cell.paragraphs[0]
                run = p.add_run(v)
                run.font.size = Pt(7)
                if i >= 4:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if idx <= 5:
                self.set_row_shading(row, '#fef9e7')
            elif idx % 2 == 0:
                self.set_row_shading(row, ROW_ALT)

        doc.add_page_break()

        # ========== SLOW MOVERS ==========
        slow_count = summary['slow_movers']
        heading = doc.add_paragraph()
        run = heading.add_run(f'SLOW MOVING ITEMS ({slow_count} items with ZERO sales)')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*self.hex_to_rgb(NAVY))

        note = doc.add_paragraph()
        run = note.add_run(
            'Items that received stock but recorded no outward movement during February. '
            'These items are potential dead stock and should be reviewed for clearance, '
            'bundling, or discontinuation to free up capital and shelf space.'
        )
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = RGBColor(*self.hex_to_rgb(GREY_TEXT))

        slow_table = doc.add_table(rows=1, cols=6)
        slow_headers = ['#', 'Code', 'Description', 'Category', 'Qty In', 'Closing']
        for i, h in enumerate(slow_headers):
            cell = slow_table.rows[0].cells[i]
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            if i >= 4:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self.set_cell_shading(cell, RED)

        for idx, item in enumerate(data['slow_movers'][:30], 1):
            cat_label = item.get('category', '')
            if not cat_label:
                for cat_name, cat_info in data['categories'].items():
                    if any(p['item_code'] == item['item_code'] for p in cat_info['products']):
                        cat_label = cat_name
                        break
            row = slow_table.add_row()
            vals = [
                str(idx), item['item_code'],
                item['item_description'][:38],
                cat_label[:18],
                f"{item['qty_in']:,.0f}",
                f"{item.get('closing_balance', 0):,.0f}"
            ]
            for i, v in enumerate(vals):
                cell = row.cells[i]
                p = cell.paragraphs[0]
                run = p.add_run(v)
                run.font.size = Pt(7)
                if i >= 4:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if idx % 2 == 0:
                self.set_row_shading(row, ROW_ALT)

        if slow_count > 30:
            remaining = slow_count - 30
            note = doc.add_paragraph()
            run = note.add_run(f'+ {remaining} more slow-moving items not shown')
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(*self.hex_to_rgb(GREY_LIGHT))

        doc.add_page_break()

        # ========== DETAILED CATEGORY BREAKDOWN ==========
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run('DETAILED MOQ BY CATEGORY')
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*self.hex_to_rgb(NAVY))

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(
            'Each category shows its top 5 moving items with calculated MOQ values. '
            'Use this section for weekly ordering decisions.'
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*self.hex_to_rgb(GREY_TEXT))

        for cat_name, cat_info in sorted_cats:
            if cat_info['total_qty_out'] == 0 and cat_info['fast_movers'] == 0:
                continue

            total_v = cat_info.get('total_weekly_value', 0)
            value_str = f'  |  Wkly Value: KES {total_v:,.0f}' if total_v > 0 else ''

            cat_header = doc.add_paragraph()
            run = cat_header.add_run(
                f"{cat_name}  |  {cat_info['total_items']} Items  |  "
                f"Fast: {cat_info['fast_movers']}  |  Slow: {cat_info['slow_movers']}  |  "
                f"Wkly MOQ: {cat_info['total_weekly_moq']:,}{value_str}"
            )
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            pPr = cat_header._p.get_or_add_pPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), NAVY.lstrip('#'))
            shading.set(qn('w:val'), 'clear')
            pPr.append(shading)

            products = [p for p in cat_info['products'] if p['qty_out'] > 0][:5]
            if not products:
                continue

            prod_table = doc.add_table(rows=1, cols=7)
            prod_headers = ['Code', 'Description', 'Qty In', 'Qty Out', 'Closing', 'Daily', 'Wkly MOQ']
            for i, h in enumerate(prod_headers):
                cell = prod_table.rows[0].cells[i]
                p = cell.paragraphs[0]
                run = p.add_run(h)
                run.font.size = Pt(7.5)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                if i >= 2:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                self.set_cell_shading(cell, GREY_HEADER)

            for idx, p_item in enumerate(products):
                row = prod_table.add_row()
                vals = [
                    p_item['item_code'],
                    p_item['item_description'][:35],
                    f"{p_item['qty_in']:,.0f}",
                    f"{p_item['qty_out']:,.0f}",
                    f"{p_item.get('closing_balance', 0):,.0f}",
                    f"{p_item['daily_average']:.1f}",
                    f"{p_item['weekly_moq']:,}"
                ]
                for i, v in enumerate(vals):
                    cell = row.cells[i]
                    pp = cell.paragraphs[0]
                    run = pp.add_run(v)
                    run.font.size = Pt(7)
                    if i >= 2:
                        pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if idx % 2 == 1:
                    self.set_row_shading(row, ROW_ALT)

            remaining = len([p for p in cat_info['products'] if p['qty_out'] > 0]) - len(products)
            if remaining > 0:
                note = doc.add_paragraph()
                run = note.add_run(f'+ {remaining} more items with movement')
                run.font.size = Pt(7)
                run.font.color.rgb = RGBColor(*self.hex_to_rgb(GREY_LIGHT))

            doc.add_paragraph()

        # ========== FOOTER ==========
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(f'Generated: {datetime.now().strftime("%d %b %Y %H:%M")}  |  Confidential — BOMAS Hardware Store')
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(*self.hex_to_rgb(GREY_LIGHT))

        try:
            doc.save(self.filename)
        except PermissionError:
            self.filename = self.filename.replace('.docx', '_v2.docx')
            doc.save(self.filename)
        print(f'Word report generated: {self.filename}')


if __name__ == '__main__':
    report = MOQReport('moq_report.pdf')
    report.generate()

    word_report = MOQReportWord('moq_report.docx')
    word_report.generate()
