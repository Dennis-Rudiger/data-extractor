from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
from datetime import datetime

def set_cell_background(cell, color_hex):
    properties = cell._element.tcPr
    if properties is None:
        return
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    properties.append(shading)

def generate_stock_take():
    doc = Document()
    
    # Title
    title = doc.add_heading('BOMAS BRANCH - STOCK TAKE REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run("Date: April 21, 2026\n").bold = True
    p.add_run("Branch Location: Bomas\n").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Executive Summary
    doc.add_heading('1. The Strategic Importance of Inventory Management', level=1)
    summary_text = (
        "Effective inventory management is the lifeblood of hardware retail profitability. A rigorous stock take "
        "reconciles physical counts with system records, identifying shrinkage, dead stock, and capital allocation errors. "
        "In highly competitive categories like cement, steel, and electricals, holding the correct minimum order "
        "quantities (MOQ) without tying up excess capital is critical. This stock take establishes a clean baseline "
        "for the Bomas branch, enabling accurate forecasting and preventing stockouts of key movers."
    )
    doc.add_paragraph(summary_text)
    
    # Methodology
    doc.add_heading('2. Methodology & Scope', level=1)
    method_text = (
        "The physical count was conducted against the system baseline for the Bomas Branch. Items across various "
        "departments-including Agricultural Tools, Cement, Plumbing, Timber, and Electrical Fittings-were mapped "
        "via their system SKU and cross-verified via floor-to-sheet counting."
    )
    doc.add_paragraph(method_text)
    
    # Load Data
    try:
        with open('inventory_count_feb.json', 'r') as f:
            inventory = json.load(f)
    except Exception as e:
        inventory = []
        doc.add_paragraph("Warning: Could not read inventory_count_feb.json.")
        
    doc.add_heading('3. Category System Baselines (Bomas)', level=1)
    
    if inventory:
        categories = {}
        for item in inventory:
            cat = item.get('category', 'UNKNOWN')
            qty = item.get('system_qty', 0)
            if cat not in categories:
                categories[cat] = {'items': 0, 'total': 0}
            categories[cat]['items'] += 1
            if isinstance(qty, (int, float)):
                categories[cat]['total'] += qty
            
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Hardware Category'
        hdr_cells[1].text = 'Unique SKUs Counted'
        hdr_cells[2].text = 'Total Expected Pieces (System)'
        
        for cell in hdr_cells:
            set_cell_background(cell, "D9D9D9")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        for c, d in sorted(categories.items(), key=lambda x: x[1]['total'], reverse=True):
            row = table.add_row().cells
            row[0].text = str(c)
            row[1].text = str(d['items'])
            row[2].text = f"{d['total']:,.1f}"
            
    # Action Plans
    doc.add_heading('4. Variances Investigation & Action Plan', level=1)
    
    actions = [
        "Cycle Counting: Implement high-frequency cyclic counts for high-risk categories (Electrical Cables/Door Locks) and bulk commodities (Cement).",
        "Shrinkage Remediation: Any variance > 2% must trigger a direct review of the respective day's receiving logs.",
        "MOQ Alignments: Shift purchasing strategies based on Bomas baseline depletion-order closer to JIT (Just-in-Time) for fast movers to free up working capital.",
        "Dead Stock Liquidation: SKUs showing zero movement over 90 days should be marked for discount or transferred if demand exists at other branches."
    ]
    
    for action in actions:
        p = doc.add_paragraph()
        p.add_run(' ').bold = True
        p.add_run(action)
    
    doc.save('Bomas_Stock_Take_Report.docx')
    print("Bomas_Stock_Take_Report.docx generated successfully.")

generate_stock_take()
