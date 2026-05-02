from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from generate_market_analysis import MARKET_BENCHMARKS

def set_cell_background(cell, color_hex):
    properties = cell._element.tcPr
    if properties is None:
        return
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    properties.append(shading)

def generate_word_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('MARKET COMPETITIVE ANALYSIS REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle / Header info
    p = doc.add_paragraph()
    p.add_run("Date: April 18, 2026\n").bold = True
    p.add_run("Focus: Top Movers Benchmarking & Sourcing Alternatives").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph("This report outlines current standard market pricing across highly competitive hardware product lines. Data is sourced from 15 active hardware stores and suppliers including: Cedar Clink, Hardware Homes, Randtech, TopTank, Polytanks Africa, Construction Kenya, Builders Kenya, Jiji Kenya, Integrum Construction, EAPC, Jumia, and Tronic.")
    
    doc.add_heading('Category Benchmarks', level=1)
    
    for category, details in MARKET_BENCHMARKS.items():
        doc.add_heading(category.replace('_', ' ').title(), level=2)
        desc = doc.add_paragraph()
        desc.add_run("Strategy/Margin Note: ").bold = True
        desc.add_run(details['category_desc'])
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Item Name'
        hdr_cells[1].text = 'Target Price (KES)'
        hdr_cells[2].text = 'Market Range (KES)'
        hdr_cells[3].text = 'Data Sources'
        
        for cell in hdr_cells:
            set_cell_background(cell, "D9D9D9")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    
        for item in details['items']:
            row_cells = table.add_row().cells
            row_cells[0].text = item['name']
            row_cells[1].text = f"KES {item['market_price']:,.2f}"
            row_cells[2].text = item['market_range']
            row_cells[3].text = item['source']
            
        doc.add_paragraph() # Spacing

    # Strategic Recommendations (copied from PDF generator)
    doc.add_page_break()
    doc.add_heading('STRATEGIC RECOMMENDATIONS', level=1)

    recommendations = [
        {
            'title': 'CEMENT & BULK MATERIALS',
            'recs': [
                'Maintain 6-8% margins to remain competitive (current market standard)',
                'Simba at KES 850, Bamburi at KES 730, Savannah at KES 695 - track weekly',
                'Consider bulk discounts for orders over 50 bags to win contractor business',
                'Use cement as a traffic driver, not primary profit center',
                'Offer delivery for bulk orders as value-add differentiation'
            ]
        },
        {
            'title': 'WATER TANKS',
            'recs': [
                'Strong margin opportunities (12-20%) due to brand differentiation',
                '1000L tanks: Market KES 9,500-14,700 (position at KES 11,000-13,000)',
                '5000L tanks: Market KES 40,000-45,000 (position at KES 42,000-43,500)',
                '10000L tanks: Market KES 90,700-146,000 (wide variance by brand)',
                'Stock multiple brands: RotoTank (budget), TopTank (mid), Premium (high-end)',
                'Offer installation services to justify higher prices',
                'Bundle with stands, taps, and delivery for complete solution'
            ]
        },
        {
            'title': 'MARINE BOARDS & PLYWOOD',
            'recs': [
                'Clear brand tiering: Budget KES 2,100-2,450 vs Premium KES 2,650-3,500',
                'Standard marine boards: Position at KES 2,350-2,500 (competitive mid-point)',
                'Premium brands (Bornwood, Honor Plex): Can command KES 2,800-3,200',
                'MDF boards market at KES 2,800-3,600 (stock quality to justify higher price)',
                'Blockboards at KES 2,900-3,200 - moderate margin potential',
                'Educate customers on quality differences to support premium pricing',
                'Offer cutting services as value-add (charge KES 50-100 per cut)'
            ]
        },
        {
            'title': 'ELECTRICAL CABLES & WIRES',
            'recs': [
                'Commodity pricing with brand premium for EA Cables',
                'Single core 1.5mm: KES 30-40/meter or KES 6,000-7,000 per 100M roll',
                'Single core 4.0mm: KES 8,500-10,000 per 100M roll',
                'Twin & Earth cables: 1.5mm at KES 2,100-2,900, 2.5mm at KES 3,200-4,000',
                'Margins typically 10-18% - stay competitive on price',
                'Stock both budget and premium brands to serve all segments',
                'Offer bulk discounts for electrician/contractor purchases'
            ]
        },
        {
            'title': 'ELECTRICAL FIXTURES & LOCKS',
            'recs': [
                'Significant margin opportunity exists (15-35% achievable)',
                'Basic door locks: Market KES 1,200-2,500 (position at KES 1,400-1,800)',
                'Premium 4-pin locks: Market KES 4,500-6,000 (target KES 5,000)',
                'Steel premium locks: KES 4,000-7,500 (position at KES 5,500-6,500)',
                'Kitchen mixers: Wide range KES 2,000-10,000 (quality matters)',
                'Bathroom accessories: Low-end KES 150-500, high-end KES 1,500-3,000',
                'Focus on mid-range quality products with good margins'
            ]
        },
        {
            'title': 'SANITARY WARE',
            'recs': [
                'Close couple toilets: Market average KES 13,500 (range KES 10,000-22,500)',
                'One piece toilets: Market average KES 28,000 (range KES 17,500-47,500)',
                'Vanity cabinets: Basic KES 7,500-25,000, Premium KES 25,000-65,000',
                'Wide price variance indicates quality/brand differentiation opportunity',
                'Position mid-range for volume, premium for margin',
                'Bundle toilet + seat + tank fittings for complete package deals',
                'Offer installation referrals or in-house service for premium positioning'
            ]
        },
        {
            'title': 'STEEL PRODUCTS',
            'recs': [
                'Binding wire 16G: KES 3,850-4,800 per 25kg (very competitive)',
                'Angle lines pricing by size: 1x1 at KES 1,040, 2x2 at KES 2,370-3,585',
                'Black sheets: 16G at KES 5,500-6,500, 18G at KES 4,400-4,500',
                'Maintain 6-8% margins maximum - this is commodity territory',
                'Consider value-adds: delivery, cutting services, bulk discounts',
                'Monitor global steel prices monthly for cost adjustments',
                'Focus on volume over margin for steel products'
            ]
        },
        {
            'title': 'TIMBER & BOARDS',
            'recs': [
                'MDF boards: Market KES 2,500-3,500 (position at KES 2,800-3,200)',
                'Plywood pricing very competitive (KES 2,000-3,500)',
                'Blockboards: KES 2,900-3,200 (moderate margin opportunity)',
                'Chipboard 18mm: KES 3,400-3,800 (less common, good margins)',
                'Gypsum boards: Market KES 700-1,200 (competitive pricing essential)',
                'Value-added processing (cutting, edging) can justify 15-20% premium',
                'Stock both imported and local to serve different price points'
            ]
        },
        {
            'title': 'BUILDING MATERIALS',
            'recs': [
                'Building stones 6x6: KES 60-68 (very tight margin, 8-12% max)',
                'Machine cut stones: KES 45-60 (price sensitive market)',
                'Ballast/Sand: Commodity pricing, focus on delivery convenience',
                'These are loss leaders - price aggressively to win project bids',
                'Upsell higher-margin items (cement, steel) once customer is committed',
                'Reliable delivery and consistent quality are key differentiators'
            ]
        },
        {
            'title': 'PAINTS',
            'recs': [
                'Crown Paints 4L: Market average KES 2,500 (range KES 2,200-2,800)',
                'Basco Paint 4L: KES 2,300 (range KES 2,000-2,600)',
                'Premium brands (Sadolin): KES 2,800-3,600 command higher prices',
                'Crown 20L: KES 10,000-12,500 (volume packaging offers better margins)',
                'Brand loyalty is strong - stock multiple brands (Crown, Basco, Sadolin)',
                'Your 8.9-10% margins align with market (appropriate for competition)',
                'Offer color mixing services and technical advice as differentiation',
                'Stock wood finishes/varnish at KES 2,400-3,200 for complete range'
            ]
        },
        {
            'title': 'PLUMBING (PPR/PVC)',
            'recs': [
                'PPR Pipes: High demand, stock PN16/PN20 for quality assurance',
                'PPR 20mm at KES 350, 25mm at KES 550 - competitive entry points',
                'PVC Pipes: Class B is standard, Heavy Gauge for drainage',
                'PVC 4" Heavy Gauge: KES 1,600-1,800 (good margin item)',
                'Stock fittings (Elbows, Tees, Sockets) - high volume, 30%+ margin',
                'Partner with plumbers for recurring sales'
            ]
        },
        {
            'title': 'FLOOR TILES',
            'recs': [
                'Ceramic 30x30/40x40: Budget friendly (KES 1,100-1,400/box)',
                'Porcelain 60x60: The growth category (KES 1,800-2,400/box)',
                'Premium Porcelain: Niche market (KES 3,000+), stock limited quantity',
                'Display is key - show installed samples to drive sales',
                'Cross-sell tile adhesive (KES 600-800) and grout'
            ]
        },
        {
            'title': 'ROOFING SHEETS',
            'recs': [
                'Dumuzas 30G: The market standard (KES 550-650/meter)',
                'Coloured sheets (Resincot): Higher value (KES 700+/meter)',
                'Stock standard lengths (2m, 2.5m, 3m) to minimize cutting waste',
                'Margins are thin (8-12%), focus on volume and project supply',
                'Offer transport for bulk orders (critical for roofing)'
            ]
        },
        {
            'title': 'AGRICULTURAL TOOLS',
            'recs': [
                'High margin category (25-40%) - "Crocodile" brand is king',
                'Jembes: Stock Crocodile (KES 1,200+) and budget options',
                'Wheelbarrows: Heavy duty Jua Kali (KES 5,500+) preferred for durability',
                'Pangas: Fast moving, keep well-stocked (KES 450-800)',
                'Target seasonal planting/harvesting times for promotions'
            ]
        }
    ]

    for rec in recommendations:
        doc.add_heading(rec['title'], level=2)
        for item in rec['recs']:
            p = doc.add_paragraph()
            p.add_run('• ').bold = True
            p.add_run(item)
        doc.add_paragraph()

    # Competitive Positioning
    doc.add_page_break()
    doc.add_heading('COMPETITIVE POSITIONING STRATEGY', level=1)

    positioning_paragraphs = [
        ('PRICE LEADERSHIP CATEGORIES (Match or beat market):', [
            'Cement and bulk materials (6% margin)',
            'Steel products & Roofing (6-8% margin)',
            'Paints (8.9-10% margin)'
        ]),
        ('BALANCED PRICING CATEGORIES (Match market, healthy margins):', [
            'Timber and boards (15-25% margin)',
            'Tiles and accessories (15-30% margin)',
            'General hardware & Plumbing (20-30% margin)'
        ]),
        ('PREMIUM MARGIN CATEGORIES (Value-add positioning):', [
            'Electrical accessories (18-36% margin)',
            'Sanitary ware (15-25% margin)',
            'Agricultural tools (37-47% margin)'
        ])
    ]

    for title, bullets in positioning_paragraphs:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        for b in bullets:
            pb = doc.add_paragraph()
            pb.add_run('• ').bold = True
            pb.add_run(b)

    # Market intelligence sources & key findings
    doc.add_heading('MARKET INTELLIGENCE SOURCES & KEY FINDINGS', level=2)
    doc.add_paragraph('Cedar Clink Hardware (Kimathi Street, Nairobi) - Premium positioning')
    doc.add_paragraph('Hardware Homes (Industrial Area, Funzi Road) - Mid-range volume')
    doc.add_paragraph('A&D Store (Eastern Bypass, Ruiru) - Competitive bulk pricing')
    doc.add_paragraph('Randtech (Ruiru Town) - Online hardware competitive pricing')
    doc.add_paragraph('Fastlane Hardware - Emerging online competitor')
    doc.add_paragraph()
    doc.add_paragraph('KEY FINDINGS:')
    findings = [
        'Your electrical margins (8-36%) are now competitive after recent adjustments',
        'Cement pricing must remain at 6% to compete with market leaders',
        'Sanitary ware has significant price variance - quality matters more than price',
        'Specialty categories (Agricultural, Jua Kali) offer best margin opportunities',
        'Service differentiation is crucial in commodity categories'
    ]
    for f in findings:
        pf = doc.add_paragraph()
        pf.add_run('• ').bold = True
        pf.add_run(f)

    # Action Items
    doc.add_page_break()
    doc.add_heading('IMMEDIATE ACTION ITEMS', level=1)
    actions = [
        'Verify cement pricing weekly against Simba (KES 850) and Bamburi (KES 730) market rates',
        'Review electrical item pricing quarterly - market is dynamic',
        'Implement price monitoring for top 50 items across 3-5 competitors',
        'Develop value-add services: delivery, cutting, technical consultation',
        'Create bundle offerings for common project needs (e.g., bathroom packages)',
        'Train staff on market positioning and value proposition for each category',
        'Consider loyalty program for repeat customers (bulk buyers, contractors)',
        'Monitor online competitors monthly for pricing shifts'
    ]
    for a in actions:
        pa = doc.add_paragraph()
        pa.add_run('• ').bold = True
        pa.add_run(a)

    # Conclusion
    doc.add_page_break()
    doc.add_heading('CONCLUSION', level=1)
    conclusion = ('The Nairobi hardware market is highly competitive but offers strategic opportunities for '
                  'differentiation. Your current pricing strategy aligns well with market dynamics, particularly after the recent '
                  'electrical margin adjustments. Success depends on: (1) Aggressive pricing on commodity items to drive traffic, '
                  '(2) Healthy margins on specialty items to ensure profitability, and (3) Service excellence to justify premium '
                  'positioning where appropriate. Regular market monitoring (monthly for commodities, quarterly for specialty items) '
                  'will ensure continued competitiveness. Focus on the total customer experience - price is just one factor in a complex buying decision.')
    doc.add_paragraph(conclusion)

    doc.save('market_competitive_analysis.docx')
    print("Report generated successfully: market_competitive_analysis.docx")

if __name__ == '__main__':
    generate_word_report()
