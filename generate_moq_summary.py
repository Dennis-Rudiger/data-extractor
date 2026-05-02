import json
from docx import Document
from docx.shared import Pt, Inches

with open('bomas_average_moq.json', 'r', encoding='utf-8') as f:
    bomas = json.load(f)

with open('march_april_moq.json', 'r', encoding='utf-8') as f:
    march_april = json.load(f)

# Map Bomas MOQ for lookup
b_map = {}
for cat_name, cat_data in bomas['categories'].items():
    for p in cat_data['products']:
        b_map[p['item_code']] = p.get('weekly_moq', 0)

summary = {
    'Increased': [],
    'Dropped': [],
    'Unchanged': []
}

category_summary = {}

for cat_name, cat_data in march_april['categories'].items():
    if not cat_data['products']: continue
    
    if cat_name not in category_summary:
        category_summary[cat_name] = {'Increased': [], 'Dropped': [], 'Unchanged': []}

    for m_p in cat_data['products']:
        code = m_p['item_code']
        desc = m_p['item_description']
        new_1_moq = m_p['1_week_moq']
        new_2_moq = m_p['2_week_moq']
        old_moq = b_map.get(code, 0)
        
        variance = new_1_moq - old_moq
        
        item_data = {
            'code': code,
            'desc': desc,
            'old_moq': old_moq,
            'new_1_moq': new_1_moq,
            'new_2_moq': new_2_moq,
            'variance': variance,
            'cat': cat_name
        }
        
        if variance > 0:
            summary['Increased'].append(item_data)
            category_summary[cat_name]['Increased'].append(item_data)
        elif variance < 0:
            summary['Dropped'].append(item_data)
            category_summary[cat_name]['Dropped'].append(item_data)
        else:
            summary['Unchanged'].append(item_data)
            category_summary[cat_name]['Unchanged'].append(item_data)

# Create Document
doc = Document()
doc.add_heading('MOQ Weekly Averages & Variance Summary', 0)
doc.add_paragraph('Comparing Bomas Q1 (74 Days) vs March-April (36 Days)')

# Overall Summary
doc.add_heading('1. Overall Summary', level=1)
total_items = len(summary['Increased']) + len(summary['Dropped']) + len(summary['Unchanged'])
doc.add_paragraph(f'Total Products Evaluated: {total_items}')
doc.add_paragraph(f"Increased 1-Week MOQs: {len(summary['Increased'])} items")
doc.add_paragraph(f"Dropped 1-Week MOQs: {len(summary['Dropped'])} items")
doc.add_paragraph(f"Unchanged 1-Week MOQs: {len(summary['Unchanged'])} items")

# Category Breakdown
doc.add_heading('2. Category Breakdown', level=1)

for cat_name, cat_data in category_summary.items():
    doc.add_heading(cat_name, level=2)
    doc.add_paragraph(f"Increased: {len(cat_data['Increased'])} | Dropped: {len(cat_data['Dropped'])} | Unchanged: {len(cat_data['Unchanged'])}")
    
    # Show Top Increases if any
    if cat_data['Increased']:
        doc.add_heading('Top MOQ Increases', level=3)
        top_inc = sorted(cat_data['Increased'], key=lambda x: x['variance'], reverse=True)[:5]
        
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Code'
        hdr_cells[1].text = 'Desc'
        hdr_cells[2].text = 'Old 1-Wk MOQ'
        hdr_cells[3].text = 'New 1-Wk MOQ'
        hdr_cells[4].text = 'New 2-Wk MOQ'
        hdr_cells[5].text = 'Variance'
        for cell in hdr_cells:
            for p in cell.paragraphs:
                for r in p.runs: r.bold = True
        
        for item in top_inc:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item['code'])
            row_cells[1].text = str(item['desc'])
            row_cells[2].text = str(item['old_moq'])
            row_cells[3].text = str(item['new_1_moq'])
            row_cells[4].text = str(item['new_2_moq'])
            row_cells[5].text = '+' + str(item['variance'])
        doc.add_paragraph()

    # Show Top Drops if any
    if cat_data['Dropped']:
        doc.add_heading('Top MOQ Drops', level=3)
        top_drop = sorted(cat_data['Dropped'], key=lambda x: x['variance'])[:5]
        
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Code'
        hdr_cells[1].text = 'Desc'
        hdr_cells[2].text = 'Old 1-Wk MOQ'
        hdr_cells[3].text = 'New 1-Wk MOQ'
        hdr_cells[4].text = 'New 2-Wk MOQ'
        hdr_cells[5].text = 'Variance'
        for cell in hdr_cells:
            for p in cell.paragraphs:
                for r in p.runs: r.bold = True
        
        for item in top_drop:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item['code'])
            row_cells[1].text = str(item['desc'])
            row_cells[2].text = str(item['old_moq'])
            row_cells[3].text = str(item['new_1_moq'])
            row_cells[4].text = str(item['new_2_moq'])
            row_cells[5].text = str(item['variance'])
        doc.add_paragraph()

# Append full lists at the end
doc.add_page_break()
doc.add_heading('3. Full List of Increased Items', level=1)
inc_sorted = sorted(summary['Increased'], key=lambda x: x['variance'], reverse=True)
if inc_sorted:
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Item'
    hdr_cells[2].text = 'Old 1-Wk'
    hdr_cells[3].text = 'New 1-Wk'
    hdr_cells[4].text = 'New 2-Wk'
    hdr_cells[5].text = 'Variance'
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True
    for item in inc_sorted:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item['cat'])
        row_cells[1].text = str(item['desc'])
        row_cells[2].text = str(item['old_moq'])
        row_cells[3].text = str(item['new_1_moq'])
        row_cells[4].text = str(item['new_2_moq'])
        row_cells[5].text = '+' + str(item['variance'])

doc.add_page_break()
doc.add_heading('4. Full List of Dropped Items', level=1)
drop_sorted = sorted(summary['Dropped'], key=lambda x: x['variance'])
if drop_sorted:
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Item'
    hdr_cells[2].text = 'Old 1-Wk'
    hdr_cells[3].text = 'New 1-Wk'
    hdr_cells[4].text = 'New 2-Wk'
    hdr_cells[5].text = 'Variance'
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True
    for item in drop_sorted:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item['cat'])
        row_cells[1].text = str(item['desc'])
        row_cells[2].text = str(item['old_moq'])
        row_cells[3].text = str(item['new_1_moq'])
        row_cells[4].text = str(item['new_2_moq'])
        row_cells[5].text = str(item['variance'])

doc.add_page_break()
doc.add_heading('5. Full List of Unchanged Items', level=1)
unchanged_sorted = sorted(summary['Unchanged'], key=lambda x: x['desc'])
if unchanged_sorted:
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Item'
    hdr_cells[2].text = 'Old 1-Wk'
    hdr_cells[3].text = 'New 1-Wk'
    hdr_cells[4].text = 'New 2-Wk'
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True
    for item in unchanged_sorted:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item['cat'])
        row_cells[1].text = str(item['desc'])
        row_cells[2].text = str(item['old_moq'])
        row_cells[3].text = str(item['new_1_moq'])
        row_cells[4].text = str(item['new_2_moq'])

doc.save('MOQ_Variance_Summary.docx')
print('Generated updated MOQ_Variance_Summary.docx')
