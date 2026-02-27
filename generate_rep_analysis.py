import json
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import Circle, Rectangle
from matplotlib.colors import LinearSegmentedColormap
import numpy as np
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.units import inch
from datetime import datetime
import os
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.chart import BarChart, PieChart, RadarChart, LineChart, Reference
from openpyxl.chart.label import DataLabelList
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Color palette
COLORS = {
    'GENERAL HARDWARE': '#3498db',
    'PAINTS': '#e74c3c', 
    'PLUMBING': '#2ecc71',
    'ELECTRICALS': '#f39c12'
}

CAT_ORDER = ['GENERAL HARDWARE', 'PAINTS', 'PLUMBING', 'ELECTRICALS']

# Rep colors for consistent visualization
REP_COLORS = ['#3498db', '#e74c3c', '#2ecc71', '#f39c12', '#9b59b6', '#1abc9c', '#34495e']

# Department Heads
DEPARTMENT_HEADS = {
    'PLUMBING': 'Lewis',
    'PAINTS': 'Bonface Muriu',
    'ELECTRICALS': 'Bonface Kitheka',
    'GENERAL HARDWARE': ['Gladys', 'Betha Odumo', 'Eliza', 'Stephen']
}

def get_rep_role(rep_name):
    """Get the role/department assignment for a rep"""
    if rep_name == 'Lewis':
        return 'Department Head - Plumbing'
    elif rep_name == 'Bonface Muriu':
        return 'Department Head - Paints'
    elif rep_name == 'Bonface Kitheka':
        return 'Department Head - Electricals'
    else:
        return 'General Hardware Team'

def get_primary_department(rep_name):
    """Get the primary department a rep is responsible for"""
    if rep_name == 'Lewis':
        return 'PLUMBING'
    elif rep_name == 'Bonface Muriu':
        return 'PAINTS'
    elif rep_name == 'Bonface Kitheka':
        return 'ELECTRICALS'
    else:
        return 'GENERAL HARDWARE'

def analyze_performance(reps_data, w1_reps, w2_reps, w3_reps=None):
    """Generate comprehensive performance analysis for each rep"""
    
    # Calculate team averages
    total_sales = sum(r['total_sales'] for r in reps_data.values())
    total_profit = sum(r['total_profit'] for r in reps_data.values())
    avg_sales = total_sales / len(reps_data)
    avg_margin = total_profit / total_sales * 100 if total_sales > 0 else 0
    
    # Category totals
    cat_totals = {}
    for cat in CAT_ORDER:
        cat_totals[cat] = {'sales': 0, 'profit': 0}
        for rep_data in reps_data.values():
            if cat in rep_data['categories']:
                cat_totals[cat]['sales'] += rep_data['categories'][cat]['sales_incl']
                cat_totals[cat]['profit'] += rep_data['categories'][cat]['profit']
    
    # Sorted lists for ranking
    sorted_by_sales = sorted(reps_data.items(), key=lambda x: x[1]['total_sales'], reverse=True)
    sorted_by_margin = sorted(reps_data.items(), key=lambda x: x[1]['overall_margin'], reverse=True)
    sorted_by_profit = sorted(reps_data.items(), key=lambda x: x[1]['total_profit'], reverse=True)
    
    # Calculate week-over-week improvements
    week_improvements_w1_w2 = {}
    week_improvements_w2_w3 = {}
    week_improvements_overall = {}
    for rep_name in reps_data.keys():
        w1_sales = w1_reps.get(rep_name, {}).get('total_sales', 0)
        w2_sales = w2_reps.get(rep_name, {}).get('total_sales', 0)
        w3_sales = w3_reps.get(rep_name, {}).get('total_sales', 0) if w3_reps else 0
        
        if w1_sales > 0:
            week_improvements_w1_w2[rep_name] = ((w2_sales - w1_sales) / w1_sales) * 100
        else:
            week_improvements_w1_w2[rep_name] = 100 if w2_sales > 0 else 0
        
        if w3_reps:
            if w2_sales > 0:
                week_improvements_w2_w3[rep_name] = ((w3_sales - w2_sales) / w2_sales) * 100
            else:
                week_improvements_w2_w3[rep_name] = 100 if w3_sales > 0 else 0
            if w1_sales > 0:
                week_improvements_overall[rep_name] = ((w3_sales - w1_sales) / w1_sales) * 100
            else:
                week_improvements_overall[rep_name] = 100 if w3_sales > 0 else 0
    
    most_improved_w1w2 = max(week_improvements_w1_w2.items(), key=lambda x: x[1])
    most_improved_w2w3 = max(week_improvements_w2_w3.items(), key=lambda x: x[1]) if w3_reps else (None, 0)
    
    # Build performance analysis for each rep
    analysis = {}
    
    for idx, (rep_name, rep_data) in enumerate(sorted_by_sales):
        role = get_rep_role(rep_name)
        primary_dept = get_primary_department(rep_name)
        
        # Determine rankings
        sales_rank = idx + 1
        margin_rank = [r[0] for r in sorted_by_margin].index(rep_name) + 1
        profit_rank = [r[0] for r in sorted_by_profit].index(rep_name) + 1
        
        # Badges/Awards
        badges = []
        if sales_rank == 1:
            badges.append("🏆 TOP SELLER")
        if margin_rank == 1:
            badges.append("💎 BEST MARGIN")
        if profit_rank == 1:
            badges.append("💰 TOP PROFIT")
        if w3_reps and rep_name == most_improved_w2w3[0] and most_improved_w2w3[1] > 20:
            badges.append("📈 MOST IMPROVED")
        elif not w3_reps and rep_name == most_improved_w1w2[0] and most_improved_w1w2[1] > 20:
            badges.append("📈 MOST IMPROVED")
        
        # Performance notes
        notes = []
        
        # Overall performance vs average
        sales_vs_avg = ((rep_data['total_sales'] - avg_sales) / avg_sales) * 100
        if sales_vs_avg > 50:
            notes.append(f"Outstanding! Sales {sales_vs_avg:.0f}% above team average")
        elif sales_vs_avg > 0:
            notes.append(f"Sales {sales_vs_avg:.0f}% above team average - solid performance")
        else:
            notes.append(f"Sales {abs(sales_vs_avg):.0f}% below team average - needs improvement")
        
        # Margin performance
        if rep_data['overall_margin'] > avg_margin + 3:
            notes.append(f"Excellent margin at {rep_data['overall_margin']:.1f}% (Team avg: {avg_margin:.1f}%)")
        elif rep_data['overall_margin'] < avg_margin - 2:
            notes.append(f"Margin at {rep_data['overall_margin']:.1f}% below team average of {avg_margin:.1f}%")
        
        # Weekly progress commentary
        w1_s = w1_reps.get(rep_name, {}).get('total_sales', 0)
        w2_s = w2_reps.get(rep_name, {}).get('total_sales', 0)
        w3_s = w3_reps.get(rep_name, {}).get('total_sales', 0) if w3_reps else 0
        
        imp_w1w2 = week_improvements_w1_w2.get(rep_name, 0)
        if imp_w1w2 > 30:
            notes.append(f"W1→W2: Strong growth +{imp_w1w2:.1f}% (KES {w1_s/1000:,.0f}K → {w2_s/1000:,.0f}K)")
        elif imp_w1w2 < -10:
            notes.append(f"W1→W2: Declined {imp_w1w2:.1f}% (KES {w1_s/1000:,.0f}K → {w2_s/1000:,.0f}K)")
        else:
            notes.append(f"W1→W2: Steady {imp_w1w2:+.1f}% (KES {w1_s/1000:,.0f}K → {w2_s/1000:,.0f}K)")
        
        if w3_reps:
            imp_w2w3 = week_improvements_w2_w3.get(rep_name, 0)
            if imp_w2w3 > 30:
                notes.append(f"W2→W3: Strong growth +{imp_w2w3:.1f}% (KES {w2_s/1000:,.0f}K → {w3_s/1000:,.0f}K)")
            elif imp_w2w3 < -10:
                notes.append(f"W2→W3: Declined {imp_w2w3:.1f}% (KES {w2_s/1000:,.0f}K → {w3_s/1000:,.0f}K)")
            else:
                notes.append(f"W2→W3: Steady {imp_w2w3:+.1f}% (KES {w2_s/1000:,.0f}K → {w3_s/1000:,.0f}K)")
            
            # Overall trajectory
            imp_overall = week_improvements_overall.get(rep_name, 0)
            if imp_w1w2 > 0 and imp_w2w3 > 0:
                notes.append(f"📈 Consistent upward trend over 3 weeks (+{imp_overall:.1f}% overall)")
            elif imp_w1w2 < 0 and imp_w2w3 < 0:
                notes.append(f"📉 Declining trend over 3 weeks ({imp_overall:.1f}% overall) - urgent review")
            elif imp_w1w2 > 0 and imp_w2w3 < 0:
                notes.append(f"⚠️ Week 3 reversal after Week 2 growth - monitor closely")
            elif imp_w1w2 < 0 and imp_w2w3 > 0:
                notes.append(f"✓ Recovery in Week 3 after Week 2 dip - positive momentum")
        
        # Department head specific analysis
        if primary_dept != 'GENERAL HARDWARE' and primary_dept in rep_data['categories']:
            dept_sales = rep_data['categories'][primary_dept]['sales_incl']
            dept_total = cat_totals[primary_dept]['sales']
            dept_share = (dept_sales / dept_total * 100) if dept_total > 0 else 0
            dept_margin = rep_data['categories'][primary_dept]['margin_pct']
            
            if primary_dept == 'PLUMBING':
                notes.append(f"Plumbing Dept Head: {dept_share:.1f}% of dept sales, {dept_margin:.1f}% margin")
                if dept_share >= 40:
                    notes.append(f"✓ Strong plumbing leadership with {dept_share:.1f}% of dept sales")
                else:
                    notes.append(f"⚠️ Should lead with higher dept contribution (currently {dept_share:.1f}%)")
            elif primary_dept == 'PAINTS':
                notes.append(f"Paints Dept Head: {dept_share:.1f}% of dept sales, {dept_margin:.1f}% margin")
                if dept_share >= 30:
                    notes.append(f"✓ Good paints leadership with {dept_share:.1f}% of dept sales")
                else:
                    notes.append(f"⚠️ Should lead with higher dept contribution (currently {dept_share:.1f}%)")
            elif primary_dept == 'ELECTRICALS':
                notes.append(f"Electricals Dept Head: {dept_share:.1f}% of dept sales, {dept_margin:.1f}% margin")
                if dept_share >= 50:
                    notes.append(f"✓ Strong leadership - controls {dept_share:.1f}% of department sales")
                else:
                    notes.append(f"⚠️ Should lead with higher dept contribution (currently {dept_share:.1f}%)")
        
        # General hardware team - evaluate GH performance
        if primary_dept == 'GENERAL HARDWARE' and 'GENERAL HARDWARE' in rep_data['categories']:
            gh_sales = rep_data['categories']['GENERAL HARDWARE']['sales_incl']
            gh_total = cat_totals['GENERAL HARDWARE']['sales']
            gh_share = (gh_sales / gh_total * 100) if gh_total > 0 else 0
            gh_margin = rep_data['categories']['GENERAL HARDWARE']['margin_pct']
            notes.append(f"General Hardware: {gh_share:.1f}% of dept, {gh_margin:.1f}% margin")
        
        # Cross-selling analysis
        cats_sold = len(rep_data['categories'])
        if cats_sold == 4:
            notes.append("✓ Selling across all 4 categories - excellent product range")
        elif cats_sold == 3:
            missing = [c for c in CAT_ORDER if c not in rep_data['categories']]
            notes.append(f"Consider expanding into {', '.join(missing)}")
        
        # Best performing category
        if rep_data['categories']:
            best_margin_cat = max(rep_data['categories'].items(), 
                                  key=lambda x: x[1].get('margin_pct', 0))
            notes.append(f"Best margin in {best_margin_cat[0]}: {best_margin_cat[1]['margin_pct']:.1f}%")
        
        analysis[rep_name] = {
            'role': role,
            'primary_department': primary_dept,
            'badges': badges,
            'notes': notes,
            'sales_rank': sales_rank,
            'margin_rank': margin_rank,
            'profit_rank': profit_rank,
            'week_improvement': week_improvements_overall.get(rep_name, imp_w1w2) if w3_reps else imp_w1w2,
            'week_improvement_w1w2': imp_w1w2,
            'week_improvement_w2w3': week_improvements_w2_w3.get(rep_name, 0) if w3_reps else 0,
            'sales_vs_avg': sales_vs_avg
        }
    
    return analysis

# Week 1 Data (Feb 1-7, 2026) - 6 working days
week1_data = {
    "period": "February 1-7, 2026",
    "working_days": 6,
    "categories": {
        "PLUMBING": {
            "Betha Odumo": {"qty": 25, "sales_incl": 5790, "cost": 3563.36, "profit": 1428.02},
            "Bonface Kitheka": {"qty": 149, "sales_incl": 74120, "cost": 50318.18, "profit": 13578.36},
            "Bonface Muriu": {"qty": 59, "sales_incl": 21120, "cost": 14755.56, "profit": 3451.33},
            "Eliza": {"qty": 392, "sales_incl": 100035, "cost": 63234.79, "profit": 23002.29},
            "Gladys": {"qty": 99, "sales_incl": 58630, "cost": 38174.12, "profit": 12368.98},
            "Lewis": {"qty": 538, "sales_incl": 208330, "cost": 135995.12, "profit": 43599.69},
            "Stephen": {"qty": 36, "sales_incl": 156730, "cost": 114305.37, "profit": 20806.71}
        },
        "ELECTRICALS": {
            "Betha Odumo": {"qty": 29, "sales_incl": 4640, "cost": 3270.29, "profit": 729.71},
            "Bonface Kitheka": {"qty": 421, "sales_incl": 37635, "cost": 21481.84, "profit": 10962.11},
            "Bonface Muriu": {"qty": 103, "sales_incl": 26760, "cost": 12548.86, "profit": 10520.1},
            "Eliza": {"qty": 339, "sales_incl": 38940, "cost": 26449.59, "profit": 7119.38},
            "Lewis": {"qty": 21, "sales_incl": 4840, "cost": 3229.89, "profit": 942.51}
        },
        "PAINTS": {
            "Betha Odumo": {"qty": 135, "sales_incl": 109900, "cost": 74676.1, "profit": 20065.3},
            "Bonface Kitheka": {"qty": 113, "sales_incl": 142420, "cost": 107533.97, "profit": 15241.91},
            "Bonface Muriu": {"qty": 329, "sales_incl": 323270, "cost": 229016.28, "profit": 49664.73},
            "Eliza": {"qty": 523, "sales_incl": 595550, "cost": 438908.57, "profit": 74496.57},
            "Gladys": {"qty": 18, "sales_incl": 33800, "cost": 23375.42, "profit": 5762.51},
            "Lewis": {"qty": 103, "sales_incl": 71660, "cost": 47183.19, "profit": 14592.67},
            "Stephen": {"qty": 67, "sales_incl": 74360, "cost": 51984.76, "profit": 12118.68}
        },
        "GENERAL HARDWARE": {
            "Betha Odumo": {"qty": 2996, "sales_incl": 2163840, "cost": 1703316.1, "profit": 162063.16},
            "Bonface Kitheka": {"qty": 420, "sales_incl": 232485, "cost": 171841.89, "profit": 28576.2},
            "Bonface Muriu": {"qty": 1175, "sales_incl": 1128619, "cost": 864689.69, "profit": 108257.76},
            "Eliza": {"qty": 2292, "sales_incl": 1636135, "cost": 1271278.16, "profit": 139182.99},
            "Gladys": {"qty": 6661, "sales_incl": 5538955, "cost": 4512274.33, "profit": 262686.92},
            "Lewis": {"qty": 709.25, "sales_incl": 552707, "cost": 423471.03, "profit": 53000.48},
            "Stephen": {"qty": 1392.5, "sales_incl": 925821, "cost": 726157.7, "profit": 71963.83}
        }
    }
}

# Week 2 Data (Feb 9-14, 2026) - 6 working days
week2_raw = {
    "period": "February 9-14, 2026",
    "working_days": 6,
    "categories": {
        "PLUMBING": {
            "Betha Odumo": {"qty": 134, "sales_incl": 25045, "cost": 15102.88, "profit": 6487.66},
            "Bonface Kitheka": {"qty": 167, "sales_incl": 27070, "cost": 14623.81, "profit": 8712.39},
            "Bonface Muriu": {"qty": 143, "sales_incl": 59480, "cost": 37416.36, "profit": 13859.50},
            "Eliza": {"qty": 40, "sales_incl": 16110, "cost": 7124.59, "profit": 6763.33},
            "Gladys": {"qty": 111, "sales_incl": 26890, "cost": 16815.69, "profit": 6365.35},
            "Lewis": {"qty": 597, "sales_incl": 157320, "cost": 99649.36, "profit": 35971.24},
            "Magdalene": {"qty": 65, "sales_incl": 17415, "cost": 11412.29, "profit": 3600.65},
            "Stephen": {"qty": 37, "sales_incl": 20040, "cost": 12542.66, "profit": 4733.22},
            "WALK IN-BOMAS": {"qty": 38, "sales_incl": 2361, "cost": 1783.32, "profit": 252.03}
        },
        "ELECTRICALS": {
            "Betha Odumo": {"qty": 117, "sales_incl": 10570, "cost": 5908.46, "profit": 3203.59},
            "Bonface Kitheka": {"qty": 1379, "sales_incl": 94545, "cost": 61012.40, "profit": 20491.88},
            "Bonface Muriu": {"qty": 3, "sales_incl": 430, "cost": 163.67, "profit": 207.02},
            "Eliza": {"qty": 10, "sales_incl": 1960, "cost": 938.58, "profit": 751.08},
            "Gladys": {"qty": 5, "sales_incl": 650, "cost": 340.75, "profit": 219.59},
            "Lewis": {"qty": 26, "sales_incl": 18715, "cost": 13545.01, "profit": 2588.62},
            "WALK IN-BOMAS": {"qty": 7, "sales_incl": 241, "cost": 205.30, "profit": 2.46}
        },
        "PAINTS": {
            "Betha Odumo": {"qty": 88, "sales_incl": 134210, "cost": 96950.09, "profit": 18748.18},
            "Bonface Kitheka": {"qty": 156, "sales_incl": 102030, "cost": 71573.19, "profit": 16383.70},
            "Bonface Muriu": {"qty": 540, "sales_incl": 477620, "cost": 335232.25, "profit": 76509.09},
            "Eliza": {"qty": 284, "sales_incl": 432050, "cost": 317026.68, "profit": 55430.12},
            "Gladys": {"qty": 30, "sales_incl": 61410, "cost": 47161.35, "profit": 5778.31},
            "Lewis": {"qty": 69, "sales_incl": 47540, "cost": 31638.86, "profit": 9343.92},
            "Stephen": {"qty": 94, "sales_incl": 117500, "cost": 84581.88, "profit": 16711.20},
            "WALK IN-BOMAS": {"qty": 1, "sales_incl": 3106, "cost": 2677.50, "profit": 0.09}
        },
        "GENERAL HARDWARE": {
            "Betha Odumo": {"qty": 4656, "sales_incl": 3331665, "cost": 2674886.71, "profit": 197238.23},
            "Bonface Kitheka": {"qty": 679.25, "sales_incl": 525460, "cost": 412309.97, "profit": 40672.78},
            "Bonface Muriu": {"qty": 1268.5, "sales_incl": 753970, "cost": 585266.62, "profit": 64707.50},
            "Eliza": {"qty": 2344, "sales_incl": 2232275, "cost": 1763019.67, "profit": 161355.25},
            "Gladys": {"qty": 9480, "sales_incl": 8927070, "cost": 7248623.49, "profit": 447126.56},
            "Lewis": {"qty": 1370.75, "sales_incl": 1002130, "cost": 800004.34, "profit": 63900.82},
            "Magdalene": {"qty": 3, "sales_incl": 1120, "cost": 585.78, "profit": 379.74},
            "Stephen": {"qty": 1972, "sales_incl": 1423935, "cost": 1108610.31, "profit": 118919.83},
            "WALK IN-BOMAS": {"qty": 36, "sales_incl": 13327, "cost": 6905.34, "profit": 4583.44}
        }
    }
}


# Week 3 Data (Feb 16-21, 2026) - 6 working days
week3_raw = {
    "period": "February 16-21, 2026",
    "working_days": 6,
    "categories": {
        "PLUMBING": {
            "Betha Odumo": {"qty": 55, "sales_incl": 25330, "cost": 14933.85, "profit": 6902.35},
            "Bonface Kitheka": {"qty": 118, "sales_incl": 39625, "cost": 25840.81, "profit": 8318.65},
            "Bonface Muriu": {"qty": 16, "sales_incl": 10150, "cost": 5665.85, "profit": 3084.15},
            "Eliza": {"qty": 192, "sales_incl": 64080, "cost": 34432.65, "profit": 20808.71},
            "Gladys": {"qty": 24, "sales_incl": 17080, "cost": 9388.57, "profit": 5335.56},
            "Lewis": {"qty": 767, "sales_incl": 253960, "cost": 161899.91, "profit": 57031.15},
            "Stephen": {"qty": 27, "sales_incl": 34950, "cost": 25984.77, "profit": 4144.53}
        },
        "ELECTRICALS": {
            "Betha Odumo": {"qty": 6, "sales_incl": 1250, "cost": 545.17, "profit": 532.41},
            "Bonface Kitheka": {"qty": 553, "sales_incl": 54340, "cost": 36305.01, "profit": 10539.80},
            "Eliza": {"qty": 51, "sales_incl": 9150, "cost": 5850.41, "profit": 2037.52},
            "Gladys": {"qty": 2, "sales_incl": 140, "cost": 66.86, "profit": 53.82},
            "Lewis": {"qty": 127, "sales_incl": 19330, "cost": 13105.12, "profit": 3558.67}
        },
        "PAINTS": {
            "Betha Odumo": {"qty": 229, "sales_incl": 151850, "cost": 107513.05, "profit": 23392.10},
            "Bonface Kitheka": {"qty": 118, "sales_incl": 50180, "cost": 32675.91, "profit": 10582.73},
            "Bonface Muriu": {"qty": 422, "sales_incl": 734470, "cost": 528075.84, "profit": 105087.98},
            "Eliza": {"qty": 429, "sales_incl": 455820, "cost": 331031.91, "profit": 61916.40},
            "Gladys": {"qty": 66, "sales_incl": 60230, "cost": 42141.40, "profit": 9781.00},
            "Lewis": {"qty": 58, "sales_incl": 35800, "cost": 24088.69, "profit": 6773.36},
            "Stephen": {"qty": 104, "sales_incl": 74140, "cost": 50732.39, "profit": 13181.39}
        },
        "GENERAL HARDWARE": {
            "Betha Odumo": {"qty": 4389, "sales_incl": 3470400, "cost": 2778307.33, "profit": 213416.77},
            "Bonface Kitheka": {"qty": 625, "sales_incl": 390480, "cost": 307724.97, "profit": 28895.74},
            "Bonface Muriu": {"qty": 829.5, "sales_incl": 413020, "cost": 316847.72, "profit": 39204.00},
            "Eliza": {"qty": 3740, "sales_incl": 2874155, "cost": 2304639.44, "profit": 173080.36},
            "Gladys": {"qty": 6110, "sales_incl": 4614165, "cost": 3742604.58, "profit": 235123.91},
            "Lewis": {"qty": 405, "sales_incl": 242150, "cost": 177858.38, "profit": 30891.59},
            "Stephen": {"qty": 1953.5, "sales_incl": 1239840, "cost": 974209.68, "profit": 94617.85},
            "WALK IN-BOMAS": {"qty": 2, "sales_incl": 1300, "cost": 1091.00, "profit": 29.69}
        }
    }
}

def merge_reps(data):
    """Merge Magdalene into Betha Odumo, and WALK IN-BOMAS into Bonface Kitheka"""
    merged = {"period": data["period"], "working_days": data["working_days"], "categories": {}}
    
    for cat, reps in data["categories"].items():
        merged["categories"][cat] = {}
        
        for rep, vals in reps.items():
            if rep == "Magdalene":
                target = "Betha Odumo"
            elif rep == "WALK IN-BOMAS":
                target = "Bonface Kitheka"
            else:
                target = rep
            
            if target not in merged["categories"][cat]:
                merged["categories"][cat][target] = {"qty": 0, "sales_incl": 0, "cost": 0, "profit": 0}
            
            merged["categories"][cat][target]["qty"] += vals["qty"]
            merged["categories"][cat][target]["sales_incl"] += vals["sales_incl"]
            merged["categories"][cat][target]["cost"] += vals["cost"]
            merged["categories"][cat][target]["profit"] += vals["profit"]
    
    return merged

# Apply merges: Magdalene → Betha Odumo, WALK IN-BOMAS → Bonface Kitheka
week1_merged = merge_reps(week1_data)
week2_data = merge_reps(week2_raw)
week3_data = merge_reps(week3_raw)

def combine_data(*weeks):
    """Combine data from multiple weeks"""
    combined = {"categories": {}}
    for w in weeks:
        if not w.get("categories"):
            continue
        for cat in w["categories"]:
            if cat not in combined["categories"]:
                combined["categories"][cat] = {}
            for rep, vals in w["categories"][cat].items():
                if rep not in combined["categories"][cat]:
                    combined["categories"][cat][rep] = {"qty": 0, "sales_incl": 0, "cost": 0, "profit": 0}
                combined["categories"][cat][rep]["qty"] += vals["qty"]
                combined["categories"][cat][rep]["sales_incl"] += vals["sales_incl"]
                combined["categories"][cat][rep]["cost"] += vals["cost"]
                combined["categories"][cat][rep]["profit"] += vals["profit"]
    return combined

def get_rep_data(combined):
    """Process rep data and add 16% VAT to profits (profits are exclusive of VAT in source)"""
    VAT_RATE = 1.16  # 16% VAT
    reps = {}
    for cat, rep_data in combined["categories"].items():
        for rep, vals in rep_data.items():
            if rep not in reps:
                reps[rep] = {"total_sales": 0, "total_profit": 0, "total_cost": 0, "total_qty": 0, "categories": {}}
            
            # Add 16% VAT to profit (profit in source is VAT-exclusive)
            profit_incl_vat = vals["profit"] * VAT_RATE
            
            reps[rep]["categories"][cat] = vals.copy()
            reps[rep]["categories"][cat]["profit"] = profit_incl_vat  # Update profit with VAT
            reps[rep]["categories"][cat]["margin_pct"] = (profit_incl_vat / vals["sales_incl"] * 100) if vals["sales_incl"] > 0 else 0
            reps[rep]["total_sales"] += vals["sales_incl"]
            reps[rep]["total_profit"] += profit_incl_vat
            reps[rep]["total_cost"] += vals["cost"]
            reps[rep]["total_qty"] += vals["qty"]
    for rep, data in reps.items():
        data["overall_margin"] = (data["total_profit"] / data["total_sales"] * 100) if data["total_sales"] > 0 else 0
    return reps

# ========== ENHANCED CHART FUNCTIONS ==========

def create_donut_chart(rep_name, rep_data, output_dir="charts"):
    """Donut chart showing sales distribution by category"""
    os.makedirs(output_dir, exist_ok=True)
    categories = []
    sales = []
    chart_colors = []
    
    for cat in CAT_ORDER:
        if cat in rep_data["categories"]:
            categories.append(cat)
            sales.append(rep_data["categories"][cat]["sales_incl"])
            chart_colors.append(COLORS[cat])
    
    if not sales or sum(sales) == 0:
        return None
    
    fig, ax = plt.subplots(figsize=(8, 5))
    
    # Create donut chart
    pie_result = ax.pie(sales, labels=None, autopct='%1.1f%%', 
                        colors=chart_colors, startangle=90,
                        pctdistance=0.75, explode=[0.02]*len(sales),
                        wedgeprops=dict(width=0.5, edgecolor='white'))
    wedges = pie_result[0]
    autotexts = pie_result[2] if len(pie_result) > 2 else []
    
    # Center text
    ax.text(0, 0, f'KES\n{sum(sales)/1000000:.2f}M', ha='center', va='center', 
            fontsize=12, fontweight='bold', color='#2c3e50')
    
    for autotext in autotexts:
        autotext.set_fontsize(9)
        autotext.set_fontweight('bold')
        autotext.set_color('white')
    
    legend_labels = [f"{cat}: KES {s:,.0f}" for cat, s in zip(categories, sales)]
    ax.legend(wedges, legend_labels, title="Categories", loc="center left", 
              bbox_to_anchor=(1, 0, 0.5, 1), fontsize=8)
    ax.set_title(f"{rep_name} - Sales Distribution", fontsize=12, fontweight='bold', pad=10)
    
    plt.tight_layout()
    filename = f"{output_dir}/{rep_name.replace(' ', '_')}_donut.png"
    plt.savefig(filename, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return filename

def create_sales_vs_profit_bar(rep_name, rep_data, output_dir="charts"):
    """Grouped bar chart showing sales vs profit by category"""
    os.makedirs(output_dir, exist_ok=True)
    categories = []
    sales = []
    profits = []
    
    for cat in CAT_ORDER:
        if cat in rep_data["categories"] and rep_data["categories"][cat]["sales_incl"] > 0:
            short_cat = cat[:12] if len(cat) > 12 else cat
            categories.append(short_cat)
            sales.append(rep_data["categories"][cat]["sales_incl"] / 1000)
            profits.append(rep_data["categories"][cat]["profit"] / 1000)
    
    if not categories:
        return None
    
    fig, ax = plt.subplots(figsize=(9, 4))
    x = np.arange(len(categories))
    width = 0.35
    
    bars1 = ax.bar(x - width/2, sales, width, label='Sales', color="#0494f5", edgecolor='white')
    bars2 = ax.bar(x + width/2, profits, width, label='Profit', color="#08f169", edgecolor='white')
    
    # Add value labels
    for bar in bars1:
        height = bar.get_height()
        ax.annotate(f'{height:.0f}K', xy=(bar.get_x() + bar.get_width() / 2, height),
                    xytext=(0, 3), textcoords="offset points", ha='center', va='bottom', fontsize=8)
    for bar in bars2:
        height = bar.get_height()
        ax.annotate(f'{height:.0f}K', xy=(bar.get_x() + bar.get_width() / 2, height),
                    xytext=(0, 3), textcoords="offset points", ha='center', va='bottom', fontsize=8)
    
    ax.set_ylabel('Amount (Thousands KES)', fontsize=10, fontweight='bold')
    ax.set_title(f"{rep_name} - Sales vs Profit by Category", fontsize=12, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(categories, fontsize=9)
    ax.legend()
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    
    plt.tight_layout()
    filename = f"{output_dir}/{rep_name.replace(' ', '_')}_sales_profit.png"
    plt.savefig(filename, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return filename

def create_margin_gauge(rep_name, rep_data, output_dir="charts"):
    """Semi-circular gauge showing overall margin"""
    os.makedirs(output_dir, exist_ok=True)
    margin = rep_data["overall_margin"]
    
    fig, ax = plt.subplots(figsize=(5, 3))
    
    # Create semi-circle
    theta = np.linspace(0, np.pi, 100)
    r = 1
    
    # Background arc
    ax.fill_between(theta, 0.5, 1, alpha=0.2, color='grey')
    
    # Colored sections
    sections = [
        (0, np.pi/3, '#e74c3c'),  # Red: 0-33% of arc (0-10% margin)
        (np.pi/3, 2*np.pi/3, '#f39c12'),  # Orange: 33-66% of arc (10-20% margin)
        (2*np.pi/3, np.pi, '#27ae60')  # Green: 66-100% of arc (20%+ margin)
    ]
    
    for start, end, color in sections:
        theta_section = np.linspace(start, end, 50)
        for i in range(len(theta_section)-1):
            ax.fill_between([theta_section[i], theta_section[i+1]], [0.6, 0.6], [0.95, 0.95], 
                          color=color, alpha=0.3, transform=ax.transData)
    
    # Calculate needle position (0% = left, 30%+ = right)
    needle_angle = np.pi - (min(margin, 30) / 30) * np.pi
    ax.annotate('', xy=(0.5 + 0.4*np.cos(needle_angle), 0.4*np.sin(needle_angle)), 
                xytext=(0.5, 0), arrowprops=dict(arrowstyle='->', color='#2c3e50', lw=3))
    
    # Center dot
    circle = Circle((0.5, 0), 0.05, color='#2c3e50')
    ax.add_patch(circle)
    
    # Labels
    ax.text(0.5, -0.15, f'{margin:.1f}%', ha='center', va='center', fontsize=16, fontweight='bold', color='#2c3e50')
    ax.text(0.5, -0.35, 'Overall Margin', ha='center', va='center', fontsize=10, color='#7f8c8d')
    ax.text(0.05, 0.05, '0%', ha='center', fontsize=8, color='#7f8c8d')
    ax.text(0.95, 0.05, '30%', ha='center', fontsize=8, color='#7f8c8d')
    
    ax.set_xlim(-0.1, 1.1)
    ax.set_ylim(-0.5, 0.6)
    ax.set_aspect('equal')
    ax.axis('off')
    ax.set_title(f"{rep_name}", fontsize=12, fontweight='bold', pad=10)
    
    plt.tight_layout()
    filename = f"{output_dir}/{rep_name.replace(' ', '_')}_gauge.png"
    plt.savefig(filename, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return filename

def create_radar_chart(reps_data, output_dir="charts"):
    """Radar chart comparing all reps across multiple metrics"""
    os.makedirs(output_dir, exist_ok=True)
    
    sorted_reps = sorted(reps_data.items(), key=lambda x: x[1]['total_sales'], reverse=True)
    rep_names = [r[0] for r in sorted_reps]
    
    # Normalize metrics (0-1 scale)
    max_sales = max(r[1]['total_sales'] for r in sorted_reps)
    max_profit = max(r[1]['total_profit'] for r in sorted_reps)
    max_qty = max(r[1]['total_qty'] for r in sorted_reps)
    max_margin = max(r[1]['overall_margin'] for r in sorted_reps)
    
    categories = ['Sales', 'Profit', 'Items Sold', 'Margin']
    N = len(categories)
    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    angles += angles[:1]
    
    # Color palette for all 7 reps - distinct and vibrant colors
    RADAR_COLORS = [
        '#2E86AB',  # Steel Blue - Gladys
        '#A23B72',  # Magenta/Berry - Betha Odumo
        '#F18F01',  # Vivid Orange - Eliza
        '#C73E1D',  # Vermillion Red - Bonface Muriu
        '#6B4C9A',  # Purple - Stephen
        '#2ECC71',  # Emerald Green - Lewis
        '#1ABC9C',  # Teal - Bonface Kitheka
    ]
    
    fig, ax = plt.subplots(figsize=(12, 9), subplot_kw=dict(polar=True))
    
    # Set background color for better visibility
    ax.set_facecolor('#f8f9fa')
    
    # Add grid styling - use plt.rgrids for polar axes
    plt.rgrids([0.2, 0.4, 0.6, 0.8, 1.0], labels=['20%', '40%', '60%', '80%', '100%'], 
               fontsize=8, color='#666666')
    ax.spines['polar'].set_color('#cccccc')
    ax.grid(color='#cccccc', linestyle='-', linewidth=0.5, alpha=0.7)
    
    for i, (rep_name, rep_data) in enumerate(sorted_reps):  # All reps
        values = [
            rep_data['total_sales'] / max_sales,
            rep_data['total_profit'] / max_profit,
            rep_data['total_qty'] / max_qty,
            rep_data['overall_margin'] / max_margin
        ]
        values += values[:1]
        
        color = RADAR_COLORS[i % len(RADAR_COLORS)]
        ax.plot(angles, values, 'o-', linewidth=2.5, label=rep_name, color=color, markersize=7)
        ax.fill(angles, values, alpha=0.15, color=color)
    
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, fontsize=11, fontweight='bold', color='#2c3e50')
    ax.set_ylim(0, 1.1)
    ax.set_title('Performance Radar - All Sales Reps', fontsize=14, fontweight='bold', pad=25, color='#2c3e50')
    
    # Better legend positioning
    legend = ax.legend(loc='upper right', bbox_to_anchor=(1.4, 1.0), fontsize=10, 
                       frameon=True, fancybox=True, shadow=True)
    legend.get_frame().set_facecolor('white')
    legend.get_frame().set_edgecolor('#cccccc')
    
    plt.tight_layout()
    filename = f"{output_dir}/summary_radar.png"
    plt.savefig(filename, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return filename

def create_heatmap(reps_data, output_dir="charts"):
    """Heatmap showing category performance by rep"""
    os.makedirs(output_dir, exist_ok=True)
    
    sorted_reps = sorted(reps_data.items(), key=lambda x: x[1]['total_sales'], reverse=True)
    rep_names = [r[0] for r in sorted_reps]
    
    # Build margin matrix
    data_matrix = []
    for rep_name, rep_data in sorted_reps:
        row = []
        for cat in CAT_ORDER:
            if cat in rep_data["categories"]:
                row.append(rep_data["categories"][cat]["margin_pct"])
            else:
                row.append(0)
        data_matrix.append(row)
    
    data_matrix = np.array(data_matrix)
    
    fig, ax = plt.subplots(figsize=(11, 7))
    
    # Create custom colormap - Light to Dark Teal
    custom_colors = ['#E8F4F8', '#B8DCE8', '#7BC8DC', '#4ABDCF', '#2E9AB8', '#1B7A9C', '#0D5C7D']
    custom_cmap = LinearSegmentedColormap.from_list('custom_teal', custom_colors, N=256)
    
    im = ax.imshow(data_matrix, cmap=custom_cmap, aspect='auto', vmin=0, vmax=35)
    
    # Add cell borders
    for i in range(len(rep_names) + 1):
        ax.axhline(i - 0.5, color='white', linewidth=2)
    for j in range(len(CAT_ORDER) + 1):
        ax.axvline(j - 0.5, color='white', linewidth=2)
    
    # Labels with better styling
    ax.set_xticks(np.arange(len(CAT_ORDER)))
    ax.set_yticks(np.arange(len(rep_names)))
    ax.set_xticklabels([c[:12] for c in CAT_ORDER], fontsize=10, fontweight='bold', color='#2c3e50')
    ax.set_yticklabels(rep_names, fontsize=10, fontweight='bold', color='#2c3e50')
    
    # Move x labels to top
    ax.xaxis.set_ticks_position('top')
    ax.xaxis.set_label_position('top')
    
    # Add values in cells with better contrast
    for i in range(len(rep_names)):
        for j in range(len(CAT_ORDER)):
            val = data_matrix[i, j]
            # Dark text for light backgrounds, white for dark
            text_color = 'white' if val > 18 else '#2c3e50'
            ax.text(j, i, f'{val:.1f}%', ha='center', va='center', 
                   fontsize=11, color=text_color, fontweight='bold')
    
    # Colorbar with better styling
    cbar = ax.figure.colorbar(im, ax=ax, shrink=0.8, pad=0.02)
    cbar.ax.set_ylabel('Profit Margin %', rotation=-90, va="bottom", fontsize=11, fontweight='bold', color='#2c3e50')
    cbar.ax.tick_params(labelsize=9)
    
    # Add margin indicators on colorbar
    cbar.ax.axhline(y=10, color='#e74c3c', linewidth=2, linestyle='--')
    cbar.ax.text(1.5, 10, 'Target', fontsize=8, va='center', color='#e74c3c', fontweight='bold')
    
    ax.set_title('Profit Margin Heatmap by Rep & Category', fontsize=14, fontweight='bold', 
                 pad=15, color='#2c3e50')
    
    # Remove spines
    for spine in ax.spines.values():
        spine.set_visible(False)
    
    plt.tight_layout()
    filename = f"{output_dir}/summary_heatmap.png"
    plt.savefig(filename, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return filename

def create_rep_contribution_chart(reps_data, output_dir="charts"):
    """Bullet chart showing each rep's contribution to total sales with category breakdown"""
    os.makedirs(output_dir, exist_ok=True)
    
    sorted_reps = sorted(reps_data.items(), key=lambda x: x[1]['total_sales'], reverse=True)
    total_sales = sum(r[1]['total_sales'] for r in sorted_reps)
    
    fig, axes = plt.subplots(len(sorted_reps), 1, figsize=(12, 10))
    
    for idx, (rep_name, rep_data) in enumerate(sorted_reps):
        ax = axes[idx]
        
        # Calculate percentages
        rep_pct = (rep_data['total_sales'] / total_sales) * 100
        
        # Draw background bar (100%)
        ax.barh(0, 100, height=0.8, color='#ecf0f1', edgecolor='none')
        
        # Draw category segments
        left = 0
        for cat in CAT_ORDER:
            if cat in rep_data["categories"]:
                cat_pct = (rep_data["categories"][cat]["sales_incl"] / total_sales) * 100
                ax.barh(0, cat_pct, height=0.6, left=left, color=COLORS[cat], edgecolor='white', linewidth=0.5)
                left += cat_pct
        
        # Rep name and value
        ax.text(-2, 0, rep_name, va='center', ha='right', fontsize=11, fontweight='bold', color='#2c3e50')
        ax.text(left + 1, 0, f'KES {rep_data["total_sales"]/1000000:.2f}M ({rep_pct:.1f}%)', 
               va='center', ha='left', fontsize=10, fontweight='bold', color='#2c3e50')
        
        # Clean up axis
        ax.set_xlim(-25, 60)
        ax.set_ylim(-0.6, 0.6)
        ax.axis('off')
    
    # Add title
    fig.suptitle('Sales Contribution by Rep\n(Colored by Category)', fontsize=14, fontweight='bold', y=0.98)
    
    # Legend at bottom
    legend_elements = [Rectangle((0,0),1,1, facecolor=COLORS[cat], label=cat) for cat in CAT_ORDER]
    fig.legend(handles=legend_elements, loc='lower center', ncol=4, fontsize=10, 
               bbox_to_anchor=(0.5, 0.02), frameon=True)
    
    plt.tight_layout()
    plt.subplots_adjust(top=0.92, bottom=0.1, hspace=0.3)
    filename = f"{output_dir}/rep_contribution.png"
    plt.savefig(filename, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return filename

def create_week_grouped_bar(w1_reps, w2_reps, w3_reps=None, output_dir="charts"):
    """Grouped bar chart comparing Week 1 vs Week 2 vs Week 3 per rep"""
    os.makedirs(output_dir, exist_ok=True)
    
    all_keys = set(w1_reps.keys()) | set(w2_reps.keys())
    if w3_reps:
        all_keys |= set(w3_reps.keys())
    all_reps = sorted(all_keys, 
                      key=lambda r: (w1_reps.get(r, {"total_sales": 0})["total_sales"] + 
                                     w2_reps.get(r, {"total_sales": 0})["total_sales"] +
                                     (w3_reps.get(r, {"total_sales": 0})["total_sales"] if w3_reps else 0)),
                      reverse=True)
    
    w1_sales = [w1_reps.get(r, {"total_sales": 0})["total_sales"]/1000000 for r in all_reps]
    w2_sales = [w2_reps.get(r, {"total_sales": 0})["total_sales"]/1000000 for r in all_reps]
    w3_sales = [w3_reps.get(r, {"total_sales": 0})["total_sales"]/1000000 for r in all_reps] if w3_reps else []
    
    fig, ax = plt.subplots(figsize=(14, 7))
    x = np.arange(len(all_reps))
    
    if w3_reps:
        width = 0.25
        bars1 = ax.bar(x - width, w1_sales, width, label='Week 1 (Feb 1-7)', color='#3498db', edgecolor='white')
        bars2 = ax.bar(x, w2_sales, width, label='Week 2 (Feb 9-14)', color='#e74c3c', edgecolor='white')
        bars3 = ax.bar(x + width, w3_sales, width, label='Week 3 (Feb 16-21)', color='#2ecc71', edgecolor='white')
        
        # Value labels with W2→W3 growth indicator
        for i in range(len(all_reps)):
            s2, s3 = w2_sales[i], w3_sales[i]
            growth = ((s3 - s2) / s2 * 100) if s2 > 0 else 0
            color = '#27ae60' if growth > 0 else '#e74c3c'
            symbol = '▲' if growth > 0 else '▼'
            peak = max(w1_sales[i], s2, s3)
            ax.annotate(f'{symbol}{abs(growth):.0f}%', xy=(i, peak + 0.1), 
                       ha='center', fontsize=7, fontweight='bold', color=color)
    else:
        width = 0.35
        bars1 = ax.bar(x - width/2, w1_sales, width, label='Week 1 (Feb 1-7)', color='#3498db', edgecolor='white')
        bars2 = ax.bar(x + width/2, w2_sales, width, label='Week 2 (Feb 9-14)', color='#e74c3c', edgecolor='white')
        
        for i, (s1, s2) in enumerate(zip(w1_sales, w2_sales)):
            growth = ((s2 - s1) / s1 * 100) if s1 > 0 else 0
            color = '#27ae60' if growth > 0 else '#e74c3c'
            symbol = '▲' if growth > 0 else '▼'
            ax.annotate(f'{symbol}{abs(growth):.0f}%', xy=(i, max(s1, s2) + 0.1), 
                       ha='center', fontsize=8, fontweight='bold', color=color)
    
    ax.set_ylabel('Sales (Millions KES)', fontsize=10, fontweight='bold')
    ax.set_title('Week-over-Week Sales Comparison by Rep', fontsize=14, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(all_reps, fontsize=9, rotation=15, ha='right')
    ax.legend(loc='upper right')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    
    plt.tight_layout()
    filename = f"{output_dir}/week_comparison_grouped.png"
    plt.savefig(filename, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return filename

def create_profit_waterfall(reps_data, output_dir="charts"):
    """Waterfall chart showing profit contribution by rep"""
    os.makedirs(output_dir, exist_ok=True)
    
    sorted_reps = sorted(reps_data.items(), key=lambda x: x[1]['total_profit'], reverse=True)
    rep_names = [r[0] for r in sorted_reps]
    profits = [r[1]['total_profit'] for r in sorted_reps]
    
    fig, ax = plt.subplots(figsize=(12, 6))
    
    # Calculate positions
    cumulative = 0
    for i, (name, profit) in enumerate(zip(rep_names, profits)):
        color = REP_COLORS[i % len(REP_COLORS)]
        ax.bar(i, profit/1000, bottom=cumulative/1000, color=color, edgecolor='white', width=0.7)
        
        # Label on bar
        ax.text(i, (cumulative + profit/2)/1000, f'{profit/1000:.0f}K', 
               ha='center', va='center', fontsize=9, fontweight='bold', color='white')
        cumulative += profit
    
    # Total bar
    ax.bar(len(rep_names), cumulative/1000, color='#2c3e50', edgecolor='white', width=0.7)
    ax.text(len(rep_names), cumulative/1000 + 30, f'Total: {cumulative/1000000:.2f}M', 
           ha='center', va='bottom', fontsize=10, fontweight='bold', color='#2c3e50')
    
    ax.set_ylabel('Profit (Thousands KES)', fontsize=10, fontweight='bold')
    ax.set_title('Profit Contribution Waterfall by Rep', fontsize=14, fontweight='bold')
    ax.set_xticks(range(len(rep_names) + 1))
    ax.set_xticklabels(rep_names + ['TOTAL'], fontsize=9, rotation=15, ha='right')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    
    plt.tight_layout()
    filename = f"{output_dir}/profit_waterfall.png"
    plt.savefig(filename, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return filename

def create_category_donut(reps_data, output_dir="charts"):
    """Donut chart showing overall category distribution"""
    os.makedirs(output_dir, exist_ok=True)
    
    cat_totals = {cat: 0 for cat in CAT_ORDER}
    for rep_name, rep_data in reps_data.items():
        for cat in CAT_ORDER:
            if cat in rep_data["categories"]:
                cat_totals[cat] += rep_data["categories"][cat]["sales_incl"]
    
    categories = [cat for cat in CAT_ORDER if cat_totals[cat] > 0]
    sales = [cat_totals[cat] for cat in categories]
    chart_colors = [COLORS[cat] for cat in categories]
    
    fig, ax = plt.subplots(figsize=(9, 7))
    
    pie_result = ax.pie(sales, labels=None, autopct='%1.1f%%',
                        colors=chart_colors, startangle=90,
                        pctdistance=0.75, explode=[0.03]*len(sales),
                        wedgeprops=dict(width=0.5, edgecolor='white'))
    wedges = pie_result[0]
    autotexts = pie_result[2] if len(pie_result) > 2 else []
    
    # Center text
    total = sum(sales)
    ax.text(0, 0, f'KES\n{total/1000000:.1f}M', ha='center', va='center', 
            fontsize=14, fontweight='bold', color='#2c3e50')
    
    for autotext in autotexts:
        autotext.set_fontsize(11)
        autotext.set_fontweight('bold')
        autotext.set_color('white')
    
    legend_labels = [f"{cat}: KES {s/1000000:.2f}M" for cat, s in zip(categories, sales)]
    ax.legend(wedges, legend_labels, title="Categories", loc="center left",
              bbox_to_anchor=(1, 0, 0.5, 1), fontsize=10)
    ax.set_title("Total Sales by Category", fontsize=14, fontweight='bold')
    
    plt.tight_layout()
    filename = f"{output_dir}/category_donut.png"
    plt.savefig(filename, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return filename

def create_margin_line_chart(reps_data, output_dir="charts"):
    """Line chart showing margin trend across categories for each rep"""
    os.makedirs(output_dir, exist_ok=True)
    
    sorted_reps = sorted(reps_data.items(), key=lambda x: x[1]['total_sales'], reverse=True)
    
    fig, ax = plt.subplots(figsize=(10, 6))
    
    x = np.arange(len(CAT_ORDER))
    
    for i, (rep_name, rep_data) in enumerate(sorted_reps):
        margins = []
        for cat in CAT_ORDER:
            if cat in rep_data["categories"]:
                margins.append(rep_data["categories"][cat]["margin_pct"])
            else:
                margins.append(0)
        ax.plot(x, margins, 'o-', linewidth=2, markersize=8, label=rep_name, 
               color=REP_COLORS[i % len(REP_COLORS)])
    
    ax.set_xlabel('Category', fontsize=10, fontweight='bold')
    ax.set_ylabel('Margin %', fontsize=10, fontweight='bold')
    ax.set_title('Profit Margin Comparison Across Categories', fontsize=14, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels([c[:12] for c in CAT_ORDER], fontsize=9)
    ax.legend(loc='upper right', fontsize=8)
    ax.axhline(y=10, color='red', linestyle='--', linewidth=1.5, alpha=0.5, label='Target 10%')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.grid(True, alpha=0.3)
    
    plt.tight_layout()
    filename = f"{output_dir}/margin_line_chart.png"
    plt.savefig(filename, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return filename

def create_summary_sales_chart(reps_data, output_dir="charts"):
    """Horizontal bar chart with sales bars and profit overlay - much clearer visualization"""
    os.makedirs(output_dir, exist_ok=True)
    sorted_reps = sorted(reps_data.items(), key=lambda x: x[1]['total_sales'], reverse=True)
    rep_names = [r[0] for r in sorted_reps]
    sales = [r[1]['total_sales']/1000000 for r in sorted_reps]
    profits = [r[1]['total_profit']/1000000 for r in sorted_reps]
    margins = [r[1]['overall_margin'] for r in sorted_reps]
    
    fig, ax = plt.subplots(figsize=(12, 7))
    
    y = np.arange(len(rep_names))
    height = 0.6
    
    # Sales bars (full width, lighter color)
    bars_sales = ax.barh(y, sales, height, label='Total Sales', color='#85C1E9', edgecolor='white', alpha=0.9)
    
    # Profit bars (shorter, darker color, overlaid)
    bars_profit = ax.barh(y, profits, height * 0.5, label='Profit', color='#1E8449', edgecolor='white')
    
    # Add value labels
    for i, (s, p, m) in enumerate(zip(sales, profits, margins)):
        # Sales label at end of bar
        ax.text(s + 0.15, i, f'KES {s:.2f}M', va='center', ha='left', fontsize=10, fontweight='bold', color='#2c3e50')
        # Profit and margin inside the sales bar
        ax.text(s * 0.5, i, f'Profit: KES {p:.2f}M ({m:.1f}%)', va='center', ha='center', 
                fontsize=9, fontweight='bold', color='white')
    
    ax.set_yticks(y)
    ax.set_yticklabels(rep_names, fontsize=11, fontweight='bold')
    ax.set_xlabel('Amount (Millions KES)', fontsize=11, fontweight='bold')
    ax.set_title('Sales & Profit Performance by Rep\nFebruary 1-14, 2026', fontsize=14, fontweight='bold', pad=15)
    ax.legend(loc='lower right', fontsize=10, frameon=True)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.set_xlim(0, max(sales) * 1.25)
    ax.invert_yaxis()  # Top performer at top
    
    plt.tight_layout()
    filename = f"{output_dir}/summary_sales_profit.png"
    plt.savefig(filename, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return filename

def create_margin_comparison_chart(reps_data, output_dir="charts"):
    os.makedirs(output_dir, exist_ok=True)
    sorted_reps = sorted(reps_data.items(), key=lambda x: x[1]['overall_margin'], reverse=True)
    rep_names = [r[0] for r in sorted_reps]
    margins = [r[1]['overall_margin'] for r in sorted_reps]
    colors_margin = ['#27ae60' if m > 10 else '#f39c12' if m > 7 else '#e74c3c' for m in margins]
    
    fig, ax = plt.subplots(figsize=(10, 5))
    bars = ax.barh(rep_names[::-1], margins[::-1], color=colors_margin[::-1], edgecolor='white', height=0.6)
    for bar, margin in zip(bars, margins[::-1]):
        width = bar.get_width()
        ax.annotate(f'{margin:.1f}%', xy=(width, bar.get_y() + bar.get_height()/2),
                    xytext=(5, 0), textcoords="offset points", ha='left', va='center', fontsize=10, fontweight='bold')
    ax.set_xlabel('Profit Margin (%)', fontsize=10, fontweight='bold')
    ax.set_title('Profit Margin Ranking by Rep', fontsize=12, fontweight='bold')
    ax.axvline(x=10, color='green', linestyle='--', linewidth=1.5, alpha=0.7, label='Target 10%')
    ax.legend(loc='lower right', fontsize=8)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    filename = f"{output_dir}/summary_margin_rank.png"
    plt.savefig(filename, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return filename

def create_category_breakdown_chart(reps_data, output_dir="charts"):
    """Horizontal stacked bar chart - much easier to read"""
    os.makedirs(output_dir, exist_ok=True)
    sorted_reps = sorted(reps_data.items(), key=lambda x: x[1]['total_sales'], reverse=True)
    rep_names = [r[0] for r in sorted_reps]
    
    fig, ax = plt.subplots(figsize=(14, 8))
    
    y = np.arange(len(rep_names))
    height = 0.7
    
    # Build category data
    left = np.zeros(len(rep_names))
    
    for cat in CAT_ORDER:
        cat_sales = []
        for rep_name, rep_data in sorted_reps:
            if cat in rep_data["categories"]:
                cat_sales.append(rep_data["categories"][cat]["sales_incl"] / 1000000)
            else:
                cat_sales.append(0)
        
        bars = ax.barh(y, cat_sales, height, left=left, label=cat, color=COLORS[cat], edgecolor='white', linewidth=0.5)
        
        # Add value labels inside bars (only if significant)
        for i, (val, l) in enumerate(zip(cat_sales, left)):
            if val > 0.3:  # Only label if > 300K
                ax.text(l + val/2, i, f'{val:.1f}M', va='center', ha='center', 
                       fontsize=9, fontweight='bold', color='white')
        
        left += np.array(cat_sales)
    
    # Add total labels at end
    for i, total in enumerate(left):
        ax.text(total + 0.1, i, f'Total: KES {total:.2f}M', va='center', ha='left', 
               fontsize=10, fontweight='bold', color='#2c3e50')
    
    ax.set_yticks(y)
    ax.set_yticklabels(rep_names, fontsize=11, fontweight='bold')
    ax.set_xlabel('Sales (Millions KES)', fontsize=11, fontweight='bold')
    ax.set_title('Category Sales Breakdown by Rep\nFebruary 1-21, 2026', fontsize=14, fontweight='bold', pad=15)
    
    # Legend at bottom
    ax.legend(title='Categories', loc='upper center', bbox_to_anchor=(0.5, -0.08), 
              ncol=4, fontsize=10, title_fontsize=11, frameon=True)
    
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.set_xlim(0, max(left) * 1.18)
    ax.invert_yaxis()  # Top performer at top
    
    plt.tight_layout()
    plt.subplots_adjust(bottom=0.15)
    filename = f"{output_dir}/summary_cat_breakdown.png"
    plt.savefig(filename, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return filename

def generate_pdf_report(reps_data, w1_reps, w2_reps, w3_reps, output_filename):
    print("Creating enhanced charts for PDF...")
    
    # Summary charts
    summary_sales = create_summary_sales_chart(reps_data)
    summary_margin = create_margin_comparison_chart(reps_data)
    summary_cat = create_category_breakdown_chart(reps_data)
    
    # NEW enhanced charts
    radar = create_radar_chart(reps_data)
    heatmap = create_heatmap(reps_data)
    rep_contribution = create_rep_contribution_chart(reps_data)
    week_grouped = create_week_grouped_bar(w1_reps, w2_reps, w3_reps)
    waterfall = create_profit_waterfall(reps_data)
    cat_donut = create_category_donut(reps_data)
    margin_line = create_margin_line_chart(reps_data)
    
    # Per-rep charts
    rep_charts = {}
    sorted_reps = sorted(reps_data.items(), key=lambda x: x[1]['total_sales'], reverse=True)
    for rep_name, rep_data in sorted_reps:
        donut = create_donut_chart(rep_name, rep_data)
        sales_profit = create_sales_vs_profit_bar(rep_name, rep_data)
        gauge = create_margin_gauge(rep_name, rep_data)
        rep_charts[rep_name] = {"donut": donut, "sales_profit": sales_profit, "gauge": gauge}
        print(f"  Created charts for {rep_name}")
    
    print("Building PDF...")
    
    doc = SimpleDocTemplate(output_filename, pagesize=A4, topMargin=0.4*inch, bottomMargin=0.4*inch, 
                           leftMargin=0.5*inch, rightMargin=0.5*inch)
    elements = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, textColor=colors.HexColor('#1a1a1a'), 
                                  spaceAfter=8, alignment=TA_CENTER, fontName='Helvetica-Bold')
    subtitle_style = ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#666666'), 
                                     spaceAfter=12, alignment=TA_CENTER)
    heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor('#2c3e50'), 
                                    spaceAfter=10, spaceBefore=10, fontName='Helvetica-Bold')
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#2c3e50'), spaceAfter=4)
    highlight_style = ParagraphStyle('Highlight', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#27ae60'), 
                                      spaceAfter=6, fontName='Helvetica-Bold')
    alert_style = ParagraphStyle('Alert', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#e74c3c'), 
                                  spaceAfter=6, fontName='Helvetica-Bold')
    
    total_sales = sum(r['total_sales'] for r in reps_data.values())
    total_profit = sum(r['total_profit'] for r in reps_data.values())
    overall_margin = (total_profit / total_sales * 100) if total_sales > 0 else 0
    
    # ===== PAGE 1: Title and Executive Summary =====
    elements.append(Spacer(1, 0.2*inch))
    elements.append(Paragraph("SALES REP PERFORMANCE ANALYSIS", title_style))
    elements.append(Paragraph("BOMAS Hardware Store - February 2026", subtitle_style))
    elements.append(Paragraph("February 1-21, 2026 (18 Working Days)", subtitle_style))
    elements.append(Spacer(1, 0.15*inch))
    
    elements.append(Paragraph("EXECUTIVE SUMMARY", heading_style))
    
    w1_total = sum(r['total_sales'] for r in w1_reps.values())
    w2_total = sum(r['total_sales'] for r in w2_reps.values())
    w3_total = sum(r['total_sales'] for r in w3_reps.values())
    w1_profit = sum(r['total_profit'] for r in w1_reps.values())
    w2_profit = sum(r['total_profit'] for r in w2_reps.values())
    w3_profit = sum(r['total_profit'] for r in w3_reps.values())
    w1_margin = (w1_profit / w1_total * 100) if w1_total > 0 else 0
    w2_margin = (w2_profit / w2_total * 100) if w2_total > 0 else 0
    w3_margin = (w3_profit / w3_total * 100) if w3_total > 0 else 0
    
    summary_data = [
        ["Metric", "Week 1 (Feb 1-7)", "Week 2 (Feb 9-14)", "Week 3 (Feb 16-21)", "Combined"],
        ["Total Sales", f"KES {w1_total:,.0f}", f"KES {w2_total:,.0f}", f"KES {w3_total:,.0f}", f"KES {total_sales:,.0f}"],
        ["Total Profit", f"KES {w1_profit:,.0f}", f"KES {w2_profit:,.0f}", f"KES {w3_profit:,.0f}", f"KES {total_profit:,.0f}"],
        ["Daily Average", f"KES {w1_total/6:,.0f}", f"KES {w2_total/6:,.0f}", f"KES {w3_total/6:,.0f}", f"KES {total_sales/18:,.0f}"],
        ["Overall Margin", f"{w1_margin:.1f}%", f"{w2_margin:.1f}%", f"{w3_margin:.1f}%", f"{overall_margin:.1f}%"],
    ]
    
    summary_table = Table(summary_data, colWidths=[1.1*inch, 1.3*inch, 1.3*inch, 1.3*inch, 1.3*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#ecf0f1')]),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
    ]))
    elements.append(summary_table)
    elements.append(Spacer(1, 0.15*inch))
    
    # Weekly progress commentary
    daily_w1w2 = ((w2_total/6 - w1_total/6) / (w1_total/6)) * 100
    daily_w2w3 = ((w3_total/6 - w2_total/6) / (w2_total/6)) * 100
    
    if daily_w1w2 > 0:
        elements.append(Paragraph(f"▲ W1→W2: Daily sales increased by {daily_w1w2:.1f}%", highlight_style))
    else:
        elements.append(Paragraph(f"▼ W1→W2: Daily sales decreased by {abs(daily_w1w2):.1f}%", alert_style))
    
    if daily_w2w3 > 0:
        elements.append(Paragraph(f"▲ W2→W3: Daily sales increased by {daily_w2w3:.1f}%", highlight_style))
    else:
        elements.append(Paragraph(f"▼ W2→W3: Daily sales decreased by {abs(daily_w2w3):.1f}%", alert_style))
    
    # Overall trend
    overall_daily_change = ((w3_total/6 - w1_total/6) / (w1_total/6)) * 100
    if overall_daily_change > 0:
        elements.append(Paragraph(f"📈 Overall 3-week trend: Daily sales UP {overall_daily_change:.1f}% from Week 1 to Week 3", highlight_style))
    else:
        elements.append(Paragraph(f"📉 Overall 3-week trend: Daily sales DOWN {abs(overall_daily_change):.1f}% from Week 1 to Week 3", alert_style))
    
    elements.append(Spacer(1, 0.1*inch))
    elements.append(Paragraph("CATEGORY SALES DISTRIBUTION", heading_style))
    if cat_donut and os.path.exists(cat_donut):
        elements.append(Image(cat_donut, width=5.5*inch, height=4*inch))
    
    elements.append(PageBreak())
    
    # ===== PAGE 2: Week Comparison =====
    elements.append(Paragraph("WEEK-OVER-WEEK COMPARISON", heading_style))
    if week_grouped and os.path.exists(week_grouped):
        elements.append(Image(week_grouped, width=6.5*inch, height=3.5*inch))
    
    elements.append(Spacer(1, 0.15*inch))
    elements.append(Paragraph("SALES VS PROFIT COMPARISON", heading_style))
    if summary_sales and os.path.exists(summary_sales):
        elements.append(Image(summary_sales, width=6.5*inch, height=3.5*inch))
    
    elements.append(PageBreak())
    
    # ===== PAGE 3: Category Analysis =====
    elements.append(Paragraph("CATEGORY BREAKDOWN BY REP", heading_style))
    if summary_cat and os.path.exists(summary_cat):
        elements.append(Image(summary_cat, width=6.5*inch, height=3.5*inch))
    
    elements.append(Spacer(1, 0.15*inch))
    elements.append(Paragraph("REP CONTRIBUTION ANALYSIS", heading_style))
    if rep_contribution and os.path.exists(rep_contribution):
        elements.append(Image(rep_contribution, width=6.5*inch, height=3.8*inch))
    
    elements.append(PageBreak())
    
    # ===== PAGE 4: Performance Analysis =====
    elements.append(Paragraph("PERFORMANCE RADAR - ALL SALES REPS", heading_style))
    if radar and os.path.exists(radar):
        elements.append(Image(radar, width=6*inch, height=4.5*inch))
    
    elements.append(Spacer(1, 0.15*inch))
    elements.append(Paragraph("PROFIT MARGIN RANKING", heading_style))
    if summary_margin and os.path.exists(summary_margin):
        elements.append(Image(summary_margin, width=6*inch, height=3*inch))
    
    elements.append(PageBreak())
    
    # ===== PAGE 5: Margin Analysis =====
    elements.append(Paragraph("PROFIT MARGIN HEATMAP", heading_style))
    if heatmap and os.path.exists(heatmap):
        elements.append(Image(heatmap, width=6*inch, height=3.5*inch))
    
    elements.append(Spacer(1, 0.15*inch))
    elements.append(Paragraph("MARGIN TREND BY CATEGORY", heading_style))
    if margin_line and os.path.exists(margin_line):
        elements.append(Image(margin_line, width=6*inch, height=3.5*inch))
    
    elements.append(PageBreak())
    
    # ===== PAGE 6: Profit Waterfall =====
    elements.append(Paragraph("PROFIT CONTRIBUTION WATERFALL", heading_style))
    if waterfall and os.path.exists(waterfall):
        elements.append(Image(waterfall, width=6.5*inch, height=3.8*inch))
    
    elements.append(Spacer(1, 0.2*inch))
    
    # Quick insights
    elements.append(Paragraph("KEY INSIGHTS", heading_style))
    insights = []
    top_rep = sorted_reps[0]
    top_margin_rep = max(reps_data.items(), key=lambda x: x[1]['overall_margin'])
    insights.append(f"◆ Top Seller: {top_rep[0]} with KES {top_rep[1]['total_sales']/1000000:.2f}M ({top_rep[1]['total_sales']/total_sales*100:.1f}% of total)")
    insights.append(f"◆ Highest Margin: {top_margin_rep[0]} at {top_margin_rep[1]['overall_margin']:.1f}%")
    insights.append(f"◆ W1→W2 daily change: {daily_w1w2:+.1f}% | W2→W3 daily change: {daily_w2w3:+.1f}%")
    insights.append(f"◆ Overall 3-week daily trend: {overall_daily_change:+.1f}%")
    insights.append(f"◆ Total profit: KES {total_profit/1000000:.2f}M at {overall_margin:.1f}% margin")
    
    for insight in insights:
        elements.append(Paragraph(insight, body_style))
    
    elements.append(PageBreak())
    
    # ===== PAGE 7: PERFORMANCE COMMENTARY =====
    # Generate performance analysis
    perf_analysis = analyze_performance(reps_data, w1_reps, w2_reps, w3_reps)
    
    elements.append(Paragraph("PERFORMANCE COMMENTARY & AWARDS", title_style))
    elements.append(Spacer(1, 0.15*inch))
    
    # Department heads section
    elements.append(Paragraph("DEPARTMENT LEADERSHIP", heading_style))
    
    dept_note_style = ParagraphStyle('DeptNote', parent=styles['Normal'], fontSize=9, 
                                      textColor=colors.HexColor('#2c3e50'), spaceAfter=3, leftIndent=20)
    badge_style = ParagraphStyle('Badge', parent=styles['Normal'], fontSize=11, 
                                  textColor=colors.HexColor('#8e44ad'), spaceAfter=4, fontName='Helvetica-Bold')
    role_style = ParagraphStyle('Role', parent=styles['Normal'], fontSize=10, 
                                 textColor=colors.HexColor('#34495e'), spaceAfter=2, fontName='Helvetica-Oblique')
    
    dept_heads_data = [
        ["Department", "Head", "Dept Sales", "Dept Margin", "Dept Share"],
    ]
    
    for dept in ['PLUMBING', 'PAINTS', 'ELECTRICALS']:
        head = DEPARTMENT_HEADS.get(dept)
        if head and head in reps_data:
            rep_data = reps_data[head]
            if dept in rep_data['categories']:
                dept_sales = rep_data['categories'][dept]['sales_incl']
                dept_margin = rep_data['categories'][dept]['margin_pct']
                # Calculate department total
                dept_total = sum(r['categories'].get(dept, {}).get('sales_incl', 0) for r in reps_data.values())
                dept_share = (dept_sales / dept_total * 100) if dept_total > 0 else 0
                dept_heads_data.append([
                    dept, head, f"KES {dept_sales:,.0f}", f"{dept_margin:.1f}%", f"{dept_share:.1f}%"
                ])
    
    if len(dept_heads_data) > 1:
        dept_table = Table(dept_heads_data, colWidths=[1.5*inch, 1.5*inch, 1.5*inch, 1.2*inch, 1.2*inch])
        dept_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8e44ad')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f0fa')]),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        elements.append(dept_table)
    
    elements.append(PageBreak())
    
    # ===== PAGE 8: REP-BY-REP COMMENTARY =====
    elements.append(Paragraph("REP-BY-REP COMMENTARY", heading_style))
    elements.append(Spacer(1, 0.1*inch))
    
    note_style = ParagraphStyle('Note', parent=styles['Normal'], fontSize=8, 
                                 textColor=colors.HexColor('#555555'), spaceAfter=2, leftIndent=15)
    rep_title_style = ParagraphStyle('RepTitle', parent=styles['Normal'], fontSize=11, 
                                      textColor=colors.HexColor('#2c3e50'), spaceAfter=4, fontName='Helvetica-Bold')
    
    for rep_name, rep_data in sorted_reps:
        analysis = perf_analysis.get(rep_name, {})
        
        # Rep header with role
        role = analysis.get('role', 'Sales Rep')
        badges_str = ' '.join(analysis.get('badges', []))
        
        elements.append(Paragraph(f"{rep_name.upper()} - {role}", rep_title_style))
        if badges_str:
            elements.append(Paragraph(badges_str, badge_style))
        
        # Performance notes
        for note in analysis.get('notes', [])[:4]:  # Limit to 4 notes per rep
            elements.append(Paragraph(f"• {note}", note_style))
        
        elements.append(Spacer(1, 0.08*inch))
    
    elements.append(PageBreak())
    
    # ===== INDIVIDUAL REP PAGES =====
    for rep_name, rep_data in sorted_reps:
        analysis = perf_analysis.get(rep_name, {})
        role = analysis.get('role', 'Sales Rep')
        
        elements.append(Paragraph(f"SALES REP: {rep_name.upper()}", heading_style))
        elements.append(Paragraph(role, role_style))
        
        # Badges
        badges = analysis.get('badges', [])
        if badges:
            badges_str = '  '.join(badges)
            elements.append(Paragraph(badges_str, badge_style))
        
        # Stats table
        rep_table_data = [
            ["Total Sales", "Total Profit", "Margin", "Items Sold", "Rank"],
            [f"KES {rep_data['total_sales']:,.0f}", f"KES {rep_data['total_profit']:,.0f}",
             f"{rep_data['overall_margin']:.1f}%", f"{rep_data['total_qty']:,.0f}",
             f"#{analysis.get('sales_rank', '-')}"]
        ]
        rep_table = Table(rep_table_data, colWidths=[1.4*inch, 1.4*inch, 1*inch, 1*inch, 0.8*inch])
        rep_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34495e')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ('TOPPADDING', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        elements.append(rep_table)
        elements.append(Spacer(1, 0.1*inch))
        
        # Donut chart
        if rep_charts[rep_name]["donut"] and os.path.exists(rep_charts[rep_name]["donut"]):
            elements.append(Image(rep_charts[rep_name]["donut"], width=5*inch, height=3*inch))
        
        elements.append(Spacer(1, 0.1*inch))
        
        # Sales vs Profit bar
        if rep_charts[rep_name]["sales_profit"] and os.path.exists(rep_charts[rep_name]["sales_profit"]):
            elements.append(Image(rep_charts[rep_name]["sales_profit"], width=5.5*inch, height=2.5*inch))
        
        elements.append(Spacer(1, 0.1*inch))
        
        # Category breakdown table
        cat_data = [["Category", "Sales", "Profit", "Margin", "Qty"]]
        for cat in CAT_ORDER:
            if cat in rep_data["categories"]:
                c = rep_data["categories"][cat]
                cat_data.append([cat, f"KES {c['sales_incl']:,.0f}", f"KES {c['profit']:,.0f}",
                                f"{c['margin_pct']:.1f}%", f"{c['qty']:,.0f}"])
        
        if len(cat_data) > 1:
            cat_table = Table(cat_data, colWidths=[1.8*inch, 1.4*inch, 1.2*inch, 1*inch, 1*inch])
            cat_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#ecf0f1')]),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            elements.append(cat_table)
        
        # Performance notes for this rep
        elements.append(Spacer(1, 0.1*inch))
        elements.append(Paragraph("PERFORMANCE NOTES", ParagraphStyle('NotesHead', parent=styles['Normal'], 
                                   fontSize=10, textColor=colors.HexColor('#8e44ad'), fontName='Helvetica-Bold', spaceAfter=4)))
        
        rep_notes = analysis.get('notes', [])
        for note in rep_notes[:5]:  # Limit to 5 notes
            note_color = '#27ae60' if '✓' in note or 'Excellent' in note or 'Outstanding' in note else (
                         '#e74c3c' if '⚠️' in note or 'below' in note or 'declined' in note else '#555555')
            elements.append(Paragraph(f"• {note}", ParagraphStyle('NoteItem', parent=styles['Normal'], 
                            fontSize=8, textColor=colors.HexColor(note_color), spaceAfter=2, leftIndent=10)))
        
        elements.append(PageBreak())
    
    doc.build(elements)
    print(f"PDF generated: {output_filename}")

def export_to_excel(reps_data, w1_reps, w2_reps, w3_reps, perf_analysis, output_filename="sales_rep_analysis_feb2026.xlsx"):
    """Export sales analysis data to Excel with multiple sheets"""
    wb = Workbook()
    
    # Styles
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")
    header_fill_green = PatternFill(start_color="27AE60", end_color="27AE60", fill_type="solid")
    header_fill_blue = PatternFill(start_color="3498DB", end_color="3498DB", fill_type="solid")
    header_fill_purple = PatternFill(start_color="8E44AD", end_color="8E44AD", fill_type="solid")
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    currency_format = '#,##0'
    percent_format = '0.0%'
    
    # ===== SHEET 1: Summary =====
    ws_summary = wb.active
    ws_summary.title = "Summary"
    
    # Title
    ws_summary['A1'] = "SALES REP PERFORMANCE ANALYSIS"
    ws_summary['A1'].font = Font(bold=True, size=16)
    ws_summary['A2'] = "BOMAS  - February 1-21, 2026"
    ws_summary['A2'].font = Font(size=12, italic=True)
    ws_summary['A3'] = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    
    # Overall metrics
    total_sales = sum(r['total_sales'] for r in reps_data.values())
    total_profit = sum(r['total_profit'] for r in reps_data.values())
    overall_margin = (total_profit / total_sales * 100) if total_sales > 0 else 0
    
    ws_summary['A5'] = "OVERALL METRICS"
    ws_summary['A5'].font = Font(bold=True, size=12)
    
    metrics = [
        ["Metric", "Value"],
        ["Total Sales", total_sales],
        ["Total Profit", total_profit],
        ["Overall Margin", overall_margin / 100],
        ["Daily Average (18 days)", total_sales / 18],
        ["Total Reps", len(reps_data)]
    ]
    
    for row_idx, row_data in enumerate(metrics, start=6):
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws_summary.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            if row_idx == 6:
                cell.font = header_font
                cell.fill = header_fill
            elif col_idx == 2 and row_idx in [7, 8, 10]:
                cell.number_format = currency_format
            elif col_idx == 2 and row_idx == 9:
                cell.number_format = percent_format
    
    # Week comparison
    ws_summary['A14'] = "WEEK COMPARISON"
    ws_summary['A14'].font = Font(bold=True, size=12)
    
    w1_total = sum(r['total_sales'] for r in w1_reps.values())
    w2_total = sum(r['total_sales'] for r in w2_reps.values())
    w3_total = sum(r['total_sales'] for r in w3_reps.values())
    w1_profit = sum(r['total_profit'] for r in w1_reps.values())
    w2_profit = sum(r['total_profit'] for r in w2_reps.values())
    w3_profit = sum(r['total_profit'] for r in w3_reps.values())
    
    week_data = [
        ["Week", "Sales", "Profit", "Margin", "Daily Avg"],
        ["Week 1 (Feb 1-7)", w1_total, w1_profit, w1_profit/w1_total if w1_total > 0 else 0, w1_total/6],
        ["Week 2 (Feb 9-14)", w2_total, w2_profit, w2_profit/w2_total if w2_total > 0 else 0, w2_total/6],
        ["Week 3 (Feb 16-21)", w3_total, w3_profit, w3_profit/w3_total if w3_total > 0 else 0, w3_total/6],
        ["W1→W2 Change", w2_total - w1_total, w2_profit - w1_profit, (w2_profit/w2_total - w1_profit/w1_total) if w1_total > 0 and w2_total > 0 else 0, (w2_total - w1_total)/6],
        ["W2→W3 Change", w3_total - w2_total, w3_profit - w2_profit, (w3_profit/w3_total - w2_profit/w2_total) if w2_total > 0 and w3_total > 0 else 0, (w3_total - w2_total)/6],
    ]
    
    for row_idx, row_data in enumerate(week_data, start=15):
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws_summary.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            if row_idx == 15:
                cell.font = header_font
                cell.fill = header_fill_blue
            elif col_idx in [2, 3, 5]:
                cell.number_format = currency_format
            elif col_idx == 4:
                cell.number_format = percent_format
    
    # Set column widths
    ws_summary.column_dimensions['A'].width = 22
    ws_summary.column_dimensions['B'].width = 18
    ws_summary.column_dimensions['C'].width = 15
    ws_summary.column_dimensions['D'].width = 12
    ws_summary.column_dimensions['E'].width = 15
    
    # ===== SHEET 2: Rep Performance =====
    ws_reps = wb.create_sheet("Rep Performance")
    
    sorted_reps = sorted(reps_data.items(), key=lambda x: x[1]['total_sales'], reverse=True)
    
    rep_headers = ["Rank", "Rep Name", "Role", "Total Sales", "Total Profit", "Margin %", 
                   "Items Sold", "Margin Rank", "Week Improvement %"]
    
    for col_idx, header in enumerate(rep_headers, start=1):
        cell = ws_reps.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = Alignment(horizontal='center')
    
    for row_idx, (rep_name, rep_data) in enumerate(sorted_reps, start=2):
        analysis = perf_analysis.get(rep_name, {})
        row_values = [
            analysis.get('sales_rank', row_idx - 1),
            rep_name,
            analysis.get('role', 'Sales Rep'),
            rep_data['total_sales'],
            rep_data['total_profit'],
            rep_data['overall_margin'] / 100,
            rep_data['total_qty'],
            analysis.get('margin_rank', '-'),
            analysis.get('week_improvement', 0) / 100
        ]
        for col_idx, value in enumerate(row_values, start=1):
            cell = ws_reps.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            if col_idx in [4, 5]:
                cell.number_format = currency_format
            elif col_idx in [6, 9]:
                cell.number_format = percent_format
            elif col_idx == 7:
                cell.number_format = '#,##0.0'
    
    # Set column widths
    col_widths = [6, 18, 25, 15, 15, 12, 12, 12, 18]
    for idx, width in enumerate(col_widths, start=1):
        ws_reps.column_dimensions[get_column_letter(idx)].width = width
    
    # ===== SHEET 3: Category Breakdown =====
    ws_cat = wb.create_sheet("Category Breakdown")
    
    cat_headers = ["Rep Name", "Category", "Sales", "Profit", "Margin %", "Qty"]
    
    for col_idx, header in enumerate(cat_headers, start=1):
        cell = ws_cat.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill_green
        cell.border = thin_border
    
    row_idx = 2
    for rep_name, rep_data in sorted_reps:
        for cat in CAT_ORDER:
            if cat in rep_data['categories']:
                cat_data = rep_data['categories'][cat]
                row_values = [
                    rep_name,
                    cat,
                    cat_data['sales_incl'],
                    cat_data['profit'],
                    cat_data['margin_pct'] / 100,
                    cat_data['qty']
                ]
                for col_idx, value in enumerate(row_values, start=1):
                    cell = ws_cat.cell(row=row_idx, column=col_idx, value=value)
                    cell.border = thin_border
                    if col_idx == 3 or col_idx == 4:
                        cell.number_format = currency_format
                    elif col_idx == 5:
                        cell.number_format = percent_format
                    elif col_idx == 6:
                        cell.number_format = '#,##0.0'
                row_idx += 1
    
    # Set column widths
    cat_col_widths = [18, 20, 15, 15, 12, 10]
    for idx, width in enumerate(cat_col_widths, start=1):
        ws_cat.column_dimensions[get_column_letter(idx)].width = width
    
    # ===== SHEET 4: Department Heads =====
    ws_dept = wb.create_sheet("Department Heads")
    
    dept_headers = ["Department", "Head", "Dept Sales", "Dept Profit", "Dept Margin", "Dept Share %"]
    
    for col_idx, header in enumerate(dept_headers, start=1):
        cell = ws_dept.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill_purple
        cell.border = thin_border
    
    row_idx = 2
    for dept in ['PLUMBING', 'PAINTS', 'ELECTRICALS']:
        head = DEPARTMENT_HEADS.get(dept)
        if head and head in reps_data:
            rep_data = reps_data[head]
            if dept in rep_data['categories']:
                dept_sales = rep_data['categories'][dept]['sales_incl']
                dept_profit = rep_data['categories'][dept]['profit']
                dept_margin = rep_data['categories'][dept]['margin_pct']
                # Calculate department total
                dept_total = sum(r['categories'].get(dept, {}).get('sales_incl', 0) for r in reps_data.values())
                dept_share = (dept_sales / dept_total * 100) if dept_total > 0 else 0
                
                row_values = [dept, head, dept_sales, dept_profit, dept_margin / 100, dept_share / 100]
                for col_idx, value in enumerate(row_values, start=1):
                    cell = ws_dept.cell(row=row_idx, column=col_idx, value=value)
                    cell.border = thin_border
                    if col_idx in [3, 4]:
                        cell.number_format = currency_format
                    elif col_idx in [5, 6]:
                        cell.number_format = percent_format
                row_idx += 1
    
    # General Hardware team
    ws_dept.cell(row=row_idx + 1, column=1, value="GENERAL HARDWARE TEAM").font = Font(bold=True)
    gh_headers = ["Rep", "GH Sales", "GH Profit", "GH Margin", "GH Share %"]
    row_idx += 2
    for col_idx, header in enumerate(gh_headers, start=1):
        cell = ws_dept.cell(row=row_idx, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill_blue
        cell.border = thin_border
    
    row_idx += 1
    gh_total = sum(r['categories'].get('GENERAL HARDWARE', {}).get('sales_incl', 0) for r in reps_data.values())
    
    for rep_name in DEPARTMENT_HEADS.get('GENERAL HARDWARE', []):
        if rep_name in reps_data and 'GENERAL HARDWARE' in reps_data[rep_name]['categories']:
            gh_data = reps_data[rep_name]['categories']['GENERAL HARDWARE']
            gh_share = (gh_data['sales_incl'] / gh_total * 100) if gh_total > 0 else 0
            row_values = [rep_name, gh_data['sales_incl'], gh_data['profit'], 
                         gh_data['margin_pct'] / 100, gh_share / 100]
            for col_idx, value in enumerate(row_values, start=1):
                cell = ws_dept.cell(row=row_idx, column=col_idx, value=value)
                cell.border = thin_border
                if col_idx in [2, 3]:
                    cell.number_format = currency_format
                elif col_idx in [4, 5]:
                    cell.number_format = percent_format
            row_idx += 1
    
    # Set column widths
    dept_col_widths = [20, 18, 15, 15, 15, 15]
    for idx, width in enumerate(dept_col_widths, start=1):
        ws_dept.column_dimensions[get_column_letter(idx)].width = width
    
    # ===== SHEET 5: Performance Notes =====
    ws_notes = wb.create_sheet("Performance Notes")
    
    note_headers = ["Rep Name", "Role", "Sales Rank", "Margin Rank", "Week Growth", "Performance Notes"]
    
    for col_idx, header in enumerate(note_headers, start=1):
        cell = ws_notes.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
    
    row_idx = 2
    for rep_name, rep_data in sorted_reps:
        analysis = perf_analysis.get(rep_name, {})
        notes = analysis.get('notes', [])
        notes_text = " | ".join(notes[:4])  # Join first 4 notes
        
        row_values = [
            rep_name,
            analysis.get('role', 'Sales Rep'),
            analysis.get('sales_rank', '-'),
            analysis.get('margin_rank', '-'),
            f"{analysis.get('week_improvement', 0):+.1f}%",
            notes_text
        ]
        for col_idx, value in enumerate(row_values, start=1):
            cell = ws_notes.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            if col_idx == 6:
                cell.alignment = Alignment(wrap_text=True)
        row_idx += 1
    
    # Set column widths
    note_col_widths = [18, 25, 12, 12, 12, 80]
    for idx, width in enumerate(note_col_widths, start=1):
        ws_notes.column_dimensions[get_column_letter(idx)].width = width
    
    # ===== SHEET 6: Charts =====
    ws_charts = wb.create_sheet("Charts")
    
    # Create data for charts - Rep Sales and Profit
    ws_charts['A1'] = "SALES REP CHARTS DATA"
    ws_charts['A1'].font = Font(bold=True, size=14)
    
    # Rep data table for charts
    chart_headers = ["Rep Name", "Sales", "Profit", "Margin %", "Items Sold"]
    for col_idx, header in enumerate(chart_headers, start=1):
        cell = ws_charts.cell(row=3, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
    
    for row_idx, (rep_name, rep_data) in enumerate(sorted_reps, start=4):
        ws_charts.cell(row=row_idx, column=1, value=rep_name).border = thin_border
        ws_charts.cell(row=row_idx, column=2, value=rep_data['total_sales']).border = thin_border
        ws_charts.cell(row=row_idx, column=2).number_format = currency_format
        ws_charts.cell(row=row_idx, column=3, value=rep_data['total_profit']).border = thin_border
        ws_charts.cell(row=row_idx, column=3).number_format = currency_format
        ws_charts.cell(row=row_idx, column=4, value=rep_data['overall_margin'] / 100).border = thin_border
        ws_charts.cell(row=row_idx, column=4).number_format = percent_format
        ws_charts.cell(row=row_idx, column=5, value=rep_data['total_qty']).border = thin_border
    
    # Set column widths
    chart_col_widths = [18, 15, 15, 12, 12]
    for idx, width in enumerate(chart_col_widths, start=1):
        ws_charts.column_dimensions[get_column_letter(idx)].width = width
    
    # ----- CHART 1: Sales Bar Chart -----
    sales_chart = BarChart()
    sales_chart.type = "col"
    sales_chart.style = 10
    sales_chart.title = "Sales by Rep"
    sales_chart.y_axis.title = "Sales (KES)"
    sales_chart.x_axis.title = "Sales Rep"
    
    sales_data = Reference(ws_charts, min_col=2, min_row=3, max_row=3 + len(sorted_reps))
    sales_cats = Reference(ws_charts, min_col=1, min_row=4, max_row=3 + len(sorted_reps))
    sales_chart.add_data(sales_data, titles_from_data=True)
    sales_chart.set_categories(sales_cats)
    sales_chart.width = 15
    sales_chart.height = 10
    
    ws_charts.add_chart(sales_chart, "G3")
    
    # ----- CHART 2: Profit Bar Chart -----
    profit_chart = BarChart()
    profit_chart.type = "col"
    profit_chart.style = 11
    profit_chart.title = "Profit by Rep"
    profit_chart.y_axis.title = "Profit (KES)"
    profit_chart.x_axis.title = "Sales Rep"
    
    profit_data = Reference(ws_charts, min_col=3, min_row=3, max_row=3 + len(sorted_reps))
    profit_chart.add_data(profit_data, titles_from_data=True)
    profit_chart.set_categories(sales_cats)
    profit_chart.width = 15
    profit_chart.height = 10
    
    ws_charts.add_chart(profit_chart, "G20")
    
    # ----- CHART 3: Margin Bar Chart -----
    margin_chart = BarChart()
    margin_chart.type = "bar"
    margin_chart.style = 12
    margin_chart.title = "Profit Margin by Rep"
    margin_chart.x_axis.title = "Margin %"
    margin_chart.y_axis.title = "Sales Rep"
    
    margin_data = Reference(ws_charts, min_col=4, min_row=3, max_row=3 + len(sorted_reps))
    margin_chart.add_data(margin_data, titles_from_data=True)
    margin_chart.set_categories(sales_cats)
    margin_chart.width = 15
    margin_chart.height = 10
    
    ws_charts.add_chart(margin_chart, "G37")
    
    # ----- Category Data for Pie Chart -----
    ws_charts['A15'] = "CATEGORY BREAKDOWN"
    ws_charts['A15'].font = Font(bold=True, size=12)
    
    cat_headers = ["Category", "Total Sales", "Total Profit"]
    for col_idx, header in enumerate(cat_headers, start=1):
        cell = ws_charts.cell(row=16, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill_green
        cell.border = thin_border
    
    cat_row = 17
    for cat in CAT_ORDER:
        cat_sales = sum(r['categories'].get(cat, {}).get('sales_incl', 0) for r in reps_data.values())
        cat_profit = sum(r['categories'].get(cat, {}).get('profit', 0) for r in reps_data.values())
        ws_charts.cell(row=cat_row, column=1, value=cat).border = thin_border
        ws_charts.cell(row=cat_row, column=2, value=cat_sales).border = thin_border
        ws_charts.cell(row=cat_row, column=2).number_format = currency_format
        ws_charts.cell(row=cat_row, column=3, value=cat_profit).border = thin_border
        ws_charts.cell(row=cat_row, column=3).number_format = currency_format
        cat_row += 1
    
    # ----- CHART 4: Category Pie Chart -----
    pie_chart = PieChart()
    pie_chart.title = "Sales by Category"
    
    pie_labels = Reference(ws_charts, min_col=1, min_row=17, max_row=20)
    pie_data = Reference(ws_charts, min_col=2, min_row=16, max_row=20)
    pie_chart.add_data(pie_data, titles_from_data=True)
    pie_chart.set_categories(pie_labels)
    pie_chart.width = 12
    pie_chart.height = 10
    
    # Add data labels
    pie_chart.dataLabels = DataLabelList()
    pie_chart.dataLabels.showPercent = True
    pie_chart.dataLabels.showCatName = True
    pie_chart.dataLabels.showVal = False
    
    ws_charts.add_chart(pie_chart, "A22")
    
    # ----- Week Comparison Data -----
    ws_charts['A35'] = "WEEK COMPARISON"
    ws_charts['A35'].font = Font(bold=True, size=12)
    
    week_headers = ["Rep", "Week 1 Sales", "Week 2 Sales", "Week 3 Sales"]
    for col_idx, header in enumerate(week_headers, start=1):
        cell = ws_charts.cell(row=36, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill_blue
        cell.border = thin_border
    
    week_row = 37
    for rep_name, _ in sorted_reps:
        w1_sales = w1_reps.get(rep_name, {}).get('total_sales', 0)
        w2_sales = w2_reps.get(rep_name, {}).get('total_sales', 0)
        w3_sales = w3_reps.get(rep_name, {}).get('total_sales', 0)
        ws_charts.cell(row=week_row, column=1, value=rep_name).border = thin_border
        ws_charts.cell(row=week_row, column=2, value=w1_sales).border = thin_border
        ws_charts.cell(row=week_row, column=2).number_format = currency_format
        ws_charts.cell(row=week_row, column=3, value=w2_sales).border = thin_border
        ws_charts.cell(row=week_row, column=3).number_format = currency_format
        ws_charts.cell(row=week_row, column=4, value=w3_sales).border = thin_border
        ws_charts.cell(row=week_row, column=4).number_format = currency_format
        week_row += 1
    
    # ----- CHART 5: Week Comparison Bar Chart -----
    week_chart = BarChart()
    week_chart.type = "col"
    week_chart.style = 10
    week_chart.title = "Week 1 vs Week 2 vs Week 3 Sales"
    week_chart.y_axis.title = "Sales (KES)"
    week_chart.x_axis.title = "Sales Rep"
    week_chart.grouping = "clustered"
    
    week_data = Reference(ws_charts, min_col=2, max_col=4, min_row=36, max_row=36 + len(sorted_reps))
    week_cats = Reference(ws_charts, min_col=1, min_row=37, max_row=36 + len(sorted_reps))
    week_chart.add_data(week_data, titles_from_data=True)
    week_chart.set_categories(week_cats)
    week_chart.width = 18
    week_chart.height = 10
    
    ws_charts.add_chart(week_chart, "A47")
    
    # ----- Radar Chart Data -----
    ws_charts['A60'] = "PERFORMANCE RADAR DATA"
    ws_charts['A60'].font = Font(bold=True, size=12)
    
    # Normalize data for radar chart (0-100 scale)
    max_sales = max(r['total_sales'] for r in reps_data.values())
    max_profit = max(r['total_profit'] for r in reps_data.values())
    max_qty = max(r['total_qty'] for r in reps_data.values())
    max_margin = max(r['overall_margin'] for r in reps_data.values())
    
    radar_headers = ["Metric"] + [rep_name for rep_name, _ in sorted_reps]
    for col_idx, header in enumerate(radar_headers, start=1):
        cell = ws_charts.cell(row=61, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill_purple
        cell.border = thin_border
    
    # Add normalized metrics
    metrics = ["Sales", "Profit", "Items Sold", "Margin"]
    for metric_idx, metric in enumerate(metrics, start=62):
        ws_charts.cell(row=metric_idx, column=1, value=metric).border = thin_border
        for col_idx, (rep_name, rep_data) in enumerate(sorted_reps, start=2):
            if metric == "Sales":
                val = (rep_data['total_sales'] / max_sales) * 100 if max_sales > 0 else 0
            elif metric == "Profit":
                val = (rep_data['total_profit'] / max_profit) * 100 if max_profit > 0 else 0
            elif metric == "Items Sold":
                val = (rep_data['total_qty'] / max_qty) * 100 if max_qty > 0 else 0
            else:
                val = (rep_data['overall_margin'] / max_margin) * 100 if max_margin > 0 else 0
            ws_charts.cell(row=metric_idx, column=col_idx, value=val).border = thin_border
    
    # ----- CHART 6: Radar Chart -----
    radar_chart = RadarChart()
    radar_chart.type = "filled"
    radar_chart.title = "Performance Radar - All Reps"
    radar_chart.style = 26
    
    radar_labels = Reference(ws_charts, min_col=1, min_row=62, max_row=65)
    radar_data = Reference(ws_charts, min_col=2, max_col=1 + len(sorted_reps), min_row=61, max_row=65)
    radar_chart.add_data(radar_data, titles_from_data=True)
    radar_chart.set_categories(radar_labels)
    radar_chart.width = 18
    radar_chart.height = 12
    
    ws_charts.add_chart(radar_chart, "A68")
    
    # ----- Sales vs Profit Line Chart Data -----
    ws_charts['A85'] = "SALES & PROFIT TREND"
    ws_charts['A85'].font = Font(bold=True, size=12)
    
    trend_headers = ["Rep", "Sales (M)", "Profit (K)"]
    for col_idx, header in enumerate(trend_headers, start=1):
        cell = ws_charts.cell(row=86, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
    
    trend_row = 87
    for rep_name, rep_data in sorted_reps:
        ws_charts.cell(row=trend_row, column=1, value=rep_name).border = thin_border
        ws_charts.cell(row=trend_row, column=2, value=rep_data['total_sales'] / 1000000).border = thin_border
        ws_charts.cell(row=trend_row, column=3, value=rep_data['total_profit'] / 1000).border = thin_border
        trend_row += 1
    
    # ----- CHART 7: Sales vs Profit Line Chart -----
    line_chart = LineChart()
    line_chart.title = "Sales vs Profit Trend"
    line_chart.style = 10
    line_chart.y_axis.title = "Value"
    line_chart.x_axis.title = "Rep (by Sales Rank)"
    
    line_data = Reference(ws_charts, min_col=2, max_col=3, min_row=86, max_row=86 + len(sorted_reps))
    line_cats = Reference(ws_charts, min_col=1, min_row=87, max_row=86 + len(sorted_reps))
    line_chart.add_data(line_data, titles_from_data=True)
    line_chart.set_categories(line_cats)
    line_chart.width = 18
    line_chart.height = 10
    
    ws_charts.add_chart(line_chart, "A97")
    
    # Save workbook
    wb.save(output_filename)
    print(f"Excel exported: {output_filename}")

def _set_cell_shading(cell, color_hex):
    """Helper to set table cell background color in Word doc"""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def _add_formatted_table_row(table, row_idx, values, bold=False, header=False, bg_color=None, font_color=None):
    """Helper to populate a row in a Word table with formatting"""
    row = table.rows[row_idx]
    for col_idx, value in enumerate(values):
        cell = row.cells[col_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(value))
        run.font.size = Pt(9)
        if bold or header:
            run.bold = True
        if font_color:
            run.font.color.rgb = RGBColor.from_string(font_color)
        if header:
            run.font.color.rgb = RGBColor(255, 255, 255)
        if bg_color:
            _set_cell_shading(cell, bg_color)

def export_to_word(reps_data, w1_reps, w2_reps, w3_reps, perf_analysis, output_filename="sales_rep_analysis_feb2026.docx"):
    """Export comprehensive sales analysis to Word document"""
    doc = DocxDocument()
    
    # --- Styles ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    
    # ===== TITLE PAGE =====
    doc.add_paragraph()
    title = doc.add_heading('SALES REP PERFORMANCE ANALYSIS', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('BOMAS Hardware Store - February 2026')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    period = doc.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = period.add_run('February 1-21, 2026 (18 Working Days)')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    
    gen_date = doc.add_paragraph()
    gen_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = gen_date.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True
    
    doc.add_page_break()
    
    # ===== EXECUTIVE SUMMARY =====
    doc.add_heading('EXECUTIVE SUMMARY', level=1)
    
    total_sales = sum(r['total_sales'] for r in reps_data.values())
    total_profit = sum(r['total_profit'] for r in reps_data.values())
    overall_margin = (total_profit / total_sales * 100) if total_sales > 0 else 0
    
    w1_total = sum(r['total_sales'] for r in w1_reps.values())
    w2_total = sum(r['total_sales'] for r in w2_reps.values())
    w3_total = sum(r['total_sales'] for r in w3_reps.values())
    w1_profit = sum(r['total_profit'] for r in w1_reps.values())
    w2_profit = sum(r['total_profit'] for r in w2_reps.values())
    w3_profit = sum(r['total_profit'] for r in w3_reps.values())
    w1_margin = (w1_profit / w1_total * 100) if w1_total > 0 else 0
    w2_margin = (w2_profit / w2_total * 100) if w2_total > 0 else 0
    w3_margin = (w3_profit / w3_total * 100) if w3_total > 0 else 0
    
    # Executive summary table
    summary_headers = ["Metric", "Week 1 (Feb 1-7)", "Week 2 (Feb 9-14)", "Week 3 (Feb 16-21)", "Combined"]
    summary_rows = [
        ["Total Sales", f"KES {w1_total:,.0f}", f"KES {w2_total:,.0f}", f"KES {w3_total:,.0f}", f"KES {total_sales:,.0f}"],
        ["Total Profit", f"KES {w1_profit:,.0f}", f"KES {w2_profit:,.0f}", f"KES {w3_profit:,.0f}", f"KES {total_profit:,.0f}"],
        ["Daily Average", f"KES {w1_total/6:,.0f}", f"KES {w2_total/6:,.0f}", f"KES {w3_total/6:,.0f}", f"KES {total_sales/18:,.0f}"],
        ["Overall Margin", f"{w1_margin:.1f}%", f"{w2_margin:.1f}%", f"{w3_margin:.1f}%", f"{overall_margin:.1f}%"],
    ]
    
    table = doc.add_table(rows=1 + len(summary_rows), cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_formatted_table_row(table, 0, summary_headers, header=True, bg_color="2C3E50")
    for i, row_data in enumerate(summary_rows, start=1):
        bg = "ECF0F1" if i % 2 == 0 else None
        _add_formatted_table_row(table, i, row_data, bg_color=bg)
    
    doc.add_paragraph()
    
    # Weekly progress commentary
    daily_w1w2 = ((w2_total/6 - w1_total/6) / (w1_total/6)) * 100
    daily_w2w3 = ((w3_total/6 - w2_total/6) / (w2_total/6)) * 100
    overall_daily_change = ((w3_total/6 - w1_total/6) / (w1_total/6)) * 100
    
    doc.add_heading('Weekly Progress', level=2)
    
    p = doc.add_paragraph()
    symbol = "▲" if daily_w1w2 > 0 else "▼"
    color = RGBColor(0x27, 0xAE, 0x60) if daily_w1w2 > 0 else RGBColor(0xE7, 0x4C, 0x3C)
    run = p.add_run(f'{symbol} W1→W2: Daily sales {"increased" if daily_w1w2 > 0 else "decreased"} by {abs(daily_w1w2):.1f}%')
    run.font.color.rgb = color
    run.bold = True
    
    p = doc.add_paragraph()
    symbol = "▲" if daily_w2w3 > 0 else "▼"
    color = RGBColor(0x27, 0xAE, 0x60) if daily_w2w3 > 0 else RGBColor(0xE7, 0x4C, 0x3C)
    run = p.add_run(f'{symbol} W2→W3: Daily sales {"increased" if daily_w2w3 > 0 else "decreased"} by {abs(daily_w2w3):.1f}%')
    run.font.color.rgb = color
    run.bold = True
    
    p = doc.add_paragraph()
    symbol = "📈" if overall_daily_change > 0 else "📉"
    run = p.add_run(f'{symbol} Overall 3-week trend: Daily sales {"UP" if overall_daily_change > 0 else "DOWN"} {abs(overall_daily_change):.1f}% from Week 1 to Week 3')
    run.bold = True
    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60) if overall_daily_change > 0 else RGBColor(0xE7, 0x4C, 0x3C)
    
    # ===== KEY INSIGHTS =====
    doc.add_heading('Key Insights', level=2)
    sorted_reps = sorted(reps_data.items(), key=lambda x: x[1]['total_sales'], reverse=True)
    top_rep = sorted_reps[0]
    top_margin_rep = max(reps_data.items(), key=lambda x: x[1]['overall_margin'])
    
    insights = [
        f"Top Seller: {top_rep[0]} with KES {top_rep[1]['total_sales']/1000000:.2f}M ({top_rep[1]['total_sales']/total_sales*100:.1f}% of total)",
        f"Highest Margin: {top_margin_rep[0]} at {top_margin_rep[1]['overall_margin']:.1f}%",
        f"W1→W2 daily change: {daily_w1w2:+.1f}% | W2→W3 daily change: {daily_w2w3:+.1f}%",
        f"Overall 3-week daily trend: {overall_daily_change:+.1f}%",
        f"Total profit: KES {total_profit/1000000:.2f}M at {overall_margin:.1f}% margin",
    ]
    for insight in insights:
        doc.add_paragraph(f"◆ {insight}", style='List Bullet')
    
    doc.add_page_break()
    
    # ===== CHARTS PAGE =====
    doc.add_heading('VISUAL ANALYSIS', level=1)
    
    # Embed charts if they exist
    chart_files = [
        ("Category Sales Distribution", "charts/summary_cat_donut.png"),
        ("Week-over-Week Comparison", "charts/week_comparison_grouped.png"),
        ("Sales vs Profit Comparison", "charts/summary_sales_chart.png"),
        ("Category Breakdown by Rep", "charts/summary_cat_breakdown.png"),
        ("Rep Contribution Analysis", "charts/rep_contribution.png"),
        ("Performance Radar", "charts/summary_radar.png"),
        ("Profit Margin Heatmap", "charts/summary_heatmap.png"),
        ("Profit Waterfall", "charts/profit_waterfall.png"),
    ]
    
    for chart_title, chart_path in chart_files:
        if os.path.exists(chart_path):
            doc.add_heading(chart_title, level=2)
            doc.add_picture(chart_path, width=Inches(6.0))
            doc.add_paragraph()
    
    doc.add_page_break()
    
    # ===== REP PERFORMANCE RANKINGS =====
    doc.add_heading('SALES REP RANKINGS', level=1)
    
    rank_headers = ["Rank", "Rep Name", "Role", "Total Sales", "Total Profit", "Margin %", "Items Sold"]
    table = doc.add_table(rows=1 + len(sorted_reps), cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_formatted_table_row(table, 0, rank_headers, header=True, bg_color="2C3E50")
    
    for idx, (rep_name, rep_data) in enumerate(sorted_reps, start=1):
        analysis = perf_analysis.get(rep_name, {})
        row_data = [
            f"#{idx}",
            rep_name,
            analysis.get('role', 'Sales Rep'),
            f"KES {rep_data['total_sales']:,.0f}",
            f"KES {rep_data['total_profit']:,.0f}",
            f"{rep_data['overall_margin']:.1f}%",
            f"{rep_data['total_qty']:,.0f}",
        ]
        bg = "ECF0F1" if idx % 2 == 0 else None
        _add_formatted_table_row(table, idx, row_data, bg_color=bg)
    
    doc.add_paragraph()
    
    # ===== WEEK COMPARISON TABLE =====
    doc.add_heading('WEEK-OVER-WEEK COMPARISON BY REP', level=1)
    
    week_headers = ["Rep", "W1 Sales", "W2 Sales", "W3 Sales", "W1→W2 %", "W2→W3 %", "Overall %"]
    table = doc.add_table(rows=1 + len(sorted_reps), cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_formatted_table_row(table, 0, week_headers, header=True, bg_color="3498DB")
    
    for idx, (rep_name, _) in enumerate(sorted_reps, start=1):
        w1_s = w1_reps.get(rep_name, {}).get('total_sales', 0)
        w2_s = w2_reps.get(rep_name, {}).get('total_sales', 0)
        w3_s = w3_reps.get(rep_name, {}).get('total_sales', 0)
        
        chg_w1w2 = ((w2_s - w1_s) / w1_s * 100) if w1_s > 0 else 0
        chg_w2w3 = ((w3_s - w2_s) / w2_s * 100) if w2_s > 0 else 0
        chg_overall = ((w3_s - w1_s) / w1_s * 100) if w1_s > 0 else 0
        
        row_data = [
            rep_name,
            f"KES {w1_s:,.0f}",
            f"KES {w2_s:,.0f}",
            f"KES {w3_s:,.0f}",
            f"{chg_w1w2:+.1f}%",
            f"{chg_w2w3:+.1f}%",
            f"{chg_overall:+.1f}%",
        ]
        bg = "ECF0F1" if idx % 2 == 0 else None
        _add_formatted_table_row(table, idx, row_data, bg_color=bg)
    
    doc.add_page_break()
    
    # ===== CATEGORY ANALYSIS =====
    doc.add_heading('CATEGORY BREAKDOWN', level=1)
    
    cat_headers = ["Category", "Total Sales", "Total Profit", "Margin %"]
    table = doc.add_table(rows=1 + len(CAT_ORDER), cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_formatted_table_row(table, 0, cat_headers, header=True, bg_color="27AE60")
    
    for idx, cat in enumerate(CAT_ORDER, start=1):
        cat_sales = sum(r['categories'].get(cat, {}).get('sales_incl', 0) for r in reps_data.values())
        cat_profit = sum(r['categories'].get(cat, {}).get('profit', 0) for r in reps_data.values())
        cat_margin = (cat_profit / cat_sales * 100) if cat_sales > 0 else 0
        row_data = [cat, f"KES {cat_sales:,.0f}", f"KES {cat_profit:,.0f}", f"{cat_margin:.1f}%"]
        bg = "ECF0F1" if idx % 2 == 0 else None
        _add_formatted_table_row(table, idx, row_data, bg_color=bg)
    
    doc.add_paragraph()
    
    # ===== DEPARTMENT LEADERSHIP =====
    doc.add_heading('DEPARTMENT LEADERSHIP', level=1)
    
    dept_headers = ["Department", "Head", "Dept Sales", "Dept Margin", "Dept Share"]
    dept_rows = []
    for dept in ['PLUMBING', 'PAINTS', 'ELECTRICALS']:
        head = DEPARTMENT_HEADS.get(dept)
        if head and head in reps_data:
            rep_data = reps_data[head]
            if dept in rep_data['categories']:
                dept_sales = rep_data['categories'][dept]['sales_incl']
                dept_margin = rep_data['categories'][dept]['margin_pct']
                dept_total = sum(r['categories'].get(dept, {}).get('sales_incl', 0) for r in reps_data.values())
                dept_share = (dept_sales / dept_total * 100) if dept_total > 0 else 0
                dept_rows.append([dept, head, f"KES {dept_sales:,.0f}", f"{dept_margin:.1f}%", f"{dept_share:.1f}%"])
    
    if dept_rows:
        table = doc.add_table(rows=1 + len(dept_rows), cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _add_formatted_table_row(table, 0, dept_headers, header=True, bg_color="8E44AD")
        for idx, row_data in enumerate(dept_rows, start=1):
            bg = "F5F0FA" if idx % 2 == 0 else None
            _add_formatted_table_row(table, idx, row_data, bg_color=bg)
    
    doc.add_page_break()
    
    # ===== REP-BY-REP COMMENTARY =====
    doc.add_heading('REP-BY-REP PERFORMANCE COMMENTARY', level=1)
    
    for rep_name, rep_data in sorted_reps:
        analysis = perf_analysis.get(rep_name, {})
        role = analysis.get('role', 'Sales Rep')
        badges = analysis.get('badges', [])
        notes = analysis.get('notes', [])
        
        # Rep heading
        doc.add_heading(f'{rep_name.upper()} — {role}', level=2)
        
        # Badges
        if badges:
            p = doc.add_paragraph()
            run = p.add_run(' '.join(badges))
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(0x8E, 0x44, 0xAD)
        
        # Stats mini-table
        stats_headers = ["Total Sales", "Total Profit", "Margin", "Items Sold", "Rank"]
        stats_values = [
            f"KES {rep_data['total_sales']:,.0f}",
            f"KES {rep_data['total_profit']:,.0f}",
            f"{rep_data['overall_margin']:.1f}%",
            f"{rep_data['total_qty']:,.0f}",
            f"#{analysis.get('sales_rank', '-')}",
        ]
        table = doc.add_table(rows=2, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _add_formatted_table_row(table, 0, stats_headers, header=True, bg_color="34495E")
        _add_formatted_table_row(table, 1, stats_values, bold=True)
        
        doc.add_paragraph()
        
        # Category breakdown for this rep
        rep_cats = [(cat, rep_data['categories'][cat]) for cat in CAT_ORDER if cat in rep_data['categories']]
        if rep_cats:
            cat_headers_rep = ["Category", "Sales", "Profit", "Margin", "Qty"]
            table = doc.add_table(rows=1 + len(rep_cats), cols=5)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            _add_formatted_table_row(table, 0, cat_headers_rep, header=True, bg_color="3498DB")
            for i, (cat, cdata) in enumerate(rep_cats, start=1):
                row_data = [cat, f"KES {cdata['sales_incl']:,.0f}", f"KES {cdata['profit']:,.0f}",
                           f"{cdata['margin_pct']:.1f}%", f"{cdata['qty']:,.0f}"]
                bg = "ECF0F1" if i % 2 == 0 else None
                _add_formatted_table_row(table, i, row_data, bg_color=bg)
        
        doc.add_paragraph()
        
        # Performance notes
        p = doc.add_paragraph()
        run = p.add_run('PERFORMANCE NOTES')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x8E, 0x44, 0xAD)
        
        for note in notes[:6]:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(note)
            run.font.size = Pt(9)
            if '✓' in note or 'Excellent' in note or 'Outstanding' in note or 'upward' in note or 'Recovery' in note:
                run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
            elif '⚠️' in note or 'below' in note or 'Declined' in note or 'Declining' in note or 'reversal' in note:
                run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
            else:
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        
        # Per-rep charts if they exist
        rep_safe = rep_name.replace(' ', '_')
        chart_files_rep = [
            f"charts/{rep_safe}_donut.png",
            f"charts/{rep_safe}_sales_profit.png",
        ]
        for cf in chart_files_rep:
            if os.path.exists(cf):
                doc.add_picture(cf, width=Inches(4.5))
        
        doc.add_page_break()
    
    # ===== DAILY AVERAGES SUMMARY =====
    doc.add_heading('DAILY SALES AVERAGES', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('Average daily sales per rep across 18 working days (6 days per week × 3 weeks)')
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    daily_headers = ["Rep", "W1 Daily Avg", "W2 Daily Avg", "W3 Daily Avg", "Overall Daily Avg"]
    table = doc.add_table(rows=1 + len(sorted_reps) + 1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_formatted_table_row(table, 0, daily_headers, header=True, bg_color="2C3E50")
    
    for idx, (rep_name, rep_data) in enumerate(sorted_reps, start=1):
        w1_s = w1_reps.get(rep_name, {}).get('total_sales', 0)
        w2_s = w2_reps.get(rep_name, {}).get('total_sales', 0)
        w3_s = w3_reps.get(rep_name, {}).get('total_sales', 0)
        row_data = [
            rep_name,
            f"KES {w1_s/6:,.0f}",
            f"KES {w2_s/6:,.0f}",
            f"KES {w3_s/6:,.0f}",
            f"KES {rep_data['total_sales']/18:,.0f}",
        ]
        bg = "ECF0F1" if idx % 2 == 0 else None
        _add_formatted_table_row(table, idx, row_data, bg_color=bg)
    
    # Team total row
    team_row = [
        "TEAM TOTAL",
        f"KES {w1_total/6:,.0f}",
        f"KES {w2_total/6:,.0f}",
        f"KES {w3_total/6:,.0f}",
        f"KES {total_sales/18:,.0f}",
    ]
    _add_formatted_table_row(table, len(sorted_reps) + 1, team_row, bold=True, bg_color="D5F5E3")
    
    doc.add_paragraph()
    
    # ===== CONCLUSION =====
    doc.add_heading('CONCLUSION & RECOMMENDATIONS', level=1)
    
    # Auto-generate conclusions based on data
    conclusions = []
    
    # Overall trend
    if overall_daily_change > 10:
        conclusions.append(f"The team has shown strong growth over the 3-week period, with daily sales increasing by {overall_daily_change:.1f}% from Week 1 to Week 3.")
    elif overall_daily_change > 0:
        conclusions.append(f"The team has shown modest growth over the 3-week period, with daily sales increasing by {overall_daily_change:.1f}% from Week 1 to Week 3.")
    else:
        conclusions.append(f"The team's daily sales have declined by {abs(overall_daily_change):.1f}% from Week 1 to Week 3. Urgent review and corrective action is needed.")
    
    # Best performer
    conclusions.append(f"{top_rep[0]} leads the team with KES {top_rep[1]['total_sales']/1000000:.2f}M in total sales ({top_rep[1]['total_sales']/total_sales*100:.1f}% of team total).")
    
    # Margin analysis
    if overall_margin > 10:
        conclusions.append(f"The overall profit margin of {overall_margin:.1f}% is healthy. {top_margin_rep[0]} leads margin performance at {top_margin_rep[1]['overall_margin']:.1f}%.")
    else:
        conclusions.append(f"The overall profit margin of {overall_margin:.1f}% is below the 10% target. Focus on higher-margin products.")
    
    # Reps who declined in W3
    declining_reps = []
    for rep_name in reps_data:
        w2_s = w2_reps.get(rep_name, {}).get('total_sales', 0)
        w3_s = w3_reps.get(rep_name, {}).get('total_sales', 0)
        if w2_s > 0 and w3_s < w2_s * 0.85:
            declining_reps.append(rep_name)
    if declining_reps:
        conclusions.append(f"The following reps showed significant decline (>15%) in Week 3: {', '.join(declining_reps)}. Individual reviews recommended.")
    
    # Improving reps
    improving_reps = []
    for rep_name in reps_data:
        w1_s = w1_reps.get(rep_name, {}).get('total_sales', 0)
        w3_s = w3_reps.get(rep_name, {}).get('total_sales', 0)
        if w1_s > 0 and w3_s > w1_s * 1.3:
            improving_reps.append(rep_name)
    if improving_reps:
        conclusions.append(f"Strong overall improvement from {', '.join(improving_reps)} — up >30% from Week 1 to Week 3.")
    
    for conclusion in conclusions:
        doc.add_paragraph(conclusion, style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('--- End of Report ---')
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save
    doc.save(output_filename)
    print(f"Word document exported: {output_filename}")

def generate_individual_rep_report(rep_name, rep_data, w1_reps, w2_reps, w3_reps, reps_data, perf_analysis, output_dir="rep_reports"):
    """Generate a standalone Word document report for a single sales rep"""
    os.makedirs(output_dir, exist_ok=True)
    
    analysis = perf_analysis.get(rep_name, {})
    role = analysis.get('role', 'Sales Rep')
    badges = analysis.get('badges', [])
    notes = analysis.get('notes', [])
    
    total_team_sales = sum(r['total_sales'] for r in reps_data.values())
    total_team_profit = sum(r['total_profit'] for r in reps_data.values())
    team_margin = (total_team_profit / total_team_sales * 100) if total_team_sales > 0 else 0
    
    w1_s = w1_reps.get(rep_name, {}).get('total_sales', 0)
    w2_s = w2_reps.get(rep_name, {}).get('total_sales', 0)
    w3_s = w3_reps.get(rep_name, {}).get('total_sales', 0)
    w1_p = w1_reps.get(rep_name, {}).get('total_profit', 0)
    w2_p = w2_reps.get(rep_name, {}).get('total_profit', 0)
    w3_p = w3_reps.get(rep_name, {}).get('total_profit', 0)
    
    chg_w1w2 = ((w2_s - w1_s) / w1_s * 100) if w1_s > 0 else 0
    chg_w2w3 = ((w3_s - w2_s) / w2_s * 100) if w2_s > 0 else 0
    chg_overall = ((w3_s - w1_s) / w1_s * 100) if w1_s > 0 else 0
    
    doc = DocxDocument()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    
    # ===== TITLE PAGE =====
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_heading(f'INDIVIDUAL PERFORMANCE REPORT', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(rep_name.upper())
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(role)
    run.font.size = Pt(14)
    run.italic = True
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    
    if badges:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(' '.join(badges))
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = RGBColor(0x8E, 0x44, 0xAD)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BOMAS Hardware Store')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('February 1-21, 2026 (18 Working Days)')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    doc.add_page_break()
    
    # ===== PAGE 2: PERFORMANCE SNAPSHOT =====
    doc.add_heading('PERFORMANCE SNAPSHOT', level=1)
    
    # Overall stats table
    stats_headers = ["Metric", "Value", "Team Comparison"]
    team_avg_sales = total_team_sales / len(reps_data)
    sales_vs_avg = ((rep_data['total_sales'] - team_avg_sales) / team_avg_sales) * 100
    
    stats_rows = [
        ["Total Sales", f"KES {rep_data['total_sales']:,.0f}", f"{'▲' if sales_vs_avg > 0 else '▼'} {abs(sales_vs_avg):.1f}% vs team avg"],
        ["Total Profit", f"KES {rep_data['total_profit']:,.0f}", f"{rep_data['total_profit']/total_team_profit*100:.1f}% of team total"],
        ["Profit Margin", f"{rep_data['overall_margin']:.1f}%", f"Team avg: {team_margin:.1f}%"],
        ["Items Sold", f"{rep_data['total_qty']:,.0f}", ""],
        ["Daily Avg Sales", f"KES {rep_data['total_sales']/18:,.0f}", ""],
        ["Daily Avg Profit", f"KES {rep_data['total_profit']/18:,.0f}", ""],
        ["Team Share", f"{rep_data['total_sales']/total_team_sales*100:.1f}%", f"#{analysis.get('sales_rank', '-')} of {len(reps_data)} reps"],
        ["Sales Rank", f"#{analysis.get('sales_rank', '-')}", ""],
        ["Margin Rank", f"#{analysis.get('margin_rank', '-')}", ""],
    ]
    
    table = doc.add_table(rows=1 + len(stats_rows), cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_formatted_table_row(table, 0, stats_headers, header=True, bg_color="2C3E50")
    for i, row_data in enumerate(stats_rows, start=1):
        bg = "ECF0F1" if i % 2 == 0 else None
        _add_formatted_table_row(table, i, row_data, bg_color=bg)
    
    doc.add_paragraph()
    
    # ===== WEEKLY BREAKDOWN =====
    doc.add_heading('WEEKLY SALES BREAKDOWN', level=1)
    
    week_headers = ["Week", "Sales", "Profit", "Margin", "Daily Avg", "Change"]
    w1_margin = (w1_p / w1_s * 100) if w1_s > 0 else 0
    w2_margin = (w2_p / w2_s * 100) if w2_s > 0 else 0
    w3_margin = (w3_p / w3_s * 100) if w3_s > 0 else 0
    
    week_rows = [
        ["Week 1 (Feb 1-7)", f"KES {w1_s:,.0f}", f"KES {w1_p:,.0f}", f"{w1_margin:.1f}%", f"KES {w1_s/6:,.0f}", "—"],
        ["Week 2 (Feb 9-14)", f"KES {w2_s:,.0f}", f"KES {w2_p:,.0f}", f"{w2_margin:.1f}%", f"KES {w2_s/6:,.0f}", f"{chg_w1w2:+.1f}%"],
        ["Week 3 (Feb 16-21)", f"KES {w3_s:,.0f}", f"KES {w3_p:,.0f}", f"{w3_margin:.1f}%", f"KES {w3_s/6:,.0f}", f"{chg_w2w3:+.1f}%"],
        ["TOTAL", f"KES {rep_data['total_sales']:,.0f}", f"KES {rep_data['total_profit']:,.0f}", f"{rep_data['overall_margin']:.1f}%", f"KES {rep_data['total_sales']/18:,.0f}", f"{chg_overall:+.1f}%"],
    ]
    
    table = doc.add_table(rows=1 + len(week_rows), cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_formatted_table_row(table, 0, week_headers, header=True, bg_color="3498DB")
    for i, row_data in enumerate(week_rows, start=1):
        bg = "D5F5E3" if i == len(week_rows) else ("ECF0F1" if i % 2 == 0 else None)
        bold = i == len(week_rows)
        _add_formatted_table_row(table, i, row_data, bold=bold, bg_color=bg)
    
    doc.add_paragraph()
    
    # Weekly trend commentary
    doc.add_heading('Weekly Progress', level=2)
    
    p = doc.add_paragraph()
    symbol = "▲" if chg_w1w2 > 0 else "▼"
    color = RGBColor(0x27, 0xAE, 0x60) if chg_w1w2 > 0 else RGBColor(0xE7, 0x4C, 0x3C)
    run = p.add_run(f'{symbol} W1→W2: Sales {"increased" if chg_w1w2 > 0 else "decreased"} by {abs(chg_w1w2):.1f}% (KES {w1_s/1000:,.0f}K → {w2_s/1000:,.0f}K)')
    run.font.color.rgb = color
    run.bold = True
    
    p = doc.add_paragraph()
    symbol = "▲" if chg_w2w3 > 0 else "▼"
    color = RGBColor(0x27, 0xAE, 0x60) if chg_w2w3 > 0 else RGBColor(0xE7, 0x4C, 0x3C)
    run = p.add_run(f'{symbol} W2→W3: Sales {"increased" if chg_w2w3 > 0 else "decreased"} by {abs(chg_w2w3):.1f}% (KES {w2_s/1000:,.0f}K → {w3_s/1000:,.0f}K)')
    run.font.color.rgb = color
    run.bold = True
    
    p = doc.add_paragraph()
    if chg_w1w2 > 0 and chg_w2w3 > 0:
        run = p.add_run(f'📈 Consistent upward trend over 3 weeks (+{chg_overall:.1f}% overall)')
        run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    elif chg_w1w2 < 0 and chg_w2w3 < 0:
        run = p.add_run(f'📉 Declining trend over 3 weeks ({chg_overall:.1f}% overall) — needs urgent attention')
        run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    elif chg_w1w2 > 0 and chg_w2w3 < 0:
        run = p.add_run(f'⚠️ Week 3 reversal after Week 2 growth — monitor closely')
        run.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
    elif chg_w1w2 < 0 and chg_w2w3 > 0:
        run = p.add_run(f'✓ Recovery in Week 3 after Week 2 dip — positive momentum')
        run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    else:
        run = p.add_run(f'Stable performance across 3 weeks ({chg_overall:+.1f}% overall)')
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.bold = True
    
    doc.add_page_break()
    
    # ===== PAGE 3: CATEGORY ANALYSIS =====
    doc.add_heading('CATEGORY PERFORMANCE', level=1)
    
    rep_cats = [(cat, rep_data['categories'][cat]) for cat in CAT_ORDER if cat in rep_data['categories']]
    
    if rep_cats:
        cat_headers = ["Category", "Sales", "Profit", "Margin %", "Qty", "% of Total"]
        table = doc.add_table(rows=1 + len(rep_cats) + 1, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _add_formatted_table_row(table, 0, cat_headers, header=True, bg_color="27AE60")
        
        for i, (cat, cdata) in enumerate(rep_cats, start=1):
            pct_total = (cdata['sales_incl'] / rep_data['total_sales'] * 100) if rep_data['total_sales'] > 0 else 0
            row_data = [cat, f"KES {cdata['sales_incl']:,.0f}", f"KES {cdata['profit']:,.0f}",
                       f"{cdata['margin_pct']:.1f}%", f"{cdata['qty']:,.0f}", f"{pct_total:.1f}%"]
            bg = "ECF0F1" if i % 2 == 0 else None
            _add_formatted_table_row(table, i, row_data, bg_color=bg)
        
        # Total row
        total_row = ["TOTAL", f"KES {rep_data['total_sales']:,.0f}", f"KES {rep_data['total_profit']:,.0f}",
                     f"{rep_data['overall_margin']:.1f}%", f"{rep_data['total_qty']:,.0f}", "100.0%"]
        _add_formatted_table_row(table, len(rep_cats) + 1, total_row, bold=True, bg_color="D5F5E3")
    
    doc.add_paragraph()
    
    # Category week-over-week breakdown per category
    doc.add_heading('Category Weekly Trends', level=2)
    
    for cat in CAT_ORDER:
        if cat not in rep_data['categories']:
            continue
        
        w1_cat_s = w1_reps.get(rep_name, {}).get('categories', {}).get(cat, {}).get('sales_incl', 0)
        w2_cat_s = w2_reps.get(rep_name, {}).get('categories', {}).get(cat, {}).get('sales_incl', 0)
        w3_cat_s = w3_reps.get(rep_name, {}).get('categories', {}).get(cat, {}).get('sales_incl', 0)
        
        p = doc.add_paragraph()
        run = p.add_run(f'{cat}: ')
        run.bold = True
        run.font.size = Pt(10)
        
        parts = [f"W1: KES {w1_cat_s:,.0f}", f"W2: KES {w2_cat_s:,.0f}", f"W3: KES {w3_cat_s:,.0f}"]
        if w1_cat_s > 0:
            cat_chg = ((w3_cat_s - w1_cat_s) / w1_cat_s * 100)
            symbol = "▲" if cat_chg > 0 else "▼"
            parts.append(f"Overall: {symbol}{abs(cat_chg):.0f}%")
        
        run = p.add_run(' → '.join(parts))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    doc.add_page_break()
    
    # ===== PAGE 4: CHARTS =====
    doc.add_heading('VISUAL ANALYSIS', level=1)
    
    rep_safe = rep_name.replace(' ', '_')
    chart_map = [
        ("Sales Distribution by Category", f"charts/{rep_safe}_donut.png"),
        ("Sales vs Profit by Category", f"charts/{rep_safe}_sales_profit.png"),
        ("Profit Margin Gauge", f"charts/{rep_safe}_gauge.png"),
    ]
    
    for chart_title, chart_path in chart_map:
        if os.path.exists(chart_path):
            doc.add_heading(chart_title, level=2)
            doc.add_picture(chart_path, width=Inches(5.5))
            doc.add_paragraph()
    
    doc.add_page_break()
    
    # ===== PAGE 5: TEAM CONTEXT =====
    doc.add_heading('TEAM CONTEXT', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run(f'How {rep_name} compares to the rest of the team:')
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()
    
    sorted_team = sorted(reps_data.items(), key=lambda x: x[1]['total_sales'], reverse=True)
    
    team_headers = ["Rank", "Rep", "Total Sales", "Margin %", "Team Share"]
    table = doc.add_table(rows=1 + len(sorted_team), cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_formatted_table_row(table, 0, team_headers, header=True, bg_color="2C3E50")
    
    for idx, (rname, rdata) in enumerate(sorted_team, start=1):
        share = (rdata['total_sales'] / total_team_sales * 100) if total_team_sales > 0 else 0
        row_data = [f"#{idx}", rname, f"KES {rdata['total_sales']:,.0f}", f"{rdata['overall_margin']:.1f}%", f"{share:.1f}%"]
        # Highlight the current rep's row
        if rname == rep_name:
            _add_formatted_table_row(table, idx, row_data, bold=True, bg_color="D4E6F1")
        else:
            bg = "ECF0F1" if idx % 2 == 0 else None
            _add_formatted_table_row(table, idx, row_data, bg_color=bg)
    
    doc.add_paragraph()
    
    # Week comparison vs team chart
    if os.path.exists("charts/week_comparison_grouped.png"):
        doc.add_heading('Team Week-over-Week Comparison', level=2)
        doc.add_picture("charts/week_comparison_grouped.png", width=Inches(6.0))
    
    doc.add_page_break()
    
    # ===== PAGE 6: DEPARTMENT ANALYSIS (for dept heads) =====
    primary_dept = get_primary_department(rep_name)
    if primary_dept != 'GENERAL HARDWARE':
        doc.add_heading(f'DEPARTMENT LEADERSHIP: {primary_dept}', level=1)
        
        # Show all reps' contributions to this department
        dept_reps = []
        dept_total_sales = 0
        for rname, rdata in reps_data.items():
            if primary_dept in rdata['categories']:
                dept_sales = rdata['categories'][primary_dept]['sales_incl']
                dept_profit = rdata['categories'][primary_dept]['profit']
                dept_margin = rdata['categories'][primary_dept]['margin_pct']
                dept_reps.append((rname, dept_sales, dept_profit, dept_margin))
                dept_total_sales += dept_sales
        
        dept_reps.sort(key=lambda x: x[1], reverse=True)
        
        dept_headers = ["Rep", "Dept Sales", "Dept Profit", "Margin", "Dept Share"]
        table = doc.add_table(rows=1 + len(dept_reps), cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _add_formatted_table_row(table, 0, dept_headers, header=True, bg_color="8E44AD")
        
        for i, (rname, ds, dp, dm) in enumerate(dept_reps, start=1):
            share = (ds / dept_total_sales * 100) if dept_total_sales > 0 else 0
            row_data = [rname, f"KES {ds:,.0f}", f"KES {dp:,.0f}", f"{dm:.1f}%", f"{share:.1f}%"]
            if rname == rep_name:
                _add_formatted_table_row(table, i, row_data, bold=True, bg_color="F5EEF8")
            else:
                bg = "ECF0F1" if i % 2 == 0 else None
                _add_formatted_table_row(table, i, row_data, bg_color=bg)
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        own_share = 0
        for rname, ds, dp, dm in dept_reps:
            if rname == rep_name:
                own_share = (ds / dept_total_sales * 100) if dept_total_sales > 0 else 0
        
        if own_share >= 40:
            run = p.add_run(f'✓ Strong department leadership — {rep_name} controls {own_share:.1f}% of {primary_dept} sales')
            run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
        else:
            run = p.add_run(f'⚠️ As Department Head, {rep_name} contributes {own_share:.1f}% of {primary_dept} sales — consider increasing dept focus')
            run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
        run.bold = True
        
        doc.add_page_break()
    
    # ===== PERFORMANCE NOTES & RECOMMENDATIONS =====
    doc.add_heading('PERFORMANCE NOTES', level=1)
    
    for note in notes:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(note)
        run.font.size = Pt(10)
        if '✓' in note or 'Excellent' in note or 'Outstanding' in note or 'upward' in note or 'Recovery' in note or 'Strong' in note:
            run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
        elif '⚠️' in note or 'below' in note or 'Declined' in note or 'Declining' in note or 'reversal' in note or 'urgent' in note:
            run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
        else:
            run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
    
    doc.add_paragraph()
    
    # Recommendations
    doc.add_heading('RECOMMENDATIONS', level=1)
    
    recommendations = []
    
    # Margin-based recommendation
    if rep_data['overall_margin'] < team_margin:
        recommendations.append(f"Focus on higher-margin products to improve margin from {rep_data['overall_margin']:.1f}% toward the team average of {team_margin:.1f}%.")
    else:
        recommendations.append(f"Maintain excellent margin discipline at {rep_data['overall_margin']:.1f}% (above team avg of {team_margin:.1f}%).")
    
    # Growth-based recommendation
    if chg_w2w3 < -10:
        recommendations.append(f"Week 3 showed a {abs(chg_w2w3):.1f}% decline. Review customer activity and pipeline for Week 4.")
    elif chg_w2w3 > 20:
        recommendations.append(f"Outstanding Week 3 growth of +{chg_w2w3:.1f}%. Maintain this momentum into Week 4.")
    
    # Cross-selling recommendation
    cats_sold = len(rep_data['categories'])
    if cats_sold < 4:
        missing = [c for c in CAT_ORDER if c not in rep_data['categories']]
        recommendations.append(f"Expand into {', '.join(missing)} to increase cross-selling coverage (currently {cats_sold}/4 categories).")
    else:
        recommendations.append("Selling across all 4 categories — excellent product range coverage.")
    
    # Volume recommendation
    if rep_data['total_sales'] < total_team_sales / len(reps_data) * 0.5:
        recommendations.append("Sales are significantly below team average. Consider additional training or customer assignment review.")
    
    # Department head recommendation
    if primary_dept != 'GENERAL HARDWARE':
        dept_cat_data = rep_data['categories'].get(primary_dept, {})
        if dept_cat_data:
            dept_pct = (dept_cat_data.get('sales_incl', 0) / rep_data['total_sales'] * 100) if rep_data['total_sales'] > 0 else 0
            if dept_pct < 30:
                recommendations.append(f"As {primary_dept} Department Head, only {dept_pct:.1f}% of your sales are in your department. Consider refocusing.")
    
    for rec in recommendations:
        doc.add_paragraph(rec, style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('--- End of Individual Report ---')
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save
    safe_name = rep_name.replace(' ', '_')
    filepath = os.path.join(output_dir, f"{safe_name}_performance_report.docx")
    doc.save(filepath)
    return filepath

def generate_all_individual_reports(reps_data, w1_reps, w2_reps, w3_reps, perf_analysis, output_dir="rep_reports"):
    """Generate individual Word reports for all reps"""
    os.makedirs(output_dir, exist_ok=True)
    sorted_reps = sorted(reps_data.items(), key=lambda x: x[1]['total_sales'], reverse=True)
    
    generated = []
    for rep_name, rep_data in sorted_reps:
        filepath = generate_individual_rep_report(
            rep_name, rep_data, w1_reps, w2_reps, w3_reps, reps_data, perf_analysis, output_dir
        )
        generated.append(filepath)
        print(f"  Generated: {filepath}")
    
    return generated

def main():
    print("=" * 60)
    print("SALES PERFORMANCE ANALYSIS - FEBRUARY 2026")
    print("=" * 60)
    
    print("\nMerging Magdalene -> Betha Odumo")
    print("Merging WALK IN-BOMAS -> Bonface Kitheka")
    
    print("\nProcessing Week 1 data (Feb 1-7)...")
    w1_combined = combine_data(week1_merged)
    w1_reps = get_rep_data(w1_combined)
    
    print("Processing Week 2 data (Feb 9-14)...")
    w2_combined = combine_data(week2_data)
    w2_reps = get_rep_data(w2_combined)
    
    print("Processing Week 3 data (Feb 16-21)...")
    w3_combined = combine_data(week3_data)
    w3_reps = get_rep_data(w3_combined)
    
    print("Combining all three weeks...")
    combined = combine_data(week1_merged, week2_data, week3_data)
    reps_data = get_rep_data(combined)
    
    total_sales = sum(r['total_sales'] for r in reps_data.values())
    total_profit = sum(r['total_profit'] for r in reps_data.values())
    overall_margin = (total_profit / total_sales * 100) if total_sales > 0 else 0
    
    print("\n" + "=" * 60)
    print("COMBINED RESULTS (Feb 1-21, 18 Working Days)")
    print("=" * 60)
    print(f"\nTotal Sales: KES {total_sales:,.0f}")
    print(f"Total Profit: KES {total_profit:,.0f}")
    print(f"Overall Margin: {overall_margin:.1f}%")
    print(f"Daily Average: KES {total_sales/18:,.0f}")
    
    # Week summaries
    w1_total = sum(r['total_sales'] for r in w1_reps.values())
    w2_total = sum(r['total_sales'] for r in w2_reps.values())
    w3_total = sum(r['total_sales'] for r in w3_reps.values())
    print(f"\nWeek 1 (Feb 1-7):   KES {w1_total:>12,.0f}  Daily: KES {w1_total/6:>10,.0f}")
    print(f"Week 2 (Feb 9-14):  KES {w2_total:>12,.0f}  Daily: KES {w2_total/6:>10,.0f}")
    print(f"Week 3 (Feb 16-21): KES {w3_total:>12,.0f}  Daily: KES {w3_total/6:>10,.0f}")
    
    print(f"\n{'Rep':<20} {'Sales':>12} {'Profit':>12} {'Margin':>8}")
    print("-" * 55)
    sorted_reps = sorted(reps_data.items(), key=lambda x: x[1]['total_sales'], reverse=True)
    for rep_name, rep_data in sorted_reps:
        print(f"{rep_name:<20} KES {rep_data['total_sales']:>10,.0f} KES {rep_data['total_profit']:>9,.0f} {rep_data['overall_margin']:>7.1f}%")
    
    print("\nGenerating PDF report with enhanced charts...")
    generate_pdf_report(reps_data, w1_reps, w2_reps, w3_reps, "sales_rep_analysis_feb2026.pdf")
    
    # Generate performance analysis for JSON
    perf_analysis = analyze_performance(reps_data, w1_reps, w2_reps, w3_reps)
    
    # Save JSON
    output_json = {
        "period": "February 1-21, 2026",
        "working_days": 18,
        "generated": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "department_heads": DEPARTMENT_HEADS,
        "week1": {
            "period": "February 1-7, 2026",
            "days": 6,
            "total_sales": sum(r['total_sales'] for r in w1_reps.values()),
            "total_profit": sum(r['total_profit'] for r in w1_reps.values())
        },
        "week2": {
            "period": "February 9-14, 2026",
            "days": 6,
            "total_sales": sum(r['total_sales'] for r in w2_reps.values()),
            "total_profit": sum(r['total_profit'] for r in w2_reps.values())
        },
        "week3": {
            "period": "February 16-21, 2026",
            "days": 6,
            "total_sales": sum(r['total_sales'] for r in w3_reps.values()),
            "total_profit": sum(r['total_profit'] for r in w3_reps.values())
        },
        "reps": {}
    }
    
    for rep_name, rep_data in sorted_reps:
        analysis = perf_analysis.get(rep_name, {})
        output_json["reps"][rep_name] = {
            "total_sales": rep_data["total_sales"],
            "total_profit": round(rep_data["total_profit"], 2),
            "total_qty": rep_data["total_qty"],
            "overall_margin": round(rep_data["overall_margin"], 2),
            "role": analysis.get('role', 'Sales Rep'),
            "primary_department": analysis.get('primary_department', 'GENERAL HARDWARE'),
            "badges": analysis.get('badges', []),
            "performance_notes": analysis.get('notes', []),
            "sales_rank": analysis.get('sales_rank'),
            "margin_rank": analysis.get('margin_rank'),
            "week_improvement_pct": round(analysis.get('week_improvement', 0), 1),
            "categories": {cat: {"sales": v["sales_incl"], "profit": round(v["profit"], 2), 
                                 "margin_pct": round(v["margin_pct"], 2), "qty": v["qty"]} 
                           for cat, v in rep_data["categories"].items()}
        }
    
    with open("sales_rep_analysis_feb2026.json", "w") as f:
        json.dump(output_json, f, indent=2)
    print("JSON saved: sales_rep_analysis_feb2026.json")
    
    # Export to Excel
    export_to_excel(reps_data, w1_reps, w2_reps, w3_reps, perf_analysis, "sales_rep_analysis_feb2026.xlsx")
    
    # Export to Word
    export_to_word(reps_data, w1_reps, w2_reps, w3_reps, perf_analysis, "sales_rep_analysis_feb2026.docx")
    
    # Generate Individual Rep Reports
    print("\n--- Generating Individual Rep Reports ---")
    generated_files = generate_all_individual_reports(reps_data, w1_reps, w2_reps, w3_reps, perf_analysis)
    print(f"\n{len(generated_files)} individual reports generated in 'rep_reports/' folder")
    
    print("\nDone!")

if __name__ == "__main__":
    main()
