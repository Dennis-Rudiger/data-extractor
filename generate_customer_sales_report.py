import json
import pandas as pd
from openpyxl.styles import Font, PatternFill
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_reports():
    with open('customer_sales_extracted.json', 'r') as f:
        data = json.load(f)
        
    df = pd.DataFrame(data)
    
    # 1. Create Excel Report
    writer = pd.ExcelWriter('Customer_Sales_Report.xlsx', engine='openpyxl')
    
    # Filter YTD out for the monthly trends
    monthly_df = df[df['Month'] != 'YTD'].copy()
    ytd_df = df[df['Month'] == 'YTD'].copy()
    
    # pivot table by customer and month (Amount Exc)
    if not monthly_df.empty:
        pivot = pd.pivot_table(monthly_df, values='Amount_Exc', index='Customer', columns='Month', aggfunc='sum', fill_value=0)
        # Order columns correctly since they're alphabetic otherwise
        month_order = [m for m in ['January', 'February', 'March', 'April'] if m in pivot.columns]
        pivot = pivot[month_order]
        pivot['Total_Monthly_Sum'] = pivot.sum(axis=1)
        pivot = pivot.sort_values('Total_Monthly_Sum', ascending=False)
        pivot.to_excel(writer, sheet_name='MoM Sales')
    
    # YTD summary
    if not ytd_df.empty:
        ytd_summary = ytd_df.groupby('Customer')[['Amount_Exc', 'Cost', 'Profit']].sum().reset_index()
        ytd_summary['Profit_Margin_%'] = (ytd_summary['Profit'] / ytd_summary['Amount_Exc']) * 100
        ytd_summary = ytd_summary.sort_values('Amount_Exc', ascending=False)
        ytd_summary.to_excel(writer, sheet_name='YTD Totals', index=False)
        
    writer.close()
    print('Created Customer_Sales_Report.xlsx')
    
    # 2. Create Word Summary Report
    doc = Document()
    head = doc.add_heading('CUSTOMER SALES PERFORMANCE REPORT', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run('Period: January - April 2026\n').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('1. Executive Overview', level=1)
    total_ytd_sales = ytd_df['Amount_Exc'].sum() if not ytd_df.empty else monthly_df['Amount_Exc'].sum()
    total_ytd_profit = ytd_df['Profit'].sum() if not ytd_df.empty else monthly_df['Profit'].sum()
    margin = (total_ytd_profit / total_ytd_sales * 100) if total_ytd_sales > 0 else 0
    
    doc.add_paragraph(f"This report consolidates customer sales performance from January to April 2026. The total sales generated across all tracked customers was KES {total_ytd_sales:,.2f}, delivering a gross profit of KES {total_ytd_profit:,.2f} (Overall Margin: {margin:.2f}%).")
    
    if not ytd_df.empty:
        doc.add_heading('2. Top 10 Customers by YTD Volume', level=1)
        top_10 = ytd_summary.head(10)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Customer Name'
        hdr[1].text = 'Sales (KES)'
        hdr[2].text = 'Profit (KES)'
        hdr[3].text = 'Margin (%)'
        
        for _, row in top_10.iterrows():
            cells = table.add_row().cells
            cells[0].text = row['Customer']
            cells[1].text = f"{row['Amount_Exc']:,.2f}"
            cells[2].text = f"{row['Profit']:,.2f}"
            cells[3].text = f"{row['Profit_Margin_%']:.1f}%"
            
        doc.add_paragraph("\nNote: Bomas - Cash Customer remains our largest generic channel. Focus on retaining high margin corporate or individual accounts like Baron Capital and Bruce.")
    
    doc.add_heading('3. Monthly Performance Trend', level=1)
    if not monthly_df.empty:
        monthly_trend = monthly_df.groupby('Month')['Amount_Exc'].sum().reindex(['January', 'February', 'March', 'April']).fillna(0)
        for mth, val in monthly_trend.items():
            if val > 0:
                doc.add_paragraph(f"{mth}: KES {val:,.2f}", style='List Bullet')
                
    doc.save('Customer_Sales_Summary.docx')
    print('Created Customer_Sales_Summary.docx')

if __name__ == '__main__':
    build_reports()
