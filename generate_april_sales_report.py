import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY = RGBColor(26, 54, 93)
DARK_BLUE = RGBColor(43, 108, 176)
WHITE = RGBColor(255, 255, 255)
TEAL = RGBColor(44, 122, 123)

def generate_report():
    with open('sales_april.json', 'r') as f:
        data = json.load(f)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('APRIL 2026 SALES ANALYSIS REPORT')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Period: {data.get('period', 'April 2026')} | Working Days: {data.get('working_days', 24)}")
    run.font.size = Pt(10)
    run.font.italic = True

    grand_sales = 0
    grand_profit = 0
    grand_qty = 0

    categories = data.get('categories', {})

    for dept, reps in categories.items():
        doc.add_heading(f"{dept} DEPARTMENT - APRIL 2026", level=1)

        dept_sales = sum(r['sales_incl'] for r in reps.values())
        dept_profit = sum(r['profit'] for r in reps.values())
        dept_qty = sum(r['qty'] for r in reps.values())
        margin = (dept_profit / dept_sales * 100) if dept_sales > 0 else 0

        grand_sales += dept_sales
        grand_profit += dept_profit
        grand_qty += dept_qty

        # Department summary
        p = doc.add_paragraph()
        p.add_run(f"Total Sales (Incl. VAT): KES {dept_sales:,.2f}\n").bold = True
        p.add_run(f"Total Profit: KES {dept_profit:,.2f} ({margin:.1f}% Margin)\n")
        p.add_run(f"Total Quantity: {dept_qty:,.0f} units\n")
        p.add_run(f"Active Reps: {len(reps)}")

        if not reps:
            doc.add_paragraph("No data extracted.")
            continue

        doc.add_heading(f"Top Performance by Transaction Rep", level=2)
        
        table = doc.add_table(rows=1, cols=5)
        try:
            table.style = 'Light Shading Accent 1'
        except:
            table.style = 'Table Grid'
            
        hdr = table.rows[0].cells
        headers = ['Transaction Rep', 'Sales Incl VAT (KES)', 'Profit (KES)', 'Margin (%)', 'Quantity']
        for i, h in enumerate(headers):
            hdr[i].text = h
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        sorted_reps = sorted(reps.items(), key=lambda x: x[1]['sales_incl'], reverse=True)
        for rep_name, vals in sorted_reps:
            row = table.add_row().cells
            row[0].text = rep_name
            row[1].text = f"{vals['sales_incl']:,.2f}"
            row[2].text = f"{vals['profit']:,.2f}"
            rep_margin = (vals['profit'] / vals['sales_incl'] * 100) if vals['sales_incl'] > 0 else 0
            row[3].text = f"{rep_margin:.1f}%"
            row[4].text = f"{vals['qty']:,.0f}"

        doc.add_page_break()

    # Grand Summary
    doc.add_heading("EXECUTIVE SUMMARY - ALL DEPARTMENTS", level=1)
    overall_margin = (grand_profit / grand_sales * 100) if grand_sales > 0 else 0
    p = doc.add_paragraph()
    p.add_run(f"Grand Total Sales: KES {grand_sales:,.2f}\n").bold = True
    p.add_run(f"Grand Total Profit: KES {grand_profit:,.2f}\n").bold = True
    p.add_run(f"Overall Profit Margin: {overall_margin:.1f}%\n").bold = True
    p.add_run(f"Total Quantity Sold: {grand_qty:,.0f} units\n").bold = True

    doc.save('April_2026_Sales_Report.docx')
    print("Successfully generated April_2026_Sales_Report.docx")

if __name__ == '__main__':
    generate_report()
