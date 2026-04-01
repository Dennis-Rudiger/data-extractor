import json
from reportlab.lib import colors
from reportlab.lib.pagesizes import landscape, A4
from reportlab.platypus import BaseDocTemplate, PageTemplate, Frame, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
from reportlab.platypus import Image
import os
import matplotlib.pyplot as plt
import numpy as np

os.makedirs('charts/q1_bomas', exist_ok=True)

# LOAD DATA
print('Loading stock movement files...')
with open('stock_movement_jan.json', 'r') as f: jan = {i['code']: i for i in json.load(f)['items']}
with open('stock_movement_feb.json', 'r') as f: feb = {i['code']: i for i in json.load(f)['items']}
with open('stock_movement_mar.json', 'r') as f: mar = {i['code']: i for i in json.load(f)['items']}

all_codes = set(jan.keys()) | set(feb.keys()) | set(mar.keys())
items = []
total_jan = 0
total_feb = 0
total_mar = 0

for code in all_codes:
    j = jan.get(code, {'qty_out': 0, 'description': ''})
    f = feb.get(code, {'qty_out': 0, 'description': ''})
    m = mar.get(code, {'qty_out': 0, 'description': ''})
    desc = m['description'] or f['description'] or j['description']
    
    total = j['qty_out'] + f['qty_out'] + m['qty_out']
    total_jan += j['qty_out']
    total_feb += f['qty_out']
    total_mar += m['qty_out']
    
    status = 'Consistent' if j['qty_out'] > 0 and f['qty_out'] > 0 and m['qty_out'] > 0 else \
             'Declining' if (j['qty_out'] > 0 or f['qty_out'] > 0) and m['qty_out'] == 0 else \
             'Emerging' if j['qty_out'] == 0 and f['qty_out'] == 0 and m['qty_out'] > 0 else \
             'Dead Stock' if total == 0 else 'Sporadic'
             
    items.append({
        'code': code.strip()[:15],
        'description': desc[:45].strip(),
        'jan': j['qty_out'],
        'feb': f['qty_out'],
        'mar': m['qty_out'],
        'total': total,
        'status': status
    })

items.sort(key=lambda x: x['total'], reverse=True)
top_items = items[:40]
dead_stock = [i for i in items if i['status'] == 'Dead Stock'][:50] # Top 50 dead stock
declining = [i for i in items if i['status'] == 'Declining'][:30]

# Write JSON
with open('quarter_one_stock_bomas.json', 'w') as f:
    json.dump({'items': items, 'top_items': top_items, 'dead_stock': dead_stock}, f, indent=2)

print('Generating charts...')

# Chart 1: Monthly Overall Trend
fig, ax = plt.subplots(figsize=(6, 4))
labels = ['Jan', 'Feb', 'Mar']
totals = [total_jan, total_feb, total_mar]
ax.bar(labels, totals, color=['#bdc3c7', '#7f8c8d', '#3498db'])
ax.set_ylabel('Total Units Moved')
ax.set_title('Overall Monthly Movement Trend', fontweight='bold')
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)
for i, v in enumerate(totals):
    ax.text(i, v + (max(totals)*0.01), f'{v:,.0f}', ha='center', va='bottom', fontweight='bold')
plt.tight_layout()
plt.savefig('charts/q1_bomas/overall_trend.png', dpi=120, bbox_inches='tight')
plt.close()

# Chart 2: Item Status Distribution
fig, ax = plt.subplots(figsize=(6, 4))
status_counts = {}
for i in items:
    status_counts[i['status']] = status_counts.get(i['status'], 0) + 1
labels = list(status_counts.keys())
sizes = list(status_counts.values())
# Make sure we have 5 colors for 5 statuses
ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90, colors=['#2ecc71', '#e74c3c', '#f1c40f', '#95a5a6', '#34495e'])
ax.set_title('SKU Status Distribution', fontweight='bold')
plt.tight_layout()
plt.savefig('charts/q1_bomas/status_pie.png', dpi=120, bbox_inches='tight')
plt.close()

# Chart 3: Top 10 items breakdown
fig, ax = plt.subplots(figsize=(10, 5))
top10 = top_items[:10]
names = [i['code'] for i in top10]
w1_vals = [i['jan'] for i in top10]
w2_vals = [i['feb'] for i in top10]
w3_vals = [i['mar'] for i in top10]

x = np.arange(len(names))
width = 0.25

ax.bar(x - width, w1_vals, width, label='Jan', color='#bdc3c7', edgecolor='white')
ax.bar(x, w2_vals, width, label='Feb', color='#7f8c8d', edgecolor='white')
ax.bar(x + width, w3_vals, width, label='Mar', color='#3498db', edgecolor='white')

ax.set_ylabel('Units Moved')
ax.set_title('Top 10 Fast Movers Breakdown', fontweight='bold')
ax.set_xticks(x)
ax.set_xticks(x, names, rotation=25, ha='right', fontsize=9)
ax.legend()
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)

plt.tight_layout()
plt.savefig('charts/q1_bomas/top10_breakdown.png', dpi=120, bbox_inches='tight')
plt.close()

print('Generating reports...')

styles = {
    'title': ParagraphStyle('Title', fontName='Helvetica-Bold', fontSize=14, textColor=colors.HexColor('#1a3a5c'), spaceAfter=15),
    'heading': ParagraphStyle('Heading', fontName='Helvetica-Bold', fontSize=12, textColor=colors.HexColor('#1a3a5c'), spaceAfter=10, spaceBefore=15),
    'text': ParagraphStyle('Text', fontName='Helvetica', fontSize=10, leading=14, spaceAfter=8, textColor=colors.HexColor('#2c3e50')),
    'recs': ParagraphStyle('Recs', fontName='Helvetica', fontSize=10, leading=14, spaceAfter=6, leftIndent=15, bulletIndent=5, textColor=colors.HexColor('#2c3e50'))
}

pdf_file = 'quarter_one_stock_bomas.pdf'
doc = BaseDocTemplate(pdf_file, pagesize=landscape(A4), leftMargin=1*inch, rightMargin=1*inch, topMargin=1*inch, bottomMargin=1*inch)
frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id='normal')
template = PageTemplate(id='base', frames=frame)
doc.addPageTemplates([template])

elements = []
elements.append(Paragraph('BOMAS HARDWARE - Q1 STOCK MOVEMENT REPORT', styles['title']))
elements.append(Paragraph('Period: January 5 - March 31, 2026 | Includes Jan, Feb, Mar Comparison', styles['text']))
elements.append(Paragraph('Total SKUs Monitored: {:,} | Total Units Moved: {:,.0f} (Jan: {:,.0f} - Feb: {:,.0f} - Mar: {:,.0f})'.format(len(items), total_jan+total_feb+total_mar, total_jan, total_feb, total_mar), styles['text']))
elements.append(Spacer(1, 10))

# Insert Charts
elements.append(Paragraph('VISUALIZED DATA SUMMARY', styles['heading']))
from reportlab.platypus import Table as rTable
images_table = rTable([
    [Image('charts/q1_bomas/overall_trend.png', width=4*inch, height=2.66*inch), 
     Image('charts/q1_bomas/status_pie.png', width=4*inch, height=2.66*inch)]
], colWidths=[4.2*inch, 4.2*inch])
elements.append(images_table)
elements.append(Spacer(1, 15))
elements.append(Image('charts/q1_bomas/top10_breakdown.png', width=8.4*inch, height=4.2*inch))
elements.append(PageBreak())

recs = [
    "Prioritize restocking investments on 'Consistent' top-moving SKUs (e.g. Simba Cement, D10) which reliably drive volume across all months.",
    "Conduct an immediate audit on the Dead Stock SKUs (recorded zero movement across Jan, Feb, and March). Consider promotions or liquidation to free up space.",
    "Investigate 'Declining' SKUs which moved well early in the quarter but recorded 0 sales in March. Verify if this drop is due to demand falloff or simply stockouts.",
    "Implement monthly MOQ re-calculations that adjust thresholds organically. Items tagged 'Emerging' that picked up pace late in Q1 require updated stock buffers."
]

elements.append(Paragraph('INFORMED RECOMMENDATIONS & INSIGHTS', styles['heading']))
for r in recs:
    elements.append(Paragraph('- '+r, styles['recs']))

elements.append(Spacer(1, 15))

headers = ['Code', 'Description', 'Jan Qty', 'Feb Qty', 'Mar Qty', 'Q1 Total', 'Status']
data = [headers]
for item in top_items:
    data.append([item['code'], item['description'], f"{item['jan']:,.0f}", f"{item['feb']:,.0f}", f"{item['mar']:,.0f}", f"{item['total']:,.0f}", item['status']])

t = Table(data, colWidths=[1.5*inch, 3*inch, 0.9*inch, 0.9*inch, 0.9*inch, 1*inch, 1.2*inch])
t.setStyle(TableStyle([
    ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1a3a5c')),
    ('TEXTCOLOR', (0,0), (-1,0), colors.white),
    ('ALIGN', (2,0), (5,-1), 'RIGHT'),
    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
    ('BOTTOMPADDING', (0,0), (-1,0), 6),
    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
    ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#d5d8dc'))
]))

elements.append(Paragraph('Q1 KEY ITEMS: TOP MOVERS (CONSISTENT DEMAND)', styles['heading']))
elements.append(t)
elements.append(PageBreak())

# Declining
data = [headers]
for item in declining:
    data.append([item['code'], item['description'], f"{item['jan']:,.0f}", f"{item['feb']:,.0f}", f"{item['mar']:,.0f}", f"{item['total']:,.0f}", item['status']])
t = Table(data, colWidths=[1.5*inch, 3*inch, 0.9*inch, 0.9*inch, 0.9*inch, 1*inch, 1.2*inch])
t.setStyle(TableStyle([
    ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1a3a5c')),
    ('TEXTCOLOR', (0,0), (-1,0), colors.white),
    ('ALIGN', (2,0), (5,-1), 'RIGHT'),
    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
    ('BOTTOMPADDING', (0,0), (-1,0), 6),
    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
    ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#d5d8dc'))
]))

elements.append(Paragraph('DECLINING ITEMS (ZERO MOVEMENT IN MARCH)', styles['heading']))
elements.append(t)
elements.append(Spacer(1, 15))

# Dead Stock
data = [headers[:6]]
for item in dead_stock:
    data.append([item['code'], item['description'], f"{item['jan']:,.0f}", f"{item['feb']:,.0f}", f"{item['mar']:,.0f}", f"{item['total']:,.0f}"])
t = Table(data, colWidths=[1.5*inch, 3.5*inch, 0.9*inch, 0.9*inch, 0.9*inch, 1.5*inch])
t.setStyle(TableStyle([
    ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#c0392b')),
    ('TEXTCOLOR', (0,0), (-1,0), colors.white),
    ('ALIGN', (2,0), (5,-1), 'RIGHT'),
    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
    ('BOTTOMPADDING', (0,0), (-1,0), 6),
    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
    ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#d5d8dc'))
]))

elements.append(Paragraph('DEAD STOCK SAMPLES (ZERO MOVEMENT ENTIRE Q1)', styles['heading']))
elements.append(t)

doc.build(elements)
print(f'Created {pdf_file}')

# WORD DOC
word_file = 'quarter_one_stock_bomas.docx'
document = Document()
document.add_heading('BOMAS HARDWARE - Q1 STOCK MOVEMENT REPORT', 0)
document.add_paragraph('Period: January 5 - March 31, 2026 | Includes Jan, Feb, Mar Comparison')
p = document.add_paragraph('Total SKUs Monitored: {:,} | Total Units Moved: {:,.0f}'.format(len(items), total_jan+total_feb+total_mar))
p.add_run(' (Jan: {:,.0f} - Feb: {:,.0f} - Mar: {:,.0f})'.format(total_jan, total_feb, total_mar)).bold = True

document.add_heading('INFORMED RECOMMENDATIONS & INSIGHTS', level=1)
for r in recs:
    document.add_paragraph(r, style='List Bullet')

document.add_heading('VISUALIZED DATA SUMMARY', level=1)
document.add_picture('charts/q1_bomas/overall_trend.png', width=Inches(6.0))
document.add_picture('charts/q1_bomas/status_pie.png', width=Inches(5.0))
document.add_picture('charts/q1_bomas/top10_breakdown.png', width=Inches(6.5))

def add_table(doc, title, items, is_dead=False):
    doc.add_heading(title, level=1)
    cols = 6 if is_dead else 7
    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Light Shading Accent 1' if not is_dead else 'Light Shading Accent 2'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Code'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Jan Qty'
    hdr_cells[3].text = 'Feb Qty'
    hdr_cells[4].text = 'Mar Qty'
    hdr_cells[5].text = 'Q1 Total'
    if not is_dead: hdr_cells[6].text = 'Status'
    
    for item in items:
        row_cells = table.add_row().cells
        row_cells[0].text = item['code']
        row_cells[1].text = item['description']
        row_cells[2].text = f"{item['jan']:,.0f}"
        row_cells[3].text = f"{item['feb']:,.0f}"
        row_cells[4].text = f"{item['mar']:,.0f}"
        row_cells[5].text = f"{item['total']:,.0f}"
        if not is_dead: row_cells[6].text = item['status']

add_table(document, 'Q1 KEY ITEMS: TOP MOVERS (CONSISTENT DEMAND)', top_items)
add_table(document, 'DECLINING ITEMS (ZERO MOVEMENT IN MARCH)', declining)
add_table(document, 'DEAD STOCK SAMPLES (ZERO MOVEMENT ENTIRE Q1)', dead_stock, True)

document.save(word_file)
print(f'Created {word_file}')

