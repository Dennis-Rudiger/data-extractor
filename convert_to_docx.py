import json
from docx import Document

with open('bomas_average_moq.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

categories_mapping = [
    ('1. Electricals', 'ELECTRICALS'),
    ('2. Plumbing', 'PLUMBING MATERIALS'),
    ('3. Paints', 'PAINTS'),
    ('4. Paints and Accessories', 'PAINT ACCESSORIES'),
    ('5. Mortar and Adhesives', 'MORTAR AND ADHESIVES')
]

doc = Document()
doc.add_heading('Top 20 Fast Movers by Category', 0)

for display_name, cat_key in categories_mapping:
    doc.add_heading(display_name, level=1)
    
    if cat_key in data['categories']:
        prods = data['categories'][cat_key].get('products', [])
        prods_sorted = sorted(prods, key=lambda x: x.get('qty_out', 0), reverse=True)[:20]
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Item Code'
        hdr_cells[1].text = 'Description'
        hdr_cells[2].text = 'Qty Out'
        hdr_cells[3].text = 'Daily Avg'
        hdr_cells[4].text = 'Weekly MOQ'
        
        for p in prods_sorted:
            row_cells = table.add_row().cells
            row_cells[0].text = str(p.get('item_code', ''))
            row_cells[1].text = str(p.get('item_description', ''))
            row_cells[2].text = '{:,.1f}'.format(p.get('qty_out', 0))
            row_cells[3].text = '{:,.2f}'.format(p.get('daily_average', 0))
            row_cells[4].text = '{:,.0f}'.format(p.get('weekly_moq', 0))
            
        doc.add_paragraph()

doc.save('top_20_fast_movers.docx')
print('Document saved as top_20_fast_movers.docx')
