"""
Generate a fillable Branch Managers Weekly Performance Report (.docx)
Uses fillable lines below each question instead of tables.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

# -- Helper functions --

def add_header_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Jenga Karen Na Bei Za Inda")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Email: zpannju@gmail.com  |  +254-733-258-419")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("WEEKLY BRANCH PERFORMANCE REPORT")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)

    # Thin line divider
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="003366"/>'
        '</w:pBdr>'
    )
    p._element.get_or_add_pPr().append(pBdr)


def add_section_header(doc, label):
    """Dark blue banner section header."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="003366"/>')
    p._element.get_or_add_pPr().append(shading)
    run = p.add_run(f"  {label}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(255, 255, 255)


def add_subsection_header(doc, label):
    """Light blue subsection header."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8EEF4"/>')
    p._element.get_or_add_pPr().append(shading)
    run = p.add_run(f"  {label}")
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = RGBColor(0, 51, 102)


def add_fill_question(doc, question, lines=3):
    """Add a numbered/bulleted question followed by blank fill lines."""
    # Question text
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(question)
    run.font.size = Pt(10.5)
    run.bold = True

    # Fill lines underneath
    for _ in range(lines):
        fp = doc.add_paragraph()
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after = Pt(0)
        fp.paragraph_format.left_indent = Cm(0.5)
        # Bottom border acts as a clean fill line
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="BFBFBF"/>'
            '</w:pBdr>'
        )
        fp._element.get_or_add_pPr().append(pBdr)
        run = fp.add_run(" ")
        run.font.size = Pt(14)  # taller line height for writing space


def add_inline_field(doc, label):
    """Add a label with an inline fill line on the same line (for short answers)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(f"{label}  ")
    run.font.size = Pt(10.5)
    run.bold = True
    # Inline fill with bottom border on entire paragraph
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="BFBFBF"/>'
        '</w:pBdr>'
    )
    p._element.get_or_add_pPr().append(pBdr)


# ==============================================================
# BUILD THE DOCUMENT
# ==============================================================

add_header_block(doc)

# -- Report Info (inline fields) --
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
run = p.add_run("Report Details")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 51, 102)

add_inline_field(doc, "Branch Name:")
add_inline_field(doc, "Report Period (Week No.):")
add_inline_field(doc, "Month / Year:")
add_inline_field(doc, "Prepared By:")
add_inline_field(doc, "Date Submitted:")

# -- Purpose / Objective --
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
run = p.add_run("Purpose / Objective")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = p.add_run(
    "Our people strategy for 2026 focuses on performance and productivity. "
    "The weekly report provides a snapshot of branch and departmental performance, "
    "to quickly identify trends and areas requiring intervention."
)
run.font.size = Pt(10)
run.italic = True

p = doc.add_paragraph()
run = p.add_run("Objectives:")
run.bold = True
run.font.size = Pt(10)

for obj in [
    "Track weekly departmental performance against KPIs.",
    "Ensure alignment with the 2026 people strategy.",
    "Provide actionable insights to support decision-making and operational improvements."
]:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(obj)
    run.font.size = Pt(10)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
run = p.add_run("Guidelines for Branch Managers:")
run.bold = True
run.font.size = Pt(10)

for g in [
    "Collaboration: Liaise with department heads (sales, finance, etc.) to gather required data.",
    "Specificity: Include names, figures, and relevant details wherever applicable.",
    "Completeness: Answer all questions. If information is unavailable, indicate clearly.",
    "Verification: Ensure all data is accurate, cross-checked, and supported by records."
]:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(g)
    run.font.size = Pt(10)


# ==============================================================
# SECTION A: SALES DEPARTMENT
# ==============================================================
doc.add_page_break()
add_section_header(doc, "A. SALES DEPARTMENT")

add_subsection_header(doc, "1. Revenue Performance")

add_fill_question(doc, "a) What was the weekly sales target? (KES)", 2)
add_fill_question(doc, "b) What was the actual revenue achieved? (KES)", 2)
add_fill_question(doc, "c) What was the percentage variance? (%)", 2)
add_fill_question(doc, "d) What were the main factors contributing to over-performance or under-performance this week?", 4)
add_fill_question(doc, "e) Were there any large or unusual sales transactions that significantly affected the weekly revenue? Please specify.", 4)

add_subsection_header(doc, "2. People Drivers")

add_fill_question(doc, "a) Which sales executives met their target? (Specify names and amounts)", 4)
add_fill_question(doc, "b) Which sales executives were below targets? (Specify names and amounts)", 4)
add_fill_question(doc, "c) What are the reasons for sales executives falling below their targets?", 4)
add_fill_question(doc, "d) What contributed towards the top-performing sales executives' performance?", 3)
add_fill_question(doc, "e) What support actions are planned for underperforming sales executives? (Names and actions)", 4)

doc.add_page_break()
add_subsection_header(doc, "3. Field Sales & Market Expansion")

add_fill_question(doc, "a) Names of clients or construction sites visited during the week:", 4)
add_fill_question(doc, "b) Names of new customer accounts opened or converted:", 4)
add_fill_question(doc, "c) How many previously dormant accounts were reactivated? (Provide names)", 3)
add_fill_question(doc, "d) What customer or market feedback was captured during field visits, and were any follow-up actions initiated?", 4)


# ==============================================================
# SECTION B: FINANCE DEPARTMENT
# ==============================================================
doc.add_page_break()
add_section_header(doc, "B. FINANCE DEPARTMENT")

add_subsection_header(doc, "1. Financial Accuracy & Discipline")

add_fill_question(doc, "a) Were all ledger entries and financial postings completed on time this week?", 3)
add_fill_question(doc, "b) Were there any posting or data entry errors identified? If yes, how many and how were they corrected?", 3)
add_fill_question(doc, "c) Is there any backlog of unposted transactions? Indicate the volume and reason for delay.", 3)
add_fill_question(doc, "d) Were all financial records reviewed and verified by the responsible officer? (Name of officer)", 2)

add_subsection_header(doc, "2. Compliance & Reporting")

add_fill_question(doc, "a) Was the weekly financial report prepared and submitted on time? (Name of officer in charge)", 3)
add_fill_question(doc, "b) Were all bank, cash, and supplier reconciliations completed for the week? (Name of officer in charge)", 3)
add_fill_question(doc, "c) Are there any missing audit or compliance documents that need follow-up?", 3)
add_fill_question(doc, "d) Were there any compliance risks or reporting gaps identified, and what actions were taken?", 3)


# ==============================================================
# SECTION C: STOCK CONTROL & INVENTORY MANAGEMENT
# ==============================================================
doc.add_page_break()
add_section_header(doc, "C. STOCK CONTROL & INVENTORY MANAGEMENT")

add_subsection_header(doc, "1. Stock Availability")

add_fill_question(doc, "a) Are all key items currently in stock? List any critical items that fell below minimum stock level.", 4)
add_fill_question(doc, "b) Were there any stock-outs this week? Specify items, duration, and impact on sales.", 4)
add_fill_question(doc, "c) List slow-moving items identified this week. Indicate whether any action plans are in place.", 4)

add_subsection_header(doc, "2. Stock Accuracy & Control")

add_fill_question(doc, "a) What was the stock vs system variance percentage this week? Comments:", 3)
add_fill_question(doc, "b) Were there discrepancies identified during stock counts? List items, quantities, and corrective actions.", 4)
add_fill_question(doc, "c) Were there any pricing errors identified in the system? List items and corrective actions.", 3)
add_fill_question(doc, "d) Is product display and shelf arrangement compliant with company standards? Comments:", 3)

doc.add_page_break()
add_subsection_header(doc, "3. Stock Risk & Cost Control")

add_fill_question(doc, "a) Was there any damage, wastage, or loss recorded this week? List items and estimated value.", 4)
add_fill_question(doc, "b) Were there any internal stock handling issues (misplacement, poor storage, expiry risks)?", 4)
add_fill_question(doc, "c) Are there any risks that could affect stock availability or profitability? Indicate risk level and mitigation.", 4)


# ==============================================================
# SECTION D: ADMIN WEEKLY REPORT
# ==============================================================
doc.add_page_break()
add_section_header(doc, "D. ADMIN WEEKLY REPORT")

add_subsection_header(doc, "1. Employee Headcount & Movements")

add_inline_field(doc, "Opening Headcount (Contracted Staff):")
add_inline_field(doc, "Closing Headcount (Contracted Staff):")
add_inline_field(doc, "Opening Headcount (Casual Workers):")
add_inline_field(doc, "Closing Headcount (Casual Workers):")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)

add_fill_question(doc, "Comments on any employee movements (new hires, exits, transfers):", 3)

add_subsection_header(doc, "2. Attendance & Leave Summary")

add_inline_field(doc, "Lateness Incidents (count):")
add_fill_question(doc, "Details of lateness incidents (names, dates, actions taken):", 3)
add_inline_field(doc, "Absenteeism Incidents (count):")
add_fill_question(doc, "Details of absenteeism incidents (names, dates, actions taken):", 3)
add_inline_field(doc, "Approved Leave Taken (count):")
add_fill_question(doc, "Details of approved leave (names, type of leave, dates):", 3)

add_subsection_header(doc, "3. Disciplinary & Employee Relations")

add_fill_question(doc, "a) Ongoing disciplinary cases (provide details):", 4)
add_fill_question(doc, "b) Employee complaints / grievances received:", 4)
add_fill_question(doc, "c) Resolutions completed:", 4)


# ==============================================================
# SIGN-OFF
# ==============================================================
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8EEF4"/>')
p._element.get_or_add_pPr().append(shading)
run = p.add_run("  Sign-Off")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)

add_inline_field(doc, "Branch Manager Name:")
add_inline_field(doc, "Signature:")
add_inline_field(doc, "Date:")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)

add_inline_field(doc, "Reviewed By (HR / MD) Name:")
add_inline_field(doc, "Signature:")
add_inline_field(doc, "Date:")

# Footer
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Jenga Karen Na Bei Za Inda!")
run.bold = True
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 51, 102)

# -- Save --
output_path = "Branch_Managers_Weekly_Performance_Report.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
print("Form uses fillable lines below each question.")
print("Users type directly on the lines in Word.")
