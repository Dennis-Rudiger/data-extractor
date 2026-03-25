"""
Generate comprehensive March Week 3 sales analysis reports.
Period: March 16-21, 2026 (6 working days)
Outputs: JSON, PDF (with charts), Excel, Word
"""
import json
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle, Paragraph, 
                                 Spacer, PageBreak, Image)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.units import inch
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.chart import BarChart, PieChart, Reference
from openpyxl.chart.label import DataLabelList
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime

# ========== CONFIGURATION ==========
PERIOD = "March 16-21, 2026"
WORKING_DAYS = 6
MONTH = "March"
YEAR = 2026
WEEK_NUM = 3

# Monthly targets (March)
MONTHLY_TARGETS = {
    'OVERALL': 60_000_000,
    'PAINTS': 7_000_000,
    'PLUMBING': 3_000_000,
    'ELECTRICALS': 1_500_000,
    'GENERAL HARDWARE': 48_500_000,
}

CAT_ORDER = ['GENERAL HARDWARE', 'PAINTS', 'PLUMBING', 'ELECTRICALS']

COLORS = {
    'GENERAL HARDWARE': '#3498db',
    'PAINTS': '#e74c3c',
    'PLUMBING': '#2ecc71',
    'ELECTRICALS': '#f39c12'
}

REP_COLORS = ['#2E86AB', '#A23B72', '#F18F01', '#C73E1D', '#6B4C9A', '#2ECC71', '#1ABC9C']

DEPARTMENT_HEADS = {
    'PLUMBING': 'Lewis',
    'PAINTS': 'Bonface Muriu',
    'ELECTRICALS': 'Bonface Kitheka',
    'GENERAL HARDWARE': ['Gladys', 'Betha Odumo', 'Eliza', 'Stephen']
}

VAT_RATE = 1.16

# ========== LOAD DATA ==========
with open("sales_march_w3.json", "r") as f:
    raw_data = json.load(f)

# ========== DATA PROCESSING ==========
def merge_reps(data):
    """Merge Magdalene -> Betha Odumo"""
    merged = {"period": data["period"], "working_days": data["working_days"], "categories": {}}
    for cat, reps in data["categories"].items():
        merged["categories"][cat] = {}
        for rep, vals in reps.items():
            target = "Betha Odumo" if rep in ["Magdalene", "WALK IN-BOMAS"] else rep
            if target not in merged["categories"][cat]:
                merged["categories"][cat][target] = {"qty": 0, "sales_incl": 0, "cost": 0, "profit": 0}
            merged["categories"][cat][target]["qty"] += vals["qty"]
            merged["categories"][cat][target]["sales_incl"] += vals["sales_incl"]
            merged["categories"][cat][target]["cost"] += vals["cost"]
            merged["categories"][cat][target]["profit"] += vals["profit"]
    return merged

week3_data = merge_reps(raw_data)

def get_rep_data(data):
    """Process rep data with 16% VAT on profits."""
    reps = {}
    for cat, rep_data in data["categories"].items():
        for rep, vals in rep_data.items():
            if rep not in reps:
                reps[rep] = {"total_sales": 0, "total_profit": 0, "total_cost": 0, "total_qty": 0, "categories": {}}
            profit_incl_vat = vals["profit"] * VAT_RATE
            reps[rep]["categories"][cat] = vals.copy()
            reps[rep]["categories"][cat]["profit"] = profit_incl_vat
            reps[rep]["categories"][cat]["margin_pct"] = (profit_incl_vat / vals["sales_incl"] * 100) if vals["sales_incl"] > 0 else 0
            reps[rep]["total_sales"] += vals["sales_incl"]
            reps[rep]["total_profit"] += profit_incl_vat
            reps[rep]["total_cost"] += vals["cost"]
            reps[rep]["total_qty"] += vals["qty"]
    for rep, data_r in reps.items():
        data_r["overall_margin"] = (data_r["total_profit"] / data_r["total_sales"] * 100) if data_r["total_sales"] > 0 else 0
    return reps

reps_data = get_rep_data(week3_data)
sorted_reps = sorted(reps_data.items(), key=lambda x: x[1]['total_sales'], reverse=True)
total_sales = sum(r['total_sales'] for r in reps_data.values())
total_profit = sum(r['total_profit'] for r in reps_data.values())
overall_margin = (total_profit / total_sales * 100) if total_sales > 0 else 0

def get_rep_role(rep_name):
    if rep_name == 'Lewis': return 'Department Head - Plumbing'
    elif rep_name == 'Bonface Muriu': return 'Department Head - Paints'
    elif rep_name == 'Bonface Kitheka': return 'Department Head - Electricals'
    else: return 'General Hardware Team'

def get_primary_dept(rep_name):
    if rep_name == 'Lewis': return 'PLUMBING'
    elif rep_name == 'Bonface Muriu': return 'PAINTS'
    elif rep_name == 'Bonface Kitheka': return 'ELECTRICALS'
    else: return 'GENERAL HARDWARE'

# Category totals
cat_totals = {}
for cat in CAT_ORDER:
    cat_sales = sum(r['categories'].get(cat, {}).get('sales_incl', 0) for r in reps_data.values())
    cat_profit = sum(r['categories'].get(cat, {}).get('profit', 0) for r in reps_data.values())
    cat_totals[cat] = {'sales': cat_sales, 'profit': cat_profit}

# Weekly target (1/4 of monthly)
weekly_target = MONTHLY_TARGETS['OVERALL'] / 4

# Load prior weeks
with open("sales_march_w1.json", "r") as f: w1_raw = json.load(f)
with open("sales_march_w2.json", "r") as f: w2_raw = json.load(f)

def get_week_sales(data_raw):
    totals_cat = {c: 0 for c in CAT_ORDER}
    totals = 0
    merged = merge_reps(data_raw)
    rep_d = get_rep_data(merged)
    for rep, r_data in rep_d.items():
        totals += r_data['total_sales']
        for cat in CAT_ORDER:
            totals_cat[cat] += r_data['categories'].get(cat, {}).get('sales_incl', 0)
    rep_sales = {rep: r_data['total_sales'] for rep, r_data in rep_d.items()}
    return totals, totals_cat, rep_sales

w1_tot, w1_cat, w1_rep = get_week_sales(w1_raw)
w2_tot, w2_cat, w2_rep = get_week_sales(w2_raw)

cumulative_total_sales = w1_tot + w2_tot + total_sales
cumulative_cat_totals = {c: cat_totals[c]['sales'] + w1_cat[c] + w2_cat[c] for c in CAT_ORDER}


# ========== CHART GENERATION ==========
def create_category_pie(output_dir="charts/march_w3"):
    os.makedirs(output_dir, exist_ok=True)
    cats = []
    sales = []
    chart_colors = []
    for cat in CAT_ORDER:
        if cat_totals[cat]['sales'] > 0:
            cats.append(cat)
            sales.append(cat_totals[cat]['sales'])
            chart_colors.append(COLORS[cat])
    fig, ax = plt.subplots(figsize=(8, 5))
    wedges, texts, autotexts = ax.pie(sales, labels=None, autopct='%1.1f%%',
                                       colors=chart_colors, startangle=90,
                                       pctdistance=0.75, explode=[0.02]*len(sales),
                                       wedgeprops=dict(width=0.5, edgecolor='white'))
    ax.text(0, 0, f'KES\n{sum(sales)/1000000:.1f}M', ha='center', va='center',
            fontsize=14, fontweight='bold', color='#2c3e50')
    for at in autotexts:
        at.set_fontsize(9)
        at.set_fontweight('bold')
        at.set_color('white')
    legend_labels = [f"{c}: KES {s/1000000:.2f}M" for c, s in zip(cats, sales)]
    ax.legend(wedges, legend_labels, title="Categories", loc="center left",
              bbox_to_anchor=(1, 0, 0.5, 1), fontsize=9)
    ax.set_title(f"Sales by Category - {MONTH} Week {WEEK_NUM}", fontsize=13, fontweight='bold')
    plt.tight_layout()
    fn = f"{output_dir}/category_pie.png"
    plt.savefig(fn, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return fn


def create_weekly_trend_chart(output_dir="charts/march_w3"):
    os.makedirs(output_dir, exist_ok=True)
    names = [r[0] for r in sorted_reps]
    w1_vals = [w1_rep.get(n, 0) / 1000000 for n in names]
    w2_vals = [w2_rep.get(n, 0) / 1000000 for n in names]
    w3_vals = [r[1]['total_sales'] / 1000000 for r in sorted_reps]

    fig, ax = plt.subplots(figsize=(12, 5))
    x = np.arange(len(names))
    width = 0.25

    bar1 = ax.bar(x - width, w1_vals, width, label='Week 1', color='#bdc3c7', edgecolor='white')
    bar2 = ax.bar(x, w2_vals, width, label='Week 2', color='#7f8c8d', edgecolor='white')
    bar3 = ax.bar(x + width, w3_vals, width, label='Week 3', color='#3498db', edgecolor='white')
    
    # Add data labels
    def add_labels(rects):
        for rect in rects:
            height = rect.get_height()
            if height > 0:
                ax.annotate(f'{height:.1f}',
                            xy=(rect.get_x() + rect.get_width() / 2, height),
                            xytext=(0, 3),  # 3 points vertical offset
                            textcoords="offset points",
                            ha='center', va='bottom', fontsize=8, rotation=90)
                
    add_labels(bar1)
    add_labels(bar2)
    add_labels(bar3)

    ax.set_ylabel('Sales (Millions KES)')
    ax.set_title('Sales Trend by Rep (W1 vs W2 vs W3)', fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticks(x, names, rotation=25, ha='right', fontsize=9)
    # Put legend outside or at a top corner so labels are not covered
    ax.legend(loc='upper right')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)

    # Adjust upper limit so labels fit
    max_val = max(max(w1_vals) if w1_vals else 0, max(w2_vals) if w2_vals else 0, max(w3_vals) if w3_vals else 0)
    ax.set_ylim(0, max_val * 1.25)

    plt.tight_layout()
    fn = f"{output_dir}/weekly_trends.png"
    plt.savefig(fn, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return fn

def create_rep_bar_chart(output_dir="charts/march_w3"):
    os.makedirs(output_dir, exist_ok=True)
    names = [r[0] for r in sorted_reps]
    sales_vals = [r[1]['total_sales']/1000000 for r in sorted_reps]
    profit_vals = [r[1]['total_profit']/1000000 for r in sorted_reps]
    fig, ax = plt.subplots(figsize=(10, 5))
    x = np.arange(len(names))
    w = 0.35
    bars1 = ax.bar(x - w/2, sales_vals, w, label='Sales', color='#3498db', edgecolor='white')
    bars2 = ax.bar(x + w/2, profit_vals, w, label='Profit', color='#2ecc71', edgecolor='white')
    for bar in bars1:
        h = bar.get_height()
        ax.annotate(f'{h:.2f}M', xy=(bar.get_x() + bar.get_width()/2, h),
                    xytext=(0, 3), textcoords="offset points", ha='center', fontsize=8)
    for bar in bars2:
        h = bar.get_height()
        ax.annotate(f'{h:.2f}M', xy=(bar.get_x() + bar.get_width()/2, h),
                    xytext=(0, 3), textcoords="offset points", ha='center', fontsize=8)
    ax.set_ylabel('Amount (Millions KES)')
    ax.set_title(f'Sales Rep Performance - {MONTH} Week {WEEK_NUM}', fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticks(x, names, rotation=25, ha='right', fontsize=9)
    ax.legend()
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    fn = f"{output_dir}/rep_bar.png"
    plt.savefig(fn, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return fn

def create_margin_comparison(output_dir="charts/march_w3"):
    os.makedirs(output_dir, exist_ok=True)
    names = [r[0] for r in sorted_reps]
    margins = [r[1]['overall_margin'] for r in sorted_reps]
    fig, ax = plt.subplots(figsize=(10, 4))
    bar_colors = ['#27ae60' if m >= 15 else '#f39c12' if m >= 10 else '#e74c3c' for m in margins]
    bars = ax.barh(names, margins, color=bar_colors, edgecolor='white', height=0.6)
    for bar, m in zip(bars, margins):
        ax.text(bar.get_width() + 0.3, bar.get_y() + bar.get_height()/2,
                f'{m:.1f}%', va='center', fontsize=10, fontweight='bold')
    ax.axvline(x=overall_margin, color='#2c3e50', linestyle='--', linewidth=1.5, label=f'Avg: {overall_margin:.1f}%')
    ax.set_xlabel('Profit Margin (%)')
    ax.set_title(f'Profit Margin by Rep - {MONTH} Week {WEEK_NUM}', fontweight='bold')
    ax.legend()
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.invert_yaxis()
    plt.tight_layout()
    fn = f"{output_dir}/margin_comparison.png"
    plt.savefig(fn, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return fn

def create_target_progress(output_dir="charts/march_w3"):
    os.makedirs(output_dir, exist_ok=True)
    fig, axes = plt.subplots(1, 5, figsize=(14, 3))
    items = [('OVERALL', cumulative_total_sales, MONTHLY_TARGETS['OVERALL'])]
    for cat in CAT_ORDER:
        items.append((cat, cumulative_cat_totals[cat], MONTHLY_TARGETS.get(cat, 0)))
    for ax, (label, actual, target) in zip(axes, items):
        pct = (actual / target * 100) if target > 0 else 0
        expected = 25  # Week 3 of 4
        color = '#27ae60' if pct >= expected else '#e74c3c'
        ax.barh([''], [pct], color=color, height=0.4, alpha=0.8)
        ax.axvline(x=expected, color='#2c3e50', linestyle='--', linewidth=1.5)
        ax.set_xlim(0, max(35, pct + 5))
        ax.set_title(label[:12], fontsize=9, fontweight='bold')
        ax.text(pct + 0.5, 0, f'{pct:.1f}%', va='center', fontsize=9, fontweight='bold', color=color)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
    fig.suptitle(f'Monthly Target Progress After Week {WEEK_NUM} (Expected: 75%)', fontweight='bold', fontsize=11)
    plt.tight_layout()
    fn = f"{output_dir}/target_progress.png"
    plt.savefig(fn, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return fn

def create_stacked_bar(output_dir="charts/march_w3"):
    os.makedirs(output_dir, exist_ok=True)
    names = [r[0] for r in sorted_reps]
    x = np.arange(len(names))
    fig, ax = plt.subplots(figsize=(10, 5))
    bottom = np.zeros(len(names))
    for cat in CAT_ORDER:
        vals = []
        for rep_name, rep_data in sorted_reps:
            vals.append(rep_data['categories'].get(cat, {}).get('sales_incl', 0) / 1000000)
        ax.bar(names, vals, bottom=bottom, label=cat, color=COLORS[cat], edgecolor='white')
        bottom += np.array(vals)
    ax.set_ylabel('Sales (Millions KES)')
    ax.set_title(f'Sales Breakdown by Category per Rep - {MONTH} Week {WEEK_NUM}', fontweight='bold')
    ax.set_xticks(x, names, rotation=25, ha='right', fontsize=9)
    ax.legend(loc='upper right', fontsize=8)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    fn = f"{output_dir}/stacked_bar.png"
    plt.savefig(fn, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return fn

def create_radar_chart(output_dir="charts/march_w3"):
    os.makedirs(output_dir, exist_ok=True)
    max_sales = max(r[1]['total_sales'] for r in sorted_reps)
    max_profit = max(r[1]['total_profit'] for r in sorted_reps)
    max_qty = max(r[1]['total_qty'] for r in sorted_reps)
    max_margin = max(r[1]['overall_margin'] for r in sorted_reps)
    categories = ['Sales', 'Profit', 'Items Sold', 'Margin']
    N = len(categories)
    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    angles += angles[:1]
    fig, ax = plt.subplots(figsize=(10, 8), subplot_kw=dict(polar=True))
    ax.set_facecolor('#f8f9fa')
    plt.rgrids([0.2, 0.4, 0.6, 0.8, 1.0], labels=['20%', '40%', '60%', '80%', '100%'], fontsize=8, color='#666666')
    for i, (rep_name, rep_data) in enumerate(sorted_reps):
        values = [
            rep_data['total_sales'] / max_sales,
            rep_data['total_profit'] / max_profit,
            rep_data['total_qty'] / max_qty,
            rep_data['overall_margin'] / max_margin
        ]
        values += values[:1]
        color = REP_COLORS[i % len(REP_COLORS)]
        ax.plot(angles, values, 'o-', linewidth=2.5, label=rep_name, color=color, markersize=7)
        ax.fill(angles, values, alpha=0.1, color=color)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, fontsize=11, fontweight='bold')
    ax.set_ylim(0, 1.1)
    ax.set_title(f'Performance Radar - {MONTH} Week {WEEK_NUM}', fontsize=14, fontweight='bold', pad=25)
    ax.legend(loc='upper right', bbox_to_anchor=(1.35, 1.0), fontsize=9, frameon=True, fancybox=True)
    plt.tight_layout()
    fn = f"{output_dir}/radar.png"
    plt.savefig(fn, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return fn

# ========== PERFORMANCE ANALYSIS ==========
def analyze_performance(reps_data, cat_totals):
    """Generate performance badges, notes, and rankings for each rep."""
    sorted_by_sales = sorted(reps_data.items(), key=lambda x: x[1]['total_sales'], reverse=True)
    sorted_by_margin = sorted(reps_data.items(), key=lambda x: x[1]['overall_margin'], reverse=True)
    sorted_by_profit = sorted(reps_data.items(), key=lambda x: x[1]['total_profit'], reverse=True)
    
    avg_sales = sum(r['total_sales'] for r in reps_data.values()) / len(reps_data) if reps_data else 0
    avg_margin = sum(r['overall_margin'] for r in reps_data.values()) / len(reps_data) if reps_data else 0
    
    analysis = {}
    
    for idx, (rep_name, rep_data) in enumerate(sorted_by_sales):
        role = get_rep_role(rep_name)
        primary_dept = get_primary_dept(rep_name)
        
        sales_rank = idx + 1
        margin_rank = [r[0] for r in sorted_by_margin].index(rep_name) + 1
        profit_rank = [r[0] for r in sorted_by_profit].index(rep_name) + 1
        
        # Badges
        badges = []
        if sales_rank == 1:
            badges.append("TOP SELLER")
        if margin_rank == 1:
            badges.append("BEST MARGIN")
        if profit_rank == 1:
            badges.append("TOP PROFIT")
        
        # Performance notes
        notes = []
        
        # Sales vs team average
        sales_vs_avg = ((rep_data['total_sales'] - avg_sales) / avg_sales) * 100 if avg_sales > 0 else 0
        if sales_vs_avg > 50:
            notes.append(f"Outstanding! Sales {sales_vs_avg:.0f}% above team average")
        elif sales_vs_avg > 0:
            notes.append(f"Sales {sales_vs_avg:.0f}% above team average - solid performance")
        else:
            notes.append(f"Sales {abs(sales_vs_avg):.0f}% below team average - needs improvement")
        
        # Margin analysis
        if rep_data['overall_margin'] > avg_margin + 3:
            notes.append(f"Excellent margin at {rep_data['overall_margin']:.1f}% (Team avg: {avg_margin:.1f}%)")
        elif rep_data['overall_margin'] < avg_margin - 2:
            notes.append(f"Margin at {rep_data['overall_margin']:.1f}% below team average of {avg_margin:.1f}%")
        
        # Department head specific analysis
        if primary_dept != 'GENERAL HARDWARE' and primary_dept in rep_data['categories']:
            dept_sales = rep_data['categories'][primary_dept]['sales_incl']
            dept_total = cat_totals[primary_dept]['sales']
            dept_share = (dept_sales / dept_total * 100) if dept_total > 0 else 0
            dept_margin = rep_data['categories'][primary_dept]['margin_pct']
            
            if primary_dept == 'PLUMBING':
                notes.append(f"Plumbing Dept Head: {dept_share:.1f}% of dept sales, {dept_margin:.1f}% margin")
                if dept_share >= 40:
                    notes.append(f"Strong plumbing leadership with {dept_share:.1f}% of dept sales")
                else:
                    notes.append(f"Should lead with higher dept contribution (currently {dept_share:.1f}%)")
            elif primary_dept == 'PAINTS':
                notes.append(f"Paints Dept Head: {dept_share:.1f}% of dept sales, {dept_margin:.1f}% margin")
                if dept_share >= 30:
                    notes.append(f"Good paints leadership with {dept_share:.1f}% of dept sales")
                else:
                    notes.append(f"Should lead with higher dept contribution (currently {dept_share:.1f}%)")
            elif primary_dept == 'ELECTRICALS':
                notes.append(f"Electricals Dept Head: {dept_share:.1f}% of dept sales, {dept_margin:.1f}% margin")
                if dept_share >= 50:
                    notes.append(f"Strong leadership - controls {dept_share:.1f}% of department sales")
                else:
                    notes.append(f"Should lead with higher dept contribution (currently {dept_share:.1f}%)")
        
        # General hardware team
        if primary_dept == 'GENERAL HARDWARE' and 'GENERAL HARDWARE' in rep_data['categories']:
            gh_sales = rep_data['categories']['GENERAL HARDWARE']['sales_incl']
            gh_total = cat_totals['GENERAL HARDWARE']['sales']
            gh_share = (gh_sales / gh_total * 100) if gh_total > 0 else 0
            gh_margin = rep_data['categories']['GENERAL HARDWARE']['margin_pct']
            notes.append(f"General Hardware: {gh_share:.1f}% of dept, {gh_margin:.1f}% margin")
        
        # Cross-selling analysis
        cats_sold = len(rep_data['categories'])
        if cats_sold == 4:
            notes.append("Selling across all 4 categories - excellent product range")
        elif cats_sold == 3:
            missing = [c for c in CAT_ORDER if c not in rep_data['categories']]
            notes.append(f"Consider expanding into {', '.join(missing)}")
        elif cats_sold <= 2:
            missing = [c for c in CAT_ORDER if c not in rep_data['categories']]
            notes.append(f"Limited to {cats_sold} categories - expand into {', '.join(missing)}")
        
        # Best margin category
        if rep_data['categories']:
            best_margin_cat = max(rep_data['categories'].items(), key=lambda x: x[1].get('margin_pct', 0))
            notes.append(f"Best margin in {best_margin_cat[0]}: {best_margin_cat[1]['margin_pct']:.1f}%")
        
        # Contribution to total
        contrib = (rep_data['total_sales'] / total_sales * 100) if total_sales > 0 else 0
        if contrib >= 30:
            notes.append(f"Key contributor at {contrib:.1f}% of total branch sales")
        elif contrib < 5:
            notes.append(f"Low contribution at {contrib:.1f}% of total branch sales - growth opportunity")
        
        analysis[rep_name] = {
            'role': role,
            'primary_department': primary_dept,
            'badges': badges,
            'notes': notes,
            'sales_rank': sales_rank,
            'margin_rank': margin_rank,
            'profit_rank': profit_rank,
            'sales_vs_avg': sales_vs_avg,
        }
    
    return analysis

# ========== PER-REP CHART FUNCTIONS ==========
def create_rep_donut(rep_name, rep_data, output_dir="charts/march_w3"):
    """Donut chart showing sales distribution by category for a single rep."""
    os.makedirs(output_dir, exist_ok=True)
    categories = []
    sales = []
    chart_colors = []
    for cat in CAT_ORDER:
        if cat in rep_data['categories'] and rep_data['categories'][cat]['sales_incl'] > 0:
            categories.append(cat)
            sales.append(rep_data['categories'][cat]['sales_incl'])
            chart_colors.append(COLORS[cat])
    if not sales or sum(sales) == 0:
        return None
    fig, ax = plt.subplots(figsize=(8, 5))
    pie_result = ax.pie(sales, labels=None, autopct='%1.1f%%',
                        colors=chart_colors, startangle=90,
                        pctdistance=0.75, explode=[0.02]*len(sales),
                        wedgeprops=dict(width=0.5, edgecolor='white'))
    wedges = pie_result[0]
    autotexts = pie_result[2] if len(pie_result) > 2 else []
    ax.text(0, 0, f'KES\n{sum(sales)/1000000:.2f}M', ha='center', va='center',
            fontsize=12, fontweight='bold', color='#2c3e50')
    for at in autotexts:
        at.set_fontsize(9)
        at.set_fontweight('bold')
        at.set_color('white')
    legend_labels = [f"{c}: KES {s:,.0f}" for c, s in zip(categories, sales)]
    ax.legend(wedges, legend_labels, title="Categories", loc="center left",
              bbox_to_anchor=(1, 0, 0.5, 1), fontsize=8)
    ax.set_title(f"{rep_name} - Sales Distribution", fontsize=12, fontweight='bold', pad=10)
    plt.tight_layout()
    fn = f"{output_dir}/{rep_name.replace(' ', '_')}_donut.png"
    plt.savefig(fn, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return fn

def create_rep_sales_profit_bar(rep_name, rep_data, output_dir="charts/march_w3"):
    """Grouped bar chart showing sales vs profit by category for a single rep."""
    os.makedirs(output_dir, exist_ok=True)
    categories = []
    sales = []
    profits = []
    for cat in CAT_ORDER:
        if cat in rep_data['categories'] and rep_data['categories'][cat]['sales_incl'] > 0:
            short_cat = cat[:12] if len(cat) > 12 else cat
            categories.append(short_cat)
            sales.append(rep_data['categories'][cat]['sales_incl'] / 1000)
            profits.append(rep_data['categories'][cat]['profit'] / 1000)
    if not categories:
        return None
    fig, ax = plt.subplots(figsize=(9, 4))
    x = np.arange(len(categories))
    width = 0.35
    bars1 = ax.bar(x - width/2, sales, width, label='Sales', color='#0494f5', edgecolor='white')
    bars2 = ax.bar(x + width/2, profits, width, label='Profit', color='#08f169', edgecolor='white')
    for bar in bars1:
        h = bar.get_height()
        ax.annotate(f'{h:.0f}K', xy=(bar.get_x() + bar.get_width()/2, h),
                    xytext=(0, 3), textcoords="offset points", ha='center', va='bottom', fontsize=8)
    for bar in bars2:
        h = bar.get_height()
        ax.annotate(f'{h:.0f}K', xy=(bar.get_x() + bar.get_width()/2, h),
                    xytext=(0, 3), textcoords="offset points", ha='center', va='bottom', fontsize=8)
    ax.set_ylabel('Amount (Thousands KES)', fontsize=10, fontweight='bold')
    ax.set_title(f"{rep_name} - Sales vs Profit by Category", fontsize=12, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(categories, fontsize=9)
    ax.legend()
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    fn = f"{output_dir}/{rep_name.replace(' ', '_')}_sales_profit.png"
    plt.savefig(fn, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close()
    return fn

# ========== GENERATE CHARTS ==========
print("Generating charts...")
chart_files = {
    'trend': create_weekly_trend_chart(),
    'category_pie': create_category_pie(),
    'rep_bar': create_rep_bar_chart(),
    'margin': create_margin_comparison(),
    'target': create_target_progress(),
    'stacked': create_stacked_bar(),
    'radar': create_radar_chart(),
}
print(f"  {len(chart_files)} summary charts generated in charts/march_w3/")

# Per-rep charts
rep_charts = {}
for rep_name, rep_data in sorted_reps:
    donut = create_rep_donut(rep_name, rep_data)
    sales_profit = create_rep_sales_profit_bar(rep_name, rep_data)
    rep_charts[rep_name] = {"donut": donut, "sales_profit": sales_profit}
    print(f"  Created charts for {rep_name}")

# Generate performance analysis
perf_analysis = analyze_performance(reps_data, cat_totals)
print(f"  Performance analysis generated for {len(perf_analysis)} reps")

# ========== PDF REPORT ==========
def generate_pdf():
    output_fn = f"sales_analysis_march_w3.pdf"
    doc = SimpleDocTemplate(output_fn, pagesize=A4,
                             topMargin=0.5*inch, bottomMargin=0.5*inch,
                             leftMargin=0.5*inch, rightMargin=0.5*inch)
    elements = []
    styles = getSampleStyleSheet()
    
    title_s = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=20,
                              textColor=colors.HexColor('#1a1a1a'), spaceAfter=8,
                              alignment=TA_CENTER, fontName='Helvetica-Bold')
    subtitle_s = ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=11,
                                 textColor=colors.HexColor('#666666'), spaceAfter=20,
                                 alignment=TA_CENTER)
    heading_s = ParagraphStyle('Heading', parent=styles['Heading2'], fontSize=14,
                                textColor=colors.HexColor('#2c3e50'), spaceAfter=12,
                                spaceBefore=16, fontName='Helvetica-Bold')
    subheading_s = ParagraphStyle('SubHeading', parent=styles['Heading3'], fontSize=11,
                                   textColor=colors.HexColor('#34495e'), spaceAfter=8,
                                   spaceBefore=12, fontName='Helvetica-Bold')
    body_s = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9,
                             textColor=colors.HexColor('#2c3e50'), spaceAfter=6)
    highlight_s = ParagraphStyle('Highlight', parent=styles['Normal'], fontSize=10,
                                  textColor=colors.HexColor('#27ae60'), spaceAfter=6,
                                  fontName='Helvetica-Bold')
    alert_s = ParagraphStyle('Alert', parent=styles['Normal'], fontSize=10,
                              textColor=colors.HexColor('#e74c3c'), spaceAfter=6,
                              fontName='Helvetica-Bold')
    
    # Title page
    elements.append(Spacer(1, 0.5*inch))
    elements.append(Paragraph("SALES PERFORMANCE ANALYSIS", title_s))
    elements.append(Paragraph(f"BOMAS Hardware Store - {MONTH} {YEAR}, Week {WEEK_NUM}", subtitle_s))
    elements.append(Paragraph(f"Period: {PERIOD} ({WORKING_DAYS} working days)", subtitle_s))
    elements.append(Paragraph(f"Report Generated: {datetime.now().strftime('%B %d, %Y')}", subtitle_s))
    elements.append(Spacer(1, 0.3*inch))
    
    # Executive Summary
    elements.append(Paragraph("EXECUTIVE SUMMARY", heading_s))
    daily_avg = total_sales / WORKING_DAYS
    daily_profit = total_profit / WORKING_DAYS
    
    summary_data = [
        ["Metric", "Value"],
        ["Total Sales (Incl. VAT)", f"KES {total_sales:,.0f}"],
        ["Total Profit (Incl. VAT)", f"KES {total_profit:,.0f}"],
        ["Overall Margin", f"{overall_margin:.1f}%"],
        ["Daily Average Sales", f"KES {daily_avg:,.0f}"],
        ["Daily Average Profit", f"KES {daily_profit:,.0f}"],
        ["Working Days", str(WORKING_DAYS)],
        ["Active Sales Reps", str(len(reps_data))],
    ]
    
    summary_table = Table(summary_data, colWidths=[2.5*inch, 3*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#ecf0f1')]),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
    ]))
    elements.append(summary_table)
    elements.append(Spacer(1, 15))
    
    # Target Progress
    elements.append(Paragraph("MONTHLY TARGET PROGRESS", subheading_s))
    target_pct = (cumulative_total_sales / MONTHLY_TARGETS['OVERALL'] * 100)
    expected_pct = 75.0  # Week 3 of 4
    pace = "AHEAD" if target_pct >= expected_pct else "BEHIND"
    pace_color = highlight_s if target_pct >= expected_pct else alert_s
    elements.append(Paragraph(f"Monthly Target: KES {MONTHLY_TARGETS['OVERALL']:,.0f}", body_s))
    elements.append(Paragraph(f"Accumulated Achievement (W1-3): KES {cumulative_total_sales:,.0f} ({target_pct:.1f}% of target)", body_s))
    elements.append(Paragraph(f"Pace: {pace} schedule (expected 75.0% after Week 3)", pace_color))
    
    target_data = [["Category", "Monthly Target", "Week 3 Actual", "% Achieved", "Expected", "Pace"]]
    for cat in CAT_ORDER:
        actual = cumulative_cat_totals[cat]
        target = MONTHLY_TARGETS.get(cat, 0)
        pct = (actual / target * 100) if target > 0 else 0
        p = "Ahead" if pct >= 75 else "Behind"
        target_data.append([cat, f"KES {target:,.0f}", f"KES {actual:,.0f}", f"{pct:.1f}%", "75.0%", p])
    target_data.append(["OVERALL", f"KES {MONTHLY_TARGETS['OVERALL']:,.0f}", f"KES {cumulative_total_sales:,.0f}",
                         f"{target_pct:.1f}%", "75.0%", pace])
    
    t = Table(target_data, colWidths=[1.3*inch, 1.2*inch, 1.2*inch, 0.8*inch, 0.8*inch, 0.7*inch])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('ALIGN', (1, 1), (-1, -1), 'CENTER'),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -2), [colors.white, colors.HexColor('#ecf0f1')]),
        ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#34495e')),
        ('TEXTCOLOR', (0, -1), (-1, -1), colors.white),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
    ]))
    elements.append(Spacer(1, 10))
    elements.append(t)
    
    # Charts
    elements.append(PageBreak())
    elements.append(Paragraph("CATEGORY BREAKDOWN", heading_s))
    if os.path.exists(chart_files['category_pie']):
        elements.append(Image(chart_files['category_pie'], width=5.5*inch, height=3.4*inch))
    elements.append(Spacer(1, 10))
    if os.path.exists(chart_files['target']):
        elements.append(Paragraph("TARGET PROGRESS", heading_s))
        elements.append(Image(chart_files['target'], width=6.5*inch, height=2*inch))
    
    elements.append(PageBreak())
    elements.append(Paragraph("SALES REP PERFORMANCE", heading_s))
    
    rep_data_rows = [["Rank", "Sales Rep", "Role", "Sales (KES)", "Profit (KES)", "Margin %", "Items"]]
    for rank, (rep_name, rep_data) in enumerate(sorted_reps, 1):
        rep_data_rows.append([
            str(rank), rep_name, get_rep_role(rep_name)[:18],
            f"{rep_data['total_sales']:,.0f}", f"{rep_data['total_profit']:,.0f}",
            f"{rep_data['overall_margin']:.1f}%", f"{rep_data['total_qty']:,.0f}"
        ])
    
    t = Table(rep_data_rows, colWidths=[0.4*inch, 1.2*inch, 1.3*inch, 1.1*inch, 1.0*inch, 0.7*inch, 0.7*inch])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('ALIGN', (0, 0), (0, -1), 'CENTER'),
        ('ALIGN', (3, 1), (-1, -1), 'RIGHT'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#ecf0f1')]),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 10))
    
    if os.path.exists(chart_files['trend']):
        elements.append(Paragraph("SALES TREND (W1 vs W2 vs W3)", heading_s))
        elements.append(Image(chart_files['trend'], width=7*inch, height=3*inch))
        elements.append(Spacer(1, 15))
        
    if os.path.exists(chart_files['rep_bar']):
        elements.append(Paragraph("SALES BY REPRESENTATIVE (WEEK 3)", heading_s))
        elements.append(Image(chart_files['rep_bar'], width=6*inch, height=3*inch))
        elements.append(Spacer(1, 15))
    
    elements.append(PageBreak())
    elements.append(Paragraph("MARGIN ANALYSIS", heading_s))
    if os.path.exists(chart_files['margin']):
        elements.append(Image(chart_files['margin'], width=6*inch, height=2.8*inch))
    elements.append(Spacer(1, 15))
    if os.path.exists(chart_files['stacked']):
        elements.append(Paragraph("CATEGORY BREAKDOWN PER REP", heading_s))
        elements.append(Image(chart_files['stacked'], width=6*inch, height=3*inch))
    
    elements.append(PageBreak())
    elements.append(Paragraph("PERFORMANCE RADAR", heading_s))
    if os.path.exists(chart_files['radar']):
        elements.append(Image(chart_files['radar'], width=5.5*inch, height=4.5*inch))
    
    # Category Detail Pages
    elements.append(PageBreak())
    elements.append(Paragraph("DEPARTMENT DETAILS", heading_s))
    
    for cat in CAT_ORDER:
        elements.append(Paragraph(cat, subheading_s))
        dept_head = DEPARTMENT_HEADS.get(cat, 'N/A')
        if isinstance(dept_head, list):
            dept_head = ', '.join(dept_head)
        elements.append(Paragraph(f"Department Head(s): {dept_head}", body_s))
        elements.append(Paragraph(f"Total Sales: KES {cat_totals[cat]['sales']:,.0f} | Profit: KES {cat_totals[cat]['profit']:,.0f}", body_s))
        
        cat_rows = [["Rep", "Sales (KES)", "Profit (KES)", "Margin %", "Qty", "% of Dept"]]
        reps_in_cat = []
        for rep_name, rep_d in sorted_reps:
            if cat in rep_d['categories']:
                c = rep_d['categories'][cat]
                share = (c['sales_incl'] / cat_totals[cat]['sales'] * 100) if cat_totals[cat]['sales'] > 0 else 0
                reps_in_cat.append((rep_name, c, share))
        reps_in_cat.sort(key=lambda x: x[1]['sales_incl'], reverse=True)
        
        for rep_name, c, share in reps_in_cat:
            cat_rows.append([rep_name, f"{c['sales_incl']:,.0f}", f"{c['profit']:,.0f}",
                              f"{c['margin_pct']:.1f}%", f"{c['qty']:,.0f}", f"{share:.1f}%"])
        
        t = Table(cat_rows, colWidths=[1.3*inch, 1.1*inch, 1.0*inch, 0.8*inch, 0.7*inch, 0.8*inch])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(COLORS.get(cat, '#2c3e50'))),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 15))
    
    # ===== PERFORMANCE COMMENTARY & AWARDS =====
    elements.append(PageBreak())
    elements.append(Paragraph("PERFORMANCE COMMENTARY &amp; AWARDS", heading_s))
    elements.append(Spacer(1, 0.1*inch))
    
    badge_s = ParagraphStyle('Badge', parent=styles['Normal'], fontSize=11,
                              textColor=colors.HexColor('#8e44ad'), spaceAfter=4, fontName='Helvetica-Bold')
    role_s = ParagraphStyle('Role', parent=styles['Normal'], fontSize=10,
                             textColor=colors.HexColor('#34495e'), spaceAfter=2, fontName='Helvetica-Oblique')
    note_s = ParagraphStyle('Note', parent=styles['Normal'], fontSize=8,
                             textColor=colors.HexColor('#555555'), spaceAfter=2, leftIndent=15)
    rep_title_s = ParagraphStyle('RepTitle', parent=styles['Normal'], fontSize=11,
                                  textColor=colors.HexColor('#2c3e50'), spaceAfter=4, fontName='Helvetica-Bold')
    
    # Department Leadership Table
    elements.append(Paragraph("DEPARTMENT LEADERSHIP", subheading_s))
    dept_heads_data = [["Department", "Head", "Dept Sales", "Dept Margin", "Dept Share"]]
    for dept in ['PLUMBING', 'PAINTS', 'ELECTRICALS']:
        head = DEPARTMENT_HEADS.get(dept)
        if head and head in reps_data:
            rd = reps_data[head]
            if dept in rd['categories']:
                dept_sales = rd['categories'][dept]['sales_incl']
                dept_margin = rd['categories'][dept]['margin_pct']
                dept_total = cat_totals[dept]['sales']
                dept_share = (dept_sales / dept_total * 100) if dept_total > 0 else 0
                dept_heads_data.append([dept, head, f"KES {dept_sales:,.0f}",
                                        f"{dept_margin:.1f}%", f"{dept_share:.1f}%"])
    if len(dept_heads_data) > 1:
        dt = Table(dept_heads_data, colWidths=[1.5*inch, 1.5*inch, 1.5*inch, 1.2*inch, 1.2*inch])
        dt.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8e44ad')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f0fa')]),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        elements.append(dt)
    elements.append(Spacer(1, 0.15*inch))
    
    # Rep-by-Rep Commentary
    elements.append(Paragraph("REP-BY-REP COMMENTARY", subheading_s))
    elements.append(Spacer(1, 0.1*inch))
    
    for rep_name, rep_data_item in sorted_reps:
        pa = perf_analysis.get(rep_name, {})
        role = pa.get('role', 'Sales Rep')
        badges_str = '  '.join(pa.get('badges', []))
        elements.append(Paragraph(f"{rep_name.upper()} - {role}", rep_title_s))
        if badges_str:
            elements.append(Paragraph(badges_str, badge_s))
        for note in pa.get('notes', [])[:4]:
            elements.append(Paragraph(f"&bull; {note}", note_s))
        elements.append(Spacer(1, 0.08*inch))
    
    # ===== INDIVIDUAL REP PAGES =====
    for rep_name, rep_data_item in sorted_reps:
        elements.append(PageBreak())
        pa = perf_analysis.get(rep_name, {})
        role = pa.get('role', 'Sales Rep')
        
        elements.append(Paragraph(f"SALES REP: {rep_name.upper()}", heading_s))
        elements.append(Paragraph(role, role_s))
        
        badges = pa.get('badges', [])
        if badges:
            elements.append(Paragraph('  '.join(badges), badge_s))
        
        # Stats table
        rep_stats = [
            ["Total Sales", "Total Profit", "Margin", "Items Sold", "Rank"],
            [f"KES {rep_data_item['total_sales']:,.0f}", f"KES {rep_data_item['total_profit']:,.0f}",
             f"{rep_data_item['overall_margin']:.1f}%", f"{rep_data_item['total_qty']:,.0f}",
             f"#{pa.get('sales_rank', '-')}"]
        ]
        st = Table(rep_stats, colWidths=[1.4*inch, 1.4*inch, 1*inch, 1*inch, 0.8*inch])
        st.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34495e')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ('TOPPADDING', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        elements.append(st)
        elements.append(Spacer(1, 0.1*inch))
        
        # Donut chart
        rc = rep_charts.get(rep_name, {})
        if rc.get('donut') and os.path.exists(rc['donut']):
            elements.append(Image(rc['donut'], width=5*inch, height=3*inch))
        elements.append(Spacer(1, 0.1*inch))
        
        # Sales vs Profit bar
        if rc.get('sales_profit') and os.path.exists(rc['sales_profit']):
            elements.append(Image(rc['sales_profit'], width=5.5*inch, height=2.5*inch))
        elements.append(Spacer(1, 0.1*inch))
        
        # Category breakdown table
        cat_detail = [["Category", "Sales", "Profit", "Margin", "Qty"]]
        for cat in CAT_ORDER:
            if cat in rep_data_item['categories']:
                c = rep_data_item['categories'][cat]
                cat_detail.append([cat, f"KES {c['sales_incl']:,.0f}", f"KES {c['profit']:,.0f}",
                                    f"{c['margin_pct']:.1f}%", f"{c['qty']:,.0f}"])
        if len(cat_detail) > 1:
            ct = Table(cat_detail, colWidths=[1.8*inch, 1.4*inch, 1.2*inch, 1*inch, 1*inch])
            ct.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#ecf0f1')]),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            elements.append(ct)
        
        # Performance notes
        elements.append(Spacer(1, 0.1*inch))
        notes_head = ParagraphStyle('NotesHead', parent=styles['Normal'],
                                     fontSize=10, textColor=colors.HexColor('#8e44ad'),
                                     fontName='Helvetica-Bold', spaceAfter=4)
        elements.append(Paragraph("PERFORMANCE NOTES", notes_head))
        
        for note in pa.get('notes', [])[:5]:
            note_color = '#27ae60' if any(w in note for w in ['Strong', 'Excellent', 'Outstanding', 'Good', 'Selling across']) else (
                         '#e74c3c' if any(w in note for w in ['below', 'needs improvement', 'Should lead', 'Low contribution', 'Limited to']) else '#555555')
            note_item_s = ParagraphStyle('NoteItem', parent=styles['Normal'],
                                          fontSize=8, textColor=colors.HexColor(note_color),
                                          spaceAfter=2, leftIndent=10)
            elements.append(Paragraph(f"&bull; {note}", note_item_s))
    
    doc.build(elements)
    print(f"  PDF: {output_fn}")
    return output_fn

# ========== EXCEL REPORT ==========
def generate_excel():
    output_fn = f"sales_analysis_march_w3.xlsx"
    wb = Workbook()
    
    header_font = Font(name='Calibri', bold=True, size=11, color='FFFFFF')
    header_fill = PatternFill(start_color='2C3E50', end_color='2C3E50', fill_type='solid')
    cat_fills = {
        'GENERAL HARDWARE': PatternFill(start_color='3498DB', end_color='3498DB', fill_type='solid'),
        'PAINTS': PatternFill(start_color='E74C3C', end_color='E74C3C', fill_type='solid'),
        'PLUMBING': PatternFill(start_color='2ECC71', end_color='2ECC71', fill_type='solid'),
        'ELECTRICALS': PatternFill(start_color='F39C12', end_color='F39C12', fill_type='solid'),
    }
    num_fmt = '#,##0'
    pct_fmt = '0.0%'
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    
    # Summary sheet
    ws = wb.active
    ws.title = "Summary"
    ws.append([f"SALES ANALYSIS - {MONTH.upper()} WEEK {WEEK_NUM}, {YEAR}"])
    ws.merge_cells('A1:G1')
    ws['A1'].font = Font(name='Calibri', bold=True, size=16, color='2C3E50')
    ws['A1'].alignment = Alignment(horizontal='center')
    ws.append([f"Period: {PERIOD} ({WORKING_DAYS} working days)"])
    ws.merge_cells('A2:G2')
    ws['A2'].font = Font(name='Calibri', size=11, color='666666')
    ws['A2'].alignment = Alignment(horizontal='center')
    ws.append([])
    
    # KPIs
    ws.append(["KEY METRICS"])
    ws['A4'].font = Font(name='Calibri', bold=True, size=12, color='2C3E50')
    kpis = [
        ("Total Sales", total_sales),
        ("Total Profit", total_profit),
        ("Overall Margin", overall_margin / 100),
        ("Daily Average Sales", total_sales / WORKING_DAYS),
        ("Daily Average Profit", total_profit / WORKING_DAYS),
    ]
    for label, val in kpis:
        if 'Margin' in label:
            ws.append([label, val])
            ws.cell(row=ws.max_row, column=2).number_format = pct_fmt
        else:
            ws.append([label, val])
            ws.cell(row=ws.max_row, column=2).number_format = num_fmt
        ws.cell(row=ws.max_row, column=1).font = Font(bold=True)
    
    ws.append([])
    ws.append(["MONTHLY TARGET PROGRESS"])
    ws.cell(row=ws.max_row, column=1).font = Font(name='Calibri', bold=True, size=12, color='2C3E50')
    headers = ["Category", "Monthly Target", "Week 3 Actual", "% Achieved", "Expected %", "Pace"]
    ws.append(headers)
    for i, h in enumerate(headers, 1):
        cell = ws.cell(row=ws.max_row, column=i)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
        cell.border = thin_border
    
    for cat in CAT_ORDER + ['OVERALL']:
        actual = cumulative_cat_totals[cat] if cat != 'OVERALL' else total_sales
        target = MONTHLY_TARGETS.get(cat, MONTHLY_TARGETS['OVERALL'])
        pct = actual / target if target > 0 else 0
        pace = "Ahead" if pct >= 0.75 else "Behind"
        ws.append([cat, target, actual, pct, 0.75, pace])
        row = ws.max_row
        ws.cell(row=row, column=2).number_format = num_fmt
        ws.cell(row=row, column=3).number_format = num_fmt
        ws.cell(row=row, column=4).number_format = pct_fmt
        ws.cell(row=row, column=5).number_format = pct_fmt
        for c in range(1, 7):
            ws.cell(row=row, column=c).border = thin_border
    
    ws.column_dimensions['A'].width = 22
    ws.column_dimensions['B'].width = 18
    ws.column_dimensions['C'].width = 18
    ws.column_dimensions['D'].width = 14
    ws.column_dimensions['E'].width = 14
    ws.column_dimensions['F'].width = 10
    
    # Rep Performance sheet
    ws2 = wb.create_sheet("Rep Performance")
    ws2.append([f"SALES REP PERFORMANCE - {MONTH.upper()} WEEK {WEEK_NUM}"])
    ws2.merge_cells('A1:H1')
    ws2['A1'].font = Font(name='Calibri', bold=True, size=14, color='2C3E50')
    ws2.append([])
    
    headers2 = ["Rank", "Sales Rep", "Role", "Total Sales", "Total Profit", "Margin %", "Total Qty", "Contribution %"]
    ws2.append(headers2)
    for i, h in enumerate(headers2, 1):
        cell = ws2.cell(row=3, column=i)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
        cell.border = thin_border
    
    for rank, (rep_name, rep_data) in enumerate(sorted_reps, 1):
        contrib = (rep_data['total_sales'] / total_sales * 100) if total_sales > 0 else 0
        ws2.append([rank, rep_name, get_rep_role(rep_name), rep_data['total_sales'],
                     rep_data['total_profit'], rep_data['overall_margin'] / 100,
                     rep_data['total_qty'], contrib / 100])
        row = ws2.max_row
        ws2.cell(row=row, column=4).number_format = num_fmt
        ws2.cell(row=row, column=5).number_format = num_fmt
        ws2.cell(row=row, column=6).number_format = pct_fmt
        ws2.cell(row=row, column=7).number_format = '#,##0'
        ws2.cell(row=row, column=8).number_format = pct_fmt
        for c in range(1, 9):
            ws2.cell(row=row, column=c).border = thin_border
    
    for col_letter, width in [('A', 6), ('B', 20), ('C', 22), ('D', 16), ('E', 16), ('F', 10), ('G', 10), ('H', 14)]:
        ws2.column_dimensions[col_letter].width = width
    
    # Add bar chart to Rep Performance sheet
    chart = BarChart()
    chart.type = "col"
    chart.title = "Sales vs Profit by Rep"
    chart.y_axis.title = "Amount (KES)"
    chart.x_axis.title = "Sales Rep"
    chart.style = 10
    
    data_ref = Reference(ws2, min_col=4, max_col=5, min_row=3, max_row=3 + len(sorted_reps))
    cats_ref = Reference(ws2, min_col=2, min_row=4, max_row=3 + len(sorted_reps))
    chart.add_data(data_ref, titles_from_data=True)
    chart.set_categories(cats_ref)
    chart.shape = 4
    chart.width = 20
    chart.height = 12
    ws2.add_chart(chart, f"A{5 + len(sorted_reps)}")
    
    # Category Detail sheets
    for cat in CAT_ORDER:
        ws_cat = wb.create_sheet(cat[:15])
        ws_cat.append([f"{cat} - {MONTH.upper()} WEEK {WEEK_NUM}"])
        ws_cat.merge_cells('A1:G1')
        ws_cat['A1'].font = Font(name='Calibri', bold=True, size=14, color='FFFFFF')
        ws_cat['A1'].fill = cat_fills.get(cat, header_fill)
        
        dept_head = DEPARTMENT_HEADS.get(cat, 'N/A')
        if isinstance(dept_head, list):
            dept_head = ', '.join(dept_head)
        ws_cat.append([f"Department Head(s): {dept_head}"])
        ws_cat.append([f"Total Sales: KES {cat_totals[cat]['sales']:,.0f}  |  Total Profit: KES {cat_totals[cat]['profit']:,.0f}"])
        ws_cat.append([])
        
        headers_cat = ["Rep", "Sales (KES)", "Profit (KES)", "Margin %", "Qty", "Share of Dept %"]
        ws_cat.append(headers_cat)
        for i, h in enumerate(headers_cat, 1):
            cell = ws_cat.cell(row=5, column=i)
            cell.font = header_font
            cell.fill = cat_fills.get(cat, header_fill)
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border
        
        reps_in_cat = []
        for rep_name, rep_d in sorted_reps:
            if cat in rep_d['categories']:
                c = rep_d['categories'][cat]
                share = (c['sales_incl'] / cat_totals[cat]['sales'] * 100) if cat_totals[cat]['sales'] > 0 else 0
                reps_in_cat.append((rep_name, c, share))
        reps_in_cat.sort(key=lambda x: x[1]['sales_incl'], reverse=True)
        
        for rep_name, c, share in reps_in_cat:
            ws_cat.append([rep_name, c['sales_incl'], c['profit'], c['margin_pct'] / 100, c['qty'], share / 100])
            row = ws_cat.max_row
            ws_cat.cell(row=row, column=2).number_format = num_fmt
            ws_cat.cell(row=row, column=3).number_format = num_fmt
            ws_cat.cell(row=row, column=4).number_format = pct_fmt
            ws_cat.cell(row=row, column=5).number_format = '#,##0'
            ws_cat.cell(row=row, column=6).number_format = pct_fmt
            for c_idx in range(1, 7):
                ws_cat.cell(row=row, column=c_idx).border = thin_border
        
        for col_letter, width in [('A', 20), ('B', 16), ('C', 16), ('D', 10), ('E', 10), ('F', 14)]:
            ws_cat.column_dimensions[col_letter].width = width
        
        # Pie chart
        pie = PieChart()
        pie.title = f"{cat} - Sales Share by Rep"
        pie.style = 10
        data_ref = Reference(ws_cat, min_col=2, min_row=5, max_row=5 + len(reps_in_cat))
        cats_ref = Reference(ws_cat, min_col=1, min_row=6, max_row=5 + len(reps_in_cat))
        pie.add_data(data_ref, titles_from_data=True)
        pie.set_categories(cats_ref)
        pie.width = 16
        pie.height = 10
        ws_cat.add_chart(pie, f"A{7 + len(reps_in_cat)}")
    
    # ===== Individual Rep Detail Sheets =====
    badge_fill = PatternFill(start_color='8E44AD', end_color='8E44AD', fill_type='solid')
    note_font_green = Font(name='Calibri', size=10, color='27AE60')
    note_font_red = Font(name='Calibri', size=10, color='E74C3C')
    note_font_grey = Font(name='Calibri', size=10, color='555555')
    
    for rep_name, rep_data_item in sorted_reps:
        pa = perf_analysis.get(rep_name, {})
        safe_name = rep_name[:15].replace(' ', '_')
        ws_rep = wb.create_sheet(safe_name)
        
        # Header
        ws_rep.append([f"{rep_name} - {pa.get('role', 'Sales Rep')}"])
        ws_rep.merge_cells('A1:F1')
        ws_rep['A1'].font = Font(name='Calibri', bold=True, size=14, color='2C3E50')
        
        # Badges
        badges = pa.get('badges', [])
        if badges:
            ws_rep.append(['  '.join(badges)])
            ws_rep.merge_cells('A2:F2')
            ws_rep['A2'].font = Font(name='Calibri', bold=True, size=11, color='8E44AD')
            ws_rep.append([])
        else:
            ws_rep.append([])
        
        # Stats row
        stats_headers = ["Total Sales", "Total Profit", "Margin %", "Items Sold", "Sales Rank"]
        ws_rep.append(stats_headers)
        row_num = ws_rep.max_row
        for i, h in enumerate(stats_headers, 1):
            cell = ws_rep.cell(row=row_num, column=i)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border
        
        ws_rep.append([rep_data_item['total_sales'], rep_data_item['total_profit'],
                       rep_data_item['overall_margin'] / 100, rep_data_item['total_qty'],
                       pa.get('sales_rank', '-')])
        row_num = ws_rep.max_row
        ws_rep.cell(row=row_num, column=1).number_format = num_fmt
        ws_rep.cell(row=row_num, column=2).number_format = num_fmt
        ws_rep.cell(row=row_num, column=3).number_format = pct_fmt
        ws_rep.cell(row=row_num, column=4).number_format = '#,##0'
        for c in range(1, 6):
            ws_rep.cell(row=row_num, column=c).border = thin_border
        
        ws_rep.append([])
        
        # Category breakdown
        cat_headers = ["Category", "Sales (KES)", "Profit (KES)", "Margin %", "Qty"]
        ws_rep.append(cat_headers)
        row_num = ws_rep.max_row
        for i, h in enumerate(cat_headers, 1):
            cell = ws_rep.cell(row=row_num, column=i)
            cell.font = Font(name='Calibri', bold=True, size=11, color='FFFFFF')
            cell.fill = PatternFill(start_color='3498DB', end_color='3498DB', fill_type='solid')
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border
        
        for cat in CAT_ORDER:
            if cat in rep_data_item['categories']:
                c = rep_data_item['categories'][cat]
                ws_rep.append([cat, c['sales_incl'], c['profit'], c['margin_pct'] / 100, c['qty']])
                row_num = ws_rep.max_row
                ws_rep.cell(row=row_num, column=2).number_format = num_fmt
                ws_rep.cell(row=row_num, column=3).number_format = num_fmt
                ws_rep.cell(row=row_num, column=4).number_format = pct_fmt
                ws_rep.cell(row=row_num, column=5).number_format = '#,##0'
                for ci in range(1, 6):
                    ws_rep.cell(row=row_num, column=ci).border = thin_border
        
        ws_rep.append([])
        
        # Performance notes
        ws_rep.append(["Performance Notes"])
        ws_rep.cell(row=ws_rep.max_row, column=1).font = Font(name='Calibri', bold=True, size=11, color='8E44AD')
        
        for note in pa.get('notes', [])[:5]:
            ws_rep.append([f"  {note}"])
            row_num = ws_rep.max_row
            if any(w in note for w in ['Strong', 'Excellent', 'Outstanding', 'Good', 'Selling across']):
                ws_rep.cell(row=row_num, column=1).font = note_font_green
            elif any(w in note for w in ['below', 'needs improvement', 'Should lead', 'Low contribution', 'Limited to']):
                ws_rep.cell(row=row_num, column=1).font = note_font_red
            else:
                ws_rep.cell(row=row_num, column=1).font = note_font_grey
        
        # Column widths
        for col_letter, width in [('A', 22), ('B', 16), ('C', 16), ('D', 12), ('E', 12), ('F', 12)]:
            ws_rep.column_dimensions[col_letter].width = width
    
    wb.save(output_fn)
    print(f"  Excel: {output_fn}")
    return output_fn

# ========== WORD REPORT ==========
def generate_word():
    output_fn = f"sales_analysis_march_w3.docx"
    doc = DocxDocument()
    
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SALES PERFORMANCE ANALYSIS")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"BOMAS Hardware Store - {MONTH} {YEAR}, Week {WEEK_NUM}")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Period: {PERIOD} ({WORKING_DAYS} working days)")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    # Executive Summary
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("EXECUTIVE SUMMARY")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    
    # Summary table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    metrics = [
        ("Total Sales (Incl. VAT)", f"KES {total_sales:,.0f}"),
        ("Total Profit (Incl. VAT)", f"KES {total_profit:,.0f}"),
        ("Overall Margin", f"{overall_margin:.1f}%"),
        ("Daily Average Sales", f"KES {total_sales/WORKING_DAYS:,.0f}"),
        ("Daily Average Profit", f"KES {total_profit/WORKING_DAYS:,.0f}"),
        ("Active Sales Reps", str(len(reps_data))),
    ]
    
    for i, (label, val) in enumerate(metrics):
        cell_l = table.cell(i, 0)
        cell_v = table.cell(i, 1)
        p_l = cell_l.paragraphs[0]
        p_v = cell_v.paragraphs[0]
        run_l = p_l.add_run(label)
        run_l.bold = True
        run_l.font.size = Pt(10)
        run_v = p_v.add_run(val)
        run_v.font.size = Pt(10)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
        cell_l._element.get_or_add_tcPr().append(shading)
    
    # Target Progress
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("MONTHLY TARGET PROGRESS")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    
    target_pct = (cumulative_total_sales / MONTHLY_TARGETS['OVERALL'] * 100)
    p = doc.add_paragraph()
    pace_text = "AHEAD of" if target_pct >= 75 else "BEHIND"
    run = p.add_run(f"After Week 3, the branch has achieved {target_pct:.1f}% of the monthly target (KES {MONTHLY_TARGETS['OVERALL']/1000000:.0f}M). This puts us {pace_text} the expected 75% pace.")
    run.font.size = Pt(10)
    
    table = doc.add_table(rows=len(CAT_ORDER) + 2, cols=5)
    table.style = 'Table Grid'
    headers = ["Category", "Monthly Target", "Week 3 Actual", "% Achieved", "Pace"]
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2C3E50"/>')
        cell._element.get_or_add_tcPr().append(shading)
    
    for i, cat in enumerate(CAT_ORDER):
        actual = cumulative_cat_totals[cat]
        target = MONTHLY_TARGETS.get(cat, 0)
        pct = (actual / target * 100) if target > 0 else 0
        pace = "Ahead" if pct >= 75 else "Behind"
        vals = [cat, f"KES {target:,.0f}", f"KES {actual:,.0f}", f"{pct:.1f}%", pace]
        for j, v in enumerate(vals):
            cell = table.cell(i + 1, j)
            p = cell.paragraphs[0]
            run = p.add_run(v)
            run.font.size = Pt(9)
    
    # Overall row
    overall_vals = ["OVERALL", f"KES {MONTHLY_TARGETS['OVERALL']:,.0f}", f"KES {cumulative_total_sales:,.0f}",
                     f"{target_pct:.1f}%", "Ahead" if target_pct >= 75 else "Behind"]
    for j, v in enumerate(overall_vals):
        cell = table.cell(len(CAT_ORDER) + 1, j)
        p = cell.paragraphs[0]
        run = p.add_run(v)
        run.bold = True
        run.font.size = Pt(9)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="34495E"/>')
        cell._element.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    # Rep Performance
    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run("SALES REP PERFORMANCE")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    
    table = doc.add_table(rows=len(sorted_reps) + 1, cols=6)
    table.style = 'Table Grid'
    headers = ["Rank", "Sales Rep", "Sales (KES)", "Profit (KES)", "Margin", "Contribution"]
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2C3E50"/>')
        cell._element.get_or_add_tcPr().append(shading)
    
    for rank, (rep_name, rep_data) in enumerate(sorted_reps, 1):
        contrib = (rep_data['total_sales'] / total_sales * 100) if total_sales > 0 else 0
        vals = [str(rank), rep_name, f"{rep_data['total_sales']:,.0f}",
                 f"{rep_data['total_profit']:,.0f}", f"{rep_data['overall_margin']:.1f}%", f"{contrib:.1f}%"]
        for j, v in enumerate(vals):
            cell = table.cell(rank, j)
            p = cell.paragraphs[0]
            run = p.add_run(v)
            run.font.size = Pt(9)
            if rank % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="ECF0F1"/>')
                cell._element.get_or_add_tcPr().append(shading)
    
    # Add charts
    doc.add_paragraph()
    for chart_key in ['trend', 'rep_bar', 'margin', 'stacked', 'category_pie', 'radar']:
        if os.path.exists(chart_files[chart_key]):
            doc.add_picture(chart_files[chart_key], width=Inches(6))
            p = doc.paragraphs[-1]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
    
    # Category Details
    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run("DEPARTMENT DETAILS")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    
    for cat in CAT_ORDER:
        p = doc.add_paragraph()
        run = p.add_run(cat)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
        
        dept_head = DEPARTMENT_HEADS.get(cat, 'N/A')
        if isinstance(dept_head, list):
            dept_head = ', '.join(dept_head)
        p = doc.add_paragraph()
        run = p.add_run(f"Department Head(s): {dept_head}")
        run.font.size = Pt(10)
        run.italic = True
        
        cat_margin = (cat_totals[cat]['profit'] / cat_totals[cat]['sales'] * 100) if cat_totals[cat]['sales'] > 0 else 0
        p = doc.add_paragraph()
        run = p.add_run(f"Sales: KES {cat_totals[cat]['sales']:,.0f}  |  Profit: KES {cat_totals[cat]['profit']:,.0f}  |  Margin: {cat_margin:.1f}%")
        run.font.size = Pt(10)
        
        reps_in_cat = []
        for rep_name, rep_d in sorted_reps:
            if cat in rep_d['categories']:
                c = rep_d['categories'][cat]
                share = (c['sales_incl'] / cat_totals[cat]['sales'] * 100) if cat_totals[cat]['sales'] > 0 else 0
                reps_in_cat.append((rep_name, c, share))
        reps_in_cat.sort(key=lambda x: x[1]['sales_incl'], reverse=True)
        
        table = doc.add_table(rows=len(reps_in_cat) + 1, cols=5)
        table.style = 'Table Grid'
        for j, h in enumerate(["Rep", "Sales", "Profit", "Margin", "Share"]):
            cell = table.cell(0, j)
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2C3E50"/>')
            cell._element.get_or_add_tcPr().append(shading)
        
        for i, (rep_name, c, share) in enumerate(reps_in_cat):
            vals = [rep_name, f"KES {c['sales_incl']:,.0f}", f"KES {c['profit']:,.0f}",
                     f"{c['margin_pct']:.1f}%", f"{share:.1f}%"]
            for j, v in enumerate(vals):
                cell = table.cell(i + 1, j)
                p = cell.paragraphs[0]
                run = p.add_run(v)
                run.font.size = Pt(9)
        
        doc.add_paragraph()
    
    # ===== INDIVIDUAL REP PERFORMANCE SECTIONS =====
    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run("INDIVIDUAL REP PERFORMANCE")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    
    for rep_name, rep_data_item in sorted_reps:
        pa = perf_analysis.get(rep_name, {})
        role = pa.get('role', 'Sales Rep')
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(f"{rep_name.upper()} - {role}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
        
        # Badges
        badges = pa.get('badges', [])
        if badges:
            p = doc.add_paragraph()
            run = p.add_run('  '.join(badges))
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x8e, 0x44, 0xad)
        
        # Stats table
        table = doc.add_table(rows=2, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        stats_headers = ["Total Sales", "Total Profit", "Margin", "Items Sold", "Rank"]
        stats_vals = [f"KES {rep_data_item['total_sales']:,.0f}",
                      f"KES {rep_data_item['total_profit']:,.0f}",
                      f"{rep_data_item['overall_margin']:.1f}%",
                      f"{rep_data_item['total_qty']:,.0f}",
                      f"#{pa.get('sales_rank', '-')}"]
        for j, h in enumerate(stats_headers):
            cell = table.cell(0, j)
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="34495E"/>')
            cell._element.get_or_add_tcPr().append(shading)
        for j, v in enumerate(stats_vals):
            cell = table.cell(1, j)
            p = cell.paragraphs[0]
            run = p.add_run(v)
            run.bold = True
            run.font.size = Pt(9)
        
        # Category breakdown
        rep_cats = []
        for cat in CAT_ORDER:
            if cat in rep_data_item['categories']:
                c = rep_data_item['categories'][cat]
                rep_cats.append((cat, c))
        
        if rep_cats:
            doc.add_paragraph()
            table = doc.add_table(rows=len(rep_cats) + 1, cols=5)
            table.style = 'Table Grid'
            for j, h in enumerate(["Category", "Sales", "Profit", "Margin", "Qty"]):
                cell = table.cell(0, j)
                p = cell.paragraphs[0]
                run = p.add_run(h)
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="3498DB"/>')
                cell._element.get_or_add_tcPr().append(shading)
            for i, (cat, c) in enumerate(rep_cats):
                vals = [cat, f"KES {c['sales_incl']:,.0f}", f"KES {c['profit']:,.0f}",
                         f"{c['margin_pct']:.1f}%", f"{c['qty']:,.0f}"]
                for j, v in enumerate(vals):
                    cell = table.cell(i + 1, j)
                    p = cell.paragraphs[0]
                    run = p.add_run(v)
                    run.font.size = Pt(9)
        
        # Charts
        rc = rep_charts.get(rep_name, {})
        if rc.get('donut') and os.path.exists(rc['donut']):
            doc.add_paragraph()
            doc.add_picture(rc['donut'], width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if rc.get('sales_profit') and os.path.exists(rc['sales_profit']):
            doc.add_paragraph()
            doc.add_picture(rc['sales_profit'], width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Performance notes
        p = doc.add_paragraph()
        run = p.add_run("Performance Notes:")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x8e, 0x44, 0xad)
        
        for note in pa.get('notes', [])[:5]:
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            if any(w in note for w in ['Strong', 'Excellent', 'Outstanding', 'Good', 'Selling across']):
                color = RGBColor(0x27, 0xae, 0x60)
            elif any(w in note for w in ['below', 'needs improvement', 'Should lead', 'Low contribution', 'Limited to']):
                color = RGBColor(0xe7, 0x4c, 0x3c)
            else:
                color = RGBColor(0x55, 0x55, 0x55)
            run = p.add_run(note)
            run.font.size = Pt(9)
            run.font.color.rgb = color
        
        doc.add_paragraph()
    
    doc.save(output_fn)
    print(f"  Word: {output_fn}")
    return output_fn

# ========== JSON REPORT ==========
def generate_json():
    output_fn = f"sales_analysis_march_w3.json"
    
    output = {
        "report": f"{MONTH} {YEAR} - Week {WEEK_NUM} Sales Analysis",
        "period": PERIOD,
        "working_days": WORKING_DAYS,
        "generated": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "department_heads": {k: v if isinstance(v, str) else v for k, v in DEPARTMENT_HEADS.items()},
        "summary": {
            "total_sales": total_sales,
            "total_profit": round(total_profit, 2),
            "overall_margin": round(overall_margin, 2),
            "daily_avg_sales": round(total_sales / WORKING_DAYS, 2),
            "daily_avg_profit": round(total_profit / WORKING_DAYS, 2),
            "active_reps": len(reps_data),
        },
        "monthly_targets": {
            "overall": {
                "target": MONTHLY_TARGETS['OVERALL'],
                "accumulated_actual": cumulative_total_sales,
                "pct_achieved": round(cumulative_total_sales / MONTHLY_TARGETS['OVERALL'] * 100, 1),
                "expected_pct": 25.0,
                "pace": "Ahead" if cumulative_total_sales / MONTHLY_TARGETS['OVERALL'] >= 0.75 else "Behind",
            },
            "categories": {}
        },
        "categories": {},
        "reps": {},
    }
    
    for cat in CAT_ORDER:
        actual = cumulative_cat_totals[cat]
        target = MONTHLY_TARGETS.get(cat, 0)
        pct = (actual / target * 100) if target > 0 else 0
        output["monthly_targets"]["categories"][cat] = {
            "target": target,
            "week3_actual": actual,
            "pct_achieved": round(pct, 1),
            "pace": "Ahead" if pct >= 75 else "Behind",
            "department_head": DEPARTMENT_HEADS.get(cat, "N/A"),
        }
        
        cat_margin = (cat_totals[cat]['profit'] / cat_totals[cat]['sales'] * 100) if cat_totals[cat]['sales'] > 0 else 0
        output["categories"][cat] = {
            "total_sales": cat_totals[cat]['sales'],
            "total_profit": round(cat_totals[cat]['profit'], 2),
            "margin": round(cat_margin, 2),
            "share_of_total": round(cat_totals[cat]['sales'] / total_sales * 100, 1) if total_sales > 0 else 0,
        }
    
    for rep_name, rep_data in sorted_reps:
        contrib = (rep_data['total_sales'] / total_sales * 100) if total_sales > 0 else 0
        pa = perf_analysis.get(rep_name, {})
        output["reps"][rep_name] = {
            "role": get_rep_role(rep_name),
            "primary_department": get_primary_dept(rep_name),
            "total_sales": rep_data["total_sales"],
            "total_profit": round(rep_data["total_profit"], 2),
            "overall_margin": round(rep_data["overall_margin"], 2),
            "total_qty": rep_data["total_qty"],
            "contribution_pct": round(contrib, 1),
            "sales_rank": pa.get('sales_rank', 0),
            "margin_rank": pa.get('margin_rank', 0),
            "profit_rank": pa.get('profit_rank', 0),
            "badges": pa.get('badges', []),
            "performance_notes": pa.get('notes', []),
            "categories": {
                cat: {
                    "sales": v["sales_incl"],
                    "profit": round(v["profit"], 2),
                    "margin_pct": round(v["margin_pct"], 2),
                    "qty": v["qty"],
                    "share_of_dept": round(v["sales_incl"] / cat_totals[cat]["sales"] * 100, 1) if cat_totals[cat]["sales"] > 0 else 0,
                }
                for cat, v in rep_data["categories"].items()
            }
        }
    
    with open(output_fn, "w") as f:
        json.dump(output, f, indent=2)
    print(f"  JSON: {output_fn}")
    return output_fn

# ========== MAIN ==========
def main():
    print("=" * 60)
    print(f"SALES ANALYSIS - {MONTH.upper()} {YEAR}, WEEK {WEEK_NUM}")
    print(f"Period: {PERIOD} ({WORKING_DAYS} working days)")
    print("=" * 60)
    
    print(f"\nMerged Magdalene -> Betha Odumo")
    print(f"Active Reps: {len(reps_data)}")
    
    print(f"\n{'Rep':<20} {'Sales':>12} {'Profit':>12} {'Margin':>8} {'Contrib':>8}")
    print("-" * 62)
    for rep_name, rep_data in sorted_reps:
        contrib = (rep_data['total_sales'] / total_sales * 100) if total_sales > 0 else 0
        print(f"{rep_name:<20} KES {rep_data['total_sales']:>10,.0f} KES {rep_data['total_profit']:>9,.0f} {rep_data['overall_margin']:>7.1f}% {contrib:>6.1f}%")
    
    print(f"\n{'Category':<20} {'Sales':>14} {'Profit':>14} {'Margin':>8} {'Share':>8}")
    print("-" * 66)
    for cat in CAT_ORDER:
        margin = (cat_totals[cat]['profit'] / cat_totals[cat]['sales'] * 100) if cat_totals[cat]['sales'] > 0 else 0
        share = (cat_totals[cat]['sales'] / total_sales * 100) if total_sales > 0 else 0
        print(f"{cat:<20} KES {cat_totals[cat]['sales']:>11,.0f} KES {cat_totals[cat]['profit']:>11,.0f} {margin:>7.1f}% {share:>6.1f}%")
    
    print(f"\n{'GRAND TOTAL':<20} KES {total_sales:>11,.0f} KES {total_profit:>11,.0f} {overall_margin:>7.1f}%")
    print(f"{'Daily Average':<20} KES {total_sales/WORKING_DAYS:>11,.0f} KES {total_profit/WORKING_DAYS:>11,.0f}")
    
    # Target progress
    target_pct = cumulative_total_sales / MONTHLY_TARGETS['OVERALL'] * 100
    print(f"\nMonthly Target Progress: {target_pct:.1f}% (expected 75% after Week 3)")
    
    print("\nGenerating reports...")
    generate_json()
    generate_pdf()
    generate_excel()
    generate_word()
    
    print("\nAll reports generated successfully!")
    print("=" * 60)

if __name__ == "__main__":
    main()
