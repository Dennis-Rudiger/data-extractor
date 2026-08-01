import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from pathlib import Path

ROOT = Path(__file__).resolve().parent
DATA_DIR = ROOT / 'Thus far' / 'data'


def safe_read_json(path):
    try:
        return pd.read_json(path)
    except Exception:
        return pd.DataFrame()


def safe_read_csv(path):
    try:
        return pd.read_csv(path, encoding='utf-8', skiprows=0)
    except Exception:
        return pd.DataFrame()


def build_doc(df_cust, df_inv, out_path):
    doc = Document()
    doc.add_heading('Half 1 Report — Customers & Inventory', 0)

    # Executive summary
    total_revenue = df_cust['Amount_Incl'].sum() if not df_cust.empty else 0
    total_cost = df_cust['Cost'].sum() if 'Cost' in df_cust.columns else 0
    total_profit = df_cust['Profit'].sum() if 'Profit' in df_cust.columns else 0
    active_clients = df_cust['Customer'].nunique() if not df_cust.empty else 0

    doc.add_heading('Executive Summary', level=1)
    p = doc.add_paragraph()
    p.add_run(f'Total Revenue (Incl): KES {total_revenue:,.2f}\n').bold = True
    p.add_run(f'Total Cost: KES {total_cost:,.2f}\n').bold = True
    p.add_run(f'Total Profit: KES {total_profit:,.2f}\n').bold = True
    p.add_run(f'Active Clients: {active_clients}\n').bold = True

    # Top customers
    doc.add_heading('Top Customers (by Revenue)', level=1)
    if not df_cust.empty:
        overall = df_cust.groupby('Customer')[['Amount_Incl', 'Profit', 'Quantity']].sum().sort_values('Amount_Incl', ascending=False)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Customer'
        hdr[1].text = 'Revenue (Incl)'
        hdr[2].text = 'Profit'
        hdr[3].text = 'Quantity'
        for idx, row in overall.head(20).iterrows():
            r = table.add_row().cells
            r[0].text = str(idx)
            r[1].text = f"{row['Amount_Incl']:,.2f}"
            r[2].text = f"{row['Profit']:,.2f}"
            r[3].text = f"{row['Quantity']:,.0f}"
    else:
        doc.add_paragraph('No customer sales data found.')

    # Top products
    doc.add_heading('Top Products (Inventory)', level=1)
    if not df_inv.empty:
        # attempt to find columns
        cols = df_inv.columns.str.strip()
        # try to detect Item Code, Description, Quantity, Profit
        code_col = next((c for c in df_inv.columns if 'Item Code' in c or 'Item' in c), df_inv.columns[0])
        desc_col = next((c for c in df_inv.columns if 'Description' in c or 'Item Description' in c), df_inv.columns[1] if len(df_inv.columns) > 1 else df_inv.columns[0])
        qty_col = next((c for c in df_inv.columns if 'Quantity' in c or 'Qty' in c), None)
        profit_col = next((c for c in df_inv.columns if 'Profit' in c), None)

        # coerce numeric
        if qty_col:
            df_inv[qty_col] = pd.to_numeric(df_inv[qty_col].astype(str).str.replace(',', ''), errors='coerce').fillna(0)
        if profit_col:
            df_inv[profit_col] = pd.to_numeric(df_inv[profit_col].astype(str).str.replace(',', ''), errors='coerce').fillna(0)

        df_inv_top_qty = df_inv.sort_values(by=qty_col if qty_col else df_inv.columns[0], ascending=False).head(20)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Item Code'
        hdr[1].text = 'Description'
        hdr[2].text = 'Quantity'
        hdr[3].text = 'Profit'
        for _, r in df_inv_top_qty.iterrows():
            row = table.add_row().cells
            row[0].text = str(r.get(code_col, ''))
            row[1].text = str(r.get(desc_col, ''))
            row[2].text = f"{r.get(qty_col, 0):,.0f}" if qty_col else ''
            row[3].text = f"{r.get(profit_col, 0):,.2f}" if profit_col else ''
    else:
        doc.add_paragraph('No inventory data found.')

    # Notes
    doc.add_heading('Notes & Next Steps', level=1)
    doc.add_paragraph('- Include sales-rep breakdown if a sales-rep file is provided.')
    doc.add_paragraph('- Generate monthly trend charts and embed in report for visual context.')

    doc.save(out_path)


def main():
    cust_json = DATA_DIR / 'customer_sales_extracted.json'
    inv_csv = DATA_DIR / 'Inventory Sales Analysis_20260713_123405.csv'
    out_path = ROOT / 'Half1_Report.docx'

    df_cust = safe_read_json(cust_json) if cust_json.exists() else pd.DataFrame()
    # normalize expected columns
    if not df_cust.empty:
        for col in ['Amount_Incl', 'Profit', 'Quantity']:
            if col in df_cust.columns:
                df_cust[col] = pd.to_numeric(df_cust[col], errors='coerce').fillna(0)

    df_inv = safe_read_csv(inv_csv) if inv_csv.exists() else pd.DataFrame()

    build_doc(df_cust, df_inv, out_path)
    print('Wrote', out_path)


if __name__ == '__main__':
    main()
