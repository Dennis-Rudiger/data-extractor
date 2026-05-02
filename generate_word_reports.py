import json
from docx import Document
from docx.shared import Pt, Inches

# Load Data
with open('bomas_average_moq.json', 'r', encoding='utf-8') as f:
    bomas = json.load(f)

with open('march_april_moq.json', 'r', encoding='utf-8') as f:
    march_april = json.load(f)

# 1. Build Basic MOQ Report DOCX
doc1 = Document()
doc1.add_heading('MOQ Report: March - April (36 Days)', 0)

for cat_name, cat_data in march_april['categories'].items():
    if not cat_data['products']: continue
    doc1.add_heading(cat_name, level=1)
    
    # Sort products by Qty Out descending
    prods = sorted(cat_data['products'], key=lambda x: x.get('qty_out', 0), reverse=True)
    
    table = doc1.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item Code'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Qty Out'
    hdr_cells[3].text = 'Daily Avg'
    hdr_cells[4].text = 'Weekly MOQ'
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True

    for p in prods:
        row_cells = table.add_row().cells
        row_cells[0].text = str(p.get('item_code', ''))
        row_cells[1].text = str(p.get('item_description', ''))
        row_cells[2].text = '{:,.1f}'.format(p.get('qty_out', 0))
        row_cells[3].text = '{:,.2f}'.format(p.get('daily_average', 0))
        row_cells[4].text = '{:,.0f}'.format(p.get('weekly_moq', 0))
    doc1.add_paragraph()

doc1.save('MOQ_Report_March_April.docx')
print('Generated MOQ_Report_March_April.docx')

# 2. Build Comparison Report DOCX
b_map = {}
for cat_name, cat_data in bomas['categories'].items():
    for p in cat_data['products']:
        b_map[p['item_code']] = {
            'b_qty_out': p.get('qty_out', 0),
            'b_daily_avg': p.get('daily_average', 0),
            'b_moq': p.get('weekly_moq', 0)
        }

doc2 = Document()
doc2.add_heading('MOQ Comparison: Q1 (74 Days) vs March-April (36 Days)', 0)

for cat_name, cat_data in march_april['categories'].items():
    if not cat_data['products']: continue
    
    comp_lines = []
    for m_p in cat_data['products']:
        code = m_p['item_code']
        b_p = b_map.get(code, {'b_qty_out': 0, 'b_daily_avg': 0, 'b_moq': 0})
        m_daily = m_p['daily_average']
        b_daily = b_p['b_daily_avg']
        variance = m_daily - b_daily
        status = 'Increased' if variance > 0.1 else ('Dropped' if variance < -0.1 else 'Unchanged')
        
        comp_lines.append({
            'code': code,
            'desc': m_p['item_description'],
            'old_avg': b_daily,
            'new_avg': m_daily,
            'var_daily': variance,
            'status': status,
            'new_moq': m_p['weekly_moq']
        })
    
    comp_lines = sorted(comp_lines, key=lambda x: abs(x['var_daily']), reverse=True)
    
    doc2.add_heading(cat_name, level=1)
    
    table = doc2.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Code'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Old Avg'
    hdr_cells[3].text = 'New Avg'
    hdr_cells[4].text = 'Variance'
    hdr_cells[5].text = 'Status'
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True

    for row_data in comp_lines:
        row_cells = table.add_row().cells
        row_cells[0].text = str(row_data['code'])
        row_cells[1].text = str(row_data['desc'])
        row_cells[2].text = '{:,.2f}'.format(row_data['old_avg'])
        row_cells[3].text = '{:,.2f}'.format(row_data['new_avg'])
        row_cells[4].text = '{:,.2f}'.format(row_data['var_daily'])
        row_cells[5].text = row_data['status']
    doc2.add_paragraph()

doc2.save('MOQ_Comparison_Bomas.docx')
print('Generated MOQ_Comparison_Bomas.docx')

