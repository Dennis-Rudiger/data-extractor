import json
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image, PageBreak        
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.units import inch

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT# Configuration
MONTHLY_TARGETS = {
    'OVERALL': 60_000_000,
    'GENERAL HARDWARE': 48_500_000,
    'PAINTS': 7_000_000,
    'PLUMBING': 3_000_000,
    'ELECTRICALS': 1_500_000,
}

VAT_RATE = 1.16

def load_data(filepath):
    with open(filepath, 'r') as f:
        return json.load(f)

def process_data(raw_data):
    overall_sales = 0
    overall_profit = 0
    cats = {}
    reps = {}
    
    for cat, cat_data in raw_data["categories"].items():
        if cat not in cats:
            cats[cat] = {'sales': 0, 'profit': 0}
            
        for rep, data in cat_data.items():
            rep_name = "Betha Odumo" if rep in ("Magdalene", "WALK IN-BOMAS") else rep
            sales = data['sales_incl']
            profit = data['profit'] * VAT_RATE
            
            if rep_name not in reps:
                reps[rep_name] = {'sales': 0, 'profit': 0}
                
            reps[rep_name]['sales'] += sales
            reps[rep_name]['profit'] += profit
            
            cats[cat]['sales'] += sales
            cats[cat]['profit'] += profit
            
            overall_sales += sales
            overall_profit += profit
            
    return {
        'sales': overall_sales,
        'profit': overall_profit,
        'margin': (overall_profit / overall_sales * 100) if overall_sales > 0 else 0,
        'categories': cats,
        'reps': reps
    }

def main():
    print("Generating Mid-Month Performance Report...")
    w1 = process_data(load_data("sales_march_w1.json"))
    w2 = process_data(load_data("sales_march_w2.json"))
    
    comb_sales = w1['sales'] + w2['sales']
    comb_profit = w1['profit'] + w2['profit']
    comb_margin = (comb_profit / comb_sales * 100) if comb_sales > 0 else 0
    
    w1_w2_diff = w2['sales'] - w1['sales']
    w1_w2_growth = (w1_w2_diff / w1['sales'] * 100) if w1['sales'] > 0 else 0
    
    overall_pace = comb_sales / MONTHLY_TARGETS['OVERALL'] * 100
    
    cats = [c for c in MONTHLY_TARGETS.keys() if c != 'OVERALL']
    w1_cat_sales = [w1['categories'].get(c, {}).get('sales', 0) for c in cats]
    w2_cat_sales = [w2['categories'].get(c, {}).get('sales', 0) for c in cats]
    comb_cat_sales = [w1_cat_sales[i] + w2_cat_sales[i] for i in range(len(cats))]
    
    # CHARTS GENERATION
    os.makedirs("charts/mid_month", exist_ok=True)
    width = 0.35
    
    # 1. Categories W1 vs W2
    fig, ax = plt.subplots(figsize=(8, 4))
    x = np.arange(len(cats))
    ax.bar(x - width/2, w1_cat_sales, width, label='Week 1', color='#3498db')
    ax.bar(x + width/2, w2_cat_sales, width, label='Week 2', color='#2ecc71')
    ax.set_ylabel('Sales (KES)')
    ax.set_title('Week 1 vs Week 2 Sales by Category')
    ax.set_xticks(x)
    ax.set_xticklabels(cats)
    ax.legend()
    plt.tight_layout()
    plt.savefig("charts/mid_month/cat_comparison.png")
    plt.close()
    
    # 2. Targets Comparison
    fig, ax = plt.subplots(figsize=(8, 4))
    targets = [MONTHLY_TARGETS[c] for c in cats]
    ax.bar(x - width/2, targets, width, label='Monthly Target', color='#95a5a6')
    ax.bar(x + width/2, comb_cat_sales, width, label='Achieved (Mid-Month)', color='#e74c3c')
    ax.set_ylabel('Sales (KES)')
    ax.set_title('Mid-Month Validated Sales vs Monthly Targets')
    ax.set_xticks(x)
    ax.set_xticklabels(cats)
    ax.legend()
    plt.tight_layout()
    plt.savefig("charts/mid_month/target_comparison.png")
    plt.close()
    
    # 3. Rep Comparison
    all_reps = set(w1['reps'].keys()) | set(w2['reps'].keys())
    rep_list = []
    for r in all_reps:
        r_w1 = w1['reps'].get(r, {}).get('sales', 0)
        r_w2 = w2['reps'].get(r, {}).get('sales', 0)
        rep_list.append((r, r_w1, r_w2, r_w1+r_w2))
    
    rep_list.sort(key=lambda x: x[3], reverse=True)
    top_reps = rep_list[:8] # Top 8
    
    r_names = [r[0] for r in top_reps]
    r_w1_s = [r[1] for r in top_reps]
    r_w2_s = [r[2] for r in top_reps]
    
    fig, ax = plt.subplots(figsize=(9, 4.5))
    x_r = np.arange(len(r_names))
    ax.bar(x_r - width/2, r_w1_s, width, label='Week 1', color='#9b59b6')
    ax.bar(x_r + width/2, r_w2_s, width, label='Week 2', color='#f1c40f')
    ax.set_ylabel('Sales (KES)')
    ax.set_title('Top Sales Reps: Week 1 vs Week 2')
    ax.set_xticks(x_r)
    ax.set_xticklabels(r_names, rotation=30, ha='right')
    ax.legend()
    plt.tight_layout()
    plt.savefig("charts/mid_month/rep_comparison.png")
    plt.close()

    # DOCUMENT GENERATION
    doc = SimpleDocTemplate("mid_month_performance.pdf", pagesize=A4, rightMargin=40, leftMargin=40)
    elements = []
    styles = getSampleStyleSheet()
    
    title_style = styles['Heading1']
    title_style.alignment = TA_CENTER
    h2 = styles['Heading2']
    h2.textColor = colors.HexColor('#2c3e50')
    h3 = styles['Heading3']
    h3.textColor = colors.HexColor('#2980b9')
    normal = styles['Normal']
    normal.alignment = TA_JUSTIFY
    
    # HEADING
    elements.append(Paragraph("Mid-Month Performance Report", title_style))
    elements.append(Paragraph("Period: March 2 - March 14, 2026", ParagraphStyle('Sub', parent=styles['Normal'], alignment=TA_CENTER)))
    elements.append(Spacer(1, 0.4*inch))
    
    # 1. EXECUTIVE SUMMARY
    elements.append(Paragraph("1. Executive Summary & Week 2 Breakdown", h2))
    
    exec_summary = f"""Overall, the first half of March concluded with a Grand Total Sales figure of <b>KES {comb_sales:,.0f}</b>, 
    representing <b>{overall_pace:.1f}%</b> of the overarching KES 60,000,000 monthly target. Week 2 showcased strong 
    momentum with total sales reaching KES {w2['sales']:,.0f}, a growth of <b>{w1_w2_growth:.1f}%</b> compared to Week 1 
    (KES {w1['sales']:,.0f}). Average combined profit margin tracks at {comb_margin:.1f}%."""
    elements.append(Paragraph(exec_summary, normal))
    elements.append(Spacer(1, 0.2*inch))
    
    w2_kpi = [
        ["Week 1 Total Sales", f"KES {w1['sales']:,.0f}", "Mid-Month Cumulative", f"KES {comb_sales:,.0f}"],
        ["Week 2 Total Sales", f"KES {w2['sales']:,.0f}", "Monthly Target", f"KES {MONTHLY_TARGETS['OVERALL']:,.0f}"],
        ["Growth (W1 -> W2)", f"{w1_w2_growth:.1f}%", "Pacing (% Achieved)", f"{overall_pace:.1f}%"]
    ]
    t1 = Table(w2_kpi, colWidths=[1.8*inch, 1.8*inch, 1.8*inch, 1.8*inch])
    t1.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f8f9fa')),
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
        ('FONTNAME', (1,0), (1,-1), 'Helvetica-Bold'),
        ('FONTNAME', (3,0), (3,-1), 'Helvetica-Bold'),
        ('TEXTCOLOR', (1,2), (1,2), colors.green if w1_w2_growth > 0 else colors.red),
        ('TEXTCOLOR', (3,2), (3,2), colors.green if overall_pace >= 50 else colors.orange),
        ('PADDING', (0,0), (-1,-1), 8),
    ]))
    elements.append(t1)
    elements.append(Spacer(1, 0.4*inch))
    
    # 2. COMPARISON CHARTS
    elements.append(Paragraph("2. Week 1 vs Week 2 Comparisons", h2))
    elements.append(Image("charts/mid_month/cat_comparison.png", width=6.5*inch, height=3.25*inch))
    elements.append(Spacer(1, 0.1*inch))
    elements.append(Image("charts/mid_month/rep_comparison.png", width=6.5*inch, height=3.25*inch))
    elements.append(PageBreak())
    
    # 3. TARGETS AFTER WEEK 2
    elements.append(Paragraph("3. Target Pacing After Week 2", h2))
    
    rem_sales = MONTHLY_TARGETS['OVERALL'] - comb_sales
    req_daily = rem_sales / 13 # Approx 13 working days remaining in March
    
    pacing_text = f"""At the halfway mark, the branch has achieved {overall_pace:.1f}% of the target. To hit the 
    KES 60M goal, an additional KES {rem_sales:,.0f} is required, translating to an average required run rate of 
    approximately KES {req_daily:,.0f} per day for the remaining weeks."""
    elements.append(Paragraph(pacing_text, normal))
    elements.append(Spacer(1, 0.1*inch))
    elements.append(Image("charts/mid_month/target_comparison.png", width=6.5*inch, height=3.25*inch))
    elements.append(Spacer(1, 0.2*inch))
    
    # Detailed target table
    cat_comb_data = [["Category", "W1 + W2 Achieved", "Monthly Target", "% Achieved"]]
    for i, cat in enumerate(cats):
        targ = MONTHLY_TARGETS[cat]
        ach = comb_cat_sales[i]
        pct = (ach / targ) * 100
        cat_comb_data.append([cat, f"KES {ach:,.0f}", f"KES {targ:,.0f}", f"{pct:.1f}%"])
        
    t_cat = Table(cat_comb_data, colWidths=[2.2*inch, 1.6*inch, 1.6*inch, 1*inch])
    t_cat.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#2c3e50')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
        ('ALIGN', (1,1), (-1,-1), 'RIGHT'),
        ('PADDING', (0,0), (-1,-1), 6),
    ]))
    elements.append(t_cat)
    elements.append(Spacer(1, 0.4*inch))
    
    # 4. COMMENTARY & RECOMMENDATIONS
    elements.append(Paragraph("4. Informed Commentary & Recommendations", h2))
    elements.append(Paragraph("Commentary", h3))
    
    hw_pct = comb_cat_sales[CATS_IDX['GENERAL HARDWARE']] / MONTHLY_TARGETS['GENERAL HARDWARE'] * 100 if 'GENERAL HARDWARE' in CATS_IDX else 0
    el_pct = comb_cat_sales[CATS_IDX['ELECTRICALS']] / MONTHLY_TARGETS['ELECTRICALS'] * 100 if 'ELECTRICALS' in CATS_IDX else 0
    
    commentary1 = """<bullet>&bull;</bullet> <b>Momentum:</b> Revenue saw a substantial upward spike in Week 2, jumping from 12.2M to 15.5M (+26%). 
    This strong growth trajectory establishes an excellent foundation for the rest of the month."""
    
    commentary2 = """<bullet>&bull;</bullet> <b>Category Performance:</b> Most departments are pacing just under the 50% waterline, with General Hardware remaining the heavyweight contributor. However, divisions like <i>Electricals</i> and <i>Plumbing</i> may need stimulation to hit their respective strict KES 1.5M and 3M thresholds, as their run rates are slacking slightly compared to Hardware."""
    
    commentary3 = """<bullet>&bull;</bullet> <b>Rep Dominance:</b> The gap between the top tier (Gladys, Betha Odumo) and mid-tier reps continues to widen, showing high dependency on a few key individuals for large ticket hardware sales."""
    
    elements.append(Paragraph(commentary1, normal))
    elements.append(Spacer(1, 0.1*inch))
    elements.append(Paragraph(commentary2, normal))
    elements.append(Spacer(1, 0.1*inch))
    elements.append(Paragraph(commentary3, normal))
    elements.append(Spacer(1, 0.2*inch))
    
    elements.append(Paragraph("Recommendations to Achieve KES 60,000,000", h3))
    
    recs1 = f"""<bullet>1.</bullet> <b>Boost the Run-Rate:</b> The daily required run-rate rests at roughly KES {req_daily/1000000:.2f} Million per day. Management should introduce daily targets to departmental heads to maintain the robust momentum seen in Week 2."""
    recs2 = """<bullet>2.</bullet> <b>Cross-Selling in Technical Departments:</b> Focus reps on cross-selling Electrical and Plumbing stock alongside bulk General Hardware purchases, to pull up their pacing percentages."""
    recs3 = """<bullet>3.</bullet> <b>Activate Mid-Tier Reps:</b> With top sales reps handling bulk orders, mid-tier members should be aggressively positioned to close fast-moving, high-margin items like Paints to diversify contribution risk."""
    recs4 = """<bullet>4.</bullet> <b>Promotional Push:</b> A quick flash promotion or discount strategy for slow-moving Electrical items in Week 3 could yield the volume needed to push the category back on track."""
    
    elements.append(Paragraph(recs1, normal))
    elements.append(Spacer(1, 0.05*inch))
    elements.append(Paragraph(recs2, normal))
    elements.append(Spacer(1, 0.05*inch))
    elements.append(Paragraph(recs3, normal))
    elements.append(Spacer(1, 0.05*inch))
    elements.append(Paragraph(recs4, normal))
    
    doc.build(elements)
    print("Report successfully saved as: mid_month_performance.pdf")

    # ========================== WORD DOCUMENT GENERATION ==========================
    doc_word = Document()
    
    # Title
    p_title = doc_word.add_heading("Mid-Month Performance Report", 0)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub = doc_word.add_paragraph("Period: March 2 - March 14, 2026")
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. EXECUTIVE SUMMARY
    doc_word.add_heading("1. Executive Summary & Week 2 Breakdown", level=1)
    
    doc_word.add_paragraph(f"Overall, the first half of March concluded with a Grand Total Sales figure of KES {comb_sales:,.0f}, representing {overall_pace:.1f}% of the overarching KES 60,000,000 monthly target. Week 2 showcased strong momentum with total sales reaching KES {w2['sales']:,.0f}, a growth of {w1_w2_growth:.1f}% compared to Week 1 (KES {w1['sales']:,.0f}). Average combined profit margin tracks at {comb_margin:.1f}%.")
    
    # KPI Table
    table = doc_word.add_table(rows=3, cols=4)
    table.style = 'Light Grid Accent 1'
    
    kpis = [
        ["Week 1 Total Sales", f"KES {w1['sales']:,.0f}", "Mid-Month Cumulative", f"KES {comb_sales:,.0f}"],
        ["Week 2 Total Sales", f"KES {w2['sales']:,.0f}", "Monthly Target", f"KES {MONTHLY_TARGETS['OVERALL']:,.0f}"],
        ["Growth (W1 -> W2)", f"{w1_w2_growth:.1f}%", "Pacing (% Achieved)", f"{overall_pace:.1f}%"]
    ]
    
    for i, row in enumerate(kpis):
        for j, val in enumerate(row):
            table.cell(i, j).text = val
    
    doc_word.add_page_break()
    
    # 2. CHARTS
    doc_word.add_heading("2. Week 1 vs Week 2 Comparisons", level=1)
    doc_word.add_picture("charts/mid_month/cat_comparison.png", width=Inches(6.0))
    doc_word.add_picture("charts/mid_month/rep_comparison.png", width=Inches(6.0))
    
    doc_word.add_page_break()
    
    # 3. TARGETS
    doc_word.add_heading("3. Target Pacing After Week 2", level=1)
    doc_word.add_paragraph(f"At the halfway mark, the branch has achieved {overall_pace:.1f}% of the target. To hit the KES 60M goal, an additional KES {rem_sales:,.0f} is required, translating to an average required run rate of approximately KES {req_daily:,.0f} per day for the remaining weeks.")
    
    doc_word.add_picture("charts/mid_month/target_comparison.png", width=Inches(6.0))
    
    cat_table = doc_word.add_table(rows=1, cols=4)
    cat_table.style = 'Light Grid Accent 1'
    hdr_cells = cat_table.rows[0].cells
    hdr_cells[0].text = "Category"
    hdr_cells[1].text = "W1 + W2 Achieved"
    hdr_cells[2].text = "Monthly Target"
    hdr_cells[3].text = "% Achieved"
    
    for i, cat in enumerate(cats):
        row_cells = cat_table.add_row().cells
        row_cells[0].text = cat
        row_cells[1].text = f"KES {comb_cat_sales[i]:,.0f}"
        row_cells[2].text = f"KES {MONTHLY_TARGETS[cat]:,.0f}"
        row_cells[3].text = f"{(comb_cat_sales[i] / MONTHLY_TARGETS[cat] * 100):.1f}%"
        
    # 4. COMMENTARY
    doc_word.add_heading("4. Informed Commentary & Recommendations", level=1)
    
    doc_word.add_heading("Commentary", level=2)
    doc_word.add_paragraph("Momentum: Revenue saw a substantial upward spike in Week 2, jumping from 12.2M to 15.5M (+26%). This strong growth trajectory establishes an excellent foundation for the rest of the month.", style='List Bullet')
    doc_word.add_paragraph("Category Performance: Most departments are pacing just under the 50% waterline, with General Hardware remaining the heavyweight contributor. However, divisions like Electricals and Plumbing may need stimulation to hit their respective strict KES 1.5M and 3M thresholds.", style='List Bullet')
    doc_word.add_paragraph("Rep Dominance: The gap between the top tier (Gladys, Betha Odumo) and mid-tier reps continues to widen, showing high dependency on a few key individuals for large ticket hardware sales.", style='List Bullet')
    
    doc_word.add_heading("Recommendations to Achieve KES 60,000,000", level=2)
    doc_word.add_paragraph(f"Boost the Run-Rate: The daily required run-rate rests at roughly KES {req_daily/1000000:.2f} Million per day. Management should introduce daily targets to departmental heads to maintain the robust momentum seen in Week 2.", style='List Number')
    doc_word.add_paragraph("Cross-Selling in Technical Departments: Focus reps on cross-selling Electrical and Plumbing stock alongside bulk General Hardware purchases, to pull up their pacing percentages.", style='List Number')
    doc_word.add_paragraph("Activate Mid-Tier Reps: With top sales reps handling bulk orders, mid-tier members should be aggressively positioned to close fast-moving, high-margin items like Paints to diversify contribution risk.", style='List Number')
    doc_word.add_paragraph("Promotional Push: A quick flash promotion or discount strategy for slow-moving Electrical items in Week 3 could yield the volume needed to push the category back on track.", style='List Number')
    
    doc_word.save("mid_month_performance.docx")
    print("Report successfully saved as: mid_month_performance.docx")

CATS_IDX = {k: i for i, k in enumerate(['GENERAL HARDWARE', 'PAINTS', 'PLUMBING', 'ELECTRICALS'])}

if __name__ == "__main__":
    main()