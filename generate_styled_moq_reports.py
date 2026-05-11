import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

NAVY = '#1a365d'
DARK_BLUE = '#2b6cb0'
ACCENT_BLUE = '#ebf8ff'
TEAL = '#2c7a7b'
TEAL_LIGHT = '#e6fffa'
ACCENT_GREEN = '#2f855a'
ORANGE = '#c05621'
RED = '#c53030'
GREY_HEADER = '#718096'
GREY_TEXT = '#4a5568'
GREY_LIGHT = '#a0aec0'
INFO_BG = '#f7fafc'
ROW_ALT = '#f8fafc'
CAT_HEADER_BG = '#2d3748'

def hex_to_rgb(hex_color):
    h = hex_color.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def set_cell_shading(cell, hex_color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), hex_color.lstrip('#'))
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_row_shading(row, hex_color):
    for cell in row.cells:
        set_cell_shading(cell, hex_color)

def generate_beautiful_report(title, json_file, out_file):
    with open(json_file, 'r') as f:
        data = json.load(f)
        
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # TITLE
    doc_title = doc.add_paragraph()
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_title.add_run(f'{title} — MOQ UPDATE & DEMAND ANALYSIS')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*hex_to_rgb(NAVY))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f'Period: Jan-Apr 2026 | Working Days: {data.get("working_days", 96)} | Formula: {data.get("formula", "")}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(*hex_to_rgb(GREY_TEXT))

    # KPI SUMMARY
    summary = data.get('summary', {})
    heading = doc.add_paragraph()
    run = heading.add_run('SUMMARY')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*hex_to_rgb(NAVY))

    kpi_table = doc.add_table(rows=1, cols=6)
    kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpi_headers = ['Total Items', 'Fast Movers', 'Slow Movers', 'Total Qty Out', 'Weekly MOQ', 'Weekly Value']
    kpi_values = [
        f"{summary.get('total_items', 0):,}", f"{summary.get('fast_movers', 0):,}", f"{summary.get('slow_movers', 0):,}",
        f"{summary.get('total_qty_out', 0):,.0f}", f"{summary.get('total_weekly_moq', 0):,.0f}",
        f"KES {summary.get('total_weekly_value', 0):,.0f}"
    ]

    for i, h in enumerate(kpi_headers):
        cell = kpi_table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, TEAL)

    val_row = kpi_table.add_row()
    for i, v in enumerate(kpi_values):
        cell = val_row.cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(v)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*hex_to_rgb(DARK_BLUE))
        set_cell_shading(cell, TEAL_LIGHT)

    doc.add_paragraph()

    # EXECUTIVE SUMMARY
    sorted_cats = sorted(data['categories'].items(), key=lambda x: x[1]['total_qty_out'], reverse=True)
    top_cat_by_vol = sorted_cats[0][0] if sorted_cats else "N/A"
    sorted_by_val = sorted(data['categories'].items(), key=lambda x: x[1].get('total_weekly_value', 0), reverse=True)
    top_cat_by_val = sorted_by_val[0][0] if sorted_by_val else "N/A"
    
    total_items = summary.get('total_items', 0)
    fast_movers = summary.get('fast_movers', 0)
    slow_movers = summary.get('slow_movers', 0)
    fast_pct = (fast_movers / total_items * 100) if total_items else 0
    weekly_val = summary.get('total_weekly_value', 0)

    exec_heading = doc.add_paragraph()
    run = exec_heading.add_run('EXECUTIVE SUMMARY & BUSINESS INSIGHTS')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*hex_to_rgb(NAVY))

    exec_body = doc.add_paragraph()
    exec_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    insights = [
        f"• Volume & Value: The branch recorded a total movement of {summary.get('total_qty_out', 0):,.0f} units during the Jan-Apr period. To sustain this demand, a weekly procurement value of KES {weekly_val:,.0f} is required.",
        f"• Inventory Velocity: Out of {total_items:,} items analyzed, {fast_movers:,} ({fast_pct:.1f}%) are classified as fast movers driving the majority of the turnover, while {slow_movers:,} items are moving slowly. Procurement should prioritize the fast-moving stock to optimize working capital.",
        f"• Category Concentration: The highest volume category is {top_cat_by_vol}, making it the primary driver of physical stock movement. Meanwhile, {top_cat_by_val} represents the most capital-intensive category based on weekly MOQ value requirements.",
        f"• Strategic Action: Focus immediate replenishment on the Top 50 Fast Moving items. Consider targeted promotions or price adjustments to clear capital tied up in the {slow_movers:,} slow-moving inventory items."
    ]

    for insight in insights:
        run = exec_body.add_run(insight + '\n')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*hex_to_rgb(GREY_TEXT))

    doc.add_paragraph()

    # CATEGORY BREAKDOWN
    heading = doc.add_paragraph()
    run = heading.add_run('MOQ BY CATEGORY')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*hex_to_rgb(NAVY))

    cat_table = doc.add_table(rows=1, cols=7)
    cat_headers = ['Category', 'Items', 'Fast', 'Slow', 'Qty Out', 'Wkly MOQ', 'Wkly Value (KES)']
    for i, h in enumerate(cat_headers):
        cell = cat_table.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        if i >= 1: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_shading(cell, DARK_BLUE)

    sorted_cats = sorted(data['categories'].items(), key=lambda x: x[1]['total_qty_out'], reverse=True)
    for idx, (cat_name, cat_info) in enumerate(sorted_cats):
        if cat_info['total_qty_out'] > 0 or cat_info['total_items'] > 0:
            row = cat_table.add_row()
            vals = [
                cat_name[:25],
                str(cat_info['total_items']), str(cat_info['fast_movers']), str(cat_info['slow_movers']),
                f"{cat_info['total_qty_out']:,.0f}", f"{cat_info['total_weekly_moq']:,.0f}",
                f"{cat_info.get('total_weekly_value', 0):,.0f}"
            ]
            for i, v in enumerate(vals):
                cell = row.cells[i]
                p = cell.paragraphs[0]
                run = p.add_run(v)
                run.font.size = Pt(7.5)
                if i >= 1: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if idx % 2 == 1:
                set_row_shading(row, ROW_ALT)

    doc.add_page_break()

    # ALL ITEMS (Top 50 Fast Movers + High MOQs)
    heading = doc.add_paragraph()
    run = heading.add_run('TOP 50 FAST MOVING INVENTORY ITEMS')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*hex_to_rgb(NAVY))
    
    # Collect all products and sort by qty out over Jan-Apr
    all_prods = []
    for cat_info in data['categories'].values():
        all_prods.extend(cat_info['products'])
    all_prods.sort(key=lambda x: x['qty_out'], reverse=True)

    fast_table = doc.add_table(rows=1, cols=7)
    fast_headers = ['#', 'Code', 'Description', 'Category', 'Qty Out (Jan-Apr)', 'Avg Daily', 'New Wkly MOQ']
    for i, h in enumerate(fast_headers):
        cell = fast_table.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        if i >= 4: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_shading(cell, ORANGE)

    for idx, item in enumerate(all_prods[:50], 1):
        row = fast_table.add_row()
        vals = [
            str(idx), str(item['item_code']), str(item['item_description'])[:38],
            item.get('category', 'UNKNOWN')[:16],
            f"{item['qty_out']:,.0f}", f"{item['daily_average']:.1f}", f"{item['weekly_moq']:,}"
        ]
        for i, v in enumerate(vals):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            run = p.add_run(v)
            run.font.size = Pt(7)
            if i >= 4: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if idx % 2 == 0:
            set_row_shading(row, ROW_ALT)

    doc.add_page_break()
    
    # DETAILED CATEGORY BREAKDOWN
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run('DETAILED MOQ BY CATEGORY (TOP ITEMS)')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*hex_to_rgb(NAVY))

    for cat_name, cat_info in sorted_cats:
        if cat_info['total_qty_out'] == 0: continue

        total_v = cat_info.get('total_weekly_value', 0)
        value_str = f'  |  Wkly Value: KES {total_v:,.0f}'

        cat_header = doc.add_paragraph()
        run = cat_header.add_run(
            f"{cat_name}  |  {cat_info['total_items']} Items  |  "
            f"Fast: {cat_info['fast_movers']}  |  Slow: {cat_info['slow_movers']}  |  "
            f"Wkly MOQ: {cat_info['total_weekly_moq']:,.0f}{value_str}"
        )
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), NAVY.lstrip('#'))
        shading.set(qn('w:val'), 'clear')
        cat_header._p.get_or_add_pPr().append(shading)

        products = sorted([p for p in cat_info['products'] if p['qty_out'] > 0], key=lambda x: x['qty_out'], reverse=True)[:5]
        if not products: continue

        prod_table = doc.add_table(rows=1, cols=6)
        prod_headers = ['Code', 'Description', 'Qty Out (Jan-Apr)', 'Avg Daily', 'Buying Prc', 'Wkly MOQ']
        for i, h in enumerate(prod_headers):
            cell = prod_table.rows[0].cells[i]
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.font.size = Pt(7.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            if i >= 2: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_cell_shading(cell, GREY_HEADER)

        for idx, p_item in enumerate(products):
            row = prod_table.add_row()
            vals = [
                str(p_item['item_code']), p_item['item_description'][:35],
                f"{p_item['qty_out']:,.0f}", f"{p_item['daily_average']:.1f}",
                f"{p_item['buying_price']:.1f}", f"{p_item['weekly_moq']:,.1f}"
            ]
            for i, v in enumerate(vals):
                cell = row.cells[i]
                pp = cell.paragraphs[0]
                run = pp.add_run(v)
                run.font.size = Pt(7)
                if i >= 2: pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if idx % 2 == 1:
                set_row_shading(row, ROW_ALT)
        doc.add_paragraph()

    doc.save(out_file)
    print(f"Generated beautifully styled report: {out_file}")

if __name__ == '__main__':
    try:
        generate_beautiful_report('BOMAS BRANCH', 'bomas_updated_moqs.json', 'Bomas_Styled_MOQ_Jan_Apr.docx')
        generate_beautiful_report('KAREN BRANCH', 'karen_updated_moqs.json', 'Karen_Styled_MOQ_Jan_Apr.docx')
    except PermissionError as e:
        print(f"\n[ERROR] Could not save the document. Please ensure that the file '{e.filename}' is CLOSED in Microsoft Word and try again.\n")
