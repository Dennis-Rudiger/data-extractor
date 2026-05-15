import json
import pandas as pd
import matplotlib.pyplot as plt
import os
from openpyxl.styles import Font, PatternFill
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_reports():
    with open('customer_sales_karen.json', 'r') as f:
        data = json.load(f)

    df = pd.DataFrame(data)

    # 1. Create Excel Report
    writer = pd.ExcelWriter('Customer_Sales_Report_Karen.xlsx', engine='openpyxl')
    
    monthly_df = df[df['Month'] != 'YTD'].copy()
    ytd_df = df[df['Month'] == 'YTD'].copy()

    # pivot table by customer and month
    if not monthly_df.empty:
        pivot = pd.pivot_table(monthly_df, values='Amount_Exc', index='Customer', columns='Month', aggfunc='sum', fill_value=0)
        month_order = [m for m in ['January', 'February', 'March', 'April'] if m in pivot.columns]
        pivot = pivot[month_order]
        pivot['Total_Monthly_Sum'] = pivot.sum(axis=1)
        pivot = pivot.sort_values('Total_Monthly_Sum', ascending=False)
        pivot.to_excel(writer, sheet_name='MoM Sales')
    
    if not ytd_df.empty:
        ytd_summary = ytd_df.groupby('Customer')[['Amount_Exc', 'Cost', 'Profit']].sum().reset_index()
        ytd_summary['Profit_Margin_%'] = (ytd_summary['Profit'] / ytd_summary['Amount_Exc']) * 100
        ytd_summary = ytd_summary.sort_values('Amount_Exc', ascending=False)
        ytd_summary.to_excel(writer, sheet_name='YTD Totals', index=False)

    writer.close()
    print('Created Customer_Sales_Report_Karen.xlsx')

    # 2. GENERATE CHARTS
    monthly_trend_path = 'karen_monthly_sales_trend.png'
    if not monthly_df.empty:
        monthly_trend = monthly_df.groupby('Month')['Amount_Exc'].sum().reindex(['January', 'February', 'March', 'April']).fillna(0)
        plt.figure(figsize=(8, 4))
        plt.plot(monthly_trend.index, monthly_trend.values, marker='o', linestyle='-', color='#1f77b4', linewidth=2)
        plt.title('Monthly Sales Trend (Karen Branch)', fontsize=14)
        plt.ylabel('Sales (KES)', fontsize=12)
        plt.grid(True, linestyle='--', alpha=0.7)
        plt.tight_layout()
        plt.savefig(monthly_trend_path)
        plt.close()

    top_10_path = 'karen_top_10_customers.png'
    if not ytd_df.empty:
        top_10 = ytd_summary.head(10)
        plt.figure(figsize=(8, 5))
        plt.barh(top_10['Customer'], top_10['Amount_Exc'], color='#2ca02c')
        plt.title('Top 10 Customers by YTD Sales (Karen Branch)', fontsize=14)
        plt.xlabel('Sales (KES)', fontsize=12)
        plt.gca().invert_yaxis()
        plt.tight_layout()
        plt.savefig(top_10_path)
        plt.close()

    # 3. Create Word Summary Report
    doc = Document()
    head = doc.add_heading('KAREN BRANCH - CUSTOMER SALES PERFORMANCE REPORT', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run('Period: January - April 2026\n').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('1. Executive Overview', level=1)
    total_ytd_sales = ytd_df['Amount_Exc'].sum() if not ytd_df.empty else monthly_df['Amount_Exc'].sum()
    total_ytd_profit = ytd_df['Profit'].sum() if not ytd_df.empty else monthly_df['Profit'].sum()
    margin = (total_ytd_profit / total_ytd_sales * 100) if total_ytd_sales > 0 else 0
    
    doc.add_paragraph(f"This report consolidates customer sales performance for the Karen Branch from January to April 2026. The total sales generated across all tracked customers was KES {total_ytd_sales:,.2f}, delivering a gross profit of KES {total_ytd_profit:,.2f} (Overall Margin: {margin:.2f}%).")
    
    if os.path.exists(monthly_trend_path):
        doc.add_picture(monthly_trend_path, width=Inches(6.0))

    if not ytd_df.empty:
        doc.add_heading('2. Top 10 Customers by YTD Volume', level=1)
        if os.path.exists(top_10_path):
            doc.add_picture(top_10_path, width=Inches(6.0))

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Customer Name'
        hdr[1].text = 'Sales (KES)'
        hdr[2].text = 'Profit (KES)'
        hdr[3].text = 'Margin (%)'

        for _, row in top_10.iterrows():
            cells = table.add_row().cells
            cells[0].text = row['Customer']
            cells[1].text = f"{row['Amount_Exc']:,.2f}"
            cells[2].text = f"{row['Profit']:,.2f}"
            cells[3].text = f"{row['Profit_Margin_%']:.1f}%"

    doc.add_heading('3. Monthly Performance Breakdown', level=1)
    if not monthly_df.empty:
        for mth, val in monthly_trend.items():
            if val > 0:
                doc.add_paragraph(f"{mth}: KES {val:,.2f}", style='List Bullet')

    doc.save('Customer_Sales_Summary_Karen.docx')
    print('Created Customer_Sales_Summary_Karen.docx')

if __name__ == '__main__':
    build_reports()
